"""Reviewer feedback #6: "Section 1.6 should be the thesis structure."

Current Chapter I (auto-numbered Heading-2, numId=8) ends with two
literature-review sections that do NOT describe the thesis:

  1.5 Paper Organization  -> a roadmap of *only* the literature review (2.1->2.9)
  1.6 Survey Methodology  -> the PRISMA literature-search method + Fig 1.1

There is no section describing how the thesis (Ch I-V) is organized.

This script (decisions confirmed with the author):
  1. MOVE both lit-review sections' bodies + the PRISMA figure into the
     Chapter II intro (under the empty "Literature Review" lead heading).
  2. RENUMBER the PRISMA figure 1.1 -> 2.1 (Fig 2.1 is free; Ch II figures
     are 2.2 and 2.7) and fix its in-text callout.
  3. FIX the stray "Section 1.6 describes the survey methodology..."
     cross-reference (now broken) and retitle the roadmap's opening from
     "The paper is organized" -> "This review is organized" for its new home.
  4. DELETE the now-orphaned "Paper Organization" / "Survey Methodology"
     Heading-2 paragraphs (and their blank spacers) from Chapter I.
  5. INSERT a fresh "Thesis Structure" Heading-2 (1.5) at the end of
     Chapter I with a clearly-marked PLACEHOLDER body (author writes the
     prose from the Phase-3 points; a follow-up script replaces it).

All look-ups are by visible text, never by paragraph index. TOC / List of
Figures are live Word fields (TOC \\o "1-3") -> NOT edited here; refresh in
OnlyOffice (select all, F9 / Update Fields) after opening.
"""

from __future__ import annotations

import copy
import sys

from docx import Document
from docx.oxml.ns import qn

DOCX = "docs/research/THESIS.docx"

PLACEHOLDER = "[[ THESIS STRUCTURE — TO BE WRITTEN (see Phase-3 points) ]]"


def norm(s: str) -> str:
    return " ".join(s.split()).strip()


def wt_nodes(p_el):
    return p_el.findall(".//" + qn("w:t"))


def set_flat_text(p_el, text: str) -> None:
    """Set paragraph text by writing it all into the first <w:t> and clearing
    the rest. Safe only for uniformly-formatted runs (verified beforehand)."""
    ts = wt_nodes(p_el)
    if not ts:
        raise SystemExit("no <w:t> in paragraph to set text")
    ts[0].text = text
    # keep xml:space so leading/trailing spaces survive
    ts[0].set(qn("xml:space"), "preserve")
    for t in ts[1:]:
        t.text = ""


def strip_bookmarks(p_el) -> None:
    """Remove TOC bookmark anchors from a cloned heading so we don't duplicate
    _Toc names; the TOC field rebuilds them on refresh."""
    for tag in ("w:bookmarkStart", "w:bookmarkEnd"):
        for bm in p_el.findall(".//" + qn(tag)):
            bm.getparent().remove(bm)


def main() -> None:
    d = Document(DOCX)
    ps = list(d.paragraphs)

    def find_idx(pred, what):
        for i, p in enumerate(ps):
            if pred(p):
                return i
        raise SystemExit(f"anchor not found: {what}")

    def is_h2(p, text):
        return p.style.name == "Heading 2" and norm(p.text) == text

    # --- capture anchors (by text) ----------------------------------------
    po_h = find_idx(lambda p: is_h2(p, "Paper Organization"), "Paper Organization H2")
    sm_h = find_idx(lambda p: is_h2(p, "Survey Methodology"), "Survey Methodology H2")
    contrib_h = find_idx(lambda p: is_h2(p, "Contributions"), "Contributions H2")
    litrev_h = find_idx(lambda p: is_h2(p, "Literature Review"), "Literature Review H2")

    # roadmap body = first non-empty para after Paper Organization heading
    roadmap_i = find_idx(
        lambda p: norm(p.text).startswith("The paper is organized into different sections"),
        "roadmap body",
    )
    survey_i = find_idx(
        lambda p: norm(p.text).startswith("To conduct this survey"),
        "survey methodology body",
    )
    cap_i = find_idx(
        lambda p: norm(p.text).startswith("Figure 1.1:"),
        "PRISMA caption",
    )
    # PRISMA image = the paragraph carrying a <w:drawing> between survey body and caption
    img_i = None
    for j in range(survey_i + 1, cap_i):
        if ps[j]._p.findall(".//" + qn("w:drawing")):
            img_i = j
            break
    if img_i is None:
        raise SystemExit("PRISMA image paragraph (w:drawing) not found")

    # blank spacers that belong to the deleted headings
    po_blank_i = po_h + 1
    sm_blank_i = sm_h + 1
    assert not norm(ps[po_blank_i].text), "expected blank spacer after Paper Organization"
    assert not norm(ps[sm_blank_i].text), "expected blank spacer after Survey Methodology"

    # Chapter II insertion anchor = blank spacer right after "Literature Review"
    litrev_blank_i = litrev_h + 1
    assert not norm(ps[litrev_blank_i].text), "expected blank spacer after Literature Review"

    # Contributions: heading / blank / body (templates + insertion anchor)
    contrib_blank_i = contrib_h + 1
    contrib_body_i = contrib_h + 2
    assert not norm(ps[contrib_blank_i].text), "expected blank spacer after Contributions"
    assert norm(ps[contrib_body_i].text).startswith("This research aims to make contributions")

    # grab element refs (immune to list staleness after mutation)
    roadmap_el = ps[roadmap_i]._p
    survey_el = ps[survey_i]._p
    img_el = ps[img_i]._p
    cap_el = ps[cap_i]._p
    po_h_el = ps[po_h]._p
    po_blank_el = ps[po_blank_i]._p
    sm_h_el = ps[sm_h]._p
    sm_blank_el = ps[sm_blank_i]._p
    litrev_blank_el = ps[litrev_blank_i]._p
    contrib_h_el = ps[contrib_h]._p
    contrib_blank_el = ps[contrib_blank_i]._p
    contrib_body_el = ps[contrib_body_i]._p

    log = []

    # --- 1. text edits on moved paragraphs (uniform formatting verified) ----
    # caption: Figure 1.1 -> Figure 2.1
    cap_text = "".join(t.text or "" for t in wt_nodes(cap_el))
    new_cap = cap_text.replace("Figure 1.1", "Figure 2.1", 1)
    set_flat_text(cap_el, new_cap)
    log.append(f"caption: {cap_text[:40]!r} -> {new_cap[:40]!r}")

    # survey body: (Figure 1.1) -> (Figure 2.1)
    survey_text = "".join(t.text or "" for t in wt_nodes(survey_el))
    new_survey = survey_text.replace("Figure 1.1", "Figure 2.1")
    set_flat_text(survey_el, new_survey)
    _hits = survey_text.count("Figure 1.1")
    log.append(f"survey callout: 'Figure 1.1' -> 'Figure 2.1' ({_hits} hit)")

    # roadmap: retitle opening + drop the broken 'Section 1.6' sentence
    rt = "".join(t.text or "" for t in wt_nodes(roadmap_el))
    rt2 = rt.replace(
        "The paper is organized into different sections",
        "This review is organized into different sections",
        1,
    )
    stray = (
        " Section 1.6 describes the survey methodology employed for "
        "literature identification and selection."
    )
    if stray in rt2:
        rt2 = rt2.replace(stray, "")
        log.append("roadmap: removed broken 'Section 1.6 describes...' cross-ref")
    else:
        # tolerate spacing variants
        import re

        rt2b = re.sub(
            r"\s*Section 1\.6 describes the survey methodology employed for "
            r"literature identification and selection\.",
            "",
            rt2,
        )
        if rt2b != rt2:
            rt2 = rt2b
            log.append("roadmap: removed broken 'Section 1.6 describes...' cross-ref (regex)")
        else:
            log.append("WARNING: stray 'Section 1.6' sentence not found in roadmap")
    set_flat_text(roadmap_el, rt2)
    log.append("roadmap: 'The paper is organized' -> 'This review is organized'")

    # --- 2. MOVE roadmap/survey/image/caption into Ch II intro -------------
    # insert after the blank spacer following 'Literature Review', in order
    litrev_blank_el.addnext(cap_el)
    litrev_blank_el.addnext(img_el)
    litrev_blank_el.addnext(survey_el)
    litrev_blank_el.addnext(roadmap_el)
    log.append("moved roadmap + survey + PRISMA image + caption -> Chapter II intro")

    # --- 3. DELETE orphaned Ch I headings + their blank spacers ------------
    for el, name in (
        (po_h_el, "Paper Organization heading"),
        (po_blank_el, "Paper Organization spacer"),
        (sm_h_el, "Survey Methodology heading"),
        (sm_blank_el, "Survey Methodology spacer"),
    ):
        el.getparent().remove(el)
        log.append(f"deleted: {name}")

    # --- 4. INSERT fresh 'Thesis Structure' section at end of Ch I ----------
    h2 = copy.deepcopy(contrib_h_el)
    strip_bookmarks(h2)
    set_flat_text(h2, "Thesis Structure")

    blank = copy.deepcopy(contrib_blank_el)
    strip_bookmarks(blank)

    body = copy.deepcopy(contrib_body_el)
    strip_bookmarks(body)
    set_flat_text(body, PLACEHOLDER)

    contrib_body_el.addnext(body)
    contrib_body_el.addnext(blank)
    contrib_body_el.addnext(h2)
    log.append("inserted: 'Thesis Structure' Heading-2 (1.5) + placeholder body")

    d.save(DOCX)

    print("RESTRUCTURE SUMMARY")
    for line in log:
        print("  -", line)

    # --- verification ------------------------------------------------------
    d2 = Document(DOCX)
    ps2 = list(d2.paragraphs)
    print("\nVERIFY  Chapter I numId=8 headings (should be 1.1-1.5):")
    for p in ps2:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            continue
        n = numPr.find(qn("w:numId"))
        if n is not None and n.get(qn("w:val")) == "8":
            print("   *", norm(p.text)[:45])

    print("\nVERIFY  Chapter II head (Literature Review -> intro -> 2.1):")
    started = False
    shown = 0
    for p in ps2:
        if norm(p.text) == "Literature Review":
            started = True
        if started:
            txt = norm(p.text)
            tag = "[img]" if p._p.findall(".//" + qn("w:drawing")) else ""
            print(f"   ({p.style.name:9}) {tag}{txt[:60]!r}")
            shown += 1
            if norm(p.text).startswith("2.1 Foundation") or shown > 8:
                break

    # leftover Figure 1.1 check
    leftover = [norm(p.text)[:50] for p in ps2 if "Figure 1.1" in p.text]
    print(f"\nVERIFY  remaining 'Figure 1.1' mentions: {leftover if leftover else 'none'}")
    # Survey/Paper Org should no longer be Ch I headings
    bad = [norm(p.text) for p in ps2 if p.style.name == "Heading 2" and norm(p.text) in
           ("Paper Organization", "Survey Methodology")]
    print(f"VERIFY  orphan Ch I headings remaining: {bad if bad else 'none'}")


if __name__ == "__main__":
    sys.exit(main())
