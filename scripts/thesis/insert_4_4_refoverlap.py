"""Insert §4.4 'Reference-Overlap Metrics' into THESIS.docx and renumber the
sections it displaces.

Order of operations (single save):
  1. Renumber the displaced sections by TEXT (so paragraph reindexing is moot):
       - Ablation  §4.4 -> §4.5   (H2 + 4.4.1-4.4.4 H3 headings)
       - Synthesis §4.5 -> §4.6   (H2 + 4.5.1-4.5.3 H3 headings)
       - Ablation figure captions  Figure 4.4.1/4.4.2 -> 4.5.1/4.5.2
         (visible <w:t> + TC instrText + Fig_4_4_x -> Fig_4_5_x bookmark)
       - In-text callouts  (Figure 4.4.1)/(Figure 4.4.2) -> 4.5.1/4.5.2
  2. Insert the new §4.4 (H2 + three H3s + six body paragraphs) right after the
     MAUVE Figure 4.3.2 caption.

Figures are embedded afterwards by scripts/thesis/insert_figure.py (TC mode),
anchored on the prose paragraphs this script writes.

Renumber guards: heading targets must be Heading-styled (skips the live TOC
cache); caption/callout targets must be style-943 body paragraphs (skips the
live List-of-Figures cache). The whole run is a no-op if §4.4 already exists.
"""
from __future__ import annotations

from copy import deepcopy

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOC = "docs/research/THESIS.docx"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

# --- §4.4 prose (author-written, Phase 5/6 polished) -----------------------

SETUP = (
    "The third evaluation arm represents the reference-overlap metrics that measure the "
    "overlap of configuration inferences compared to the golden reference. This evaluation "
    "arm assumes the golden reference pool represented by our dataset is the only winning ad "
    "from that brief. Despite this limitation, these metrics will provide insight into model "
    "performance on how their output style overlaps with the real-world market-validated ad "
    "set. We use five metrics: BLEU, chrF, ROUGE-L, METEOR, and BERTScore. Our setup involves "
    "comparing any inference to two references, including the single gold winning ad as well "
    "as a pool of five nearest winning ads from the same platform. This is to mitigate the "
    "effect of overpenalizing and bias for a single style. Similar to the other evaluation "
    "arms, we use the 215-brief test set with configs A, B, C, and GOLD."
)

RANK_HEADLINE = (
    "The headline result of the calculated metrics shows that C leads all five gold-reference "
    "metrics. C gold means — BLEU 0.138, ROUGE-L 0.342, METEOR 0.300, chrF 0.386, "
    "BERTScore 0.893; A gold means — BLEU 0.051, ROUGE-L 0.226, METEOR 0.261, chrF 0.356, "
    "BERTScore 0.868; B gold means — BLEU 0.035, ROUGE-L 0.189, METEOR 0.215, chrF 0.322, "
    "BERTScore 0.858; ranking C > A > B on every metric."
)

RANK_FIGREAD = (
    "C's confidence intervals clear A and B on BLEU, ROUGE-L, and METEOR, showing a clear "
    "separation between C and the other configurations. chrF and BERTScore scores are high for "
    "all three configurations, which means these metrics barely tell the models apart; despite "
    "that, C still ranks the highest of the three. The baseline GOLD scoring 1.0 against itself "
    "proves the validity of the evaluation pipeline. Moreover, comparing GOLD against the "
    "five-ad reference pool gives a multi-reference ceiling (BLEU 0.361, ROUGE-L 0.453, METEOR "
    "0.418, chrF 0.485, BERTScore 0.900), which represents the ceiling the other configurations "
    "should aim for (Figure 4.4.1)."
)

RANK_MEMO = (
    "To check whether configuration C really has a winning style or copies ads verbatim, we run "
    "a simple check comparing each generated ad's overlap with the one actual winner against its "
    "overlap with the other five similar winners. A small gap indicates the model has learned "
    "style, not memorized verbatim. C's gap is small (0.120); GOLD's is huge (0.547) because it "
    "literally is the ad — therefore the small gap indicates C writes like winners, it does "
    "not copy them."
)

GROUND_UPWORTHY = (
    "With these results, it is important to validate the accuracy of these metrics. We run a "
    "200-pair real A/B headline test on the Upworthy dataset (Matias et al., 2021) to see which "
    "metric actually manages to pick the right winner. The results show only METEOR beats chance "
    "for real (0.613, p = 0.007); chrF, ROUGE-L, and BLEU do not (Figure 4.4.2). This is a "
    "critical reason why the third evaluation arm is provided as supplementary evidence of the "
    "limitations of using traditional overlap metrics in creative marketing tasks."
)

GROUND_CONCL = (
    "We conclude that the results of this evaluation are consistent with the other two, where C "
    "ranked the highest: this arm ranks C > A > B, while the learned scorer ranked C > B > A. "
    "This consistency strengthens the evaluation validity of the other arms, but is not enough "
    "to decide by itself."
)

HEADING_MAP = {
    "4.4 Fine-Tuning and Agent Ablation": "4.5 Fine-Tuning and Agent Ablation",
    "4.4.1 2×2 Design": "4.5.1 2×2 Design",
    "4.4.2 Paired Contrasts": "4.5.2 Paired Contrasts",
    "4.4.3 Brief Shape": "4.5.3 Brief Shape",
    "4.4.4 Agent Scoring Caveats": "4.5.4 Agent Scoring Caveats",
    "4.5 Synthesis and Limitations": "4.6 Synthesis and Limitations",
    "4.5.1 Recap": "4.6.1 Recap",
    "4.5.2 Interpretation": "4.6.2 Interpretation",
    "4.5.3 Limitations": "4.6.3 Limitations",
}


def find_by_text(doc, *, equals=None, startswith=None, style_id=None, heading=None, last=False):
    hit = None
    for p in doc.paragraphs:
        txt = p.text.strip()
        if equals is not None and txt != equals:
            continue
        if startswith is not None and not txt.startswith(startswith):
            continue
        if style_id is not None and (not p.style or p.style.style_id != style_id):
            continue
        if heading is not None:
            is_h = bool(p.style and p.style.name.startswith("Heading"))
            if is_h != heading:
                continue
        if not last:
            return p
        hit = p
    if hit is None:
        raise RuntimeError(f"paragraph not found: equals={equals!r} startswith={startswith!r}")
    return hit


def clean_clone(tpl_p, text: str):
    """Clone a template paragraph, keep only its <w:pPr>, add one clean text run."""
    new = deepcopy(tpl_p._p)
    for child in list(new):
        if child.tag != qn("w:pPr"):
            new.remove(child)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(XML_SPACE, "preserve")
    t.text = text
    r.append(t)
    new.append(r)
    return new


def set_first_t(p_elem, text: str) -> None:
    ts = p_elem.findall(".//" + qn("w:t"))
    if not ts:
        raise RuntimeError("paragraph has no <w:t> to set")
    ts[0].text = text
    for t in ts[1:]:
        t.text = ""


def renumber_caption(p_elem, old_num, new_num, old_bm, new_bm) -> None:
    for t in p_elem.findall(".//" + qn("w:t")):
        if t.text and old_num in t.text:
            t.text = t.text.replace(old_num, new_num)
    for it in p_elem.findall(".//" + qn("w:instrText")):
        if it.text and old_num in it.text:
            it.text = it.text.replace(old_num, new_num)
    for bs in p_elem.findall(".//" + qn("w:bookmarkStart")):
        if bs.get(qn("w:name")) == old_bm:
            bs.set(qn("w:name"), new_bm)


def main() -> None:
    doc = docx.Document(DOC)

    if any(p.text.strip() == "4.4 Reference-Overlap Metrics" for p in doc.paragraphs):
        print("§4.4 already present — nothing to do.")
        return

    # --- templates (unaffected by renumber) ---
    h2_tpl = find_by_text(doc, equals="4.3 MAUVE Distribution Matching", heading=True)
    h3_tpl = find_by_text(doc, equals="4.3.1 Arm Setup", heading=True)
    body_tpl = find_by_text(doc, startswith="The second arm, which is the MAUVE arm")
    anchor = find_by_text(doc, startswith="Figure 4.3.2:", style_id="943", last=True)

    # --- 1. renumber headings ---
    n_head = 0
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith("Heading"):
            txt = p.text.strip()
            if txt in HEADING_MAP:
                set_first_t(p._p, HEADING_MAP[txt])
                n_head += 1
                print(f"  heading: {txt!r} -> {HEADING_MAP[txt]!r}")

    # --- renumber ablation captions + callouts (style-943 body only) ---
    n_cap = n_call = 0
    for p in doc.paragraphs:
        if not (p.style and p.style.style_id == "943"):
            continue
        txt = p.text.strip()
        if txt.startswith("Figure 4.4.1"):
            renumber_caption(p._p, "Figure 4.4.1", "Figure 4.5.1", "Fig_4_4_1", "Fig_4_5_1")
            n_cap += 1
            print("  caption: Figure 4.4.1 -> Figure 4.5.1")
        elif txt.startswith("Figure 4.4.2"):
            renumber_caption(p._p, "Figure 4.4.2", "Figure 4.5.2", "Fig_4_4_2", "Fig_4_5_2")
            n_cap += 1
            print("  caption: Figure 4.4.2 -> Figure 4.5.2")
        elif not txt.startswith("Figure"):
            for old, new in (("Figure 4.4.2", "Figure 4.5.2"), ("Figure 4.4.1", "Figure 4.5.1")):
                if old in txt:
                    for t in p._p.findall(".//" + qn("w:t")):
                        if t.text and old in t.text:
                            t.text = t.text.replace(old, new)
                    n_call += 1
                    print(f"  callout: ({old}) -> ({new})")

    # --- 2. insert new §4.4 after the MAUVE 4.3.2 caption ---
    new_paras = [
        clean_clone(h2_tpl, "4.4 Reference-Overlap Metrics"),
        clean_clone(h3_tpl, "4.4.1 Arm Setup"),
        clean_clone(body_tpl, SETUP),
        clean_clone(h3_tpl, "4.4.2 Overlap Rankings"),
        clean_clone(body_tpl, RANK_HEADLINE),
        clean_clone(body_tpl, RANK_FIGREAD),
        clean_clone(body_tpl, RANK_MEMO),
        clean_clone(h3_tpl, "4.4.3 Grounding Against Real Outcomes"),
        clean_clone(body_tpl, GROUND_UPWORTHY),
        clean_clone(body_tpl, GROUND_CONCL),
    ]
    prev = anchor._p
    for np in new_paras:
        prev.addnext(np)
        prev = np

    doc.save(DOC)
    print(
        f"\nDONE: renumbered {n_head} headings, {n_cap} captions, {n_call} callouts; "
        f"inserted {len(new_paras)} new §4.4 paragraphs."
    )


if __name__ == "__main__":
    main()
