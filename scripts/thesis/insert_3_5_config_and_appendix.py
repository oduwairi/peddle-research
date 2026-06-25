"""Insert the Ch. III fine-tuning config table + prompt-template appendix.

Addresses professor feedback #3 (weak methodology description, sec 3.5):
  - Table 3.1 "Fine-tuning configuration" at the end of sec 3.5 (renumbers the
    existing eval-methods table 3.1 -> 3.2, body caption + List-of-Tables).
  - sec 3.5.4 numeric fills: peak LR 2x10^-4, warmup ratio 0.03, training time.
  - Appendix A "Prompt Templates" (A.1 system prompt, A.2 teacher prompt,
    A.3 brief registers, A.4 representative example), imported verbatim from
    source so they cannot drift.
  - In-text pointers from sec 3.4.3 / sec 3.4.5 to the appendix.

Index-free: every anchor is located by visible text. Idempotent: re-running
skips blocks already present. Whole-file python-docx save -- do not run while
another writer has THESIS.docx open.
"""

from __future__ import annotations

import copy

import docx
from docx.oxml.ns import qn
from docx.shared import Pt

# --- verbatim prompt constants, imported from the live source ---------------
from draper.construction.bundle import BACKTRANSLATION_STYLE_RULES
from draper.construction.formats.copywriting.constructor import CopywritingConstructor
from draper.construction.formats.copywriting.dice import (
    ConversationRegister,
    _REGISTER_BODIES,
)

DOC = "docs/research/THESIS.docx"
SYSTEM_PROMPT = CopywritingConstructor.SYSTEM_PROMPT

CONFIG_ROWS: list[tuple[str, str]] = [
    ("Base model", "Qwen3-8B (unsloth/Qwen3-8B), 4-bit"),
    ("Quantization", "4-bit NF4 (QLoRA)"),
    ("Max sequence length", "4,096 tokens"),
    ("Compute precision", "bfloat16"),
    ("LoRA rank (r)", "32"),
    ("LoRA alpha (α)", "64"),
    ("LoRA dropout", "0.05"),
    ("Target modules", "q, k, v, o, gate, up, down projections"),
    ("rsLoRA / DoRA", "both enabled"),
    ("Trainable parameters", "≈48 M (≈0.58% of 8.2 B)"),
    ("Optimizer", "8-bit AdamW"),
    ("Peak learning rate", "2×10⁻⁴"),
    ("LR schedule", "cosine decay"),
    ("Warmup ratio", "0.03"),
    ("Weight decay", "0.0"),
    ("Epochs", "3"),
    ("Batch size (per-device × accumulation)", "2 × 8 = 16 effective"),
    ("Loss masking", "assistant-only"),
    ("Random seed", "42"),
    ("Evaluation / checkpoint interval", "every 50 steps"),
    ("Model-selection metric", "validation loss (lower is better)"),
    ("Early-stopping patience", "2"),
    ("Dataset (train / val / test)", "2,442 / 215 / 215"),
    ("Training hardware", "single cloud RTX 4090 (24 GB)"),
    ("Training wall-clock", "≈1.5–3 hours"),
    ("Compute cost", "≈US$1.50"),
    ("Merge / serving", "bfloat16 merge → Modal vLLM"),
]

log: list[str] = []


def find_para(d: docx.Document, pred):
    for p in d.paragraphs:
        if pred(p):
            return p
    return None


def by_contains(d, sub, style=None):
    return find_para(
        d, lambda p: sub in p.text and (style is None or p.style.name == style)
    )


def by_eq(d, s, style=None):
    return find_para(
        d, lambda p: p.text.strip() == s and (style is None or p.style.name == style)
    )


def replace_first_run(p, old, new):
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return True
    full = "".join(r.text for r in p.runs)
    if old in full and p.runs:
        p.runs[0].text = full.replace(old, new)
        for r in p.runs[1:]:
            r.text = ""
        return True
    raise SystemExit(f"ANCHOR NOT FOUND for replace: {old!r}")


def set_element_text(el, new_text):
    """Set first <w:t> of an element to new_text, clear the rest (preserves
    pPr/tab stops; leader dots come from the paragraph tab stop, not the run)."""
    ts = el.findall(".//" + qn("w:t"))
    if not ts:
        raise SystemExit("no <w:t> in cloned element")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def mono(run):
    run.font.size = Pt(9)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), "Courier New")


def add_heading(d, text, style_obj):
    p = d.add_paragraph()
    p.add_run(text)
    p.style = style_obj
    return p


def add_code_block(d, text, style_obj):
    for line in text.split("\n"):
        p = d.add_paragraph()
        p.style = style_obj
        mono(p.add_run(line))
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0


def add_label(d, text, style_obj):
    p = d.add_paragraph()
    p.style = style_obj
    r = p.add_run(text)
    r.bold = True
    p.paragraph_format.space_after = Pt(0)


def main() -> None:
    d = docx.Document(DOC)

    # ---- 1. sec 3.5.4 numeric fills -------------------------------------
    opt = by_contains(d, "cosine schedule (Loshchilov & Hutter, 2017)")
    if opt is None:
        raise SystemExit("sec 3.5.4 Optimization paragraph not found")
    if "peak learning rate" not in opt.text:
        replace_first_run(
            opt,
            "cosine schedule (Loshchilov & Hutter, 2017),",
            "cosine schedule (Loshchilov & Hutter, 2017) with a peak learning "
            "rate of 2×10⁻⁴,",
        )
        replace_first_run(
            opt,
            "first three percent of the training loop,",
            "first three percent of the training loop (a warmup ratio of 0.03),",
        )
        log.append("EDIT 3.5.4: inserted peak LR 2x10^-4 + warmup ratio 0.03")
    else:
        log.append("SKIP 3.5.4 LR/warmup (already present)")
    if "is estimated to take from one to two hours" in opt.text:
        replace_first_run(
            opt,
            "is estimated to take from one to two hours",
            "takes from one and a half to three hours",
        )
        log.append("EDIT 3.5.4: training time 1-2h -> 1.5-3h")
    else:
        log.append("SKIP 3.5.4 training time (already updated)")

    # ---- 2. Table 3.1 (config) + renumber existing -> 3.2 ----------------
    h36 = by_eq(d, "3.6 Agent Architecture", style="Heading 2")
    if h36 is None:
        raise SystemExit("'3.6 Agent Architecture' heading not found")

    eval_cap = by_contains(
        d, "Comparison of the two evaluation methods", style="Body Text"
    )  # body caption
    if eval_cap is None:
        raise SystemExit("eval-table body caption not found")

    if by_contains(d, "Table 3.1: Fine-tuning configuration") is None:
        # new caption: clone the eval body caption element for identical style
        cap_el = copy.deepcopy(eval_cap._p)
        set_element_text(cap_el, "Table 3.1: Fine-tuning configuration.")
        h36._p.addprevious(cap_el)

        # new table
        tbl = d.add_table(rows=len(CONFIG_ROWS) + 1, cols=2)
        eval_tbl = next(
            (t for t in d.tables if t.rows and t.rows[0].cells[0].text.strip() == "Aspect"),
            None,
        )
        try:
            tbl.style = eval_tbl.style if eval_tbl else d.styles["Table Grid"]
        except KeyError:
            tbl.style = d.styles["Table Grid"]
        hdr = tbl.rows[0].cells
        hdr[0].text, hdr[1].text = "Parameter", "Value"
        for c in hdr:
            for r in c.paragraphs[0].runs:
                r.bold = True
        for i, (k, v) in enumerate(CONFIG_ROWS, start=1):
            tbl.rows[i].cells[0].text = k
            tbl.rows[i].cells[1].text = v
        h36._p.addprevious(tbl._tbl)
        log.append(f"INSERT: Table 3.1 (config, {len(CONFIG_ROWS)} rows) before sec 3.6")
    else:
        log.append("SKIP Table 3.1 config (already present)")

    # renumber existing eval table body caption + ToF entry
    if "Table 3.1:" in eval_cap.text:
        replace_first_run(eval_cap, "Table 3.1", "Table 3.2")
        log.append("RENUMBER: body caption Table 3.1 -> 3.2")
    tof_eval = by_contains(
        d, "Comparison of the two evaluation methods", style="table of figures"
    )
    if tof_eval is not None and "Table 3.1:" in tof_eval.text:
        replace_first_run(tof_eval, "Table 3.1", "Table 3.2")
        log.append("RENUMBER: List-of-Tables entry Table 3.1 -> 3.2")
        tof_eval = by_contains(
            d, "Comparison of the two evaluation methods", style="table of figures"
        )

    # new ToF entry for Table 3.1 (clone the renumbered eval entry's styling)
    if (
        tof_eval is not None
        and by_contains(
            d, "Fine-tuning configuration", style="table of figures"
        )
        is None
    ):
        tof_el = copy.deepcopy(tof_eval._p)
        full = tof_eval.text  # e.g. "Table 3.2: Comparison...\t1"
        new_full = full.replace(
            "Table 3.2: Comparison of the two evaluation methods",
            "Table 3.1: Fine-tuning configuration",
        )
        set_element_text(tof_el, new_full)
        tof_eval._p.addprevious(tof_el)
        log.append("INSERT: List-of-Tables entry for Table 3.1")

    # ---- 3. In-text pointers (sec 3.4.3 teacher, sec 3.4.5 schema) -------
    teacher = by_contains(
        d, "the raw ads are sent to the teacher as clean, unlabeled text"
    )
    if teacher is not None and "Appendix A.2" not in teacher.text:
        teacher.add_run(
            " The full teacher prompt template is reproduced in Appendix A.2; "
            "the brief is written in one of three registers (conversational, "
            "structured, or imperative), reproduced in Appendix A.3."
        )
        log.append("POINTER: sec 3.4.3 -> Appendix A.2 / A.3 (register note)")

    schema = by_contains(d, "You are an ad copywriter")
    if schema is not None and "Appendix A" not in schema.text:
        schema.add_run(
            " The complete prompt templates and a representative training "
            "example are provided in Appendix A."
        )
        log.append("POINTER: sec 3.4.5 -> Appendix A")

    # ---- 4. Appendix A: Prompt Templates --------------------------------
    if by_eq(d, "Appendix A: Prompt Templates") is None:
        h2_style = h36.style
        h3_style = find_para(d, lambda p: p.style.name == "Heading 3").style
        normal_style = by_eq(d, "APPENDIX").style

        add_heading(d, "Appendix A: Prompt Templates", h2_style)

        add_heading(d, "A.1 Student System Prompt", h3_style)
        add_code_block(d, SYSTEM_PROMPT, normal_style)

        add_heading(d, "A.2 Teacher Backtranslation Prompt", h3_style)
        add_code_block(d, BACKTRANSLATION_STYLE_RULES, normal_style)

        add_heading(d, "A.3 Brief Registers", h3_style)
        for reg in (
            ConversationRegister.CONVERSATIONAL,
            ConversationRegister.STRUCTURED,
            ConversationRegister.IMPERATIVE,
        ):
            add_label(d, reg.value.capitalize(), normal_style)
            add_code_block(d, _REGISTER_BODIES[reg], normal_style)

        add_heading(d, "A.4 Representative Training Example", h3_style)
        from datasets import load_from_disk

        ex = load_from_disk("data/final")["test"][0]
        roles = {m["role"]: m["content"] for m in ex["messages"]}
        for role, label in (
            ("system", "System"),
            ("user", "User brief"),
            ("assistant", "Assistant response"),
        ):
            add_label(d, label, normal_style)
            add_code_block(d, roles.get(role, ""), normal_style)
        log.append("INSERT: Appendix A (A.1 system, A.2 teacher, A.3 registers, A.4 example)")
    else:
        log.append("SKIP Appendix A (already present)")

    # ---- save + summary --------------------------------------------------
    d.save(DOC)
    print("CHANGES:")
    for ln in log:
        print("  -", ln)


if __name__ == "__main__":
    main()
