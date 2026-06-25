"""Reviewer feedback #5 — Chapter III technical-depth / reproducibility pass.

Scoped, idempotent edits to docs/research/THESIS.docx:

  1. Table 3.1 (fine-tuning config): correct LoRA rank/alpha/param-count to the
     ACTUAL shipped+evaluated v1 run (#002: r=16, alpha=32, ~24M params), and
     add a Framework row. The table had drifted to a later config (r=32).
  2. Insert a new Table 3.2 "Scoring-predictor configuration" at the end of
     section 3.7 (mirrors Table 3.1's placement at the end of 3.5).
  3. Renumber the existing "Table 3.2: Comparison ..." -> 3.3 (visible run text,
     hidden TC field, bookmark, and the in-text callout).
  4. Add two loss-function display equations (3.5.3 SFT; 3.7.4 predictor).
  5. Add Appendix B "Algorithm Listings" with three pseudocode blocks, plus
     minimal parenthetical cross-ref callouts in 3.3 / 3.4 / 3.5.

Everything is located by visible text (never a hardcoded paragraph index) and
clones only from our own safe paragraphs (no <w:drawing> text-box headings).
Re-runnable: each step checks for its own marker and skips if already applied.
"""

from __future__ import annotations

import copy
import sys

from docx import Document
from docx.oxml.ns import qn

DOC = "docs/research/THESIS.docx"

W_T = qn("w:t")
W_R = qn("w:r")
W_P = qn("w:p")
W_TBL = qn("w:tbl")
W_TR = qn("w:tr")
W_TC = qn("w:tc")
W_PPR = qn("w:pPr")
W_RPR = qn("w:rPr")
W_JC = qn("w:jc")
W_IND = qn("w:ind")
W_INSTR = qn("w:instrText")
W_BMS = qn("w:bookmarkStart")
W_SECTPR = qn("w:sectPr")
XML_SPACE = qn("xml:space")


def para_text(p_el) -> str:
    return "".join(t.text or "" for t in p_el.iter(W_T)).strip()


def find_para(paras, *, equals=None, startswith=None, contains=None, skip_toc=True):
    """First paragraph element matching a visible-text predicate (skips TOC)."""
    for p in paras:
        t = p.text.strip()
        if skip_toc and "\t" in t:
            continue
        if equals is not None and t == equals:
            return p._element
        if startswith is not None and t.startswith(startswith):
            return p._element
        if contains is not None and contains in t:
            return p._element
    return None


def tbl_after(caption_el):
    """The <w:tbl> immediately following a caption paragraph."""
    sib = caption_el.getnext()
    while sib is not None and sib.tag != W_TBL:
        sib = sib.getnext()
    return sib


def rows(tbl):
    return tbl.findall(W_TR)


def row_left(tr) -> str:
    tcs = tr.findall(W_TC)
    return para_text(tcs[0]) if tcs else ""


def set_row_text(tr, left, right):
    """Set the two cells of a 2-col row, preserving the first run's rPr/font."""
    for tc, txt in zip(tr.findall(W_TC), [left, right]):
        ps = tc.findall(W_P)
        first_p = ps[0]
        # drop extra paragraphs in the cell
        for extra in ps[1:]:
            tc.remove(extra)
        runs = first_p.findall(W_R)
        done = False
        for r in runs:
            t = r.find(W_T)
            if t is not None and not done:
                t.text = txt
                t.set(XML_SPACE, "preserve")
                done = True
            elif t is not None:
                first_p.remove(r)
        if not done:
            # no run with text -> build one (clone rPr from any existing run)
            new_r = runs[0] if runs else first_p.makeelement(W_R, {})
            if not runs:
                first_p.append(new_r)
            t = new_r.makeelement(W_T, {XML_SPACE: "preserve"})
            t.text = txt
            new_r.append(t)


def clone_para(template_el, text, *, center=False, mono=False):
    """Clone a paragraph element, replace its text with a single run."""
    new = copy.deepcopy(template_el)
    # strip bookmarks / fields that came along
    for tag in (W_BMS, qn("w:bookmarkEnd")):
        for el in new.findall(tag):
            new.remove(el)
    # keep one run as the rPr donor
    runs = new.findall(W_R)
    donor_rpr = None
    for r in runs:
        rpr = r.find(W_RPR)
        if rpr is not None:
            donor_rpr = copy.deepcopy(rpr)
        new.remove(r)
    # also remove any field runs (instrText/fldChar) left behind
    for r in new.findall(W_R):
        new.remove(r)
    run = new.makeelement(W_R, {})
    if mono:
        rpr = run.makeelement(W_RPR, {})
        fonts = rpr.makeelement(qn("w:rFonts"), {})
        fonts.set(qn("w:ascii"), "Consolas")
        fonts.set(qn("w:hAnsi"), "Consolas")
        fonts.set(qn("w:cs"), "Consolas")
        rpr.append(fonts)
        sz = rpr.makeelement(qn("w:sz"), {})
        sz.set(qn("w:val"), "18")
        rpr.append(sz)
        run.append(rpr)
    elif donor_rpr is not None:
        run.append(donor_rpr)
    t = run.makeelement(W_T, {XML_SPACE: "preserve"})
    t.text = text
    run.append(t)
    new.append(run)
    # paragraph-level formatting
    ppr = new.find(W_PPR)
    if ppr is None:
        ppr = new.makeelement(W_PPR, {})
        new.insert(0, ppr)
    if center or mono:
        # left/centre align + kill first-line indent for equations / code
        for el in ppr.findall(W_IND):
            ppr.remove(el)
        jc = ppr.find(W_JC)
        if jc is None:
            jc = ppr.makeelement(W_JC, {})
            ppr.append(jc)
        jc.set(qn("w:val"), "center" if center else "left")
    return new


def append_callout(paras, anchor_text, callout):
    """Append ' (callout)' to the body paragraph whose text starts with anchor.

    Mirrors add_figure_callouts: a trailing parenthetical, idempotent.
    """
    # find the heading, then its first body paragraph
    p_list = list(paras)
    for i, p in enumerate(p_list):
        t = p.text.strip()
        if "\t" in t:
            continue
        if t.startswith(anchor_text):
            # walk forward to first non-empty, non-heading body paragraph
            for q in p_list[i + 1 :]:
                qt = q.text.strip()
                if not qt or "\t" in qt:
                    continue
                if q.style.name.startswith("Heading"):
                    return False  # no body paragraph before next heading
                if callout in qt:
                    return True  # already present
                runs = q._element.findall(W_R)
                # append a run carrying the parenthetical
                rpr_donor = None
                for r in runs:
                    if r.find(W_RPR) is not None:
                        rpr_donor = copy.deepcopy(r.find(W_RPR))
                new_r = q._element.makeelement(W_R, {})
                if rpr_donor is not None:
                    new_r.append(rpr_donor)
                tt = new_r.makeelement(W_T, {XML_SPACE: "preserve"})
                tt.text = f" ({callout})"
                new_r.append(tt)
                q._element.append(new_r)
                return True
    return False


# ---------------------------------------------------------------------------
PREDICTOR_ROWS = [
    ("Base model", "DeBERTa-v3-base (microsoft/deberta-v3-base), ≈140 M params"),
    ("Architecture", "mean-pooled encoder → linear head → sigmoid; 4 scalar outputs"),
    ("Output heads", "composite, survivability, engagement volume, engagement velocity"),
    ("Max sequence length", "256 tokens"),
    ("Dropout", "0.1"),
    ("Backbone learning rate", "2×10⁻⁵"),
    ("Head learning rate", "1×10⁻⁴ (5× backbone)"),
    ("Optimizer", "AdamW"),
    ("Weight decay", "0.01"),
    ("LR schedule", "linear with warmup"),
    ("Warmup ratio", "0.06"),
    ("Epochs", "4"),
    ("Batch size (per-device × accumulation)", "8 × 4 = 32 effective"),
    ("Compute precision", "bf16 autocast on fp32 weights"),
    ("Loss", "masked, quality-weighted MSE summed over heads (equal head weights)"),
    ("Sample weighting", "by training-quality rating"),
    ("Evaluation / checkpoint interval", "every 250 steps"),
    ("Model-selection metric", "Spearman ρ (composite head)"),
    ("Early-stopping patience", "3"),
    ("Random seed", "42"),
    (
        "Splits",
        "80 / 10 / 10 random (stratified by platform) + hold-out-platform + hold-out-vertical",
    ),
    ("Training corpus", "≈55k v3-scored ads"),
    ("Training hardware", "consumer RTX 3060 (6 GB)"),
    ("Framework", "Hugging Face Transformers Trainer"),
]

SFT_EQ = (
    "ℒ_SFT(θ) = − (1 / Σᵢ |Aᵢ|) · "
    "Σᵢ Σ_{t∈Aᵢ} log p_θ(y_t | y_<t)"
)
SFT_WHERE = (
    "where Aᵢ is the set of assistant-response token positions of training "
    "example i; system-prompt and user-brief tokens are masked and contribute "
    "no loss (assistant-only loss masking)."
)
PRED_EQ = (
    "ℒ(θ) = [ Σ_b Σ_h m_{b,h} · q_b · α_h · "
    "(ŷ_{b,h} − y_{b,h})² ] / [ Σ_b Σ_h m_{b,h} · q_b · α_h ]"
)
PRED_WHERE = (
    "where b indexes ads in the batch and h ∈ {composite, survivability, "
    "engagement volume, engagement velocity}; ŷ_{b,h} = σ(·) is the "
    "sigmoid-bounded prediction; m_{b,h} ∈ {0,1} masks the two engagement "
    "heads for weak-engagement platforms (Reddit, other); q_b is the per-ad "
    "training-quality weight; α_h is the per-head weight (all 1.0)."
)

ALGO_B1 = """Algorithm B.1  Engagement-based scoring (v3 hybrid scorer)
Input: corpus C of ads; weights w = {surv: 0.50, vol: 0.25, vel: 0.25}
for each ad a in C:
    vol_raw  = log(1 + sum_f weight_f * engagement_f(a))   # shares:5 comments:3 likes:1 reactions:1 views:0.1
    vel_raw  = vol_raw / days_running(a)
    surv(a)  = KM_survival_percentile(a)                   # per-platform Kaplan-Meier, 7-day censoring
# percentile-normalise: engagement per platform, survivability globally
vol  = percentile_normalise(vol_raw,  by = platform)
vel  = percentile_normalise(vel_raw,  by = platform)
surv = percentile_normalise(surv,     global = True)
for each ad a in C:
    S = signals available for a            # weak platforms (Reddit, other) drop vol, vel
    composite(a) = sum_{s in S} (w_s / sum_{s in S} w_s) * s     # renormalise over available signals
    composite(a) = clip(composite(a), 0, 1)
# corpus-wide percentile tiers of composite
high   = top 20%      (>= 80th pct)
medium = middle 50%   (30th - 80th pct)
low    = bottom 30%   (< 30th pct)"""

ALGO_B2 = """Algorithm B.2  Training-data construction via backtranslation
Input: scored corpus C; teacher mix T = {Haiku 0.40, GPT 0.35, Gemini 0.25};
       registers R = {conversational, structured, imperative}; final size N
# pre-teacher selection
E = [ a in C if composite(a) >= 0.70 and is_english(a)
              and copy_len(a) >= 60 and quality_label(a) >= 3 ]
E = cap_per_vertical(E); cap_per_advertiser(E, max = 50)        # diversity caps
target = round(N * 1.25)                                        # overgeneration buffer
examples = []
for a in sample(E, target):
    teacher  = choose(T); register = random(R)
    brief    = teacher.reverse_engineer_brief(a, register)      # product facts only; ad held verbatim
    if gates_pass(brief, a):
        examples.append( (system_prompt, brief, ad_copy(a) + rationale) )
# ingestion gates (all must hold)
gates_pass(brief, a):
    structured tags present
    and fidelity(ad) == verbatim                 # no paraphrase of the real ad
    and no schema / field leakage into brief
    and TF-IDF_dedup(brief, kept) < 0.80
split examples -> 85% train / 7.5% val / 7.5% test   (stratified by platform)"""

ALGO_B3 = """Algorithm B.3  QLoRA fine-tuning setup
base  = load("Qwen3-8B", quantization = 4-bit NF4)             # base weights frozen
model = attach_LoRA(base, r = 16, alpha = 32, dropout = 0.05,
                    variants = [DoRA, rsLoRA],
                    target = [q, k, v, o, gate, up, down]_proj) # ~24M trainable (~0.29%)
data  = format_chatml(examples)                                # system + user brief + assistant
mask  = assistant_tokens_only                                  # system + brief masked from loss
opt   = AdamW_8bit(lr = 2e-4, weight_decay = 0.0)
sched = cosine(warmup_ratio = 0.03)
train(model, data, epochs = 3, eff_batch = 16, max_len = 4096, seed = 42,
      eval_every = 50 steps, early_stop(patience = 2, metric = val_loss))
adapter = best_checkpoint(model)                               # ~285 MB
merge_to_bf16(base, adapter) -> push_to_hub -> serve(vLLM)"""


def main() -> None:
    d = Document(DOC)
    paras = d.paragraphs
    log: list[str] = []

    # --- 1. Table 3.1 corrections + Framework row ---------------------------
    cap31 = find_para(paras, equals="Table 3.1: Fine-tuning configuration.")
    assert cap31 is not None, "Table 3.1 caption not found"
    t31 = tbl_after(cap31)
    assert t31 is not None, "Table 3.1 grid not found"
    PRED_TBL_TEMPLATE = copy.deepcopy(t31)  # clone BEFORE editing 3.1

    fixes = {
        "LoRA rank (r)": "16",
        "LoRA alpha (α)": "32",
        "Trainable parameters": "≈24 M (≈0.29% of 8.2 B)",
    }
    seen = set()
    for tr in rows(t31):
        lk = row_left(tr)
        if lk in fixes:
            set_row_text(tr, lk, fixes[lk])
            seen.add(lk)
            log.append(f"  Table 3.1: {lk} -> {fixes[lk]}")
    missing = set(fixes) - seen
    assert not missing, f"Table 3.1 rows not found: {missing}"

    if not any(row_left(tr) == "Framework" for tr in rows(t31)):
        # insert Framework after the Compute precision row
        anchor = next(tr for tr in rows(t31) if row_left(tr) == "Compute precision")
        new_tr = copy.deepcopy(anchor)
        set_row_text(new_tr, "Framework", "Unsloth + TRL (PyTorch / Hugging Face)")
        anchor.addnext(new_tr)
        log.append("  Table 3.1: + Framework row")

    # --- 3. Renumber existing Table 3.2 (comparison) -> 3.3 -----------------
    cap_cmp = find_para(paras, equals="Table 3.2: Comparison of the two evaluation methods.")
    if cap_cmp is not None:
        for t in cap_cmp.iter(W_T):
            if t.text and "Table 3.2" in t.text:
                t.text = t.text.replace("Table 3.2", "Table 3.3")
        for instr in cap_cmp.iter(W_INSTR):
            if instr.text and "Table 3.2" in instr.text:
                instr.text = instr.text.replace("Table 3.2", "Table 3.3")
        for bm in cap_cmp.findall(W_BMS):
            if bm.get(qn("w:name")) == "Tbl_3_2":
                bm.set(qn("w:name"), "Tbl_3_3")
        log.append("  Renumbered caption Table 3.2 -> Table 3.3 (+ TC field, bookmark)")
        # in-text callout
        for p in paras:
            for t in p._element.iter(W_T):
                if t.text and "Table 3.2" in t.text and not p.text.strip().startswith("Table 3."):
                    t.text = t.text.replace("Table 3.2", "Table 3.3")
                    log.append(f"  In-text callout: Table 3.2 -> Table 3.3 (para '{p.text.strip()[:40]}...')")
    else:
        log.append("  (comparison table already renumbered — skipped)")

    # --- 2. Insert new Table 3.2 scoring-predictor config -------------------
    if find_para(paras, equals="Table 3.2: Scoring-predictor configuration.") is None:
        sec38 = find_para(paras, equals="3.8 Evaluation Methodology")
        assert sec38 is not None, "section 3.8 heading not found"
        # caption: clone Table 3.1 caption, swap text + TC + bookmark
        new_cap = copy.deepcopy(cap31)
        for t in new_cap.iter(W_T):
            if t.text and "Table 3.1: Fine-tuning configuration." in t.text:
                t.text = "Table 3.2: Scoring-predictor configuration."
        for instr in new_cap.iter(W_INSTR):
            if instr.text and "Table 3.1" in instr.text:
                instr.text = ' TC "Table 3.2: Scoring-predictor configuration." \\f T \\l 1 '
        for bm in new_cap.findall(W_BMS):
            bm.set(qn("w:name"), "Tbl_3_2")
            bm.set(qn("w:id"), "10052")
        for bme in new_cap.findall(qn("w:bookmarkEnd")):
            bme.set(qn("w:id"), "10052")
        # grid: reuse cloned Table-3.1 grid, overwrite rows, trim surplus
        grid = PRED_TBL_TEMPLATE
        all_rows = rows(grid)
        header, data_rows = all_rows[0], all_rows[1:]
        set_row_text(header, "Parameter", "Value")
        for i, (lk, rv) in enumerate(PREDICTOR_ROWS):
            set_row_text(data_rows[i], lk, rv)
        for surplus in data_rows[len(PREDICTOR_ROWS):]:
            grid.remove(surplus)
        sec38.addprevious(new_cap)
        new_cap.addnext(grid)
        log.append(f"  + Table 3.2 'Scoring-predictor configuration' ({len(PREDICTOR_ROWS)} rows) before 3.8")
    else:
        log.append("  (Table 3.2 predictor config already present — skipped)")

    # --- 4. Loss equations --------------------------------------------------
    def body_after(heading_text):
        p_list = list(d.paragraphs)
        for i, p in enumerate(p_list):
            if "\t" in p.text.strip():
                continue
            if p.text.strip().startswith(heading_text):
                for q in p_list[i + 1 :]:
                    qt = q.text.strip()
                    if qt and "\t" not in qt and not q.style.name.startswith("Heading"):
                        return q._element
        return None

    body353 = body_after("3.5.3 Training Inputs")
    if body353 is not None and SFT_EQ not in para_text(body353.getparent()) and find_para(d.paragraphs, contains="ℒ_SFT") is None:
        where_p = clone_para(body353, SFT_WHERE)
        eq_p = clone_para(body353, SFT_EQ, center=True)
        body353.addnext(where_p)
        body353.addnext(eq_p)
        log.append("  + SFT loss equation after 3.5.3")

    body374 = body_after("3.7.4 Model Size and Hardware")
    if body374 is not None and find_para(d.paragraphs, contains="m_{b,h}") is None:
        where_p = clone_para(body374, PRED_WHERE)
        eq_p = clone_para(body374, PRED_EQ, center=True)
        body374.addnext(where_p)
        body374.addnext(eq_p)
        log.append("  + predictor loss equation after 3.7.4")

    # --- 5. Appendix B + pseudocode ----------------------------------------
    if find_para(d.paragraphs, contains="Appendix B: Algorithm Listings") is None:
        # safe heading templates from our own Appendix A
        h2 = find_para(d.paragraphs, equals="Appendix A: Prompt Templates")
        h3 = find_para(d.paragraphs, startswith="A.1 ")
        assert h2 is not None and h3 is not None, "Appendix A heading templates missing"
        assert h2.find(".//" + qn("w:drawing")) is None, "Appendix A H2 carries a drawing — unsafe to clone"
        body_code_template = body353 if body353 is not None else h3

        # insertion point: before final sectPr (end of document)
        bodyel = d.element.body
        last = bodyel[-1]
        anchor = last if last.tag == W_SECTPR else None

        def add(el):
            if anchor is not None:
                anchor.addprevious(el)
            else:
                bodyel.append(el)

        add(clone_para(h2, "Appendix B: Algorithm Listings"))
        for h3_text, code in [
            ("B.1 Engagement-Based Scoring", ALGO_B1),
            ("B.2 Training-Data Construction", ALGO_B2),
            ("B.3 QLoRA Fine-Tuning", ALGO_B3),
        ]:
            add(clone_para(h3, h3_text))
            for line in code.split("\n"):
                add(clone_para(body_code_template, line if line else " ", mono=True))
        log.append("  + Appendix B: Algorithm Listings (B.1, B.2, B.3 pseudocode)")
    else:
        log.append("  (Appendix B already present — skipped)")

    # --- cross-ref callouts -------------------------------------------------
    if append_callout(d.paragraphs, "3.7.4 Model Size and Hardware", "Table 3.2"):
        log.append("  + callout (Table 3.2) in 3.7.4")
    if append_callout(d.paragraphs, "3.3.4 Composite Score and Tiers", "Algorithm B.1"):
        log.append("  + callout (Algorithm B.1) in 3.3.4")
    if append_callout(d.paragraphs, "3.4.4 Ingestion Gates", "Algorithm B.2"):
        log.append("  + callout (Algorithm B.2) in 3.4.4")
    if append_callout(d.paragraphs, "3.5.4 Optimization", "Algorithm B.3"):
        log.append("  + callout (Algorithm B.3) in 3.5.4")

    d.save(DOC)
    print("CHANGES APPLIED:")
    print("\n".join(log) if log else "  (none)")


if __name__ == "__main__":
    try:
        main()
    except AssertionError as e:
        print(f"ABORTED (no save): {e}", file=sys.stderr)
        sys.exit(1)
