"""Rebuild the List-of-Tables entry for Table 3.1 with clean tab runs.

The insert script jammed "\t1" as literal text into a cloned run, leaving a
stray <w:tab/> ("...configuration\t1\t"). Rebuild the entry from the clean
eval entry: keep its pPr (tab stop + leader), set fresh runs description + tab
+ page number. Idempotent: skips if already clean.
"""

from __future__ import annotations

import copy

import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DOC = "docs/research/THESIS.docx"


def main() -> None:
    d = docx.Document(DOC)

    def tof(sub):
        return next(
            (p for p in d.paragraphs
             if p.style.name == "table of figures" and sub in p.text),
            None,
        )

    bad = tof("Fine-tuning configuration")
    evt = tof("Comparison of the two evaluation methods")
    if bad is None or evt is None:
        raise SystemExit("ToF entries not found")
    if bad.text == "Table 3.1: Fine-tuning configuration\t1":
        print("already clean, nothing to do")
        return

    # rebuild from the clean eval entry's pPr
    clone_el = copy.deepcopy(evt._p)
    for r in clone_el.findall(qn("w:r")):
        clone_el.remove(r)
    par = Paragraph(clone_el, evt._parent)
    par.add_run("Table 3.1: Fine-tuning configuration")
    par.add_run().add_tab()
    par.add_run("1")

    evt._p.addprevious(clone_el)
    bad._p.getparent().remove(bad._p)
    d.save(DOC)
    print("rebuilt ToF entry:", repr(par.text))


if __name__ == "__main__":
    main()
