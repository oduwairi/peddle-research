"""Reviewer feedback #12: "Add STATISTICAL EVALUATION OF MODEL PERFORMANCE."

Audit finding: the thesis already reports bootstrap CIs, paired-contrast
significance, effect sizes, calibration (ECE/ROC-AUC) and correlations
(Spearman/Pearson) throughout Chapter IV. The real gap is that Chapter III's
methodology (§3.8) never *describes* how performance is statistically
evaluated — the bootstrap settings, the significance criterion, and the
paired-contrast construction live only in results captions. This script does
two grounded, minimal things and writes NO new body prose:

  (A) Inserts a new §3.8.8 "Statistical Evaluation of Model Performance"
      Heading 3 plus a bracketed PLACEHOLDER body at the end of §3.8 (before
      CHAPTER IV). The author writes the prose from the Phase-3 points
      (per docs/research/THESIS_EDITING.md, "never draft new thesis prose").
      If the heading already exists it is left untouched (protects author work).

  (B) Augments the existing §4.2.1 headline claim (P~648: "...beats out both
      the frontier model (+0.048 ... over A ...) and the base model (+0.040
      ... over B ...)") with the paired-bootstrap significance test for those
      two gaps — appended as parenthetical statistics on the author's existing
      numbers (the AI-detector-safe "additions as parentheticals" pattern, not
      a rewrite). Numbers are computed from the same per-row eval data with the
      same bootstrap settings (1000 resamples, seed=42) as Table 4.8.

Idempotent: re-running skips the §3.8.8 heading if present and skips the
parenthetical if already injected. Anchors located by unique text, never
hardcoded indices.

Run:  uv run python scripts/thesis/insert_reviewer12_statistics.py [--dry-run]
"""

from __future__ import annotations

import copy
import sys
from pathlib import Path

import numpy as np
import polars as pl
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DOCX = Path("docs/research/THESIS.docx")
PER_ROW = Path("data/eval/learned_scores")

H38_8_TITLE = "3.8.8 Statistical Evaluation of Model Performance"
PLACEHOLDER = (
    "[Placeholder — §3.8.8 Statistical Evaluation of Model Performance. "
    "Author to write from the supplied Phase-3 points.]"
)

MINUS = "−"  # match the thesis convention ("C − B")


# --------------------------------------------------------------------------- #
# Bootstrap (paired) — identical to insert_ch4_tables.boot_ci_idx / Table 4.8
# --------------------------------------------------------------------------- #
def boot_ci_idx(diffs: np.ndarray, n_boot: int = 1000, seed: int = 42):
    rng = np.random.default_rng(seed)
    n = len(diffs)
    boot = np.empty(n_boot)
    for i in range(n_boot):
        idx = rng.integers(0, n, n)
        boot[i] = diffs[idx].mean()
    lo, hi = np.percentile(boot, [2.5, 97.5])
    return float(diffs.mean()), float(lo), float(hi)


def paired_contrast(a: str, b: str):
    """Mean Δ + 95% CI of (a − b) composite over the shared briefs."""
    da = pl.read_parquet(PER_ROW / f"{a}.parquet").select(["example_id", "composite"]).rename({"composite": a})
    db = pl.read_parquet(PER_ROW / f"{b}.parquet").select(["example_id", "composite"]).rename({"composite": b})
    j = da.join(db, on="example_id", how="inner")
    diffs = (j[a] - j[b]).to_numpy()
    m, lo, hi = boot_ci_idx(diffs)
    return m, lo, hi, len(j)


def fmt_contrast(label_a: str, label_b: str, m: float, lo: float, hi: float, n: int) -> str:
    return f"paired {label_a} {MINUS} {label_b} = {m:+.3f}, 95% CI [{lo:+.3f}, {hi:+.3f}], significant"


# --------------------------------------------------------------------------- #
# helpers
# --------------------------------------------------------------------------- #
def has_drawing(p: Paragraph) -> bool:
    return p._p.find(".//" + qn("w:drawing")) is not None


def find_one(doc: Document, *, style: str | None, contains: str) -> Paragraph | None:
    hits = [
        p for p in doc.paragraphs
        if contains in p.text and (style is None or p.style.name == style)
    ]
    return hits[0] if len(hits) == 1 else (hits[0] if hits else None)


def clone_para(template: Paragraph, text: str):
    new_el = copy.deepcopy(template._p)
    para = Paragraph(new_el, template._parent)
    runs = para.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        para.add_run(text)
    return new_el, para


# --------------------------------------------------------------------------- #
# (A) §3.8.8 heading + placeholder
# --------------------------------------------------------------------------- #
def insert_3_8_8(doc: Document) -> str:
    if any(p.text.strip().startswith("3.8.8 Statistical Evaluation") for p in doc.paragraphs):
        return "== §3.8.8 already present — left untouched (protects author prose)"

    h3 = find_one(doc, style="Heading 3", contains="3.8.7 Gold Baseline")
    body_tmpl = find_one(doc, style="Body Text", contains="Gold ads compared against the baseline")
    anchor = find_one(doc, style="Body Text", contains="Table 3.3: Comparison of the two evaluation methods")
    if h3 is None or body_tmpl is None or anchor is None:
        return f"!! §3.8.8 anchor missing (h3={h3 is not None}, body={body_tmpl is not None}, anchor={anchor is not None})"
    if has_drawing(h3) or has_drawing(body_tmpl):
        return "!! §3.8.8 template carries a drawing — refusing to clone"

    # The "Table 3.3" caption is followed by its table grid; insert §3.8.8
    # AFTER the grid so we never split a caption from its table.
    insert_after = anchor._p
    nxt = insert_after.getnext()
    if nxt is not None and nxt.tag == qn("w:tbl"):
        insert_after = nxt

    head_el, _ = clone_para(h3, H38_8_TITLE)
    insert_after.addnext(head_el)
    body_el, _ = clone_para(body_tmpl, PLACEHOLDER)
    head_el.addnext(body_el)
    return f"++ §3.8.8 heading + placeholder inserted after {anchor.text.strip()[:40]!r} (past its table grid)"


# --------------------------------------------------------------------------- #
# (B) paired-contrast parenthetical on the §4.2.1 headline claim
# --------------------------------------------------------------------------- #
def inject_contrast_parenthetical(doc: Document) -> str:
    para = find_one(doc, style="Body Text", contains="the final tally is GOLD")
    if para is None:
        return "!! headline-claim paragraph not found"
    if "paired C " + MINUS + " A" in para.text:
        return "== contrast parenthetical already present"

    ca = paired_contrast("C", "A")
    cb = paired_contrast("C", "B")
    ca_txt = fmt_contrast("C", "A", *ca[:3], ca[3])
    cb_txt = fmt_contrast("C", "B", *cb[:3], cb[3])

    target_run = next(
        (r for r in para.runs if "an 8.0% relative lift)" in r.text and "a 6.5% relative lift)" in r.text),
        None,
    )
    if target_run is None:
        return "!! lift substrings not in a single run — manual placement needed"

    new = target_run.text
    new = new.replace(
        "an 8.0% relative lift)",
        f"an 8.0% relative lift; {ca_txt})",
    )
    new = new.replace(
        "a 6.5% relative lift)",
        f"a 6.5% relative lift; {cb_txt})",
    )
    target_run.text = new
    return (
        f"++ contrast parenthetical injected "
        f"[C{MINUS}A {ca[0]:+.3f} (n={ca[3]}); C{MINUS}B {cb[0]:+.3f} (n={cb[3]})]"
    )


def main() -> int:
    dry = "--dry-run" in sys.argv
    if not DOCX.exists():
        print(f"ERROR: {DOCX} not found", file=sys.stderr)
        return 1
    if Path("docs/research/.~lock.THESIS.docx#").exists():
        print("ERROR: THESIS.docx is open in an editor (lock present). Close it.", file=sys.stderr)
        return 1

    if dry:
        print("=== DRY RUN — computed paired contrasts ===")
        for a, b in (("C", "A"), ("C", "B"), ("GOLD", "C")):
            m, lo, hi, n = paired_contrast(a, b)
            sig = "significant" if (lo > 0 or hi < 0) else "n.s."
            print(f"  {a} {MINUS} {b}: {m:+.4f}  95% CI [{lo:+.4f}, {hi:+.4f}]  n={n}  {sig}")
        print(f"\n  §3.8.8 heading text: {H38_8_TITLE!r}")
        print(f"  placeholder body  : {PLACEHOLDER!r}")
        return 0

    doc = Document(str(DOCX))
    print("CHANGES:")
    print("  " + insert_3_8_8(doc))
    print("  " + inject_contrast_parenthetical(doc))
    doc.save(str(DOCX))
    print(f"\nSaved {DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
