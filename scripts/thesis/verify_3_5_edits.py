"""Verify the Ch. III config-table + appendix edits landed on disk."""

from __future__ import annotations

import docx
from docx.oxml.ns import qn

DOC = "docs/research/THESIS.docx"


def main() -> None:
    d = docx.Document(DOC)
    ps = d.paragraphs

    print("=== 3.5.4 Optimization paragraph (edited) ===")
    opt = next(p for p in ps if "cosine schedule (Loshchilov" in p.text)
    t = opt.text
    for probe in ("peak learning rate of 2×10⁻⁴", "warmup ratio of 0.03",
                  "takes from one and a half to three hours"):
        print(f"  [{'OK' if probe in t else 'MISSING'}] {probe}")
    print("  ...snippet:", t[t.index("cosine"):t.index("cosine") + 170])
    print("  ...time:", t[t.index("Training takes" if "Training takes" in t else "Training"):][:120])

    print("\n=== Tables: captions + numbering ===")
    for p in ps:
        s = p.text.strip()
        if s.startswith("Table 3.") and ("configuration" in s or "Comparison" in s):
            print(f"  body caption | {p.style.name} | {s}")
    print("  -- List of Tables entries --")
    for p in ps:
        if p.style.name == "table of figures" and "Table 3." in p.text:
            print(f"  ToF | {p.text!r}")

    print("\n=== New config table content (find by header 'Parameter') ===")
    cfg = next((t for t in d.tables if t.rows and t.rows[0].cells[0].text.strip() == "Parameter"), None)
    if cfg is None:
        print("  MISSING config table!")
    else:
        print(f"  rows={len(cfg.rows)} cols={len(cfg.columns)}  style={cfg.style.name if cfg.style else None}")
        for r in cfg.rows[:4]:
            print("   ", r.cells[0].text, "::", r.cells[1].text)
        print("    ... (last 3) ...")
        for r in cfg.rows[-3:]:
            print("   ", r.cells[0].text, "::", r.cells[1].text)
    print(f"  total tables in doc: {len(d.tables)}")

    print("\n=== Order check: caption -> table -> 3.6 heading ===")
    body = list(d.element.body)
    tags = []
    for ch in body:
        tag = ch.tag.split('}')[-1]
        if tag == 'p':
            txt = ''.join(t.text or '' for t in ch.findall('.//' + qn('w:t')))
            if 'Table 3.1: Fine-tuning' in txt or '3.6 Agent Architecture' in txt or '3.5.5 Adapter' in txt:
                tags.append(('p', txt[:42]))
        elif tag == 'tbl':
            first = ch.find('.//' + qn('w:t'))
            if first is not None and first.text == 'Parameter':
                tags.append(('TBL', 'config'))
    print("  sequence:", tags)

    print("\n=== Pointers ===")
    for sub, want in (("the raw ads are sent to the teacher", "Appendix A.2"),
                      ("You are an ad copywriter", "Appendix A")):
        p = next((p for p in ps if sub in p.text), None)
        ok = p is not None and want in p.text
        print(f"  [{'OK' if ok else 'MISSING'}] sec pointer -> {want}")

    print("\n=== Appendix A structure ===")
    seen = False
    for p in ps:
        st = p.text.strip()
        if st == "Appendix A: Prompt Templates":
            seen = True
        if seen and (st.startswith("Appendix A") or st.startswith("A.")):
            print(f"  {p.style.name} | {st}")

    print("\n=== Drawing-clone safety: count <w:drawing> in doc ===")
    n_draw = len(d.element.body.findall('.//' + qn('w:drawing')))
    print(f"  total <w:drawing> elements: {n_draw} (was 4 figures pre-edit; appendix/table add none)")

    print("\n=== Appendix sanity: system prompt + a teacher-rule line present verbatim ===")
    alltext = "\n".join(p.text for p in ps)
    for probe in ("You are an ad copywriter. When a user describes a product",
                  "Reproduce the source ad word-for-word",
                  "Get 2 FREE tacos"):
        print(f"  [{'OK' if probe in alltext else 'MISSING'}] {probe[:48]}")


if __name__ == "__main__":
    main()
