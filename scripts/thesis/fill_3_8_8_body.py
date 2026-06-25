"""Fill the §3.8.8 body: replace the placeholder with the author's polished prose.

Phase 7 of the collaborative voice workflow (docs/research/THESIS_EDITING.md).
The author wrote the §3.8.8 draft; this drops the Phase-5-polished version (with
the one author-requested MAUVE fact added) over the placeholder paragraph,
preserving the Body-Text run/formatting. Idempotent: if the placeholder is gone
and the body is already present, it skips.

Run:  uv run python scripts/thesis/fill_3_8_8_body.py [--dry-run]
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

DOCX = Path("docs/research/THESIS.docx")

PLACEHOLDER_MARK = "Placeholder — §3.8.8"

BODY = (
    "The thesis followed a rigorous results-reporting strategy to ensure the "
    "results are as close to reality as possible. This includes a confidence "
    "interval on every reported number mean (95% CI). Intervals also follow a "
    "bootstrap, where we resample per-ad scores with replacement over 1000 "
    "resamples, with the interval taken as the 2.5th/97.5th percentiles of the "
    "resampled means. This is applied to all three evaluation arms — the "
    "learned-scorer (including its per-platform and per-head breakdowns), MAUVE, "
    "and reference-overlap. Our test set contains 215 briefs in total, as the "
    "10% held-out portion of the training split; since the sample number can be "
    "considered relatively small, some uncertainty is expected. When comparing "
    "models, we always ensure the same briefs and answers are compared across "
    "configurations, and the thesis reports the number of samples for each "
    "comparison as the paired n. The headline fine-tuning gap is also reported "
    "as an effect size (Cohen's dz). Individually, the scorer was validated on a "
    "held-out split, with Spearman, Pearson, and MAE reported per head; MAUVE "
    "was grounded by anchoring its scale against real ads, with the GOLD "
    "configuration of real held-out ads setting the ceiling and random text the "
    "floor; and the reference-overlap metrics were grounded on Upworthy A/B "
    "headline test sets."
)


def set_para_text(para, text: str) -> None:
    runs = para.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        para.add_run(text)


def main() -> int:
    dry = "--dry-run" in sys.argv
    if not DOCX.exists():
        print(f"ERROR: {DOCX} not found", file=sys.stderr)
        return 1
    if Path("docs/research/.~lock.THESIS.docx#").exists():
        print("ERROR: THESIS.docx is open in an editor (lock present). Close it.", file=sys.stderr)
        return 1

    doc = Document(str(DOCX))

    # Target the placeholder, or the already-inserted body (so this is
    # re-runnable when the prose is edited).
    target = next((p for p in doc.paragraphs if PLACEHOLDER_MARK in p.text), None)
    if target is None:
        target = next(
            (p for p in doc.paragraphs if p.text.strip().startswith("The thesis followed a rigorous results-reporting")),
            None,
        )
    if target is None:
        print("!! neither placeholder nor existing §3.8.8 body found — aborting", file=sys.stderr)
        return 2

    if target.text.strip() == BODY.strip():
        print("== §3.8.8 body already up to date — nothing to do")
        return 0

    if dry:
        print("=== DRY RUN ===\nWould set §3.8.8 body to:\n")
        print(BODY)
        return 0

    set_para_text(target, BODY)
    doc.save(str(DOCX))
    print(f"++ §3.8.8 body set ({len(BODY)} chars). Saved {DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
