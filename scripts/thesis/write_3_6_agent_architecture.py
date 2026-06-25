"""Write §3.6 Agent Architecture body.

Same pattern as write_3_4_construction.py: wipes every paragraph between the
§3.6 heading and the §3.7 heading, then re-inserts the current PARAGRAPHS
list using a Body Text (style 943) template paragraph from CH3. New
bibliography entries (if any) are inserted in NEW_REFERENCES at their
alphabetical positions before APPENDIX, idempotent.

Idempotent: re-run safely as more paragraphs land in PARAGRAPHS.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"


PARAGRAPHS: list[str] = [
    # ¶1 — Motivation for agent-first architecture
    "Our fine-tuned model specializes in ad copywriting and high-performance ad generation. Other tasks such as general queries, web searching, and landing-page URL retrieval are less suited to such a model. For Draper to be a fully deployable system, the surrounding workflow must be capable of handling multi-step tasks and tool-calling sequences (function calling). This motivates an agent-first architecture, splitting the model's responsibilities across tools that handle these different tasks, such as reading uploaded documents (PDF, DOCX, and plain text), searching the internet, scraping the web, and following multi-step workflows.",
    # ¶2 — Two roles and the hard separation
    "The agent system consists of two main models, an orchestrator model and a writer model. The orchestrator model serves as a generalist frontier-model LLM via API (GPT-5.4-mini in our deployment); it handles every task except for creative generation, which is there for the writer. For example, it can talk to the user, provide follow-ups, explanations, call tools, perform web searches, and more. On the other hand, the writer model is our custom fine-tuned model (see §3.5). This separation maintains that general-purpose models are better at agentic tasks while our custom LLM is specifically suited to generating creative ads as per its training. This is a strict separation: the orchestrator model is never allowed to paraphrase or produce any ad copy. Anything of such creative nature is routed to the writer model as a tool call.",
    # ¶3 — The freeform loop and the four tool buckets
    "When the orchestrator and writer model are set, the agent workflow runs as a freeform loop. The orchestrator model sequentially reasons and executes actions with provided tools according to the information provided in the system prompt. The agent loop typically ends in a specific marketing campaign emission, but may not, depending on the nature of the task asked by the user. The loop is also bounded at fourteen steps maximum to ensure cost efficiency and avoid long loops (a typical successful campaign completes in six to nine steps). The main tools provided to the agent are classified as follows. Research tools (scrape_url, attach_file, web_search, exa_similar) are responsible for gathering facts and searching the web for information to ground the marketing campaign in facts. Writing tools (draft_campaign, ask_draper, and generate_image) are used to route requests to the Draper model for creative-related content such as headlines and artwork generation. Scoring tools (score_copy) are used to execute the scoring predictor on the generated ad copy, whether during the agent loop or at the end. Output tools (emit_campaign) are used to present the output to the user in a specific format depending on the request or platform, printed on the screen in a specific UI. The agent is a freeform agent that decides and executes which tool to use at any time, avoiding rigid, hard-coded workflows.",
    # ¶4 — Research phase and the product-only brief
    "The research phase represents the forefront of the agentic workflow. The research consists of web searching as well as scraping URLs in turn. The goal of this phase is to ground the information and facts and to polish the ads with the latest market and emerging trends as well as competitive analysis. The most common and convenient entry point is the product URL that the user can paste. The agent can call the scrape_url tool (powered by Jina Reader, which converts the page to markdown) to extract the content of the landing page into an extracted response. Another popular mode of input is attached files such as PDFs, DOCX, and plain text. The model can similarly fetch the contents into a structured response. Live web search (via the Tavily API) is used so that the agent can search the web for related content and competitor analysis as well as buyer voice cues. Once the research phase is complete and the model has enough, it combines what it learned into a single brief (product name, description, key features, unique selling points, and category) that is sent to the writer model for ad copy generation. The brief is intentionally stripped of raw data and full scraping responses, since this can clutter the writer model and inhibit its creative writing ability.",
    # ¶5 — Writer tools, best-of-N, image, emit, scorecard
    "Communication between the orchestrator and the writer is done mainly through two tools. The draft_campaign tool is a more structured workflow where the orchestrator model asks the writer model for an ad campaign in the native shape of a platform. It is also constrained by platform constraints, such as character caps (for example, 280 characters on X and 2,200 on Meta). A secondary ask_draper tool may also be invoked by the agent for smaller and more unstructured tasks such as writing variants, rewrites, improvements, and taglines that don't follow a specific construction. The proposed architecture fires multiple requests in parallel with different temperatures (six draws at temperatures 0.5 to 1.0) and scores them with the predictor, which then returns the third-best option, safer than the best option. After the draft, the agent is usually encouraged to invoke the generate_image tool, which calls the OpenAI gpt-image-1.5 model for visual generation based on the visual brief generated by the writer. The final compiled output is presented to the user using the emit_campaign tool that produces a structured output response rendered by the UI. Each emitted campaign is also accompanied by a scorecard showing the predicted score by the predictor (a 0–100 badge attached to the campaign card).",
]


# (alphabetical-anchor-prefix, full reference text)
NEW_REFERENCES: list[tuple[str, str]] = []


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def set_paragraph_text(p_elem, new_text):
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def write_section_body(doc):
    h_3_6 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("3.6 Agent Architecture"),
        "§3.6 heading",
    )
    h_3_7 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("3.7 Scoring Predictor"),
        "§3.7 heading",
    )

    # Body Text template: first style-943 paragraph inside CH3 (excluding
    # §3.6's own body, which we're about to wipe).
    body_tpl = None
    in_ch3 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.upper() == "CHAPTER III":
            in_ch3 = True
            continue
        if t.upper() == "CHAPTER IV":
            break
        if t.startswith("3.6 Agent Architecture"):
            in_ch3 = "in-36"
            continue
        if in_ch3 == "in-36" and t.startswith("3.7 "):
            in_ch3 = True
            continue
        if in_ch3 == "in-36":
            continue
        sid = p.style.style_id if p.style else ""
        if in_ch3 and sid == "943" and t and p._element.findall(f".//{QN_T}"):
            body_tpl = p._element
            break
    if body_tpl is None:
        raise RuntimeError("No style-943 Body Text paragraph found inside CH3 (outside §3.6)")

    body = h_3_6.getparent()
    children = list(body)
    i_start = children.index(h_3_6) + 1
    i_end = children.index(h_3_7)

    to_delete = [c for c in children[i_start:i_end] if c.tag == QN_P]
    print(f"DELETING: {len(to_delete)} existing §3.6 body paragraphs")
    for c in to_delete:
        body.remove(c)

    insert_after = h_3_6
    for i, text in enumerate(PARAGRAPHS, 1):
        p = deepcopy(body_tpl)
        set_paragraph_text(p, text)
        insert_after.addnext(p)
        insert_after = p
        print(f"  ¶{i}: {text[:80]}…")

    print(f"\nWROTE: {len(PARAGRAPHS)} paragraph(s) into §3.6 Agent Architecture")


def insert_references(doc):
    if not NEW_REFERENCES:
        print("(no new references for §3.6)")
        return

    refs_heading = None
    appendix_heading = None
    for p in doc.paragraphs:
        t = p.text.strip().upper()
        if t == "REFERENCES" and refs_heading is None:
            refs_heading = p._element
        elif t.startswith("APPENDIX") and appendix_heading is None:
            appendix_heading = p._element
    if refs_heading is None or appendix_heading is None:
        raise RuntimeError("References / APPENDIX bounds not found")

    body = refs_heading.getparent()
    children = list(body)
    i_refs = children.index(refs_heading)
    i_appendix = children.index(appendix_heading)
    ref_range = children[i_refs + 1 : i_appendix]

    ref_tpl = None
    for c in ref_range:
        if c.tag != QN_P:
            continue
        pPr = c.find(f"{{{W_NS}}}pPr")
        pStyle = pPr.find(f"{{{W_NS}}}pStyle") if pPr is not None else None
        sid = pStyle.get(f"{{{W_NS}}}val") if pStyle is not None else ""
        ts = c.findall(f".//{QN_T}")
        if sid == "943" and ts and any(t.text for t in ts):
            ref_tpl = c
            break
    if ref_tpl is None:
        raise RuntimeError("No style-943 reference template paragraph found")

    inserted = 0
    skipped = 0
    for anchor_prefix, ref_text in NEW_REFERENCES:
        first_words = ref_text.split("(")[0].strip().rstrip(",")
        already_present = any(
            c.tag == QN_P and "".join(t.text or "" for t in c.findall(f".//{QN_T}")).startswith(first_words)
            for c in ref_range
        )
        if already_present:
            print(f"  SKIP (already present): {ref_text[:80]}…")
            skipped += 1
            continue

        anchor = None
        for c in ref_range:
            if c.tag != QN_P:
                continue
            text = "".join(t.text or "" for t in c.findall(f".//{QN_T}"))
            if text.startswith(anchor_prefix):
                anchor = c
                break
        if anchor is None:
            raise RuntimeError(f"Alphabetical anchor not found: {anchor_prefix!r}")

        new_p = deepcopy(ref_tpl)
        set_paragraph_text(new_p, ref_text)
        anchor.addprevious(new_p)
        inserted += 1
        print(f"  INSERT before {anchor_prefix!r}: {ref_text[:80]}…")

        children = list(body)
        i_refs_new = children.index(refs_heading)
        i_appendix_new = children.index(appendix_heading)
        ref_range = children[i_refs_new + 1 : i_appendix_new]

    print(f"\nREFERENCES: inserted {inserted}, skipped {skipped} (already present)")


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    print("=== §3.6 body ===")
    write_section_body(doc)

    print("\n=== References ===")
    insert_references(doc)

    doc.save(str(path))
    print(f"\nSAVED: {path}")


if __name__ == "__main__":
    main()
