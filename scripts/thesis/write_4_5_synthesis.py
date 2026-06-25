"""Write §4.5 Synthesis and Limitations body (between §4.5 and CHAPTER V).

Same wipe-and-replace pattern as write_4_4_ablation.py: wipes every paragraph
between the "4.5 Synthesis and Limitations" heading and the "CHAPTER V"
heading, then re-inserts the current PARAGRAPHS list using a Body Text
(style 943) template paragraph cloned from Chapter III.

Idempotent — re-running produces the same output. §4.5 has no embedded
figures, but the figure-preserving predicate is kept for consistency.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"
QN_DRAWING = f"{{{W_NS}}}drawing"


def _is_figure_block(p_elem) -> bool:
    """Preserve image paragraphs and their caption paragraphs (start with
    'Figure 4.5.'). Lets us re-run this script without wiping embedded
    figures, even though §4.5 currently has none."""
    if p_elem.find(f".//{QN_DRAWING}") is not None:
        return True
    text = "".join((t.text or "") for t in p_elem.findall(f".//{QN_T}"))
    return text.strip().startswith("Figure 4.5.")


PARAGRAPHS: list[str] = [
    # ¶1 — Recap the two evaluation arms; report both rankings; flag the
    # cross-arm disagreement on the agent loop's effect as the section's
    # foundational question.
    "We recap our evaluation procedures. §4.2 uses the learned scorer, "
    "which is a DeBERTa-v3-base regressor we trained on a 55k-ad AdFlex "
    "corpus. This model outputs a number between zero and one "
    "representing the composite performance score of the individual ad, "
    "as well as three underlying scores: survivability, engagement "
    "volume, and engagement velocity. Higher composite score "
    "corresponds to better performance, and this model is used for each "
    "individual inference by all configurations. Meanwhile, §4.3 uses "
    "the MAUVE method (Pillutla et al., 2021), which is a corpus-level "
    "comparison method based on embedding similarity of text. Again, "
    "the output is a number between zero and one for each comparison "
    "between the reference ads and the ads to be tested. Higher number "
    "means higher distribution or similarity overlap. Both evaluation "
    "arms run over the same test briefs and configurations. Per the "
    "findings of §4.2 and §4.3, the learned scorer yielded results of "
    "composite mean: GOLD 0.684, C 0.651, B 0.611, C_pipe 0.607, A "
    "0.603, B_pipe 0.586; and MAUVE all-platform mean: GOLD 0.462, "
    "C_pipe 0.420, B_pipe 0.302, C 0.287, B 0.183, A 0.180. The most "
    "important fact is that both evaluation arms position GOLD at the "
    "top, which means the results are compared to the correct baseline. "
    "It is also noted that in both evaluation methods, configuration A "
    "sits at the bottom of the ranks. However, more critically, the "
    "evaluation arms disagree on the effect of agent-wrapped models on "
    "the overall performance. The learned scorer indicates that the "
    "agent RAG slightly decreases performance for both configurations "
    "(fine-tuned and base model): C beats C_pipe by 0.044, B beats "
    "B_pipe by 0.025; while the MAUVE method contradictorily indicates "
    "that the agent loop helps significantly: C_pipe beats C by 0.132, "
    "B_pipe beats B by 0.119. We think it even serves as the foundation "
    "of this section, where we explore possible causes and "
    "interpretations of this result.",
    # ¶2 — Unpack what each arm measures: learned scorer is absolute and
    # per-ad (inherits v3 biases); MAUVE is comparative and distributional
    # (blind to per-ad quality). This frames the disagreement as two
    # facets of the same competence rather than a contradiction.
    "We provide a deeper dive into what each evaluation arm actually "
    "measures in practice. The learned scorer provides a single number "
    "which represents the composite performance score. This is an "
    "absolute score, not a comparative one, unlike MAUVE, which "
    "provides only comparative results by comparing embedding "
    "distribution overlap. Moreover, the learned scorer is a model "
    "trained on our custom v3 hybrid scoring system (§3.3) that we use "
    "to label the raw ad corpus, meaning any biases or inaccuracies in "
    "the original system are inherited or learned by this model. The "
    "learned scorer has its distinct advantage of not being influenced "
    "by other inputs during production or testing, while providing a "
    "single score proxying predicted performance. It also has the "
    "ability to distinguish individual ad performance, as compared to "
    "MAUVE, which compares distributions and is blind to individual "
    "scores. Meanwhile, MAUVE is more representative of the overall "
    "shape of the data and can capture variety better than the "
    "individual scorer, making it a better candidate for pooled data.",
    # ¶3 — Limitations: text-only across both arms; single-split test
    # data drawn from the writer's own construction pool; LLM-as-judge
    # tournament implemented but not reported (judge bias toward
    # generative content); real-world deployment as the true test,
    # deferred beyond thesis scope.
    "Having discussed the interpretation of the results, it is "
    "important to outline a few limitations in our methodology as well "
    "as the ceiling of what can be tested. First of all, we acknowledge "
    "that both our evaluation arms score text only, and there is no "
    "judgment of any other metadata such as images, brand consistency, "
    "or layout, which can often have an extensive effect on campaign "
    "performance. Additionally, the held-out split of the 215 briefs "
    "is from the same construction pool as the original corpus that "
    "the writer trained on. No additional datasets or external corpora "
    "are used. Within this methodology, additional evaluation methods "
    "are implemented but not reported, such as the LLM-as-a-judge "
    "tournament. The main reason for omitting it is that AI models are "
    "unreliable at judging high-performance ads from a human "
    "perspective and would offer a great bias toward generative AI "
    "content. Finally, a real test of ad performance is real-world "
    "deployment, where ads are deployed and performance is observed in "
    "real time. However, this has great cost implications and is "
    "deferred beyond the scope of this thesis.",
]


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def set_paragraph_text(p_elem, new_text):
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def write_section_body(doc):
    h_4_5 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("4.5 Synthesis and Limitations"),
        "§4.5 heading",
    )
    h_ch5 = find_paragraph(
        doc,
        lambda p: p.text.strip().upper() == "CHAPTER V",
        "Chapter V heading",
    )

    # Body Text template: first style-943 paragraph inside Chapter III with
    # real text content. Same convention as write_4_4_ablation.py.
    body_tpl = None
    in_ch3 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.upper() == "CHAPTER III":
            in_ch3 = True
            continue
        if t.upper() == "CHAPTER IV":
            break
        sid = p.style.style_id if p.style else ""
        if in_ch3 and sid == "943" and t and p._element.findall(f".//{QN_T}"):
            body_tpl = p._element
            break
    if body_tpl is None:
        raise RuntimeError("No style-943 Body Text template found inside CH3")

    body = h_4_5.getparent()
    children = list(body)
    i_start = children.index(h_4_5) + 1
    i_end = children.index(h_ch5)

    to_delete = [
        c
        for c in children[i_start:i_end]
        if c.tag == QN_P and not _is_figure_block(c)
    ]
    print(f"DELETING: {len(to_delete)} existing §4.5 paragraph(s) (figures preserved)")
    for c in to_delete:
        body.remove(c)

    insert_after = h_4_5
    for i, text in enumerate(PARAGRAPHS, 1):
        p = deepcopy(body_tpl)
        set_paragraph_text(p, text)
        insert_after.addnext(p)
        insert_after = p
        print(f"  ¶{i}: {text[:80]}…")

    print(
        f"\nWROTE: {len(PARAGRAPHS)} paragraph(s) into §4.5 Synthesis and "
        "Limitations"
    )


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    print("=== §4.5 Synthesis and Limitations ===")
    write_section_body(doc)

    doc.save(str(path))
    print(f"\nSAVED: {path}")


if __name__ == "__main__":
    main()
