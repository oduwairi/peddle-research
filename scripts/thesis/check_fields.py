"""Determine whether the List of Tables / captions are live Word fields."""

from __future__ import annotations

import docx
from docx.oxml.ns import qn

DOC = "docs/research/THESIS.docx"


def field_info(p):
    el = p._p
    n_begin = len(el.findall(".//" + qn("w:fldChar")))
    instr = [t.text for t in el.findall(".//" + qn("w:instrText"))]
    n_seq = sum(1 for i in instr if i and "SEQ" in i)
    return n_begin, instr


def main() -> None:
    d = docx.Document(DOC)

    print("=== doc-wide field presence ===")
    body = d.element.body
    print("  <w:fldChar>:", len(body.findall(".//" + qn("w:fldChar"))))
    print("  <w:instrText>:", len(body.findall(".//" + qn("w:instrText"))))

    print("\n=== List-of-Tables paragraphs (style 'table of figures') ===")
    for p in d.paragraphs:
        if p.style.name == "table of figures" and "Table" in p.text:
            nb, instr = field_info(p)
            print(f"  fldChar={nb} instr={instr} :: {p.text[:60]!r}")

    print("\n=== Table CAPTIONS (Body Text 'Table 3.x: ...') — SEQ field? ===")
    for p in d.paragraphs:
        s = p.text.strip()
        if s.startswith("Table 3.") and p.style.name != "table of figures":
            nb, instr = field_info(p)
            print(f"  style={p.style.name} fldChar={nb} instr={instr} :: {s[:50]!r}")

    print("\n=== FIGURE captions (for comparison) ===")
    for p in d.paragraphs:
        s = p.text.strip()
        if s.startswith("Figure 3.") and p.style.name != "table of figures":
            nb, instr = field_info(p)
            print(f"  style={p.style.name} fldChar={nb} instr={instr} :: {s[:50]!r}")

    print("\n=== TOC region check (first 'toc' styled paras) ===")
    cnt = 0
    for p in d.paragraphs:
        if p.style.name.startswith("toc") and cnt < 4:
            nb, instr = field_info(p)
            print(f"  style={p.style.name} fldChar={nb} instr={instr[:1]} :: {p.text[:50]!r}")
            cnt += 1

    # is there a TOC/TOF field begin anywhere with instr text?
    print("\n=== all instrText values (unique kinds) ===")
    kinds = set()
    for it in body.findall(".//" + qn("w:instrText")):
        if it.text:
            kinds.add(it.text.strip()[:40])
    for k in sorted(kinds):
        print("   ", repr(k))


if __name__ == "__main__":
    main()
