"""Append the evaluation-justification paragraph (C.1) to sec 3.8.1 Goals.

Phase-5/6 polished author prose, added as a second paragraph in sec 3.8.1
(before the sec 3.8.2 heading) to strengthen the methodology's justification
(professor feedback #3: "described but not rigorously justified"). Index-free
anchor, idempotent. Citations (Matias 2021; Efron & Tibshirani 1993) get their
References entries + verification in the Phase-8 pass.
"""

from __future__ import annotations

import docx

DOC = "docs/research/THESIS.docx"

BODY = (
    "For a full-scope evaluation of the model performance, we need more than "
    "one signal to cross-reference and judge model performance. Our methodology "
    "uses multiple evaluation arms to answer different questions. Agreement or "
    "disagreement between arms is itself informative. Additionally, every "
    "evaluation arm is validated against real-world labels, such as real "
    "engagement or a real winning ad, which justifies the use of the evaluation "
    "arm. Specifically, our learned DeBERTa scorer is trained on true "
    "engagement signals using our proxy scoring system. The reliability of the "
    "model post-training is also checked and validated on a held-out set "
    "(rank-correlation, calibration; §3.7). Similarly, MAUVE compares embedding "
    "similarity against high-performing ads in the corpus. The reliability of "
    "this method is checked by measuring the similarity of the reference set to "
    "itself, expecting it to reach a high ceiling. The reference-overlap arm is "
    "compared against the real pool of winning ads. They are similarly "
    "validated in picking an A/B winner against the Upworthy Research Archive "
    "(Matias et al., 2021). Importantly, during evaluation, held-out test ads "
    "are removed from the reference pools, so nothing is scored against itself. "
    "When reporting results, model-to-model gaps are reported with uncertainty "
    "(bootstrap confidence intervals; Efron & Tibshirani, 1993)."
)


def main() -> None:
    d = docx.Document(DOC)

    def by_eq(s, style=None):
        for p in d.paragraphs:
            if p.text.strip() == s and (style is None or p.style.name == style):
                return p
        return None

    if any("we need more than one signal to cross-reference" in p.text
           for p in d.paragraphs):
        print("already inserted; nothing to do")
        return

    anchor = by_eq("3.8.2 Held-Out Test Split", style="Heading 3")
    if anchor is None:
        raise SystemExit("'3.8.2 Held-Out Test Split' heading not found")

    body_style = next(p for p in d.paragraphs if p.style.name == "Body Text").style

    b = d.add_paragraph()
    b.add_run(BODY)
    b.style = body_style
    anchor._p.addprevious(b._p)

    d.save(DOC)
    print("APPENDED: justification paragraph to §3.8.1 Goals (before §3.8.2)")


if __name__ == "__main__":
    main()
