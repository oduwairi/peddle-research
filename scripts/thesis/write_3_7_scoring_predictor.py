"""Write §3.7 Scoring Predictor body.

Same pattern as write_3_6_agent_architecture.py: wipes every paragraph between
the §3.7 heading and the §3.8 heading, then re-inserts the current PARAGRAPHS
list using a Body Text (style 943) template paragraph from CH3. New
bibliography entries (if any) are inserted in NEW_REFERENCES at their
alphabetical positions before APPENDIX, idempotent.

Idempotent: re-run safely as more paragraphs land in PARAGRAPHS.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"


PARAGRAPHS: list[str] = [
    # ¶1 — The gap v3 leaves behind
    "Our v3 proxy scoring system measures ad quality based on real-world signals and labels, including engagement and survivability. However, this doesn't generalize well to ads that have not run yet. A brand-new ad written by our model or a frontier model cannot be measured using this proxy scorer, which means we simply cannot use it for evaluation of our model. Instead, we need a new way to predict the performance of ads based on our pre-labeled corpus of 55k AdFlex ads. This would allow us to score any ad on the spot using a model that has trained on a ground-truth dataset.",
    # ¶2 — The idea, and why it's not circular
    "The idea is to train a small text regressor model (a fine-tuned DeBERTa-v3-base) that takes in ad copy text as input and outputs a proxy performance score. This regressor will learn to map ad copy text to the proxy scorer system, and since the label set is grounded in real-world numbers, the model would not be circular and the predictions would be grounded in what real high-performing ads would look like.",
    # ¶3 — Training data and label setup
    "To train the model we are using the same 55,000-ad corpus already obtained and scored. The model input is a pure, concatenated string with the platform tag, vertical tag, as well as the ad copy, including the headline, body, and description. The output to be predicted by the model represents the composite score, labels obtained using the original scorer, as well as the scores of the individual components contributing to the overall score, including the engagement volume, engagement velocity, and survivability. The model is trained to predict all four labels instead of relying only on the composite, which can give more insight and has many benefits with little risk. A filter that drops broken ads from the training set and penalizes clickbait by half is also applied. Some platforms with limited inputs because of API limitations are also considered, such as Reddit and the \"other\" platform bucket, which have no engagement numbers, so the engagement heads are not trained for those platforms.",
    # ¶4 — Model and training procedure
    "We use a small pretrained text model with four prediction outputs per scoring label. The model is chosen to be small enough to fit on a consumer RTX 3060 laptop GPU with six gigabytes of VRAM, and can be trained within an hour. A CPU is enough for inference on most consumer laptops. It uses standard supervised regression learning, with square error per head summed, and weighted by the training-quality scores. This ensures that high-quality ads contribute more to the model training. We use the standard AdamW optimizer algorithm. The pretrained body of the model has small learning rates (2e-5), since the model is pretrained and does not need to be tuned on the base weights; the four prediction heads get larger learning rates (1e-3) since they produce the desired output (and start from random initialisation). We use mixed precision (bfloat16 and float32) to make it fit on a six-gigabyte GPU.",
    # ¶5 — Train/val/test split
    "For training, the corpus is split 80 / 10 / 10 into training, validation, and test sets. The split is stratified by platform, so platforms are fairly distributed across the training, validation, and test sets. The three-set assignment is chosen to be deterministic and determined by the source ID of the ad, for reproducibility across re-runs. The original corpus is cleaned by dropping the rows with a low training-quality level as labeled by the GPT-4o-mini model. The remaining corpus contains around 43,000 ads. During training, the training set is used to drive the weight updates of the model, and the validation set is run every 250 steps to validate the performance during training and consider early stopping to avoid overfitting if the model stops improving. The final held-out test set is never seen by the model and is used to evaluate the model's performance at the end, after training completion.",
    # ¶6 — Evaluation methodology (closing paragraph)
    "As for evaluation of the trained regressor, a held-out set is used to calculate performance metrics. To give an overall idea of the model's performance, rank correlation (Spearman) is used to calculate the model's ability to rank ads in the same order as the original dataset. Perfect correlation corresponds to a score of one, which means the model orders ads identically to the original dataset; zero corresponds to completely random (fifty-fifty), meaning the model didn't learn anything. Mean absolute error is used to calculate how far numerically each prediction is from the real v3 label. Lower mean absolute error corresponds to better performance. AUC for the bottom and top tier ad sets (as defined in §3.3) is used to determine if the model can reliably classify high-performing and low-performing ads from the rest.",
]


# (alphabetical-anchor-prefix, full reference text)
NEW_REFERENCES: list[tuple[str, str]] = []


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def set_paragraph_text(p_elem, new_text):
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def write_section_body(doc):
    h_3_7 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("3.7 Scoring Predictor"),
        "§3.7 heading",
    )
    h_3_8 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("3.8 Evaluation Methodology"),
        "§3.8 heading",
    )

    # Body Text template: first style-943 paragraph inside CH3 (excluding
    # §3.7's own body, which we're about to wipe).
    body_tpl = None
    in_ch3 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.upper() == "CHAPTER III":
            in_ch3 = True
            continue
        if t.upper() == "CHAPTER IV":
            break
        if t.startswith("3.7 Scoring Predictor"):
            in_ch3 = "in-37"
            continue
        if in_ch3 == "in-37" and t.startswith("3.8 "):
            in_ch3 = True
            continue
        if in_ch3 == "in-37":
            continue
        sid = p.style.style_id if p.style else ""
        if in_ch3 and sid == "943" and t and p._element.findall(f".//{QN_T}"):
            body_tpl = p._element
            break
    if body_tpl is None:
        raise RuntimeError("No style-943 Body Text paragraph found inside CH3 (outside §3.7)")

    body = h_3_7.getparent()
    children = list(body)
    i_start = children.index(h_3_7) + 1
    i_end = children.index(h_3_8)

    to_delete = [c for c in children[i_start:i_end] if c.tag == QN_P]
    print(f"DELETING: {len(to_delete)} existing §3.7 body paragraphs")
    for c in to_delete:
        body.remove(c)

    insert_after = h_3_7
    for i, text in enumerate(PARAGRAPHS, 1):
        p = deepcopy(body_tpl)
        set_paragraph_text(p, text)
        insert_after.addnext(p)
        insert_after = p
        print(f"  ¶{i}: {text[:80]}…")

    print(f"\nWROTE: {len(PARAGRAPHS)} paragraph(s) into §3.7 Scoring Predictor")


def insert_references(doc):
    if not NEW_REFERENCES:
        print("(no new references for §3.7)")
        return

    refs_heading = None
    appendix_heading = None
    for p in doc.paragraphs:
        t = p.text.strip().upper()
        if t == "REFERENCES" and refs_heading is None:
            refs_heading = p._element
        elif t.startswith("APPENDIX") and appendix_heading is None:
            appendix_heading = p._element
    if refs_heading is None or appendix_heading is None:
        raise RuntimeError("References / APPENDIX bounds not found")

    body = refs_heading.getparent()
    children = list(body)
    i_refs = children.index(refs_heading)
    i_appendix = children.index(appendix_heading)
    ref_range = children[i_refs + 1 : i_appendix]

    ref_tpl = None
    for c in ref_range:
        if c.tag != QN_P:
            continue
        pPr = c.find(f"{{{W_NS}}}pPr")
        pStyle = pPr.find(f"{{{W_NS}}}pStyle") if pPr is not None else None
        sid = pStyle.get(f"{{{W_NS}}}val") if pStyle is not None else ""
        ts = c.findall(f".//{QN_T}")
        if sid == "943" and ts and any(t.text for t in ts):
            ref_tpl = c
            break
    if ref_tpl is None:
        raise RuntimeError("No style-943 reference template paragraph found")

    inserted = 0
    skipped = 0
    for anchor_prefix, ref_text in NEW_REFERENCES:
        first_words = ref_text.split("(")[0].strip().rstrip(",")
        already_present = any(
            c.tag == QN_P and "".join(t.text or "" for t in c.findall(f".//{QN_T}")).startswith(first_words)
            for c in ref_range
        )
        if already_present:
            print(f"  SKIP (already present): {ref_text[:80]}…")
            skipped += 1
            continue

        anchor = None
        for c in ref_range:
            if c.tag != QN_P:
                continue
            text = "".join(t.text or "" for t in c.findall(f".//{QN_T}"))
            if text.startswith(anchor_prefix):
                anchor = c
                break
        if anchor is None:
            raise RuntimeError(f"Alphabetical anchor not found: {anchor_prefix!r}")

        new_p = deepcopy(ref_tpl)
        set_paragraph_text(new_p, ref_text)
        anchor.addprevious(new_p)
        inserted += 1
        print(f"  INSERT before {anchor_prefix!r}: {ref_text[:80]}…")

        children = list(body)
        i_refs_new = children.index(refs_heading)
        i_appendix_new = children.index(appendix_heading)
        ref_range = children[i_refs_new + 1 : i_appendix_new]

    print(f"\nREFERENCES: inserted {inserted}, skipped {skipped} (already present)")


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    print("=== §3.7 body ===")
    write_section_body(doc)

    print("\n=== References ===")
    insert_references(doc)

    doc.save(str(path))
    print(f"\nSAVED: {path}")


if __name__ == "__main__":
    main()
