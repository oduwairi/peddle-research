"""Convert the two Chapter III loss equations to real Word math (OMML).

The plain-text unicode versions render with literal ``_{b,h}`` markup in Word
because LaTeX-style underscores are not interpreted. This rebuilds both as
``<m:oMathPara>`` display equations (render correctly in Word + OnlyOffice) and
rewords the two "where ..." gloss lines so they no longer use underscore/brace
subscript notation in body prose.

Idempotent: if a target paragraph already contains an ``<m:oMath>`` it is skipped.
"""

from __future__ import annotations

import sys

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

DOC = "docs/research/THESIS.docx"
M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
XMLSPACE = "{http://www.w3.org/XML/1998/namespace}space"


def _m(tag: str) -> str:
    return f"{{{M}}}{tag}"


def mr(s: str):
    """A math run <m:r><m:t>s</m:t></m:r>."""
    r = etree.Element(_m("r"))
    t = etree.SubElement(r, _m("t"))
    t.set(XMLSPACE, "preserve")
    t.text = s
    return r


def _wrap(tag: str, kids):
    e = etree.Element(_m(tag))
    for k in kids:
        e.append(k)
    return e


def ssub(base, sub):
    e = etree.Element(_m("sSub"))
    e.append(_wrap("e", base))
    e.append(_wrap("sub", sub))
    return e


def ssup(base, sup):
    e = etree.Element(_m("sSup"))
    e.append(_wrap("e", base))
    e.append(_wrap("sup", sup))
    return e


def frac(num, den):
    e = etree.Element(_m("f"))
    e.append(_wrap("num", num))
    e.append(_wrap("den", den))
    return e


def delim(kids, beg="(", end=")"):
    e = etree.Element(_m("d"))
    pr = etree.SubElement(e, _m("dPr"))
    b = etree.SubElement(pr, _m("begChr"))
    b.set(_m("val"), beg)
    en = etree.SubElement(pr, _m("endChr"))
    en.set(_m("val"), end)
    e.append(_wrap("e", kids))
    return e


def nary(sub_kids, body_kids):
    e = etree.Element(_m("nary"))
    pr = etree.SubElement(e, _m("naryPr"))
    etree.SubElement(pr, _m("chr")).set(_m("val"), "∑")
    etree.SubElement(pr, _m("limLoc")).set(_m("val"), "undOvr")
    etree.SubElement(pr, _m("subHide")).set(_m("val"), "0")
    etree.SubElement(pr, _m("supHide")).set(_m("val"), "1")
    e.append(_wrap("sub", sub_kids))
    e.append(_wrap("sup", []))  # empty (hidden)
    e.append(_wrap("e", body_kids))
    return e


def omath_para(children):
    """A centered <w:p> carrying an <m:oMathPara><m:oMath> equation."""
    p = etree.Element(qn("w:p"))
    ppr = etree.SubElement(p, qn("w:pPr"))
    etree.SubElement(ppr, qn("w:jc")).set(qn("w:val"), "center")
    # declare the math namespace on the oMathPara so descendants serialize as m:
    opar = etree.SubElement(p, _m("oMathPara"), nsmap={"m": M})
    omath = etree.SubElement(opar, _m("oMath"))
    for c in children:
        omath.append(c)
    return p


# --- equation 1: SFT assistant-only cross-entropy --------------------------
def build_sft():
    L = ssub([mr("ℒ")], [mr("SFT")])
    return [
        L,
        delim([mr("θ")]),
        mr(" = "),
        mr("−"),
        frac([mr("1")], [mr("|A|")]),
        mr(" "),
        nary(
            [mr("t ∈ A")],
            [
                mr("log "),
                ssub([mr("p")], [mr("θ")]),
                delim(
                    [
                        ssub([mr("y")], [mr("t")]),
                        mr(" | "),
                        ssub([mr("y")], [mr("<t")]),
                    ]
                ),
            ],
        ),
    ]


# --- equation 2: masked, quality-weighted multi-head MSE -------------------
def _m_nh():
    return ssub([mr("m")], [mr("n,h")])


def _q_n():
    return ssub([mr("q")], [mr("n")])


def _a_h():
    return ssub([mr("α")], [mr("h")])


def _yhat_nh():
    return ssub([mr("ŷ")], [mr("n,h")])


def _y_nh():
    return ssub([mr("y")], [mr("n,h")])


def _double_sum(tail_factors):
    """Σ_n Σ_h <factors>."""
    inner = nary([mr("h")], tail_factors)
    return nary([mr("n")], [inner])


def build_pred():
    sq = ssup([delim([_yhat_nh(), mr(" − "), _y_nh()])], [mr("2")])
    numer = _double_sum([_m_nh(), mr(" "), _q_n(), mr(" "), _a_h(), mr(" "), sq])
    denom = _double_sum([_m_nh(), mr(" "), _q_n(), mr(" "), _a_h()])
    return [
        mr("ℒ"),
        delim([mr("θ")]),
        mr(" = "),
        frac([numer], [denom]),
    ]


SFT_WHERE = (
    "Here A is the set of assistant-response token positions across the dataset "
    "and |A| their total count. The system-prompt and user-brief tokens are "
    "masked and contribute no loss (assistant-only loss masking), and the "
    "per-token factor is the model's next-token probability, so the loss is the "
    "mean negative log-likelihood over the assistant tokens."
)
PRED_WHERE = (
    "Here n indexes the ads in a batch and h the four prediction heads "
    "(composite, survivability, engagement volume, engagement velocity). The "
    "mask m is one for a valid ad–head pair and zero for the two engagement "
    "heads on weak-engagement platforms (Reddit and other); q is the per-ad "
    "training-quality weight; α is the per-head weight (all equal to one); and "
    "the prediction ŷ is sigmoid-bounded. The loss is therefore a masked, "
    "quality-weighted mean-squared error averaged over the valid elements."
)


def set_para_text(p_el, text):
    runs = p_el.findall(qn("w:r"))
    done = False
    for r in runs:
        t = r.find(qn("w:t"))
        if t is not None and not done:
            t.text = text
            t.set(XMLSPACE, "preserve")
            done = True
        elif t is not None:
            p_el.remove(r)
    return done


def replace_with_omath(eq_el, children) -> bool:
    if eq_el.find(".//" + _m("oMath")) is not None:
        return False  # already converted
    new_p = omath_para(children)
    eq_el.addprevious(new_p)
    eq_el.getparent().remove(eq_el)
    return True


def main() -> None:
    d = Document(DOC)
    ps = d.paragraphs
    log = []

    def find_eq(marker):
        for p in ps:
            if marker in p.text:
                return p._element
        return None

    # equation 1 + its where line (the immediate next paragraph)
    eq1 = find_eq("ℒ_SFT")
    if eq1 is not None:
        where1 = eq1.getnext()
        if replace_with_omath(eq1, build_sft()):
            log.append("  eq 3.5.3 -> OMML")
        if where1 is not None and "token positions" in "".join(
            t.text or "" for t in where1.iter(qn("w:t"))
        ):
            set_para_text(where1, SFT_WHERE)
            log.append("  reworded 3.5.3 'where' gloss")
    else:
        log.append("  (eq 3.5.3 already OMML — skipped)")

    # equation 2 + its where line
    ps = d.paragraphs  # refresh after structural change
    eq2 = find_eq("ŷ_{b,h}") or find_eq("m_{b,h}")
    if eq2 is not None:
        where2 = eq2.getnext()
        if replace_with_omath(eq2, build_pred()):
            log.append("  eq 3.7.4 -> OMML")
        if where2 is not None and "engagement heads" in "".join(
            t.text or "" for t in where2.iter(qn("w:t"))
        ):
            set_para_text(where2, PRED_WHERE)
            log.append("  reworded 3.7.4 'where' gloss")
    else:
        log.append("  (eq 3.7.4 already OMML — skipped)")

    d.save(DOC)
    print("CHANGES:")
    print("\n".join(log) if log else "  (none)")


if __name__ == "__main__":
    try:
        main()
    except AssertionError as e:
        print(f"ABORTED (no save): {e}", file=sys.stderr)
        sys.exit(1)
