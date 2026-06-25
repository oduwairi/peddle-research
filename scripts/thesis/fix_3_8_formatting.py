"""Transplant body-paragraph formatting onto the sec 3.8 inserts.

The C.1 / C.2 paragraphs were created with only a style assignment, so they
lack the first-line indent + line spacing the thesis applies as *direct*
formatting. Copy the <w:pPr> (and first plain-run <w:rPr>) from a neighbouring
real body paragraph. Field codes are NOT copied (only pPr + the first run's
rPr). Idempotent: re-running just re-applies identical formatting.
"""

from __future__ import annotations

import copy

import docx
from docx.oxml.ns import qn

DOC = "docs/research/THESIS.docx"


def first_text_run_rpr(p_el):
    for r in p_el.findall(qn("w:r")):
        if r.find(qn("w:t")) is not None:
            return r.find(qn("w:rPr"))
    return None


def transplant(target, model, *, do_rpr=True):
    tp, mp = target._p, model._p
    # pPr
    new_ppr = copy.deepcopy(mp.find(qn("w:pPr")))
    old_ppr = tp.find(qn("w:pPr"))
    if new_ppr is None:
        return
    if old_ppr is not None:
        tp.replace(old_ppr, new_ppr)
    else:
        tp.insert(0, new_ppr)
    # rPr of every run -> the model's first text-run rPr
    if do_rpr:
        mrpr = first_text_run_rpr(mp)
        if mrpr is not None:
            for r in tp.findall(qn("w:r")):
                newr = copy.deepcopy(mrpr)
                old = r.find(qn("w:rPr"))
                if old is not None:
                    r.replace(old, newr)
                else:
                    r.insert(0, newr)


def report(tag, p):
    pf = p.paragraph_format
    fr = p.runs[0].font if p.runs else None
    print(f"  {tag}: indent={pf.first_line_indent} spacing={pf.line_spacing} "
          f"align={pf.alignment} font={(fr.name, fr.size) if fr else None}")


def main() -> None:
    d = docx.Document(DOC)

    def find(sub, style=None):
        return next(
            (p for p in d.paragraphs
             if sub in p.text and (style is None or p.style.name == style)),
            None,
        )

    body_model = find("Model evaluation remains as the final and most critical step")
    head_model = next(p for p in d.paragraphs
                      if p.text.strip() == "3.8.5 MAUVE Arm")
    c1 = find("we need more than one signal to cross-reference")
    c2 = find("This marketing-domain-specific evaluation framework")
    c2h = next(p for p in d.paragraphs
               if p.text.strip() == "3.8.6 Reference-Overlap Metrics")

    if not all([body_model, head_model, c1, c2, c2h]):
        raise SystemExit("missing an anchor; aborting")

    print("BEFORE:")
    report("body_model", body_model)
    report("c1       ", c1)
    report("c2       ", c2)
    report("head_modl", head_model)
    report("c2_head  ", c2h)

    transplant(c1, body_model)
    transplant(c2, body_model)
    # heading: only fix if spacing/indent differ from the sibling heading
    if (c2h.paragraph_format.line_spacing != head_model.paragraph_format.line_spacing
            or c2h.paragraph_format.first_line_indent
            != head_model.paragraph_format.first_line_indent):
        transplant(c2h, head_model)
        print("  (heading formatting also transplanted)")
    else:
        print("  (heading formatting already matches sibling — left as is)")

    d.save(DOC)

    d2 = docx.Document(DOC)
    def find2(sub):
        return next(p for p in d2.paragraphs if sub in p.text)
    print("AFTER:")
    report("c1       ", find2("we need more than one signal to cross-reference"))
    report("c2       ", find2("This marketing-domain-specific evaluation framework"))
    report("body_model", find2("Model evaluation remains as the final"))


if __name__ == "__main__":
    main()
