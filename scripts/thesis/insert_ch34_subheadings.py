"""Insert numbered Heading 3 subheadings into Chapters III and IV of THESIS.docx.

Anchors are first-N-character prefixes of the paragraph that the subheading should
sit *immediately before*. Insertion is scoped per §3.x / §4.x section: an anchor
only fires when the previous Heading 2 matches the section title prefix, so a
prefix that happens to also appear elsewhere will not produce a stray heading.

Run from repo root:

    uv run python scripts/thesis/insert_ch34_subheadings.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from lxml import etree

THESIS = Path("docs/research/THESIS.docx")
# python-docx can't resolve "Heading 3" by name in this template (numeric styleIds);
# reference the styleId directly. 882 is the Heading 3 paragraph style in styles.xml.
HEADING3_STYLE_ID = "882"

# (section_h2_prefix, anchor_para_prefix, heading_text)
PLAN: list[tuple[str, str, str]] = [
    # 3.1
    ("3.1 Proposed System", "The proposed marketing AI agent system takes a marketing brief", "3.1.1 Input and Output"),
    ("3.1 Proposed System", "The proposed system is mainly developed over three main phases", "3.1.2 Three-Phase Pipeline"),
    ("3.1 Proposed System", "The proposed agent system architecture involves a two-role", "3.1.3 Two-Role Agent System"),
    # 3.2
    ("3.2 Data Acquisition", "For the ad collection strategy, we followed a multi-sweep", "3.2.1 Collection Strategy"),
    ("3.2 Data Acquisition", "In order to produce a fine-tuned model capable of generating", "3.2.2 Rationale for High-Performing Ads"),
    ("3.2 Data Acquisition", "The collection engine supports five main sweep configurations", "3.2.3 Sweep Configurations"),
    ("3.2 Data Acquisition", "Following the ad collection and sweeping, a normalization step", "3.2.4 Normalization"),
    # 3.3
    ("3.3 Engagement-Based Scoring", "The main limitation of AdFlex ads is the fact that they don", "3.3.1 Proxy Scoring Rationale"),
    ("3.3 Engagement-Based Scoring", "Engagement volume refers to the total bulk volume of the engagement", "3.3.2 Engagement Volume and Velocity"),
    ("3.3 Engagement-Based Scoring", "As for lifespan, we use a Kaplan", "3.3.3 Survivability"),
    ("3.3 Engagement-Based Scoring", "These three signals are then combined into a single composite", "3.3.4 Composite Score and Tiers"),
    ("3.3 Engagement-Based Scoring", "A developed scoring system uses proxies, not absolute signals", "3.3.5 Validation Against Ground Truth"),
    # 3.4
    ("3.4 Training-Data Construction via Backtranslation", "Following data collection, the final AdFlex corpus obtained", "3.4.1 Corpus and Target Size"),
    ("3.4 Training-Data Construction via Backtranslation", "Before feeding data from the collected corpus directly to the teacher", "3.4.2 Pre-Teacher Filters"),
    ("3.4 Training-Data Construction via Backtranslation", "Following the filters and finalization of the ad set", "3.4.3 Teacher Prompt"),
    ("3.4 Training-Data Construction via Backtranslation", "Before ingesting the teacher", "3.4.4 Ingestion Gates"),
    ("3.4 Training-Data Construction via Backtranslation", "The final training dataset example structure follows", "3.4.5 Training Example Schema"),
    # 3.5
    ("3.5 Fine-Tuning", "The finalized dataset contains approximately three thousand", "3.5.1 Dataset and Splits"),
    ("3.5 Fine-Tuning", "The setup phase includes shrinking the base model weights", "3.5.2 QLoRA Setup"),
    ("3.5 Fine-Tuning", "During training, each training example is given to the model", "3.5.3 Training Inputs"),
    ("3.5 Fine-Tuning", "The training is set to run for three epochs", "3.5.4 Optimization"),
    ("3.5 Fine-Tuning", "After training completes, the LoRA adapter stays separate", "3.5.5 Adapter Merging and Hosting"),
    # 3.6
    ("3.6 Agent Architecture", "Our fine-tuned model specializes in ad copywriting", "3.6.1 Motivation"),
    ("3.6 Agent Architecture", "The agent system consists of two main models", "3.6.2 Two-Role Design"),
    ("3.6 Agent Architecture", "When the orchestrator and writer model are set, the agent workflow", "3.6.3 Freeform Loop"),
    ("3.6 Agent Architecture", "The research phase represents the forefront", "3.6.4 Research Phase"),
    ("3.6 Agent Architecture", "Communication between the orchestrator and the writer", "3.6.5 Orchestrator–Writer Tools"),
    # 3.7
    ("3.7 Scoring Predictor", "Our v3 proxy scoring system measures ad quality", "3.7.1 Motivation"),
    ("3.7 Scoring Predictor", "The idea is to train a small text regressor", "3.7.2 Regressor Architecture"),
    ("3.7 Scoring Predictor", "To train the model we are using the same 55,000-ad", "3.7.3 Training Data and Input Format"),
    ("3.7 Scoring Predictor", "We use a small pretrained text model with four prediction outputs", "3.7.4 Model Size and Hardware"),
    ("3.7 Scoring Predictor", "For training, the corpus is split 80", "3.7.5 Splits"),
    ("3.7 Scoring Predictor", "As for evaluation of the trained regressor", "3.7.6 Evaluation Metrics"),
    # 3.8
    ("3.8 Evaluation Methodology", "Model evaluation remains as the final and most critical", "3.8.1 Goals"),
    ("3.8 Evaluation Methodology", "For the evaluation, the held-out test split", "3.8.2 Held-Out Test Split"),
    ("3.8 Evaluation Methodology", "Before feeding the text into the scorers", "3.8.3 Copy Normalization"),
    ("3.8 Evaluation Methodology", "With our obtained scoring predictor model, every clean ad text", "3.8.4 Learned-Scorer Arm"),
    ("3.8 Evaluation Methodology", "As a second technique to evaluate the model", "3.8.5 MAUVE Arm"),
    ("3.8 Evaluation Methodology", "Gold ads compared against the baseline", "3.8.6 Gold Baseline"),
    # 4.1
    ("4.1 Evaluation Setup", "The evaluation uses the 215 held-out briefs", "4.1.1 Held-Out Briefs"),
    ("4.1 Evaluation Setup", "The evaluation pipeline uses five main configurations", "4.1.2 Configurations"),
    ("4.1 Evaluation Setup", "In our methodology, we have two evaluation arms", "4.1.3 Evaluation Arms"),
    # 4.2
    ("4.2 Learned-Scorer Absolute Scores", "This section reports the findings of the first evaluation arm", "4.2.1 Composite Scores"),
    ("4.2 Learned-Scorer Absolute Scores", "B_pipe and C_pipe are the agent-integrated versions", "4.2.2 Agent-Integrated Variants"),
    ("4.2 Learned-Scorer Absolute Scores", "Now focusing on our fine-tuned model", "4.2.3 Per-Platform Scores"),
    ("4.2 Learned-Scorer Absolute Scores", "We have looked at the composite scores", "4.2.4 Per-Head Scores"),
    ("4.2 Learned-Scorer Absolute Scores", "Given these results, it is important to address the reliability", "4.2.5 Predictor Reliability"),
    # 4.3
    ("4.3 MAUVE Distribution Matching", "The second arm, which is the MAUVE arm", "4.3.1 Arm Setup"),
    ("4.3 MAUVE Distribution Matching", "The final score rankings are as follows", "4.3.2 Score Rankings"),
    # 4.4
    ("4.4 Fine-Tuning and Agent Ablation", "This section reports the 2", "4.4.1 2×2 Design"),
    ("4.4 Fine-Tuning and Agent Ablation", "For paired contrast, we simply cannot just subtract", "4.4.2 Paired Contrasts"),
    ("4.4 Fine-Tuning and Agent Ablation", "The test-set brief shape is similar to the training examples", "4.4.3 Brief Shape"),
    ("4.4 Fine-Tuning and Agent Ablation", "Although the agent results are consistently lower", "4.4.4 Agent Scoring Caveats"),
    # 4.5
    ("4.5 Synthesis and Limitations", "We recap our evaluation procedures", "4.5.1 Recap"),
    ("4.5 Synthesis and Limitations", "We provide a deeper dive into what each evaluation arm actually measures", "4.5.2 Interpretation"),
    ("4.5 Synthesis and Limitations", "Having discussed the interpretation of the results", "4.5.3 Limitations"),
]


def main() -> None:
    doc: DocumentObject = Document(str(THESIS))

    # Walk paragraphs once, remember the most recent Heading 2 title, and for each
    # plan entry whose section matches AND whose anchor prefix matches the paragraph,
    # collect (anchor_para, heading_text) targets. We then insert in reverse document
    # order so element-index churn doesn't matter.
    current_section: str | None = None
    targets: list[tuple[object, str]] = []  # (anchor_xml_element, heading_text)
    consumed_plan_indices: set[int] = set()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name == "Heading 2" and (text.startswith("3.") or text.startswith("4.")):
            current_section = text
            continue
        if current_section is None:
            continue
        for k, (sec_prefix, anchor_prefix, heading_text) in enumerate(PLAN):
            if k in consumed_plan_indices:
                continue
            if not current_section.startswith(sec_prefix):
                continue
            if text.startswith(anchor_prefix):
                targets.append((para._element, heading_text))
                consumed_plan_indices.add(k)
                break  # one plan entry per paragraph

    missing = [PLAN[k] for k in range(len(PLAN)) if k not in consumed_plan_indices]
    if missing:
        print("MISSING ANCHORS:")
        for sec, anc, head in missing:
            print(f"  [{sec}] {head!r}  anchor prefix not found: {anc!r}")
        raise SystemExit(1)

    # Insert in reverse order so we never invalidate an earlier target's element.
    for anchor_el, heading_text in reversed(targets):
        new_para = doc.add_paragraph(heading_text)
        pPr = new_para._element.get_or_add_pPr()
        # Drop any default style python-docx may have set, then write our own pStyle.
        for existing in pPr.findall(qn("w:pStyle")):
            pPr.remove(existing)
        pStyle = etree.SubElement(pPr, qn("w:pStyle"))
        pStyle.set(qn("w:val"), HEADING3_STYLE_ID)
        # pStyle must be the first child of pPr per the OOXML schema.
        pPr.insert(0, pStyle)
        anchor_el.addprevious(new_para._element)

    doc.save(str(THESIS))
    print(f"Inserted {len(targets)} subheadings into {THESIS}")


if __name__ == "__main__":
    main()
