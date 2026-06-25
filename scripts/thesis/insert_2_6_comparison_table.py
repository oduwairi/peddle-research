"""Insert Table 2.6 — the related-work comparison matrix — into THESIS.docx.

Reviewer feedback #11: "Add a table showing a comparison of your work. Add more
related projects that other researchers have done, at least 10 or 15 other
research works with citations. This should serve as a benchmark to prove that
your project is better than other previous researchers' projects."

Audit (docs/research/THESIS_EDITING.md workflow): Chapter II already carries five
comparison tables (2.1-2.5), but none place *this work* as a row against a broad
set of related projects. The Chapter II Conclusion states the gap in prose ("No
existing work fine-tunes an open-source sub-10B model for marketing campaign
generation") with no visual companion. This script inserts a consolidated
capability matrix (Table 2.6) immediately after that gap paragraph: 15 related
research works + a final "This work (Draper)" row, across the dimensions on which
the thesis's contribution rests. Honest framing — a capability/approach matrix,
not a fabricated numeric leaderboard (competing systems use different data,
languages, and metrics). The decisive, defensible signal is that Draper is the
only row that is BOTH an open-weights sub-10B fine-tuned model AND
marketing-specialized — exactly the existing Conclusion claim.

Every citation in the table is already present in the bibliography, so no new
References entries are required. Cell values are facts verified against the cited
papers; "—" / "Quality only" / "Partial" are used conservatively where a source
does not report a live outcome or releases no model.

The caption is cloned from the Table 2.3 caption so it inherits the Chapter II
caption style (Body, style 948) AND its embedded `TC "..." \\f T \\l 1` field —
the field is what the List-of-Tables field collects on refresh. A minimal
"(Tables 2.3 and 2.6)" cross-ref is added to the gap paragraph (reviewer #5
pattern). No new body prose is authored.

Idempotent: a prior Table 2.6 block (caption + table + spacer + note) is removed
before re-inserting; the callout is skipped if "2.6" is already present. Anchors
are located by unique text, never hardcoded indices.

Run:  uv run python scripts/thesis/insert_2_6_comparison_table.py [--dry-run]
"""

from __future__ import annotations

import copy
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DOCX = Path("docs/research/THESIS.docx")
LOCK = Path("docs/research/.~lock.THESIS.docx#")

CAPTION_SRC_PREFIX = "Table 2.3 — Domain-Specialized"
NOTE_SRC_PREFIX = "Note. “—” indicates data not reported in the original paper. RLHF"
GAP_ANCHOR = "This survey has examined the research landscape"
OLD_TITLE_2_3 = "Domain-Specialized Small Models: Training Approach and Performance"

BOOKMARK_ID = "30071"  # max existing id is 30070
BOOKMARK_NAME = "Tbl_2_6"

TITLE = ("Comparison of this work against related domain-specialized and "
         "marketing language-model systems")

NOTE = (
    "Note. “Open wts” = downloadable model weights; “— (data)” / "
    "“— (method)” = not a released model (a dataset, benchmark, or training "
    "method). “Real-perf grounded” = trained or selected on real measured ad "
    "outcomes (CTR, engagement, or longevity) rather than human or heuristic quality "
    "labels; “Quality only” = human or heuristic quality scores, not live "
    "outcomes. CTR = click-through rate; FT = fine-tuned; RAG = retrieval-augmented "
    "generation; SERP = search-engine results page."
)

HEADERS = ["Work", "Focus", "Type", "Open wts", "Mktg",
           "Real-perf grounded", "Agentic / RAG", "Reference"]

ROWS = [
    ["Stanford Alpaca", "General instruction", "Open FT model (LLaMA 7B)",
     "Yes", "No", "No", "No", "Taori et al. (2023)"],
    ["Med42", "Medical", "Open FT model (Llama-3 8/70B)",
     "Yes", "No", "No", "No", "Christophe et al. (2024)"],
    ["AlpaCare", "Medical", "Open FT model (LLaMA 7/13B)",
     "Yes", "No", "No", "No", "Zhang et al. (2023a)"],
    ["FinGPT", "Finance", "Open FT framework (7B)",
     "Yes", "No", "No", "Partial", "Wang et al. (2023)"],
    ["MarketingFM", "E-commerce ads", "Hosted FT (proprietary)",
     "No", "Yes", "Yes (live CTR)", "Yes", "Liu, H., Tahmasbi et al. (2025)"],
    ["Lops et al.", "Fashion ad copy", "Hosted FT (GPT-4o-mini)",
     "No", "Yes", "Quality only", "No", "Lops et al. (2025)"],
    ["Reisenbichler et al.", "Content / SEO marketing", "NLG pipeline (API)",
     "No", "Yes", "Yes (SERP rank)", "No", "Reisenbichler et al. (2022)"],
    ["Reisenbichler et al.", "Sponsored search ads", "LLM field study (API)",
     "No", "Yes", "Yes (field CTR/CPC)", "No", "Reisenbichler et al. (2025)"],
    ["LOLA", "Content experiments", "LLM + online learning",
     "Partial", "Yes", "Yes (A/B CTR)", "Partial", "Ye et al. (2025)"],
    ["Matz et al.", "Personalized persuasion", "NLG study (ChatGPT)",
     "No", "Yes", "Quality only", "No", "Matz et al. (2024)"],
    ["AdTEC", "Search ad text (JP)", "Public benchmark",
     "— (data)", "Yes", "Quality only", "No", "Zhang et al. (2025b)"],
    ["CAMERA", "Ad text (JP)", "Public dataset",
     "— (data)", "Yes", "Quality only", "No", "Mita et al. (2024)"],
    ["LCTG-Bench", "Controllable text (JP)", "Public benchmark",
     "— (data)", "Partial", "No", "No", "Kurihara et al. (2025)"],
    ["AD-Bench", "Ad-analytics agents", "Public benchmark (agents)",
     "— (data)", "Yes", "Yes (live platform)", "Yes", "Hu et al. (2026)"],
    ["Instruction Backtranslation (Humpback)", "General instruction",
     "Training method", "— (method)", "No", "No", "No", "Li et al. (2024)"],
    ["This work (Draper)", "Multi-platform ad copywriting",
     "Open FT model (Qwen3-8B) + agent", "Yes", "Yes",
     "Yes (engagement + survival)", "Yes (live web)", "This work"],
]


# --------------------------------------------------------------------------- #
def ptext(p_el) -> str:
    return "".join(t.text or "" for t in p_el.findall(".//" + qn("w:t")))


def pstyle(p_el) -> str | None:
    pr = p_el.find(qn("w:pPr"))
    if pr is not None:
        s = pr.find(qn("w:pStyle"))
        if s is not None:
            return s.get(qn("w:val"))
    return None


def find_para(doc: Document, prefix: str, style: str | None = None) -> Paragraph | None:
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix) and (style is None or pstyle(p._p) == style):
            return p
    return None


def set_cell(cell, text: str, bold: bool) -> None:
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(text)
    if bold:
        run.bold = True


def set_table_full_width(tbl) -> None:
    tblPr = tbl._tbl.tblPr
    for w in tblPr.findall(qn("w:tblW")):
        tblPr.remove(w)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")  # 100%
    tblPr.append(tblW)


def build_caption(src_caption: Paragraph) -> "object":
    """Clone the Table 2.3 caption element; rewrite bookmark, visible text, TC field."""
    el = copy.deepcopy(src_caption._p)

    bm_start = el.find(qn("w:bookmarkStart"))
    if bm_start is not None:
        bm_start.set(qn("w:id"), BOOKMARK_ID)
        bm_start.set(qn("w:name"), BOOKMARK_NAME)
    for bm_end in el.findall(qn("w:bookmarkEnd")):
        bm_end.set(qn("w:id"), BOOKMARK_ID)

    for r in el.findall(qn("w:r")):
        t = r.find(qn("w:t"))
        if t is not None and t.text:
            if t.text.strip().startswith("Table 2.3"):
                t.text = "Table 2.6 "
            elif t.text.strip() == OLD_TITLE_2_3:
                t.text = TITLE
        instr = r.find(qn("w:instrText"))
        if instr is not None and instr.text and "TC " in instr.text:
            instr.text = f' TC "Table 2.6 — {TITLE}" \\f T \\l 1 '
    return el


def remove_prior(doc: Document) -> int:
    """Remove a previously inserted Table 2.6 block (caption + table + spacer + note)."""
    removed = 0
    for p in list(doc.paragraphs):
        if pstyle(p._p) == "948" and p.text.strip().startswith("Table 2.6"):
            el = p._p
            # remove following siblings: tbl, empty 948, note 948 (bounded, conditional)
            for _ in range(3):
                nxt = el.getnext()
                if nxt is None:
                    break
                tag = nxt.tag.split("}")[-1]
                if tag == "tbl":
                    nxt.getparent().remove(nxt)
                elif tag == "p" and pstyle(nxt) == "948" and (
                    not ptext(nxt).strip() or ptext(nxt).strip().startswith("Note.")
                ):
                    nxt.getparent().remove(nxt)
                else:
                    break
            el.getparent().remove(el)
            removed += 1
    return removed


def add_callout(doc: Document) -> str:
    p = find_para(doc, GAP_ANCHOR)
    if p is None:
        return "!! gap paragraph not found — callout skipped"
    if "2.6" in p.text:
        return "== callout already present"
    runs = p.runs
    done = False
    for i, r in enumerate(runs):
        if r.text and r.text.rstrip().endswith("(Table 2") and "marketing campaign generation" in r.text:
            r.text = r.text.replace("(Table 2", "(Tables 2")
            nxt = runs[i + 1] if i + 1 < len(runs) else None
            if nxt is not None and nxt.text.startswith(".3)"):
                nxt.text = ".3 and 2.6)" + nxt.text[3:]
                done = True
            break
    return "++ callout: (Table 2.3) -> (Tables 2.3 and 2.6)" if done else "!! callout anchor not matched"


def main() -> int:
    dry = "--dry-run" in sys.argv
    if not DOCX.exists():
        print(f"ERROR: {DOCX} not found", file=sys.stderr)
        return 1
    if LOCK.exists():
        print("ERROR: THESIS.docx is open in an editor (lock present). Close it.", file=sys.stderr)
        return 1

    if dry:
        print("=== DRY RUN — Table 2.6 ===\n")
        print(f"Caption: Table 2.6 — {TITLE}")
        print("Anchor : after paragraph starting", repr(GAP_ANCHOR))
        print("\n  " + " | ".join(HEADERS))
        for r in ROWS:
            print("  " + " | ".join(r))
        print("\n" + NOTE)
        return 0

    doc = Document(str(DOCX))

    src_caption = find_para(doc, CAPTION_SRC_PREFIX, style="948")
    if src_caption is None:
        print("ERROR: caption template (Table 2.3) not found", file=sys.stderr)
        return 2
    if src_caption._p.find(".//" + qn("w:drawing")) is not None:
        print("ERROR: caption template carries a drawing — unsafe to clone", file=sys.stderr)
        return 2
    src_note = find_para(doc, NOTE_SRC_PREFIX, style="948")
    if src_note is None:
        print("ERROR: note template not found", file=sys.stderr)
        return 2
    anchor = find_para(doc, GAP_ANCHOR)
    if anchor is None:
        print("ERROR: gap anchor paragraph not found", file=sys.stderr)
        return 2

    removed = remove_prior(doc)
    print(f"Removed {removed} prior Table 2.6 block(s).\n")

    # caption
    cap_el = build_caption(src_caption)
    anchor._p.addnext(cap_el)

    # table (built at end of body, then moved after the caption)
    tbl = doc.add_table(rows=len(ROWS) + 1, cols=len(HEADERS))
    tbl.style = doc.styles["Table Grid"]
    tbl.autofit = True
    set_table_full_width(tbl)
    for ci, h in enumerate(HEADERS):
        set_cell(tbl.rows[0].cells[ci], h, bold=True)
    for ri, row in enumerate(ROWS, start=1):
        bold = row[0].startswith("This work")
        for ci, val in enumerate(row):
            set_cell(tbl.rows[ri].cells[ci], val, bold=bold)
    cap_el.addnext(tbl._tbl)

    # spacer (empty 948) + note (948), cloned from the Table 2.3 note paragraph
    spacer_el = copy.deepcopy(src_note._p)
    for r in list(spacer_el.findall(qn("w:r"))):
        spacer_el.remove(r)
    note_el = copy.deepcopy(src_note._p)
    runs = note_el.findall(qn("w:r"))
    if runs:
        first_t = runs[0].find(qn("w:t"))
        if first_t is None:
            first_t = OxmlElement("w:t")
            runs[0].append(first_t)
        first_t.set(qn("xml:space"), "preserve")
        first_t.text = NOTE
        for r in runs[1:]:
            note_el.remove(r)
    tbl._tbl.addnext(spacer_el)
    spacer_el.addnext(note_el)

    print("INSERTIONS:")
    print(f"  ++ Caption  : Table 2.6 — {TITLE[:50]}...")
    print(f"  ++ Table    : {len(ROWS)+1}x{len(HEADERS)} (Table Grid, header + Draper row bold)")
    print(f"  ++ Note     : {NOTE[:50]}...")
    print("  " + add_callout(doc))

    doc.save(str(DOCX))
    print(f"\nSaved {DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
