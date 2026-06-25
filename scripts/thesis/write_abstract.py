"""Write the Abstract body.

Same pattern as write_3_3_scoring.py: wipes every paragraph between the
'Abstract' heading and the 'Özet' heading, then re-inserts the current
PARAGRAPHS list using the existing style-944 abstract body paragraph as the
clone template. Idempotent — re-running produces the same result.

No reference insertions: abstracts are self-contained, no citations.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"


PARAGRAPHS: list[str] = [
    # Abstract body — single continuous paragraph (standard academic format).
    (
        "With the rapid development of AI technology and domain-specialized agents, they have seen extensive applications in marketing and advertising domains. "
        "However, most LLMs and models off the shelf often produce generic ad copy and campaigns that tend to underperform in real-world deployments. "
        "Existing research does not establish whether a fine-tuned 8B model can match or surpass large proprietary models on the marketing domain vertical, "
        "nor whether wrapping them with an agent workflow and live RAG can improve the performance of the generated ads. "
        "This research aims to address these gaps by asking whether an 8B domain-specialized model can outperform a frontier model when generating ad copy, "
        "and additionally to understand the cumulative effect of fine-tuning and live RAG and the influence of each component separately. "
        "This research uses a 55,000-ad corpus scraped from AdFlex across five platforms (Facebook, TikTok, Pinterest, Twitter, Reddit), "
        "each scored using our proxy scoring system: per-platform Kaplan-Meier survival curves over campaign longevity, combined with continuous engagement signals (likes, shares, comments). "
        "Top-performing ads in the pile are used as training material. "
        "The training construction pipeline involves creating brief-answer pairs using teacher models that the student will learn to replicate; "
        "the pipeline uses the instruction backtranslation technique. "
        "The student model is Qwen3-8B, fine-tuned with QLoRA. "
        "The resulting fine-tuned writer is wrapped in an agent workflow where an orchestrator handles general queries and tool calls, "
        "and the fine-tuned model acts as the writer which is called when creative tasks are involved. "
        "The evaluation methodology uses a 215-brief held-out test set from the construction set, stratified across platforms. "
        "Five configurations are compared as well as GOLD ads, which are the actual high-performing ads from the set; "
        "configurations include a frontier model as well as fine-tuning and RAG combinations (4 combinations). "
        "Two evaluation arms are designed, which include a DeBERTa-v3-base regressor scorer model trained on engagement labels, "
        "as well as a corpus-level MAUVE distribution comparison based on embedding overlap. "
        "All numbers are reported with 95% bootstrap confidence intervals. "
        "Results show per-ad composite scores: GOLD 0.684, C 0.651, B 0.611, C_pipe 0.607, A 0.603, B_pipe 0.586. "
        "The fine-tuned model outperforms the other four configurations and lags behind the GOLD reference, which indicates a positive answer to RQ1. "
        "Corpus-level MAUVE results show GOLD 0.462, C_pipe 0.420, B_pipe 0.302, C 0.287, B 0.183, A 0.180. "
        "Here again, our agent-wrapped fine-tuned model beats all other configurations by a considerable margin. "
        "The two evaluation arms produce contradictory results about the influence of agent-wrapped workflows. "
        "Contributions of this research include a proxy scoring system based on engagement and survivability signals, "
        "a fine-tuned open-source model trained on successful ad data, and a fully-deployed end-to-end marketing agent."
    ),
    # Keywords line.
    "Keywords: domain-specialized LLM, ad copy generation, instruction backtranslation, QLoRA fine-tuning, retrieval-augmented generation, engagement-based scoring, marketing agent.",
]


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def set_paragraph_text(p_elem, new_text):
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def write_abstract(doc):
    h_abstract = find_paragraph(
        doc,
        lambda p: p.text.strip() == "Abstract" and (p.style.style_id if p.style else "") == "945",
        "'Abstract' heading",
    )
    h_ozet = find_paragraph(
        doc,
        lambda p: p.text.strip() == "Özet" and (p.style.style_id if p.style else "") == "945",
        "'Özet' heading",
    )

    body = h_abstract.getparent()
    children = list(body)
    i_start = children.index(h_abstract) + 1
    i_end = children.index(h_ozet)

    # Body template: first style-944 paragraph with text between Abstract and Özet,
    # otherwise fall back to any style-944 paragraph with text in the doc.
    body_tpl = None
    for c in children[i_start:i_end]:
        if c.tag != QN_P:
            continue
        pPr = c.find(f"{{{W_NS}}}pPr")
        pStyle = pPr.find(f"{{{W_NS}}}pStyle") if pPr is not None else None
        sid = pStyle.get(f"{{{W_NS}}}val") if pStyle is not None else ""
        ts = c.findall(f".//{QN_T}")
        if sid == "944" and ts:
            body_tpl = deepcopy(c)
            break
    if body_tpl is None:
        # Fallback: search whole doc for any style-944 paragraph with <w:t>.
        for p in doc.paragraphs:
            sid = p.style.style_id if p.style else ""
            ts = p._element.findall(f".//{QN_T}")
            if sid == "944" and ts:
                body_tpl = deepcopy(p._element)
                break
    if body_tpl is None:
        raise RuntimeError("No style-944 body paragraph found to clone")

    # Wipe everything between Abstract and Özet headings.
    to_delete = [c for c in children[i_start:i_end] if c.tag == QN_P]
    print(f"DELETING: {len(to_delete)} existing paragraph(s) between Abstract and Özet")
    for c in to_delete:
        body.remove(c)

    insert_after = h_abstract
    for i, text in enumerate(PARAGRAPHS, 1):
        p = deepcopy(body_tpl)
        set_paragraph_text(p, text)
        insert_after.addnext(p)
        insert_after = p
        print(f"  ¶{i}: {text[:80]}…")

    print(f"\nWROTE: {len(PARAGRAPHS)} paragraph(s) into Abstract")


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    print("=== Abstract body ===")
    write_abstract(doc)

    doc.save(str(path))
    print(f"\nSAVED: {path}")


if __name__ == "__main__":
    main()
