"""Add in-text callouts for every figure/table that lacks one.

Reviewer feedback #5: "Call out or reference ALL figures and tables
accordingly in the text." Chapter II artifacts (Figs 2.2/2.7, Tables
2.1-2.5) are already referenced; every figure/table authored in
Chapters III-V (15 in total) had a caption but no in-text callout.

This script appends a minimal parenthetical cross-reference, e.g.
" (Figure 4.2)", to the *adjacent discussion paragraph* for each one.
No new sentences are written; the author's wording is untouched.

Each target paragraph is located by a unique text ANCHOR (never a
hardcoded index, per docs/research/THESIS_EDITING.md). The append is
idempotent: a paragraph that already carries its "(Figure X)" callout
is skipped, so the script is safe to re-run.

Run:  uv run python scripts/thesis/add_figure_callouts.py [--dry-run]
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

DOCX = Path("docs/research/THESIS.docx")

# (label, anchor substring identifying the discussion paragraph)
# Order = document order, so the INSERTIONS summary reads top-to-bottom.
CALLOUTS: list[tuple[str, str]] = [
    # --- Chapter II (lit-review survey methodology) ---
    ("Figure 1.1", "an AI-powered systematic search was used"),
    # --- Chapter III ---
    ("Figure 3.1", "developed over three main phases"),
    ("Figure 3.2", "we followed a multi-sweep strategy"),
    ("Figure 3.3", "The construction phase follows this"),
    ("Table 3.1", "shrinking the base model weights into a four-bit"),
    ("Figure 3.4", "The agent system consists of two main models"),
    ("Table 3.2", "As a second technique to evaluate"),
    # --- Chapter IV ---
    ("Figure 4.1", "reports the findings of the first evaluation arm"),
    ("Figure 4.2", "yields composite scores per platform"),
    ("Figure 4.3", "Looking at the three component heads"),
    ("Figure 4.4", "it is important to address the reliability of the trained predictor"),
    ("Figure 4.3.1", "uses embedding-based comparison between ad pools"),
    ("Figure 4.3.2", "The final score rankings are as follows"),
    ("Figure 4.4.1", "we simply cannot just subtract cell scores"),
    ("Figure 4.4.2", "The two-by-two design examines the four combinations"),
    # --- Chapter V ---
    ("Figure 5.1", "one important question was how it would be fine-tuned"),
]


def is_caption(text: str) -> bool:
    t = text.lstrip()
    return t.startswith("Figure") or t.startswith("Table")


def find_paragraph(doc: Document, anchor: str) -> list[tuple[int, Paragraph]]:
    hits = []
    for i, p in enumerate(doc.paragraphs):
        if anchor in p.text and not is_caption(p.text):
            hits.append((i, p))
    return hits


def append_callout(p: Paragraph, label: str) -> str:
    """Append ' (label)' before the paragraph's terminal period.

    Operates on the last run carrying non-whitespace text, so run
    properties (rPr) and the run structure are preserved.
    """
    runs = p.runs
    idx = None
    for k in range(len(runs) - 1, -1, -1):
        if runs[k].text and runs[k].text.strip():
            idx = k
            break
    insert = f" ({label})"
    if idx is None:
        p.add_run(insert)
        return "(appended new run — paragraph had no text run)"
    r = runs[idx]
    t = r.text
    stripped = t.rstrip()
    trail = t[len(stripped):]
    if stripped.endswith("."):
        r.text = stripped[:-1] + insert + "." + trail
    else:
        r.text = stripped + insert + trail
    return r.text[-60:]


def main() -> int:
    dry = "--dry-run" in sys.argv
    if not DOCX.exists():
        print(f"ERROR: {DOCX} not found", file=sys.stderr)
        return 1

    doc = Document(str(DOCX))
    print(f"INSERTIONS into {DOCX}{' (DRY RUN)' if dry else ''}:\n")

    applied = skipped = errors = 0
    for label, anchor in CALLOUTS:
        token = f"({label})"
        hits = find_paragraph(doc, anchor)
        if not hits:
            print(f"  !! {label:14s} NO MATCH for anchor {anchor!r}")
            errors += 1
            continue
        if len(hits) > 1:
            idxs = [i for i, _ in hits]
            print(f"  !! {label:14s} {len(hits)} matches {idxs} for {anchor!r} — ambiguous, skip")
            errors += 1
            continue
        i, p = hits[0]
        if token in p.text:
            print(f"  == {label:14s} [p{i}] already present — skip")
            skipped += 1
            continue
        if dry:
            print(f"  -> {label:14s} [p{i}] would append after: …{p.text.rstrip()[-50:]!r}")
            applied += 1
            continue
        tail = append_callout(p, label)
        print(f"  ++ {label:14s} [p{i}] -> …{tail!r}")
        applied += 1

    verb = "to apply" if dry else "applied"
    print(f"\nSummary: {applied} {verb}, {skipped} skipped, {errors} errors")
    if errors:
        print("Resolve errors before writing.", file=sys.stderr)
        if not dry:
            return 2
    if not dry and errors == 0:
        doc.save(str(DOCX))
        print(f"Saved {DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
