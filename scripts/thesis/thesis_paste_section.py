"""Paste a literature-review §x.y section into THESIS.docx.

Reads source prose directly from `docs/research/literature-review.docx` so we
never hardcode body text. Section structure is detected from heading text
patterns (every paragraph in the source uses the same `Body` style):

    ^\\d+\\.\\d+(\\.\\d+)?\\s        -> heading (section or sub-subsection)
    ^Table \\d+\\.\\d+               -> table caption -> placeholder line
    ^Note\\.                         -> note row of a table (skip)
    ^Figure \\d+\\.\\d+              -> figure caption -> placeholder line
    everything else                  -> body paragraph

In the destination thesis (Chapter II), we *rename* an existing style-`945`
subheading slot and then *insert* a sequence of (heading, body, ...) blocks
between it and the next style-`945` subheading. Heading blocks clone the slot;
body blocks clone an existing body-style paragraph in the section.

Usage:
    uv run python scripts/ops/thesis_paste_section.py <slot-text> <section-key>
    uv run python scripts/ops/thesis_paste_section.py APPEND <section-key>

`<section-key>` is something like "2.3" — must match a heading in
literature-review.docx beginning with that number. Use the sentinel
"APPEND" as <slot-text> to insert a brand-new top-level section at the
end of Chapter II (no slot rename); the section heading + sub-subsection
headings clone the most recent existing subheading paragraph, and body
paragraphs clone the most recent existing body paragraph.
"""

from __future__ import annotations

import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document  # type: ignore[import-untyped]
from docx.oxml.ns import qn  # type: ignore[import-untyped]

DOCX = Path("docs/research/THESIS.docx")
LIT_REVIEW = Path("docs/research/literature-review.docx")

# Heading patterns inside the lit-review doc.
SECTION_RE = re.compile(r"^(\d+)\.(\d+)\s+(.+)$")           # 2.3 Foo
SUBSEC_RE = re.compile(r"^(\d+)\.(\d+)\.(\d+)\s+(.+)$")     # 2.3.1 Bar
TABLE_RE = re.compile(r"^Table\s+\d+\.\d+\b")
FIGURE_RE = re.compile(r"^Figure\s+\d+\.\d+\b")
NOTE_RE = re.compile(r"^Note\.")


def build_spec(section_key: str) -> tuple[str, list[tuple[str, str]]]:
    """Return (rename_to_heading, [(kind, text), ...]) for the requested section.

    `kind` is one of "heading" (sub-subsection) or "body" (prose / placeholder).
    """
    src = Document(str(LIT_REVIEW))
    paras = [p.text.strip() for p in src.paragraphs]

    # Locate the section heading line.
    start = None
    section_heading = None
    for i, t in enumerate(paras):
        m = SECTION_RE.match(t)
        if m and f"{m.group(1)}.{m.group(2)}" == section_key and not SUBSEC_RE.match(t):
            start = i
            section_heading = t
            break
    if start is None:
        raise SystemExit(f"Section '{section_key}' not found in {LIT_REVIEW}.")

    # Bound: next section header that is *not* a sub-subsection of this one,
    # OR an end-of-chapter terminator ("Conclusion" / "References" /
    # "Bibliography" line), OR the first APA-style reference. Without these
    # extra terminators, a section that sits immediately before the
    # bibliography (e.g. §2.9 when followed directly by references) silently
    # consumes the bibliography as body paragraphs.
    apa_ref = re.compile(
        r"^[A-ZÀ-Ÿ][A-Za-zÀ-ÿ\-'\.]+(?:[\s\-][A-ZÀ-Ÿ][A-Za-zÀ-ÿ\-'\.]+)*,\s+[A-Z]\."
    )
    terminators = {"Conclusion", "References", "Reference", "Bibliography"}
    end = len(paras)
    for i in range(start + 1, len(paras)):
        line = paras[i]
        m = SECTION_RE.match(line)
        if m and not SUBSEC_RE.match(line):
            end = i
            break
        if line in terminators:
            end = i
            break
        if apa_ref.match(line):
            end = i
            break

    spec: list[tuple[str, str]] = []
    buffer: list[str] = []

    def _should_merge(prev: str, nxt: str) -> bool:
        # Lit-review has hard-wrapped paragraphs where each visual line is its
        # own <w:p>. Merge a continuation into the previous body when either
        # (a) the previous doesn't end in sentence-terminator, or (b) the next
        # starts with lowercase / "(" (mid-citation pattern like "(2023) introduced…").
        if not prev:
            return False
        last = prev.rstrip()[-1] if prev.rstrip() else ""
        first = nxt.lstrip()[0] if nxt.lstrip() else ""
        if last not in ".?!":
            return True
        if first == "(" or first.islower():
            return True
        return False

    def _flush() -> None:
        for entry in buffer:
            spec.append(("body", entry))
        buffer.clear()

    def _add_body(t: str) -> None:
        if buffer and _should_merge(buffer[-1], t):
            buffer[-1] = buffer[-1].rstrip() + " " + t.lstrip()
        else:
            buffer.append(t)

    for t in paras[start + 1 : end]:
        if not t:
            continue
        if SUBSEC_RE.match(t):
            _flush()
            spec.append(("heading", t))
        elif TABLE_RE.match(t) or FIGURE_RE.match(t):
            _flush()
            spec.append(("body", f"[{t} — manual paste from literature-review.docx.]"))
        elif NOTE_RE.match(t):
            continue
        else:
            _add_body(t)
    _flush()

    return section_heading, spec


def para_text(elem) -> str:
    return "".join(t.text or "" for t in elem.findall(".//" + qn("w:t"))).strip()


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def set_para_text(elem, text: str) -> None:
    """Replace paragraph contents with a single run carrying `text`.

    - Strips <w:bookmarkStart>/<w:bookmarkEnd> (their IDs are document-unique
      and conflict when paragraphs are cloned).
    - Preserves the first run that already has a <w:t> as the template
      (keeps its <w:rPr>); removes all other runs.
    - If no text-bearing run exists, creates one.
    """
    from lxml import etree  # type: ignore[import-untyped]

    for tag in ("w:bookmarkStart", "w:bookmarkEnd"):
        for el in elem.findall(qn(tag)):
            el.getparent().remove(el)

    runs = elem.findall(qn("w:r"))
    template_run = next(
        (r for r in runs if r.findall(qn("w:t"))),
        runs[0] if runs else None,
    )

    if template_run is None:
        new_run = etree.SubElement(elem, f"{{{_W_NS}}}r")
        new_t = etree.SubElement(new_run, f"{{{_W_NS}}}t")
        new_t.text = text
        new_t.set(qn("xml:space"), "preserve")
        return

    t_elems = template_run.findall(qn("w:t"))
    if t_elems:
        t_elems[0].text = text
        t_elems[0].set(qn("xml:space"), "preserve")
        for t in t_elems[1:]:
            template_run.remove(t)
    else:
        new_t = etree.SubElement(template_run, f"{{{_W_NS}}}t")
        new_t.text = text
        new_t.set(qn("xml:space"), "preserve")

    for r in runs:
        if r is not template_run:
            r.getparent().remove(r)


def _append_section(
    doc,
    paras,
    ch2_idx: int,
    ch3_idx: int,
    new_heading: str,
    spec: list[tuple[str, str]],
    section_key: str,
) -> None:
    """Insert a brand-new section (heading + spec blocks) at the end of Chapter II.

    Templates for new heading/body paragraphs are the most recent existing
    same-style paragraph found inside Chapter II.
    """
    heading_template = None
    body_template = None
    for i in range(ch3_idx - 1, ch2_idx, -1):
        p = paras[i]
        if not p.text.strip():
            continue
        sid = p.style.style_id if p.style else None
        m = SECTION_RE.match(p.text.strip()) or SUBSEC_RE.match(p.text.strip())
        if m and heading_template is None:
            heading_template = p._element
        elif not m and body_template is None and sid not in {"948"}:
            body_template = p._element
        if heading_template is not None and body_template is not None:
            break
    if heading_template is None or body_template is None:
        raise SystemExit("Could not find heading/body templates in Chapter II.")

    # Insertion anchor: last non-empty paragraph in Chapter II (i.e. just
    # before CHAPTER III, skipping trailing blanks).
    anchor_elem = None
    for i in range(ch3_idx - 1, ch2_idx, -1):
        if paras[i].text.strip():
            anchor_elem = paras[i]._element
            break
    if anchor_elem is None:
        raise SystemExit("Could not find an insertion anchor in Chapter II.")

    print(
        f"\nAppend mode: inserting §{section_key} after last non-empty Chapter II "
        f"paragraph; {len(spec) + 1} new block(s) total."
    )

    # Build the full ordered list of (kind, text) blocks: top-level heading first,
    # then the spec items.
    full_spec: list[tuple[str, str]] = [("heading", new_heading), *spec]

    cur_anchor = anchor_elem
    for kind, text in full_spec:
        if kind == "heading":
            new_p = deepcopy(heading_template)
        elif kind == "body":
            new_p = deepcopy(body_template)
        else:
            raise ValueError(f"Unknown spec kind: {kind}")
        set_para_text(new_p, text)
        cur_anchor.addnext(new_p)
        cur_anchor = new_p

    doc.save(str(DOCX))
    print(f"\nAppended §{section_key} '{new_heading}' to Chapter II.")


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: thesis_paste_section.py <slot-text> <section-key>")
    slot_text, section_key = sys.argv[1], sys.argv[2]

    new_heading, spec = build_spec(section_key)
    print(f"Section §{section_key} '{new_heading}': {len(spec)} blocks from lit-review.docx")
    for kind, text in spec:
        head = text[:80].replace("\n", " ")
        print(f"  [{kind:7}] {head}{'…' if len(text) > 80 else ''}")

    doc = Document(str(DOCX))
    paras = list(doc.paragraphs)

    ch2_idx = next((i for i, p in enumerate(paras) if p.text.strip() == "CHAPTER II"), None)
    ch3_idx = next((i for i, p in enumerate(paras) if p.text.strip() == "CHAPTER III"), None)
    if ch2_idx is None or ch3_idx is None:
        raise SystemExit("Could not locate Chapter II / III boundaries.")

    if slot_text == "APPEND":
        return _append_section(doc, paras, ch2_idx, ch3_idx, new_heading, spec, section_key)

    slot_idx = None
    for i in range(ch2_idx, ch3_idx):
        if paras[i].text.strip() == slot_text:
            slot_idx = i
            break
    if slot_idx is None:
        raise SystemExit(f"Could not find slot subheading '{slot_text}' in Chapter II.")

    slot_elem = paras[slot_idx]._element
    sub_style = paras[slot_idx].style.style_id

    end_elem = None
    for i in range(slot_idx + 1, ch3_idx):
        if paras[i].style.style_id == sub_style:
            end_elem = paras[i]._element
            break
    if end_elem is None:
        end_elem = paras[ch3_idx]._element

    template_body_elem = None
    cur = slot_elem.getnext()
    while cur is not None and cur is not end_elem:
        if cur.tag == qn("w:p") and para_text(cur):
            style_el = cur.find(qn("w:pPr") + "/" + qn("w:pStyle"))
            sid = style_el.get(qn("w:val")) if style_el is not None else None
            if sid != sub_style:
                template_body_elem = cur
                break
        cur = cur.getnext()
    if template_body_elem is None:
        raise SystemExit("No body-style template paragraph found in section.")

    body_elems: list = []
    cur = slot_elem.getnext()
    while cur is not None and cur is not end_elem:
        if cur.tag == qn("w:p"):
            body_elems.append(cur)
        cur = cur.getnext()

    print(
        f"\nSlot '{slot_text}' @ idx {slot_idx} (style {sub_style}); "
        f"{len(body_elems)} body paragraph(s) to remove; "
        f"inserting {len(spec)} new block(s)."
    )

    set_para_text(slot_elem, new_heading)

    spacer = None
    rest = body_elems
    if body_elems and not para_text(body_elems[0]):
        spacer = body_elems[0]
        rest = body_elems[1:]
    for e in rest:
        parent = e.getparent()
        if parent is not None:
            parent.remove(e)

    cur_anchor = spacer if spacer is not None else slot_elem
    for kind, text in spec:
        if kind == "heading":
            new_p = deepcopy(slot_elem)
        elif kind == "body":
            new_p = deepcopy(template_body_elem)
        else:
            raise ValueError(f"Unknown spec kind: {kind}")
        set_para_text(new_p, text)
        cur_anchor.addnext(new_p)
        cur_anchor = new_p

    doc.save(str(DOCX))
    print(f"\nPasted §{section_key} into slot '{slot_text}' -> '{new_heading}'.")


if __name__ == "__main__":
    main()
