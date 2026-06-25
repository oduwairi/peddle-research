"""Write §3.2 Data Acquisition body — running list of paragraphs.

Same pattern as write_3_1_proposed_system.py: wipes every paragraph between the
§3.2 heading and the §3.3 heading, then re-inserts the current PARAGRAPHS list
using a Body Text (style 943) template paragraph from CH3.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"


PARAGRAPHS: list[str] = [
    # ¶1 — Collection engine: sweeps, sessions, loops
    "For the ad collection strategy, we followed a multi-sweep strategy where a collection engine is given a query — a combination of platform, filtering, ordering, and geo (and optionally a search term, such as a keyword or landing-page domain) — and paginates through the result pages from the API, then deduplicates against existing ads and appends them to the corpus. The system allowed resumability by saving the cursor at the last paginated page, so it can pick up where it left off in case of interruption or adjustment. The chosen collection strategy consists of sessions and loops: a session is a single run, and a loop is a complete pass over all queries for a single target depth in that run. We chose the loop strategy simply because frequent adjustments were necessary to ensure duplication rates stayed low and the variety of the collected ad set was sufficient for training.",
    # ¶2 — Motivation: why we need a real-ad corpus, and source choice
    "In order to produce a fine-tuned model capable of generating actual high-performing ad campaigns, we need a database of high-performing ads from which we can build the training instructions, using the instruction-backtranslation technique (§2.7.3) where the brief is reverse-engineered from the response (the real performing ad). This provides the main motivation: a domain-specific model trained on successful ads will outperform a generalist model, which outputs generic AI content that is often perceived as lacking or inhuman, especially in advertising domains. Creating a data corpus of successful ads presents a significant challenge, as many providers either lack the automated tooling — such as a public API — to scrape and store ads at scale, or do not expose the engagement metadata required for generating training labels. For this reason, AdFlex was chosen as the primary source, because it provides a single large source for ads across five different platforms (Facebook, TikTok, X, Pinterest, and Reddit) with engagement metadata for our proxy scoring system (§3.3). AdFlex charges 100 credits per API call returning 18 ads, against a monthly budget of 500K credits, which constrains the sweep design described below. Complementary free and public datasets — Meta Ad Library, Google Ads Transparency Center, TikTok Ad Library, BigSpy — are used in supplementary roles, for cross-checking and validation across the pipeline.",
    # ¶3 — Sweep families and configuration
    "The collection engine supports five main sweep configurations, each aimed at scraping the dataset from a different angle to ensure maximum variety at the end for training. The five main sweep configurations include filter, range, broad, keyword, and domain sweeps. Broad sweeps are used first, with no category filter at all, generating ordering combinations (across orderings such as popularity and recency) across the entire AdFlex corpus; this is good initially because it provides good insight into the ad dataset distribution. The second type is filter sweeps, where very specific filters exposed by the API — such as industry, call-to-action type, or ad format codes — are queried per sweep; this provides a more specific scraping of the dataset, as it targets specific industry or format pools. Range sweeps refer to fields with a numerical range, such as run duration or ad spend, bucketed into tiers; this provides another slice of the ad set, focused on numerical qualities. Other sweeps, such as keyword sweeps (which search for specific words or phrases — including intent terms such as “free trial” and vertical terms such as “insurance”) and domain sweeps (which search by landing-page URL, targeting domains such as shopify or stripe), are used towards the end to provide maximum diversity and scrape very specific ads not covered by the main sweep configurations. These sweep plans are used interchangeably after each loop, depending on the resulting collected ad-set variety, quality, and deduplication rates, to ensure maximum utilization of the AdFlex corpus.",
    # ¶4 — Normalisation step
    "Following the ad collection and sweeping, a normalization step is applied to the raw collected ads from AdFlex or any other supplementary source, in order to unify the structure that the downstream operations use. This is done by converting the raw API response into a single canonical schema that normalizes all metadata fields, including engagement metadata (such as likes, shares, comments, reactions, views, active days, and first/last seen dates), platform, or language tags. The schema is automatically validated to make sure ads with missing or type-mismatched required fields are rejected and not written as broken ads. The normalized schema also carries provenance information, such as which sweep produced each ad, as well as platform and other source-specific metadata. The final output is a single append-only JSONL file, used for all downstream operations such as scoring (§3.3) or construction (§3.4). Additional LLM-generated metadata labels — such as training quality, content safety, and business vertical (43 categories, generated via batch APIs by a general-purpose LLM and later used as gates in the construction pipeline §3.4) — are also appended to the same centralized schema, yielding a downstream corpus of approximately 55,000 ads.",
]


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def set_paragraph_text(p_elem, new_text):
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    h_3_2 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("3.2 Data Acquisition"),
        "§3.2 heading",
    )
    h_3_3 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("3.3 Engagement-Based Scoring"),
        "§3.3 heading",
    )

    # Body Text template: first style-943 paragraph inside CH3. Prefer an
    # existing body paragraph over a placeholder so the run/rPr matches the
    # surrounding prose exactly.
    body_tpl = None
    in_ch3 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.upper() == "CHAPTER III":
            in_ch3 = True
            continue
        if t.upper() == "CHAPTER IV":
            break
        sid = p.style.style_id if p.style else ""
        if in_ch3 and sid == "943" and t and p._element.findall(f".//{QN_T}"):
            body_tpl = p._element
            break
    if body_tpl is None:
        raise RuntimeError("No style-943 Body Text paragraph found inside CH3")

    body = h_3_2.getparent()
    children = list(body)
    i_start = children.index(h_3_2) + 1
    i_end = children.index(h_3_3)

    to_delete = [c for c in children[i_start:i_end] if c.tag == QN_P]
    print(f"DELETING: {len(to_delete)} existing §3.2 body paragraphs")
    for c in to_delete:
        body.remove(c)

    insert_after = h_3_2
    for i, text in enumerate(PARAGRAPHS, 1):
        p = deepcopy(body_tpl)
        set_paragraph_text(p, text)
        insert_after.addnext(p)
        insert_after = p
        print(f"  ¶{i}: {text[:80]}…")

    doc.save(str(path))
    print(f"\nWROTE: {len(PARAGRAPHS)} paragraph(s) into §3.2 Data Acquisition")


if __name__ == "__main__":
    main()
