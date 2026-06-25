"""Insert the three deployment/testing screenshots into Chapter III.

Reviewer feedback #10: "Add images to prove the deployment and testing of your
model/project." Chapter IV already carries the quantitative testing evidence
(the learned-scorer / MAUVE / reference-overlap arms and the fine-tuning x agent
ablation); the gap was *visual* proof that the system is deployed and runs. This
script lands three real screenshots:

  * Fig 3.4 (NEW)  Modal dashboard            -> §3.5.5 Adapter Merging and Hosting
  * Fig 3.5        agent architecture (RENUMBERED from 3.4) -> §3.6
  * Fig 3.6 (NEW)  logged agent execution trace (the loop/process) -> §3.6.5
  * Fig 3.7 (NEW)  deployed chat campaign card + score badge (the output) -> §3.6.5

Numbering rationale: the Modal figure lands in §3.5.5 (before the existing
agent-architecture figure), so the existing Fig 3.4 renumbers to Fig 3.5. The
trace (process) is 3.6 and the campaign card (output) is 3.7 so the in-text
callouts run in ascending order (loop is referenced in §3.6.3, the emitted card
in §3.6.5).

The real agent-architecture caption is located by its ``Fig_3_4`` bookmark (NOT
by text — the List-of-Figures cache paragraph also starts with "Figure 3.4" and
sorts first in document order). Nothing cross-refs the old Fig_3_4 bookmark
except the LoF cache (a live field, refreshed in OnlyOffice), and the in-text
callout is plain text, so the renumber is safe.

Idempotent: re-running skips figures already present, skips the renumber if it
is already done, and skips callouts already appended.

Run:
    uv run python scripts/thesis/insert_deployment_figures.py
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document

# Reuse the tested figure-insertion helper that lives next to this script.
sys.path.insert(0, str(Path(__file__).resolve().parent))
from insert_figure import insert_figure  # noqa: E402

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QT = f"{{{W}}}t"
QI = f"{{{W}}}instrText"
QBMS = f"{{{W}}}bookmarkStart"
QNAME = f"{{{W}}}name"

DOC = Path("docs/research/THESIS.docx")
LOCK = Path("docs/research/.~lock.THESIS.docx#")
BACKUP = Path("docs/research/THESIS.docx.pre-deploy-figs.bak")
FIGDIR = Path("docs/research/figures")

WIDTH = 5.5  # inches — matches the existing Ch III inline figures

CAP_34 = (
    "Figure 3.4: Serverless deployment of the fine-tuned model. The Modal "
    "dashboard for the draper-vllm application shows the live vLLM endpoint "
    "serving OpenAI-compatible chat-completion requests (HTTP 200)."
)
CAP_36 = (
    "Figure 3.6: A logged agent execution trace. The freeform loop's steps — "
    "research tool calls, draft, and emit — recorded to the database; the "
    "best-of-N draft surfaces four temperature-varied candidates, each scored, "
    "before one is selected."
)
CAP_37 = (
    "Figure 3.7: The deployed chat interface returning a finished campaign. A "
    "TikTok ad generated from a single-line brief, shown with its auto-attached "
    "predicted-performance score (61/100) and generated creative."
)

ANCHOR_MODAL = "After training completes, the LoRA adapter stays separate"
ANCHOR_TOOLS = "Communication between the orchestrator and the writer"
ANCHOR_TRACE_CAP = "Figure 3.6:"  # the trace caption inserted just above the campaign
ANCHOR_LOOP = "When the orchestrator and writer model are set"

CALLOUT_MODAL = ("representing both the merged adapter and the base model", "(Figure 3.4)")
CALLOUT_TRACE = ("the agent workflow runs as a freeform loop", "(Figure 3.6)")
CALLOUT_CAMPAIGN = ("attached to the campaign card)", "(Figure 3.7)")


# --------------------------------------------------------------------------
# small XML helpers
# --------------------------------------------------------------------------
def visible_texts(p_elem):
    """Visible <w:t> run elements (excludes <w:instrText> field codes)."""
    return p_elem.findall(f".//{QT}")


def find_para(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def insert_callout(p_elem, phrase: str, callout: str) -> bool:
    """Insert ``" " + callout`` immediately after ``phrase`` in a paragraph.

    Works across run boundaries: concatenates the visible <w:t>, finds the end
    offset of ``phrase``, then splices the callout into the run that owns that
    offset. Idempotent — no-op if the callout text is already present.
    """
    full = "".join(t.text or "" for t in visible_texts(p_elem))
    if callout in full:
        return False  # already there
    idx = full.find(phrase)
    if idx < 0:
        raise RuntimeError(f"callout phrase not found: {phrase!r}")
    end = idx + len(phrase)
    ins = " " + callout
    cum = 0
    for t in visible_texts(p_elem):
        txt = t.text or ""
        if cum <= end <= cum + len(txt):
            local = end - cum
            t.text = txt[:local] + ins + txt[local:]
            return True
        cum += len(txt)
    raise RuntimeError(f"could not place callout for phrase: {phrase!r}")


def renumber_agent_arch(doc) -> bool:
    """Renumber the existing 'Draper.ai agent architecture' figure 3.4 -> 3.5.

    Located by its ``Fig_3_4`` bookmark so we never touch the List-of-Figures
    cache paragraph (which also starts with "Figure 3.4" and sorts first).
    Updates the visible caption, the TC field text, and the bookmark name.
    Returns False (no-op) if it has already been renumbered.
    """
    cap = None
    for p in doc.paragraphs:
        names = [b.get(QNAME) for b in p._element.findall(f".//{QBMS}")]
        if "Fig_3_4" in names:
            cap = p
            break
    if cap is None:
        # already renumbered (bookmark is Fig_3_5) -> verify and no-op
        for p in doc.paragraphs:
            names = [b.get(QNAME) for b in p._element.findall(f".//{QBMS}")]
            if "Fig_3_5" in names and "architecture" in p.text:
                return False
        raise RuntimeError("agent-architecture caption (Fig_3_4) not found")

    e = cap._element
    for t in visible_texts(e):
        if t.text and "Figure 3.4" in t.text:
            t.text = t.text.replace("Figure 3.4", "Figure 3.5")
    for it in e.findall(f".//{QI}"):
        if it.text and "Figure 3.4" in it.text:
            it.text = it.text.replace("Figure 3.4", "Figure 3.5")
    for bs in e.findall(f".//{QBMS}"):
        if bs.get(QNAME) == "Fig_3_4":
            bs.set(QNAME, "Fig_3_5")
    return True


# --------------------------------------------------------------------------
def main() -> None:
    if LOCK.exists():
        sys.exit(f"REFUSING: {LOCK} exists — close OnlyOffice/Word first.")
    for f in ("fig-deploy-modal-dashboard.png", "fig-deploy-chat-campaign.png",
              "fig-deploy-agent-trace.png"):
        if not (FIGDIR / f).exists():
            sys.exit(f"missing figure: {FIGDIR / f}")
    if not BACKUP.exists():
        shutil.copy2(DOC, BACKUP)
        print(f"backup -> {BACKUP}")
    else:
        print(f"backup already exists -> {BACKUP} (kept)")

    # 1) renumber existing Fig 3.4 -> 3.5 and bump its plain-text callout.
    doc = Document(str(DOC))
    did_renum = renumber_agent_arch(doc)
    bumped = False
    for p in doc.paragraphs:
        if "(Figure 3.4)" in p.text and "routed to the writer model" in p.text:
            for t in visible_texts(p._element):
                if t.text and "(Figure 3.4)" in t.text:
                    t.text = t.text.replace("(Figure 3.4)", "(Figure 3.5)")
                    bumped = True
            break
    doc.save(str(DOC))
    print(f"renumber 3.4->3.5: caption={'done' if did_renum else 'skip'} "
          f"callout={'done' if bumped else 'skip'}")

    # 2) insert the three new figures (TC mode -> listed in the LoF on refresh).
    #    Order: Modal (3.4) in §3.5.5; trace (3.6) then campaign (3.7) in §3.6.5.
    insert_figure(DOC, ANCHOR_MODAL, FIGDIR / "fig-deploy-modal-dashboard.png",
                  CAP_34, WIDTH, bookmark="Fig_3_4", bm_id=30040)
    insert_figure(DOC, ANCHOR_TOOLS, FIGDIR / "fig-deploy-agent-trace.png",
                  CAP_36, WIDTH, bookmark="Fig_3_6", bm_id=30060)
    insert_figure(DOC, ANCHOR_TRACE_CAP, FIGDIR / "fig-deploy-chat-campaign.png",
                  CAP_37, WIDTH, bookmark="Fig_3_7", bm_id=30070)

    # 3) in-text callouts for the three new figures (ascending order).
    doc = Document(str(DOC))
    results = []
    for prefix, (phrase, callout) in (
        (ANCHOR_MODAL, CALLOUT_MODAL),
        (ANCHOR_LOOP, CALLOUT_TRACE),
        (ANCHOR_TOOLS, CALLOUT_CAMPAIGN),
    ):
        p = find_para(doc, prefix)
        if p is None:
            raise RuntimeError(f"callout anchor not found: {prefix!r}")
        changed = insert_callout(p._element, phrase, callout)
        results.append((callout, "added" if changed else "skip"))
    doc.save(str(DOC))
    for c, s in results:
        print(f"callout {c}: {s}")

    print("DONE")


if __name__ == "__main__":
    main()
