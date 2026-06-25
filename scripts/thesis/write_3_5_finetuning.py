"""Write §3.5 Fine-Tuning body.

Same pattern as write_3_4_construction.py: wipes every paragraph between
the §3.5 heading and the §3.6 heading, then re-inserts the current
PARAGRAPHS list using a Body Text (style 943) template paragraph from
CH3. New bibliography entries (if any) are inserted in NEW_REFERENCES at
their alphabetical positions before APPENDIX, idempotent.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"


PARAGRAPHS: list[str] = [
    # ¶1 — Setup and the single-GPU constraint
    "The finalized dataset contains approximately three thousand examples (2,872 in total), each of which consists of a three-turn chat: system, user brief, and an assistant response. The final set is split into training, validation, and test splits (roughly 85/7.5/7.5), where the validation set is used during training for evaluation of the model during training, and the test set is held out after training for evaluation of the finished model. The student model must be balanced between cost and efficiency, in that it can be served on a single GPU with less than fifty gigabytes of VRAM, yet the model should be large enough to perform tasks such as creative writing and handling briefs. For this, we picked an eight-billion-parameter model in the Qwen3 family, which is popular for fine-tuning and instruction-tuning tasks and can fit on one consumer GPU. Since this is a fine-tuning task, we will not be training all eight billion parameter weights; instead we train a small adapter that is later merged into the base model.",
    # ¶2 — The adapter setup (QLoRA + LoRA + DoRA + rsLoRA)
    "The setup phase includes shrinking the base model weights into a four-bit (NF4) compressed form rather than the complete sixteen-bit form. This makes the model fit easier on a consumer GPU and shrinks the model from sixteen gigabytes down to about five gigabytes. The base model weights are frozen during training, and they are never touched or updated. They are used for inferencing during training. A small low-rank (LoRA) adapter (rank r=32, α=64) is attached into every attention and feed-forward projection in the base model — the combination of a four-bit-quantized base and a trainable LoRA adapter is known as QLoRA. During training, loss and optimization are applied to the adapter weights, which are the only weights updated during training. That's the total size. It's about fifty-one million trainable parameters, which is about 0.6% of the base eight-billion-parameter model. In addition to this adapter, we apply two techniques to improve training efficiency. The first, DoRA (weight-decomposed low-rank adaptation), splits the adapter into a magnitude and direction for updating weights. The second, rsLoRA (rank-stabilized LoRA), tells how strongly the adapter affects the base model output. This allows us to change the adapter size without having to change the learning rate.",
    # ¶3 — How each example is shown to the model (ChatML + assistant-only loss masking)
    "During training, each training example is given to the model in the form of a chat conversation. Using the chat template (ChatML) the base model uses, the system prompt as well as the user prompt are fed to the model as input. The assistant's response is meant to minimize the difference between the training example response and the actual response, and it's the part that the model trains to learn. Therefore, each time training loss is computed, it only applies to the assistant response. The system prompt and the user brief are masked during training (this is known as assistant-only loss masking) and do not contribute to weight change. Therefore, the model slowly learns on the responses it generates, and its ability to reproduce the successful ad is already established.",
    # ¶4 — Training schedule and compute (AdamW-8bit + cosine schedule + effective batch 16 + early stopping)
    "The training is set to run for three epochs over the entire training set. Weight updates use an eight-bit version of the AdamW optimizer in order to save memory. The learning rate is set to follow a cosine schedule, where it starts by increasing in the first three percent of the training loop, then starts to decay to zero towards the end. Examples are processed by the GPU two at a time and accumulated over eight passes until the weight update happens. Each training step is sixteen examples (2 per device × 8 accumulation steps). This is in order to balance speed of learning as well as learning stability. During training, the validation set is used for testing the model's learning and runs every 50 steps to ensure the model is learning and not overfitting; if the model stops learning for two consecutive checks, the training stops early (early stopping with patience of 2). The estimated GPU memory needed is around sixteen gigabytes at minimum, to fit the model, adapter, and training state as well as having some headroom to avoid out-of-memory (OOM) errors. Training is estimated to take from one to two hours depending on stopping time and performance on the chosen GPU.",
    # ¶5 — Output artefact and scope boundaries (adapter + merge + deployment + SFT-only scope)
    "After training completes, the LoRA adapter stays separate and is pushed separately to our Hugging Face repository, where it can optionally be merged into the base model to create the single fine-tuned model. This flexibility allows us to test before merging into the base model. The final merged model is served as the same base model with sixteen-bit (bfloat16) precision, back from the four-bit version used during training. The model is deployed on a serverless GPU on Modal for inference after training as one model file representing both the merged adapter and the base model. The scope of the work stays limited to supervised fine-tuning. Techniques such as reward modeling, preference optimization, or curriculum learning are deferred for later judgement.",
]


# (alphabetical-anchor-prefix, full reference text)
# Empty so far — LoRA/QLoRA/DoRA/rsLoRA/Unsloth/Qwen3 citations will be
# added as their owning paragraphs land.
NEW_REFERENCES: list[tuple[str, str]] = []


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def set_paragraph_text(p_elem, new_text):
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def write_section_body(doc):
    h_3_5 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("3.5 Fine-Tuning"),
        "§3.5 heading",
    )
    h_3_6 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("3.6 "),
        "§3.6 heading",
    )

    # Body Text template: first style-943 paragraph inside CH3 (excluding
    # §3.5's own body, which we're about to wipe).
    body_tpl = None
    in_ch3 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.upper() == "CHAPTER III":
            in_ch3 = True
            continue
        if t.upper() == "CHAPTER IV":
            break
        if t.startswith("3.5 Fine-Tuning"):
            in_ch3 = "in-35"
            continue
        if in_ch3 == "in-35" and t.startswith("3.6 "):
            in_ch3 = True
            continue
        if in_ch3 == "in-35":
            continue
        sid = p.style.style_id if p.style else ""
        if in_ch3 and sid == "943" and t and p._element.findall(f".//{QN_T}"):
            body_tpl = p._element
            break
    if body_tpl is None:
        raise RuntimeError("No style-943 Body Text paragraph found inside CH3 (outside §3.5)")

    body = h_3_5.getparent()
    children = list(body)
    i_start = children.index(h_3_5) + 1
    i_end = children.index(h_3_6)

    to_delete = [c for c in children[i_start:i_end] if c.tag == QN_P]
    print(f"DELETING: {len(to_delete)} existing §3.5 body paragraphs")
    for c in to_delete:
        body.remove(c)

    insert_after = h_3_5
    for i, text in enumerate(PARAGRAPHS, 1):
        p = deepcopy(body_tpl)
        set_paragraph_text(p, text)
        insert_after.addnext(p)
        insert_after = p
        print(f"  ¶{i}: {text[:80]}…")

    print(f"\nWROTE: {len(PARAGRAPHS)} paragraph(s) into §3.5 Fine-Tuning")


def insert_references(doc):
    if not NEW_REFERENCES:
        print("(no new references for §3.5)")
        return

    refs_heading = None
    appendix_heading = None
    for p in doc.paragraphs:
        t = p.text.strip().upper()
        if t == "REFERENCES" and refs_heading is None:
            refs_heading = p._element
        elif t.startswith("APPENDIX") and appendix_heading is None:
            appendix_heading = p._element
    if refs_heading is None or appendix_heading is None:
        raise RuntimeError("References / APPENDIX bounds not found")

    body = refs_heading.getparent()
    children = list(body)
    i_refs = children.index(refs_heading)
    i_appendix = children.index(appendix_heading)
    ref_range = children[i_refs + 1 : i_appendix]

    ref_tpl = None
    for c in ref_range:
        if c.tag != QN_P:
            continue
        pPr = c.find(f"{{{W_NS}}}pPr")
        pStyle = pPr.find(f"{{{W_NS}}}pStyle") if pPr is not None else None
        sid = pStyle.get(f"{{{W_NS}}}val") if pStyle is not None else ""
        ts = c.findall(f".//{QN_T}")
        if sid == "943" and ts and any(t.text for t in ts):
            ref_tpl = c
            break
    if ref_tpl is None:
        raise RuntimeError("No style-943 reference template paragraph found")

    inserted = 0
    skipped = 0
    for anchor_prefix, ref_text in NEW_REFERENCES:
        first_words = ref_text.split("(")[0].strip().rstrip(",")
        already_present = any(
            c.tag == QN_P and "".join(t.text or "" for t in c.findall(f".//{QN_T}")).startswith(first_words)
            for c in ref_range
        )
        if already_present:
            print(f"  SKIP (already present): {ref_text[:80]}…")
            skipped += 1
            continue

        anchor = None
        for c in ref_range:
            if c.tag != QN_P:
                continue
            text = "".join(t.text or "" for t in c.findall(f".//{QN_T}"))
            if text.startswith(anchor_prefix):
                anchor = c
                break
        if anchor is None:
            raise RuntimeError(f"Alphabetical anchor not found: {anchor_prefix!r}")

        new_p = deepcopy(ref_tpl)
        set_paragraph_text(new_p, ref_text)
        anchor.addprevious(new_p)
        inserted += 1
        print(f"  INSERT before {anchor_prefix!r}: {ref_text[:80]}…")

        children = list(body)
        i_refs_new = children.index(refs_heading)
        i_appendix_new = children.index(appendix_heading)
        ref_range = children[i_refs_new + 1 : i_appendix_new]

    print(f"\nREFERENCES: inserted {inserted}, skipped {skipped} (already present)")


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    print("=== §3.5 body ===")
    write_section_body(doc)

    print("\n=== References ===")
    insert_references(doc)

    doc.save(str(path))
    print(f"\nSAVED: {path}")


if __name__ == "__main__":
    main()
