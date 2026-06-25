"""Post-repair verification: field structure of LoT [314]/[315], dirty flags,
settings.xml updateFields placement, and overall doc integrity."""
from __future__ import annotations

import re

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

THESIS = "docs/research/THESIS.docx"
doc = Document(THESIS)
paras = doc.paragraphs


def show(i):
    xml = etree.tostring(paras[i]._element, pretty_print=True).decode()
    xml = re.sub(r'xmlns:\w+="[^"]+"\s*', "", xml)
    print(f"\n##### [{i}] {paras[i].text.strip()[:60]!r}\n{xml}")


show(314)
show(315)

# dirty-flag count across list regions
dirty = 0
for i in list(range(294, 308)) + list(range(309, 316)):
    for fb in paras[i]._element.findall(".//" + qn("w:fldChar")):
        if fb.get(qn("w:fldCharType")) == "begin" and fb.get(qn("w:dirty")) in ("1", "true"):
            dirty += 1
print(f"\nPAGEREF begins marked dirty in LoF/LoT: {dirty}")

# settings.xml
settings = doc.settings.element
order = [c.tag.split('}')[-1] for c in settings]
print(f"\nsettings.xml children order: {order}")
uf = settings.find(qn("w:updateFields"))
print(f"updateFields present: {uf is not None}  val={uf.get(qn('w:val')) if uf is not None else None}")

# integrity
print(f"\nparagraphs: {len(paras)}")
print(f"tables: {len(doc.tables)}")
drawings = doc.element.body.findall(".//" + qn("w:drawing"))
print(f"drawings: {len(drawings)}")
# duplicate bookmark id check
ids = [bm.get(qn('w:id')) for bm in doc.element.body.iter(qn('w:bookmarkStart'))]
dupes = {x for x in ids if ids.count(x) > 1}
print(f"duplicate bookmark ids: {dupes or 'none'}")
