"""Write §4.1 Evaluation Setup body (between the §4.1 heading and §4.2).

Same wipe-and-replace pattern as write_ch4_intro.py: wipes every paragraph
between the "4.1 Evaluation Setup" heading and the "4.2 Learned-Scorer
Absolute Scores" heading, then re-inserts the current PARAGRAPHS list using
a Body Text (style 943) template paragraph cloned from Chapter III.

Idempotent — re-running produces the same output.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"


PARAGRAPHS: list[str] = [
    # ¶1 — test set + platform breakdown
    "The evaluation uses the 215 held-out briefs from the construction "
    "training examples. These are briefs that the model has never seen "
    "during training. That way, we can ensure that the model won't get "
    "any unfair boost in score if the model overfits or memorizes, and "
    "ensure a clean comparison. Each brief in this test set contains the "
    "high-performer ad as part of the response. Of the 215 briefs: "
    "Facebook 93 (43%), Reddit 39 (18%), TikTok 28 (13%), Pinterest 28 "
    "(13%), Twitter 27 (13%). Facebook dominates the ads since the "
    "original corpus, and the API (AdFlex) was Facebook-heavy.",
    # ¶2 — all configurations (single-shot A/B/C + agentic B_pipe/C_pipe + GOLD)
    "The evaluation pipeline uses five main configurations and a "
    "reference configuration across all combinations. These "
    "configurations are classified as single-shot prompted LLMs and "
    "agent-wrapped LLMs. The first configuration (Config A) is the "
    "frontier model, which is a GPT-5.4 model (gpt-5.4-mini) that is "
    "prompted using the API and gives the response as its candidate ad. "
    "This is essential as the frontier model is the representation that "
    "our model is trying to beat (RQ1). The second configuration "
    "(Config B) represents the base 8B model (Qwen3-8B) without our "
    "fine-tuning. Similar to the frontier model, it gets the brief and "
    "writes a response. This configuration is essential to understand "
    "the effect of fine-tuning on the quality of the output. "
    "Configuration C is our fine-tuned model (Draper, served via Modal "
    "vLLM). This single-shot configuration is essential to compare "
    "against the base untuned model to understand the effect of "
    "fine-tuning and our training. Two more configurations (Config "
    "B_pipe and Config C_pipe) use the untuned and tuned models, but "
    "with agentic workflows (RAG) integrating either model in the full "
    "agent loop, which calls tools (web search, image generation) and "
    "performs research then calls the writers. These configurations "
    "together are essential to understand the effect of the existence "
    "or absence of agentic systems on the quality of the output (RQ2). "
    "The final configuration (GOLD) represents the golden reference, "
    "which is the golden held-out ads themselves, and they serve as the "
    "upper ceiling which other baselines are trying to chase.",
    # ¶3 — evaluation procedure overview (Phase 4 + 5 + 6)
    "In our methodology, we have two evaluation arms to test the "
    "performance of the model. Since marketing and advertisements are "
    "subjective tasks and there are no right or wrong answers, the best "
    "approach is to use trained models to evaluate the performance of "
    "the models based on ground truth data, giving an objective "
    "evaluation of the model performance. Our first evaluation arm is "
    "the trained DeBERTa regressor (§3.7) that performs pointwise "
    "evaluation and comparison by giving a composite score for each ad "
    "copy text as well as underlying individual component scores "
    "(survivability, engagement volume, engagement velocity). The "
    "second evaluation arm is MAUVE (§3.8), which performs cluster "
    "comparisons rather than pointwise comparisons. It has the "
    "advantage of comparing the overall shape of text based on their "
    "embedding representation (GPT-2 Large), and focuses less on "
    "individual ads, giving an overall overview of the model "
    "performance. For both these evaluation arms, outputs are "
    "normalized across providers and runs to ensure they have the same "
    "shape before being fed into the model for fair comparison. Since "
    "some models output different formats, including markdown "
    "formatting or filler text, a normalization step is important to "
    "isolate the ad copy, which is to be graded. Extraction failures "
    "are also addressed when models fail to deliver in the expected "
    "format. These are naturally flagged and excluded from scoring to "
    "ensure accurate comparisons. Critically, when evaluating paired "
    "comparisons, we always use the intersection between any two "
    "configs, meaning that the same brief has to be present in both "
    "configurations to be used, not one or the other. At the end, all "
    "reported numbers come with 95% confidence intervals (via bootstrap "
    "resampling). Both arms report CIs to ensure the sample results are "
    "properly reported without overclaiming.",
]


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def set_paragraph_text(p_elem, new_text):
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def write_section_body(doc):
    h_4_1 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("4.1 Evaluation Setup"),
        "§4.1 heading",
    )
    h_4_2 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("4.2 Learned-Scorer"),
        "§4.2 heading",
    )

    # Body Text template: first style-943 paragraph inside Chapter III with
    # real text content. Same convention as write_ch4_intro.py.
    body_tpl = None
    in_ch3 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.upper() == "CHAPTER III":
            in_ch3 = True
            continue
        if t.upper() == "CHAPTER IV":
            break
        sid = p.style.style_id if p.style else ""
        if in_ch3 and sid == "943" and t and p._element.findall(f".//{QN_T}"):
            body_tpl = p._element
            break
    if body_tpl is None:
        raise RuntimeError("No style-943 Body Text template found inside CH3")

    body = h_4_1.getparent()
    children = list(body)
    i_start = children.index(h_4_1) + 1
    i_end = children.index(h_4_2)

    to_delete = [c for c in children[i_start:i_end] if c.tag == QN_P]
    print(f"DELETING: {len(to_delete)} existing §4.1 paragraph(s)")
    for c in to_delete:
        body.remove(c)

    insert_after = h_4_1
    for i, text in enumerate(PARAGRAPHS, 1):
        p = deepcopy(body_tpl)
        set_paragraph_text(p, text)
        insert_after.addnext(p)
        insert_after = p
        print(f"  ¶{i}: {text[:80]}…")

    print(f"\nWROTE: {len(PARAGRAPHS)} paragraph(s) into §4.1 Evaluation Setup")


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    print("=== §4.1 Evaluation Setup ===")
    write_section_body(doc)

    doc.save(str(path))
    print(f"\nSAVED: {path}")


if __name__ == "__main__":
    main()
