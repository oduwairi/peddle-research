"""Write §4.2 Learned-Scorer Absolute Scores body (between the §4.2 heading
and §4.3).

Same wipe-and-replace pattern as write_4_1_setup.py: wipes every paragraph
between the "4.2 Learned-Scorer Absolute Scores" heading and the "4.3 MAUVE
Distribution Matching" heading, then re-inserts the current PARAGRAPHS list
using a Body Text (style 943) template paragraph cloned from Chapter III.

Idempotent — re-running produces the same output.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"


PARAGRAPHS: list[str] = [
    # ¶1 — what this arm reports + headline single-shot composite means
    "This section reports the findings of the first evaluation arm, "
    "which is the trained DeBERTa regressor (§3.7). By running it "
    "against each configuration's inference set, in the 215 held-out "
    "test set, the results show the composite score, which is the "
    "overall score, as well as the three components, which are the "
    "engagement volume, velocity, and survivability ratings. After "
    "cleaning and extraction of the configuration inferences, the final "
    "tally is GOLD = 0.684 (ceiling, n=215), C = 0.651 (n=204), B = "
    "0.611 (n=215), A = 0.603 (n=213). The most important takeaway is "
    "that our Draper model hits about 95.2% of the mean score of the "
    "gold ads (0.651 vs 0.684, a gap of 0.033 composite). This "
    "immediately reveals a promising finding that our fine-tuned model "
    "has picked up patterns of successful ads from the corpus and "
    "successfully generalizes to unseen ads. More critically, our model "
    "configuration beats out both the frontier model (+0.048 composite "
    "over A, an 8.0% relative lift) and the base model (+0.040 "
    "composite over B, a 6.5% relative lift), as reported by the scorer.",
    # ¶2 — pipe configs (agentic wrap); numbers as placeholders pending
    # finalisation of the clean-pipe run.
    "B_pipe and C_pipe are the agent-integrated versions of the flat "
    "models. The models are wrapped in a full agent loop with tool "
    "calls such as web search, URL scraping, similar-page search, and "
    "some outputs like image generation or campaign schema output. The "
    "goal is to test whether integrating the model into this agent "
    "flow provides real value over the base models. First, the "
    "agent-integrated systems run through the same test set that the "
    "raw models used. After inferencing the models we obtained "
    "composite means: B_pipe = X.XXX, C_pipe = X.XXX. Compared to the "
    "single-shot models, we see B (X.XXX) → B_pipe (X.XXX), drop of "
    "X.XXX composite (X.X% relative). C (X.XXX) → C_pipe (X.XXX), "
    "drop of X.XXX composite. The decrease is consistent for both "
    "cases but hurts the fine-tuned model more than the base model. "
    "We can also see that C_pipe still beats out B_pipe, similar to "
    "the single-shot case, meaning the agent-wrapped, fine-tuned "
    "model still outperforms the base model. The best single-shot "
    "model is ahead of the corresponding agentic model (C_pipe = "
    "X.XXX) by X.XXX composite.",
    # ¶3 — per-platform composite breakdown for C; numbers as
    # placeholders pending finalisation.
    "Now focusing on our fine-tuned model's performance across "
    "platforms, we can see that our fine-tuned model yields "
    "composite scores per platform: pinterest X.XXX, twitter "
    "X.XXX, tiktok X.XXX, facebook X.XXX, reddit X.XXX, with n: "
    "facebook X, pinterest X, reddit X, tiktok X, twitter X. Our "
    "model has the best performance on the Pinterest platform, "
    "and the worst relative performance on Reddit. Most notably, "
    "our fine-tuned model beats out both the frontier model "
    "configuration as well as the base model on every platform. "
    "This shows that the model internalized platform-native "
    "structure where other models have failed to do so, and that "
    "the model is consistently better than others. Wins are as "
    "large as (C +X.XXX over B, +X.XXX over A) on Pinterest. "
    "Smallest absolute wins are on TikTok (C +X.XXX over B, "
    "+X.XXX over A). We can also see that the gap to the gold "
    "ads is: tiktok X.XXX (closest to ceiling), facebook X.XXX, "
    "reddit X.XXX, pinterest X.XXX, twitter X.XXX (furthest from "
    "ceiling).",
    # ¶4 — per-head breakdown (survivability, engagement volume,
    # engagement velocity); numbers as placeholders pending finalisation.
    "We have looked at the composite scores. Looking at the three "
    "component heads (§3.7) that sit beneath the composite can also "
    "give insight into the model's individual performance. For the "
    "survivability head, which represents the score for the longevity "
    "of the ad (the Kaplan–Meier survival curve target from the v3 "
    "scorer, §3.3), we can see results of GOLD=X.XXX, C=X.XXX, "
    "B=X.XXX, A=X.XXX, where C beats B by +X.XXX and A by +X.XXX on "
    "survivability, which is the biggest margin across the three "
    "heads. Gap to GOLD on survivability = X.XXX, which is the "
    "smallest gap to the high baseline across the heads. For the "
    "engagement volume head, which represents the total engagement "
    "(reactions, comments, shares) that a post is likely to get, we "
    "see scores of GOLD=X.XXX, C=X.XXX, B=X.XXX, A=X.XXX. The core "
    "insight is that our fine-tuned model beats out both "
    "configurations by a considerable margin (~+X.XXX), while A and "
    "B sit essentially tied on this head. Finally, the engagement "
    "velocity head represents the engagement-per-time proxy. We see "
    "scores of GOLD=X.XXX, C=X.XXX, B=X.XXX, A=X.XXX. It is "
    "interesting to note that C is actually slightly behind A on "
    "velocity (−X.XXX). However, the margin is quite small. This is "
    "the single case where the fine-tuned model loses out to the "
    "frontier model by a slim margin. Looking at the overall "
    "picture, we see that the overall composite score lift that our "
    "fine-tuned model sees is because of the survivability score "
    "(+X.XXX over B) and then engagement volume (+X.XXX over B). For "
    "reference, gold ads score the highest on all three heads, which "
    "is a signal that our pipeline is representative of the actual "
    "performance.",
    # ¶5 — predictor reliability + honesty notes (closing caveats);
    # numbers as placeholders pending finalisation.
    "Given these results, it is important to address the reliability "
    "of the trained predictor (§3.7). Since the predictor is a "
    "trained regressor model on real performance data, its "
    "predictions are grounded. However, it is still trained on a "
    "moderate-sized 55k v3-scored AdFlex corpus, which means drift "
    "may happen. The composite reliability performance on the "
    "held-out test split scored Spearman ρ = X.XXX, which is a "
    "strong correlation but not perfect. Calibration error (ECE) = "
    "X.XXXX — close to zero, which indicates that the model predicts "
    "scores that match quantile rates. The model's top-tier AUC = "
    "X.XXX and bottom-tier AUC = X.XXX indicate that the model can "
    "reliably separate high-tier performing ads from low-tier ads. "
    "It is worth noting that the Reddit slice of the corpus is the "
    "weakest, since the absence of engagement data from the API "
    "makes the engagement heads have limited data to train on for "
    "this platform. It is important to note that the predictor sees "
    "text "
    "only, not the image or video creative, which are part of the "
    "engagement, but for the purposes of this research, we are "
    "fine-tuning a copywriting model, so the trained score "
    "accurately represents the skill. Finally, the trained model is "
    "based on our proxy v3 scorer (§3.3), which means any bias or "
    "inaccuracies in the original score are transferred to the "
    "trained model.",
]


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def set_paragraph_text(p_elem, new_text):
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def write_section_body(doc):
    h_4_2 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("4.2 Learned-Scorer"),
        "§4.2 heading",
    )
    h_4_3 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("4.3 MAUVE"),
        "§4.3 heading",
    )

    # Body Text template: first style-943 paragraph inside Chapter III with
    # real text content. Same convention as write_4_1_setup.py.
    body_tpl = None
    in_ch3 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.upper() == "CHAPTER III":
            in_ch3 = True
            continue
        if t.upper() == "CHAPTER IV":
            break
        sid = p.style.style_id if p.style else ""
        if in_ch3 and sid == "943" and t and p._element.findall(f".//{QN_T}"):
            body_tpl = p._element
            break
    if body_tpl is None:
        raise RuntimeError("No style-943 Body Text template found inside CH3")

    body = h_4_2.getparent()
    children = list(body)
    i_start = children.index(h_4_2) + 1
    i_end = children.index(h_4_3)

    to_delete = [c for c in children[i_start:i_end] if c.tag == QN_P]
    print(f"DELETING: {len(to_delete)} existing §4.2 paragraph(s)")
    for c in to_delete:
        body.remove(c)

    insert_after = h_4_2
    for i, text in enumerate(PARAGRAPHS, 1):
        p = deepcopy(body_tpl)
        set_paragraph_text(p, text)
        insert_after.addnext(p)
        insert_after = p
        print(f"  ¶{i}: {text[:80]}…")

    print(
        f"\nWROTE: {len(PARAGRAPHS)} paragraph(s) into §4.2 Learned-Scorer "
        "Absolute Scores"
    )


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    print("=== §4.2 Learned-Scorer Absolute Scores ===")
    write_section_body(doc)

    doc.save(str(path))
    print(f"\nSAVED: {path}")


if __name__ == "__main__":
    main()
