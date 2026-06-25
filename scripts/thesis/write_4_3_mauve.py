"""Write §4.3 MAUVE Distribution Matching body (between §4.3 and §4.4).

Same wipe-and-replace pattern as write_4_1_setup.py: wipes every paragraph
between the "4.3 MAUVE Distribution Matching" heading and the "4.4
Fine-Tuning and Agent Ablation" heading, then re-inserts the current
PARAGRAPHS list using a Body Text (style 943) template paragraph cloned
from Chapter III.

Idempotent — re-running produces the same output.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"
QN_DRAWING = f"{{{W_NS}}}drawing"


def _is_figure_block(p_elem) -> bool:
    """Preserve image paragraphs and their caption paragraphs (start with
    'Figure 4.3.'). Lets us re-run this script without wiping embedded
    figures."""
    if p_elem.find(f".//{QN_DRAWING}") is not None:
        return True
    text = "".join((t.text or "") for t in p_elem.findall(f".//{QN_T}"))
    return text.strip().startswith("Figure 4.3.")


PARAGRAPHS: list[str] = [
    # ¶1 — MAUVE methodology + GOLD ceiling framing.
    "The second arm, which is the MAUVE arm (Pillutla et al., 2021), uses "
    "embedding-based comparison between ad pools related to the original "
    "corpus. Compared to §4.2, it offers a more holistic view of the model "
    "performance rather than pointwise comparison. This arm is used to "
    "cover all six configurations including the frontier model, base "
    "model, fine-tuned model, and the agent versions. GOLD ads represent "
    "the held-out ads from the original corpus, and naturally serve as "
    "the top baseline. For this method, each ad in both pools is embedded "
    "using the GPT-2 Large model (774M parameters), which transforms the "
    "text into vector embedding representations. This allows us to "
    "measure the embedding overlap between pools rather than just normal "
    "text, which captures nuances and deeper understanding of text. The "
    "output is a single number between zero and one, where zero is zero "
    "similarity while one is perfect similarity. A good reference ad "
    "scores well below the 1.0 threshold since practically the corpus "
    "size as well as the data variety makes it impossible to match "
    "perfectly; therefore it lands at 0.462, creating a ceiling baseline "
    "for the other models where they are expected to score lower.",
    # ¶2 — corpus-level rankings, fine-tuning + agent-wrap effects,
    # cross-arm disagreement flag.
    "The final score rankings are as follows: GOLD at 0.462 (ceiling), "
    "C_pipe at 0.420 (highest non-GOLD configuration), B_pipe at 0.302, "
    "C at 0.287, B at 0.183, and A at 0.180. We observe that A and B, "
    "which are the traditional AI models, are nearly tied, showing that "
    "generalized AI models have similar performance regardless of size. "
    "The most promising result is that our fine-tuned model raises the "
    "bar from 0.183 → 0.287 alone, which shows our fine-tuning has "
    "extensive effect on the style of the output to make it match our "
    "high-performer ad corpus. Additionally, wrapping any model in an "
    "agent workflow has positive effect on all configurations: B → "
    "B_pipe is +0.119, C → C_pipe is +0.133. The best configuration is "
    "our fine-tuned model and agent-wrapped workflow, reaching about 91% "
    "of the GOLD ceiling (0.420 / 0.462). The bottom-line result is that "
    "our agent wrapper and fine-tuned model have a combined effect to "
    "improve the model from the base version to the fine-tuned version. "
    "This finding differs fundamentally from the findings from the "
    "previous section.",
]


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def set_paragraph_text(p_elem, new_text):
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def write_section_body(doc):
    h_4_3 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("4.3 MAUVE Distribution Matching"),
        "§4.3 heading",
    )
    h_4_4 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("4.4 Fine-Tuning and Agent Ablation"),
        "§4.4 heading",
    )

    # Body Text template: first style-943 paragraph inside Chapter III with
    # real text content. Same convention as write_4_1_setup.py.
    body_tpl = None
    in_ch3 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.upper() == "CHAPTER III":
            in_ch3 = True
            continue
        if t.upper() == "CHAPTER IV":
            break
        sid = p.style.style_id if p.style else ""
        if in_ch3 and sid == "943" and t and p._element.findall(f".//{QN_T}"):
            body_tpl = p._element
            break
    if body_tpl is None:
        raise RuntimeError("No style-943 Body Text template found inside CH3")

    body = h_4_3.getparent()
    children = list(body)
    i_start = children.index(h_4_3) + 1
    i_end = children.index(h_4_4)

    to_delete = [
        c
        for c in children[i_start:i_end]
        if c.tag == QN_P and not _is_figure_block(c)
    ]
    print(f"DELETING: {len(to_delete)} existing §4.3 paragraph(s) (figures preserved)")
    for c in to_delete:
        body.remove(c)

    insert_after = h_4_3
    for i, text in enumerate(PARAGRAPHS, 1):
        p = deepcopy(body_tpl)
        set_paragraph_text(p, text)
        insert_after.addnext(p)
        insert_after = p
        print(f"  ¶{i}: {text[:80]}…")

    print(f"\nWROTE: {len(PARAGRAPHS)} paragraph(s) into §4.3 MAUVE Distribution Matching")


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    print("=== §4.3 MAUVE Distribution Matching ===")
    write_section_body(doc)

    doc.save(str(path))
    print(f"\nSAVED: {path}")


if __name__ == "__main__":
    main()
