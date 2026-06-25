"""Write §4.4 Fine-Tuning and Agent Ablation body (between §4.4 and §4.5).

Same wipe-and-replace pattern as write_4_3_mauve.py: wipes every paragraph
between the "4.4 Fine-Tuning and Agent Ablation" heading and the "4.5
Synthesis and Limitations" heading, then re-inserts the current PARAGRAPHS
list using a Body Text (style 943) template paragraph cloned from Chapter
III.

Idempotent — re-running produces the same output. Figure paragraphs
(those containing a <w:drawing> or whose visible text starts with
'Figure 4.4.') are preserved across runs.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"
QN_DRAWING = f"{{{W_NS}}}drawing"


def _is_figure_block(p_elem) -> bool:
    """Preserve image paragraphs and their caption paragraphs (start with
    'Figure 4.4.'). Lets us re-run this script without wiping embedded
    figures."""
    if p_elem.find(f".//{QN_DRAWING}") is not None:
        return True
    text = "".join((t.text or "") for t in p_elem.findall(f".//{QN_T}"))
    return text.strip().startswith("Figure 4.4.")


PARAGRAPHS: list[str] = [
    # ¶1 — 2x2 ablation setup + per-cell composite means + headline that C
    # (fine-tuned, no agent) is the highest cell, contradicting §4.3.
    "This section reports the 2×2 ablation results for RQ2. The "
    "two-by-two design examines the four combinations possible from the "
    "binary factors as fine-tuning / RAG (agent), which allows us to get "
    "four different experimental results representing all combinations of "
    "the above to see their individual effects on performance. Factor one "
    "is the presence or absence of fine-tuning. We either use the base "
    "untuned Qwen3-8B model, or our fine-tuned Draper model "
    "interchangeably. Factor two is whether the writer model is wrapped "
    "in an agentic workflow with RAG capabilities. This tests the on and "
    "off effects of Retrieval-Augmented Generation on the quality of "
    "output in the context of ad generation. Crossing these four "
    "configurations gives us the four different ablations to test side "
    "by side. All these configurations are run over the same 215-brief "
    "held-out test split to ensure a fair comparison. Configuration "
    "results are: B = 0.611 (n=215), C = 0.651 "
    "(n=204), B_pipe = 0.586 (n=139), C_pipe = 0.607 (n=145). From these "
    "results, we can observe that we find through the raw arm without "
    "agentic workflow as the highest score according to our learned "
    "scorer (§3.7), contradictory to the findings of §4.3 where the "
    "agent-wrapped model had the clearest high performance.",
    # ¶2 — paired-contrast methodology + three significant effects (FT
    # helps, agent hurts FT, FT-only beats base+agent) + the three n.s.
    # contrasts as the foil; closes on the headline that fine-tuning
    # lifts and the agent does not.
    "For paired contrast, we simply cannot just subtract cell scores, "
    "simply because the total number of successfully parsed briefs for "
    "each configuration differs due to some issues like generation "
    "errors, format mismatch, or other failures during inference. "
    "Subtracting the total results from each means they are tested over "
    "different-sized datasets, which sets an unfair advantage. Instead, "
    "we only measure the intersection of briefs for the four cells in "
    "the configurations. That intersection = 94 briefs for this case. "
    "Each brief gets a single score for all four configurations that "
    "is comparable since we eliminate the brief-to-brief variance. The "
    "six contrasts (95% CIs from 1000 bootstrap resamples on the n=94 "
    "paired briefs, seed=42): C − B = +0.036 [+0.017, +0.055] — "
    "fine-tuning alone helps; C_pipe − C = −0.040 [−0.061, −0.020] — "
    "agent hurts the fine-tuned writer; C − B_pipe = +0.054 [+0.034, "
    "+0.073] — fine-tuning alone beats base with agent. Other "
    "comparisons are shown to be too little or noise to make any "
    "conclusive assessment: C_pipe − B = −0.005, B_pipe − B = −0.018, "
    "C_pipe − B_pipe = +0.013. Seeing this data, the most critical "
    "takeaway is that the fine-tuning lifts up the score from the base "
    "model, consistent with the earlier 3-arm C vs B comparison at "
    "Cohen's dz ≈ 0.41. This is directly a proof of effective fine-tuning on "
    "improving the model's ability to generate successful ad patterns "
    "rather than generic AI ones. You can also see that adding the "
    "agent, the RAG workflow, does not help any configuration. And, "
    "actually, it reduces the performance for our fine-tuned model.",
    # ¶3 — distribution-shift mechanism: test briefs mimic training-time
    # user form; agent enrichment pushes the brief off-distribution, which
    # hurts the FT writer more than the base writer.
    "The test-set brief shape is similar to the training examples and "
    "consists of naturally written briefs in user form. This was "
    "designed during construction to mimic what a user would type into "
    "a model in the front end. This brief is plain language and simple, "
    "doesn't contain platform preferences, detailed product info, or "
    "other specific information. On the other hand, the agent pipeline "
    "transforms these briefs after performing research and tool calls "
    "into one whose product fields (description, key features, USPs) "
    "have been shaped by the orchestrator's research before reaching "
    "the writer. This makes the brief slightly different from the "
    "training distribution, which would explain the performance drop "
    "especially in the fine-tuned model (C_pipe − C = −0.040, "
    "significant), where we can see the base model has a negligible "
    "difference vs the agent-wrapped one (B_pipe − B = −0.018, not "
    "significant), meaning strengthening the prompts with research data "
    "and detailed product information has negative returns when it "
    "comes to the fine-tuned model.",
    # ¶4 — closing caveats: text-only scorer underweights what the agent
    # actually delivers (grounding, accuracy, multi-step reasoning, image
    # generation) and these are framed as limitations of the offline
    # methodology rather than a verdict on agent value.
    "Although the agent results are consistently lower, it's important "
    "to note that the trained scorer (§3.7) is likely to be somewhat "
    "unreliable for measuring the effect of RAG on output quality, "
    "since RAG is about grounding in information and facts rather than "
    "engagement signals which the model was trained to measure. For "
    "that reason, the agent workflow offers other advantages in "
    "practice that may not be justifiably measured by these models. "
    "For instance, it can provide more grounded marketing "
    "advertisements as well as being aware of competitor analysis and "
    "market trends; it also has the capability for integration of "
    "tools such as image generation or others for full campaign "
    "generation. The multi-step agent also has the benefit of better "
    "interpretability (via persisted step traces) as well as multi-step reasoning and execution "
    "workflows, allowing the final campaign to be more adherent to the "
    "requirements than a single-shot generation. All of these and the "
    "above are marked as the limitations of this methodology for "
    "measuring objective quality of generated advertising by the LLMs.",
]


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def set_paragraph_text(p_elem, new_text):
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def write_section_body(doc):
    h_4_4 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("4.4 Fine-Tuning and Agent Ablation"),
        "§4.4 heading",
    )
    h_4_5 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("4.5 Synthesis and Limitations"),
        "§4.5 heading",
    )

    # Body Text template: first style-943 paragraph inside Chapter III with
    # real text content. Same convention as write_4_1_setup.py.
    body_tpl = None
    in_ch3 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.upper() == "CHAPTER III":
            in_ch3 = True
            continue
        if t.upper() == "CHAPTER IV":
            break
        sid = p.style.style_id if p.style else ""
        if in_ch3 and sid == "943" and t and p._element.findall(f".//{QN_T}"):
            body_tpl = p._element
            break
    if body_tpl is None:
        raise RuntimeError("No style-943 Body Text template found inside CH3")

    body = h_4_4.getparent()
    children = list(body)
    i_start = children.index(h_4_4) + 1
    i_end = children.index(h_4_5)

    to_delete = [
        c
        for c in children[i_start:i_end]
        if c.tag == QN_P and not _is_figure_block(c)
    ]
    print(f"DELETING: {len(to_delete)} existing §4.4 paragraph(s) (figures preserved)")
    for c in to_delete:
        body.remove(c)

    insert_after = h_4_4
    for i, text in enumerate(PARAGRAPHS, 1):
        p = deepcopy(body_tpl)
        set_paragraph_text(p, text)
        insert_after.addnext(p)
        insert_after = p
        print(f"  ¶{i}: {text[:80]}…")

    print(
        f"\nWROTE: {len(PARAGRAPHS)} paragraph(s) into §4.4 Fine-Tuning and "
        "Agent Ablation"
    )


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    print("=== §4.4 Fine-Tuning and Agent Ablation ===")
    write_section_body(doc)

    doc.save(str(path))
    print(f"\nSAVED: {path}")


if __name__ == "__main__":
    main()
