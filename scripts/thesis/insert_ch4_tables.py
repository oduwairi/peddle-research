"""Insert Chapter IV results tables (Tables 4.1-4.9) into THESIS.docx.

Reviewer feedback #9: "Add images ... to show all the evaluation metrics
(accuracy, precision, F1-score, etc.) for the comparison of the models ...
both in tables and figures."

Chapter IV currently presents every model comparison as a FIGURE only, with
no companion data table. This script tabulates the exact numbers behind each
Chapter IV figure, computed from the SAME on-disk eval data with the SAME
bootstrap settings (seed=42) the figure generators use, so table and figure
share one source of truth. No new measurements are taken and no body prose is
written.

("accuracy/precision/F1" do not apply to a generative copywriting task; the
legitimate classification-flavoured analogues — ROC-AUC, calibration error,
and the Upworthy decision accuracy — are tabulated in Tables 4.4 and 4.7.)

Each table is inserted immediately after its figure's body caption, with a
"Table 4.N: ..." caption above it (matching the Chapter III table convention),
plus a minimal "(Table 4.N)" parenthetical callout on the adjacent discussion
paragraph (reviewer #5 pattern, AI-detector-safe). Config A is labelled
GPT-5.5 (per the reviewer #6 standardisation), not gpt-5.4-mini.

Idempotent: prior "Table 4.x" insertions are removed before re-inserting, and
callouts are skipped if already present. Anchors are located by unique text,
never hardcoded indices (per docs/research/THESIS_EDITING.md).

Run:  uv run python scripts/thesis/insert_ch4_tables.py [--dry-run]
"""

from __future__ import annotations

import copy
import json
import sys
from pathlib import Path

import numpy as np
import polars as pl
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DOCX = Path("docs/research/THESIS.docx")

RUN_SINGLE = Path("data/eval/runs/2026-05-10-learned-may-smoke/aggregates")
RUN_PIPE = Path("data/eval/runs/2026-05-14-clean-pipe/aggregates")
RUN_MAUVE = Path("data/eval/runs/2026-05-22-mauve-with-pipes/aggregates")
PER_ROW = Path("data/eval/learned_scores")
REF_DIR = Path("data/eval/reference_scores")
VAL_DIR = Path("data/eval/validation")
EVAL_REPORT = Path("data/scoring_predictor/checkpoints/random/best/eval_report.json")

CONFIG_ORDER = ["A", "B", "B_pipe", "C", "C_pipe", "GOLD"]
CONFIG_DESC = {
    "A": "GPT-5.5",
    "B": "Qwen3-8B (base)",
    "B_pipe": "Qwen3-8B + agent",
    "C": "Draper (fine-tuned)",
    "C_pipe": "Draper + agent",
    "GOLD": "Real ads (ceiling)",
}
PLATFORMS = ["facebook", "pinterest", "reddit", "tiktok", "twitter"]
PLATFORM_LABEL = {
    "facebook": "Facebook",
    "pinterest": "Pinterest",
    "reddit": "Reddit",
    "tiktok": "TikTok",
    "twitter": "X (Twitter)",
}


# --------------------------------------------------------------------------- #
# Bootstrap helpers — copied verbatim from the figure generators so the table
# numbers are bit-for-bit the same as the plotted means/CIs.
# --------------------------------------------------------------------------- #
def boot_ci_choice(values: np.ndarray, n_bootstrap: int = 2000, seed: int = 42):
    """generate_figures_4_2.bootstrap_ci"""
    rng = np.random.default_rng(seed)
    n = len(values)
    means = np.empty(n_bootstrap)
    for i in range(n_bootstrap):
        means[i] = rng.choice(values, size=n, replace=True).mean()
    lo = float(np.percentile(means, 2.5))
    hi = float(np.percentile(means, 97.5))
    return float(values.mean()), lo, hi


def boot_ci_idx(diffs: np.ndarray, n_boot: int = 1000, seed: int = 42):
    """generate_figures_4_4.bootstrap_ci (paired) / refoverlap boot_ci"""
    rng = np.random.default_rng(seed)
    n = len(diffs)
    boot = np.empty(n_boot)
    for i in range(n_boot):
        idx = rng.integers(0, n, n)
        boot[i] = diffs[idx].mean()
    lo, hi = np.percentile(boot, [2.5, 97.5])
    return float(diffs.mean()), float(lo), float(hi)


def f3(x) -> str:
    return f"{x:.3f}"


def ci3(lo, hi) -> str:
    return f"[{lo:.3f}, {hi:.3f}]"


# --------------------------------------------------------------------------- #
# Per-table number computation
# --------------------------------------------------------------------------- #
def compute_tables() -> list[dict]:
    tables: list[dict] = []

    # Per-row score cache (reused for bootstrap CIs in Tables 4.2/4.3/4.9).
    per_row = {
        c: pl.read_parquet(PER_ROW / f"{c}.parquet")
        for c in CONFIG_ORDER
        if (PER_ROW / f"{c}.parquet").exists()
    }

    # ---- Table 4.1 : composite by config (Fig 4.1) -----------------------
    single = pl.read_parquet(RUN_SINGLE / "learned_scores_summary.parquet")
    pipe = pl.read_parquet(RUN_PIPE / "learned_scores_summary.parquet")
    summ = pl.concat([single, pipe.filter(pl.col("config").is_in(["B_pipe", "C_pipe"]))])
    rows = []
    for cfg in CONFIG_ORDER:
        p = PER_ROW / f"{cfg}.parquet"
        if p.exists():
            arr = pl.read_parquet(p)["composite"].drop_nulls().to_numpy()
            m, lo, hi = boot_ci_choice(arr)
            n = len(arr)
        else:
            r = summ.filter(pl.col("config") == cfg).row(0, named=True)
            m, lo, hi, n = r["composite_mean"], r["composite_p25"], r["composite_p75"], r["n"]
        rows.append([cfg, CONFIG_DESC[cfg], f3(m), ci3(lo, hi), str(n)])
    tables.append({
        "num": "4.1",
        "anchor_fig": "Figure 4.1:",
        "caption": "Table 4.1: Composite score per configuration on the 215-brief "
                   "held-out test set (mean with 95% bootstrap confidence interval). "
                   "Tabulates the values plotted in Figure 4.1.",
        "headers": ["Config", "Model", "Composite (mean)", "95% CI", "n"],
        "rows": rows,
        "callout": "reports the findings of the first evaluation arm",
    })

    # ---- Table 4.2 : composite by platform (Fig 4.2) ---------------------
    # Per-cell mean + 95% bootstrap CI from the per-row scores (same settings
    # as Table 4.1 / Fig 4.2: boot_ci_choice, 2000 resamples, seed=42).
    cfgs = ["A", "B", "C", "GOLD"]
    rows = []
    for plat in PLATFORMS:
        cells = [PLATFORM_LABEL[plat]]
        for cfg in cfgs:
            vals = per_row[cfg].filter(pl.col("platform") == plat)["composite"].drop_nulls().to_numpy()
            if len(vals) == 0:
                cells.append("—")
            else:
                m, lo, hi = boot_ci_choice(vals)
                cells.append(f"{f3(m)} {ci3(lo, hi)}")
        rows.append(cells)
    tables.append({
        "num": "4.2",
        "anchor_fig": "Figure 4.2:",
        "caption": "Table 4.2: Composite score by platform and configuration on the "
                   "215-brief held-out test set (single-shot configurations; mean with "
                   "95% bootstrap confidence interval). Tabulates Figure 4.2.",
        "headers": ["Platform", "A (GPT-5.5)", "B (Qwen3-8B)", "C (Draper)", "GOLD"],
        "rows": rows,
        "callout": "yields composite scores per platform",
    })

    # ---- Table 4.3 : per-head means by config (Fig 4.3) ------------------
    # Per-head mean + 95% bootstrap CI from the per-row scores.
    heads = ["composite", "survivability", "engagement_volume", "engagement_velocity"]
    rows = []
    for cfg in cfgs:
        df = per_row[cfg]
        cells = [f"{cfg} — {CONFIG_DESC[cfg]}"]
        for h in heads:
            m, lo, hi = boot_ci_choice(df[h].drop_nulls().to_numpy())
            cells.append(f"{f3(m)} {ci3(lo, hi)}")
        rows.append(cells)
    tables.append({
        "num": "4.3",
        "anchor_fig": "Figure 4.3:",
        "caption": "Table 4.3: Per-head learned-scorer means by configuration on the "
                   "215-brief held-out test set (mean with 95% bootstrap confidence "
                   "interval). Tabulates Figure 4.3.",
        "headers": ["Configuration", "Composite", "Survivability",
                    "Engagement volume", "Engagement velocity"],
        "rows": rows,
        "callout": "Looking at the three component heads",
    })

    # ---- Table 4.4 : predictor reliability (Fig 4.4) ---------------------
    rep = json.loads(EVAL_REPORT.read_text())
    m = rep["metrics"]
    hh = ["composite", "survivability", "engagement_volume", "engagement_velocity"]
    rows = [
        ["Spearman ρ"] + [f3(m[f"spearman_{h}"]) for h in hh],
        ["Pearson r"] + [f3(m[f"pearson_{h}"]) for h in hh],
        ["MAE"] + [f3(m[f"mae_{h}"]) for h in hh],
        ["Top-tier ROC-AUC (target ≥ 0.80)", f3(rep["composite_auc_top_tier"]), "—", "—", "—"],
        ["Bottom-tier ROC-AUC (target ≤ 0.30)", f3(rep["composite_auc_bottom_tier"]), "—", "—", "—"],
        ["Calibration error (ECE)", f"{rep['composite_ece']:.4f}", "—", "—", "—"],
    ]
    tables.append({
        "num": "4.4",
        "anchor_fig": "Figure 4.4:",
        "caption": f"Table 4.4: Scoring-predictor reliability on the held-out random "
                   f"split (n = {rep['n_test']:,}). Tabulates the quantities summarised "
                   f"in Figure 4.4; the tier ROC-AUC and ECE rows apply to the composite "
                   f"head only.",
        "headers": ["Reliability metric", "Composite", "Survivability",
                    "Engagement volume", "Engagement velocity"],
        "rows": rows,
        "callout": "it is important to address the reliability of the trained predictor",
    })

    # ---- Table 4.5 : per-platform MAUVE (Fig 4.3.2) ----------------------
    mv = pl.read_parquet(RUN_MAUVE / "mauve_scores_by_platform.parquet")
    plat_cols = PLATFORMS + ["ALL"]
    rows = []
    for cfg in CONFIG_ORDER:
        cells = [f"{cfg} — {CONFIG_DESC[cfg]}"]
        for plat in plat_cols:
            r = mv.filter((pl.col("config") == cfg) & (pl.col("platform") == plat))
            cells.append(f3(r.row(0, named=True)["mauve_mean"]) if not r.is_empty() else "—")
        rows.append(cells)
    tables.append({
        "num": "4.5",
        "anchor_fig": "Figure 4.3.2:",
        "caption": "Table 4.5: Per-platform MAUVE score by configuration (mean; higher "
                   "is closer to the high-tier reference distribution). 95% bootstrap "
                   "confidence intervals are shown in Figure 4.3.2.",
        "headers": ["Configuration"] + [PLATFORM_LABEL[p] for p in PLATFORMS] + ["ALL"],
        "rows": rows,
        "callout": "The final score rankings are as follows",
    })

    # ---- Table 4.6 : reference overlap by config (Fig 4.4.1) -------------
    metrics = [("bleu_gold", "BLEU"), ("rouge_l_gold", "ROUGE-L"),
               ("meteor_gold", "METEOR"), ("chrf_gold", "chrF"),
               ("bertscore_gold", "BERTScore")]
    data = {c: pl.read_parquet(REF_DIR / f"{c}.parquet") for c in ("A", "B", "C")}
    gold = pl.read_parquet(REF_DIR / "GOLD.parquet")
    rows = []
    for col, label in metrics:
        cells = [label]
        for cfg in ("A", "B", "C"):
            vals = np.asarray([v for v in data[cfg][col].to_list() if v is not None], float)
            vals = vals[~np.isnan(vals)]
            mm, lo, hi = boot_ci_idx(vals)
            cells.append(f"{f3(mm)} {ci3(lo, hi)}")
        ceil = float(gold[col.replace("_gold", "_multi")].drop_nulls().mean())
        cells.append(f3(ceil))
        rows.append(cells)
    tables.append({
        "num": "4.6",
        "anchor_fig": "Figure 4.4.1:",
        "caption": "Table 4.6: Per-ad reference overlap with the real winning ad by "
                   "configuration (mean with 95% bootstrap CI). BLEU and chrF are on a "
                   "0–100 scale; ROUGE-L, METEOR, and BERTScore on a 0–1 scale. The GOLD "
                   "ceiling is the real ad's overlap with the winner pool. Tabulates "
                   "Figure 4.4.1.",
        "headers": ["Overlap metric", "A (GPT-5.5)", "B (Qwen3-8B)",
                    "C (Draper)", "GOLD ceiling"],
        "rows": rows,
        "callout": None,  # resolved at runtime; see CALLOUT_FALLBACKS
    })

    # ---- Table 4.7 : Upworthy grounding (Fig 4.4.2) ----------------------
    up = [("meteor", "METEOR"), ("chrf", "chrF"), ("rouge_l", "ROUGE-L"), ("bleu", "BLEU")]
    rows = []
    for key, label in up:
        d = json.loads((VAL_DIR / f"refmetrics_{key}_upworthy.json").read_text())
        npairs, nc, nt = d["n_pairs"], d["n_correct"], d["n_ties"]
        nw = npairs - nt - nc
        rows.append([label, str(npairs), str(nc), str(nw), str(nt),
                     f3(d["accuracy"]), f"{d['binomial_p_value']:.3f}"])
    tables.append({
        "num": "4.7",
        "anchor_fig": "Figure 4.4.2:",
        "caption": "Table 4.7: How each overlap metric decides the 200 Upworthy A/B "
                   "headline tests; accuracy is computed over decided (non-tied) pairs "
                   "only. Tabulates Figure 4.4.2.",
        "headers": ["Metric", "Pairs", "Correct", "Wrong", "Tied",
                    "Accuracy", "p-value"],
        "rows": rows,
        "callout": "it is important to validate the accuracy of these metrics",
    })

    # ---- Table 4.8 : paired contrasts (Fig 4.5.1) ------------------------
    cells_cfg = ["B", "C", "B_pipe", "C_pipe"]
    parts = {c: pl.read_parquet(PER_ROW / f"{c}.parquet").select(["example_id", "composite"]).rename({"composite": c}) for c in cells_cfg}
    joined = parts["B"]
    for c in cells_cfg[1:]:
        joined = joined.join(parts[c], on="example_id", how="inner")
    n_paired = len(joined)
    arr = joined.select(cells_cfg).to_numpy()
    contrasts = [
        ("C", "B", "fine-tuning, no agent"),
        ("C_pipe", "C", "agent on fine-tuned"),
        ("C", "B_pipe", "fine-tuned vs base+agent"),
        ("C_pipe", "B", "full product vs base"),
        ("B_pipe", "B", "agent on base"),
        ("C_pipe", "B_pipe", "fine-tuning effect, both with agent"),
    ]
    rows = []
    for a, b, desc in contrasts:
        diffs = arr[:, cells_cfg.index(a)] - arr[:, cells_cfg.index(b)]
        mm, lo, hi = boot_ci_idx(diffs)
        sig = "Yes" if (lo > 0 or hi < 0) else "No"
        rows.append([f"{a} − {b}", desc, f"{mm:+.3f}", f"[{lo:+.3f}, {hi:+.3f}]", sig])
    tables.append({
        "num": "4.8",
        "anchor_fig": "Figure 4.5.1:",
        "caption": f"Table 4.8: Paired contrasts on the 2×2 ablation (n = {n_paired} "
                   f"paired briefs; 1000 bootstrap resamples, seed = 42). A contrast is "
                   f"significant when its 95% CI excludes zero. Tabulates Figure 4.5.1.",
        "headers": ["Contrast", "Interpretation", "Mean Δ", "95% CI", "Significant"],
        "rows": rows,
        "callout": "we simply cannot just subtract cell scores",
    })

    # ---- Table 4.9 : 2x2 cell means (Fig 4.5.2) --------------------------
    full = {}
    for c in cells_cfg:
        vals = per_row[c]["composite"].drop_nulls().to_numpy()
        m, lo, hi = boot_ci_choice(vals)
        full[c] = (m, lo, hi, len(vals))

    def cell9(c: str) -> str:
        m, lo, hi, n = full[c]
        return f"{c} = {f3(m)} {ci3(lo, hi)} (n = {n})"

    rows = [
        ["Base Qwen3-8B", cell9("B"), cell9("B_pipe")],
        ["Fine-tuned Draper", cell9("C"), cell9("C_pipe")],
    ]
    tables.append({
        "num": "4.9",
        "anchor_fig": "Figure 4.5.2:",
        "caption": "Table 4.9: Per-cell composite means on the 215-brief held-out test "
                   "set (2×2 writer × agent design; mean with 95% bootstrap confidence "
                   "interval). Tabulates Figure 4.5.2.",
        "headers": ["Writer \\ Agent", "Agent off (direct call)",
                    "Agent on (orchestrator)"],
        "rows": rows,
        "callout": "The two-by-two design examines the four combinations",
    })

    return tables


# Candidate callout anchors for tables whose discussion paragraph differs from
# the figure callout map (resolved best-effort, reported if missing).
CALLOUT_FALLBACKS = {
    "4.6": [
        "C leads all five gold-reference metrics",
        "clear separation between C and the other configurations",
    ],
}


# --------------------------------------------------------------------------- #
# docx manipulation
# --------------------------------------------------------------------------- #
def is_caption(text: str) -> bool:
    t = text.lstrip()
    return t.startswith("Figure") or t.startswith("Table")


def find_body_caption(doc: Document, prefix: str) -> Paragraph | None:
    """The body figure caption (style Body Text), not the List-of-Figures twin."""
    for p in doc.paragraphs:
        if p.style.name == "Body Text" and p.text.strip().startswith(prefix):
            return p
    return None


def caption_template(doc: Document) -> Paragraph:
    for p in doc.paragraphs:
        if p.style.name == "Body Text" and p.text.strip().startswith("Table 3.2:"):
            if p._p.find(".//" + qn("w:drawing")) is None:
                return p
    raise RuntimeError("could not find a safe 'Table 3.2:' caption template")


def set_caption_text(p_el, text: str) -> None:
    para = Paragraph(p_el, None)
    runs = para.runs
    if not runs:
        para.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def set_cell(cell, text: str, bold: bool) -> None:
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(text)
    if bold:
        run.bold = True


def remove_prior_tables(doc: Document) -> int:
    removed = 0
    for p in list(doc.paragraphs):
        if p.style.name == "Body Text" and p.text.strip().startswith("Table 4."):
            el = p._p
            nxt = el.getnext()
            if nxt is not None and nxt.tag == qn("w:tbl"):
                nxt.getparent().remove(nxt)
            el.getparent().remove(el)
            removed += 1
    return removed


def insert_table(doc: Document, spec: dict, tmpl: Paragraph) -> str:
    anchor = find_body_caption(doc, spec["anchor_fig"])
    if anchor is None:
        return f"!! NO ANCHOR {spec['anchor_fig']!r}"

    # caption paragraph (cloned from the Table 3.2 caption → inherits Body Text)
    cap_el = copy.deepcopy(tmpl._p)
    set_caption_text(cap_el, spec["caption"])
    anchor._p.addnext(cap_el)

    # table (built at end of body, then moved after the caption)
    headers, data = spec["headers"], spec["rows"]
    tbl = doc.add_table(rows=len(data) + 1, cols=len(headers))
    tbl.style = doc.styles["Table Grid"]
    for ci, h in enumerate(headers):
        set_cell(tbl.rows[0].cells[ci], h, bold=True)
    for ri, row in enumerate(data, start=1):
        for ci, val in enumerate(row):
            set_cell(tbl.rows[ri].cells[ci], val, bold=False)
    cap_el.addnext(tbl._tbl)
    return f"++ Table {spec['num']} ({len(data)}x{len(headers)}) after {spec['anchor_fig']!r}"


def append_callout(p: Paragraph, label: str) -> None:
    runs = p.runs
    idx = next((k for k in range(len(runs) - 1, -1, -1) if runs[k].text and runs[k].text.strip()), None)
    insert = f" ({label})"
    if idx is None:
        p.add_run(insert)
        return
    r = runs[idx]
    stripped = r.text.rstrip()
    trail = r.text[len(stripped):]
    r.text = (stripped[:-1] + insert + "." + trail) if stripped.endswith(".") else (stripped + insert + trail)


def add_callout(doc: Document, spec: dict) -> str:
    label = f"Table {spec['num']}"
    anchors = []
    if spec.get("callout"):
        anchors.append(spec["callout"])
    anchors += CALLOUT_FALLBACKS.get(spec["num"], [])
    for anchor in anchors:
        hits = [p for p in doc.paragraphs if anchor in p.text and not is_caption(p.text)]
        if len(hits) == 1:
            p = hits[0]
            if f"({label})" in p.text:
                return f"== {label} callout already present"
            append_callout(p, label)
            return f"++ {label} callout on …{p.text.rstrip()[-46:]!r}"
    return f"!! {label} callout NO unambiguous anchor (skipped)"


def main() -> int:
    dry = "--dry-run" in sys.argv
    if not DOCX.exists():
        print(f"ERROR: {DOCX} not found", file=sys.stderr)
        return 1
    if Path("docs/research/.~lock.THESIS.docx#").exists():
        print("ERROR: THESIS.docx is open in an editor (lock present). Close it.", file=sys.stderr)
        return 1

    specs = compute_tables()

    if dry:
        print("=== DRY RUN — computed table data ===\n")
        for s in specs:
            print(f"Table {s['num']}  (after {s['anchor_fig']})")
            print("  " + s["caption"])
            print("  H: " + " | ".join(s["headers"]))
            for r in s["rows"]:
                print("     " + " | ".join(r))
            print()
        return 0

    doc = Document(str(DOCX))
    tmpl = caption_template(doc)
    removed = remove_prior_tables(doc)
    print(f"Removed {removed} prior Table 4.x block(s).\n")

    print("TABLE INSERTIONS:")
    errors = 0
    for s in specs:
        msg = insert_table(doc, s, tmpl)
        print("  " + msg)
        if msg.startswith("!!"):
            errors += 1

    print("\nCALLOUTS:")
    for s in specs:
        print("  " + add_callout(doc, s))

    if errors:
        print(f"\n{errors} anchor error(s) — not saving.", file=sys.stderr)
        return 2

    doc.save(str(DOCX))
    print(f"\nSaved {DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
