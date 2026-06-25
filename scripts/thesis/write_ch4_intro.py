"""Write the Chapter IV intro paragraph (between the chapter title and §4.1).

Same wipe-and-replace pattern as write_3_3_scoring.py: wipes every paragraph
between the "Result and discussion" chapter-title paragraph and the §4.1
heading, then re-inserts the current PARAGRAPHS list using a Body Text
(style 943) template paragraph cloned from Chapter III.

Idempotent — re-running produces the same output.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"


PARAGRAPHS: list[str] = [
    # ¶1 — author voice (Phase 4), polished Phase 5, enriched Phase 6
    "Following the methodology, this chapter of this thesis delves into the "
    "evaluation of the obtained model to understand the performance of the "
    "fine-tuned model on performance data when compared to other baselines "
    "such as the frontier model (gpt-5.4-mini) or the base non-fine-tuned "
    "model. In Chapter III, we discussed two evaluation procedures — per-ad "
    "absolute scoring via the trained regressor (§3.7) and corpus-level "
    "distribution matching with MAUVE (Pillutla et al., 2021) — to measure "
    "the output quality of our model. In this chapter, we will provide "
    "results and reports side by side. Mainly, we are comparing the "
    "fine-tuned Draper model against a frontier model to answer the core "
    "research question (RQ1) whether a 7–9B fine-tuned model plus agent "
    "capabilities can compete with frontier large proprietary models. "
    "Secondary evaluation is also used to compare our fine-tuned model to "
    "the base model, as well as the effect of RAG and fine-tuning on the "
    "performance interchangeably (RQ2). The findings can be summarized as: "
    "in fine-tuning, the model closes the gap to high performance more "
    "than any other alternative (+0.036 composite over the base model, see "
    "§4.4). However, agent and RAG capabilities report mixed results when "
    "it comes to improving the fine-tuned output (−0.040 composite on the "
    "fine-tuned writer, see §4.4).",
]


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def set_paragraph_text(p_elem, new_text):
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def write_intro_body(doc):
    # Chapter title sits between CHAPTER IV (style 944) and §4.1 (style 945).
    # It is itself style 945 with text "Result and discussion".
    title = find_paragraph(
        doc,
        lambda p: p.text.strip() == "Result and discussion",
        '"Result and discussion" chapter title',
    )
    h_4_1 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("4.1 Evaluation Setup"),
        "§4.1 heading",
    )

    # Body Text template: first style-943 paragraph inside Chapter III with
    # actual text content. Same convention as write_3_3_scoring.py.
    body_tpl = None
    in_ch3 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.upper() == "CHAPTER III":
            in_ch3 = True
            continue
        if t.upper() == "CHAPTER IV":
            break
        sid = p.style.style_id if p.style else ""
        if in_ch3 and sid == "943" and t and p._element.findall(f".//{QN_T}"):
            body_tpl = p._element
            break
    if body_tpl is None:
        raise RuntimeError("No style-943 Body Text template found inside CH3")

    body = title.getparent()
    children = list(body)
    i_start = children.index(title) + 1
    i_end = children.index(h_4_1)

    to_delete = [c for c in children[i_start:i_end] if c.tag == QN_P]
    print(f"DELETING: {len(to_delete)} existing intro paragraph(s)")
    for c in to_delete:
        body.remove(c)

    insert_after = title
    for i, text in enumerate(PARAGRAPHS, 1):
        p = deepcopy(body_tpl)
        set_paragraph_text(p, text)
        insert_after.addnext(p)
        insert_after = p
        print(f"  ¶{i}: {text[:80]}…")

    print(f"\nWROTE: {len(PARAGRAPHS)} paragraph(s) into Chapter IV intro")


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    print("=== Chapter IV intro ===")
    write_intro_body(doc)

    doc.save(str(path))
    print(f"\nSAVED: {path}")


if __name__ == "__main__":
    main()
