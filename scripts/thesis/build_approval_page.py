"""Replace the Approval-page placeholder with a proper NEU approval page,
replicated from the reference template (SURE NEU THESIS FORMAT ...docx).

Known fields filled: author (OSAMA DUWAIRI), thesis title, degree, supervisor
(Prof. Dr. Fadi Al-Turjman), 2026 dates, NEU institute director. Committee
members + department head left as signature fill-in lines (known at defense).

Preserves the section break carried by the last approval paragraph; deletes only
the placeholder + empty filler paragraphs and inserts content before the break.
"""
from __future__ import annotations
import os, sys
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT

THESIS = "docs/research/THESIS.docx"
if os.path.exists("docs/research/.~lock.THESIS.docx#"):
    sys.exit("ABORT: editor lock present.")

doc = Document(THESIS)
paras = doc.paragraphs

def idx_of(pred):
    for i, p in enumerate(paras):
        if pred(p):
            return i
    return -1

ai = idx_of(lambda p: p.text.strip() == "Approval" and p.style.style_id == "945")
di = idx_of(lambda p: p.text.strip().startswith("Declaration of Ethical"))
assert ai >= 0 and di >= 0, f"headings not found ai={ai} di={di}"
# section-break paragraph between them
from docx.oxml.ns import qn
si = -1
for i in range(ai + 1, di):
    if paras[i]._element.find(qn("w:pPr") + "/" + qn("w:sectPr")) is not None or \
       paras[i]._element.find(".//" + qn("w:sectPr")) is not None:
        si = i
        break
assert si >= 0, "section-break paragraph not found in approval body"
print(f"Approval heading=[{ai}] Declaration=[{di}] sectPr para=[{si}]")

sect_para = paras[si]
# delete empties/placeholder between heading and section-break para
for i in range(ai + 1, si):
    el = paras[i]._element
    el.getparent().remove(el)

L, R = WD_TAB_ALIGNMENT.LEFT, WD_TAB_ALIGNMENT.RIGHT
TITLE = ("DOMAIN SPECIALIZED MARKETING AGENT FOR HIGH PERFORMANCE CAMPAIGN "
         "GENERATION: FINE-TUNING AND RAG-BASED APPROACH")
CERTIFY = (f"We certify that we have read the thesis submitted by OSAMA DUWAIRI "
           f"titled “{TITLE}” and that, in our combined opinion, it is "
           f"fully adequate, in scope and quality, as a thesis for the degree of "
           f"Master of Artificial Intelligence Engineering.")
DOTS_NAME = "...................................."
DOTS_SIG = ".............................."
CT = [(2.40, L), (4.40, L)]

LINES = [
    ([], None),
    ([], None),
    ([CERTIFY], None),
    ([], None),
    ([], None),
    (["Examining Committee", "\t", "Name-Surname", "\t", "Signature"], [(2.59, L), (5.30, L)]),
    ([], None),
    (["Head of the Committee:", "\t", DOTS_NAME, "\t", DOTS_SIG], CT),
    ([], None),
    (["Committee Member:", "\t", DOTS_NAME, "\t", DOTS_SIG], CT),
    ([], None),
    (["Supervisor:", "\t", "Prof. Dr. Fadi Al-Turjman", "\t", DOTS_SIG], CT),
    ([], None),
    (["Co-Supervisor:", "\t", DOTS_NAME, "\t", DOTS_SIG], CT),
    ([], None),
    ([], None),
    (["Approved by the Head of the Department"], None),
    ([], None),
    (["…../…../2026"], None),
    ([], None),
    ([], None),
    ([DOTS_SIG + "   Head of the Department"], None),
    (["Approved by the Institute of Graduate Studies"], None),
    ([], None),
    (["……/…../ 2026"], None),
    ([], None),
    (["Prof. Dr. Kemal Hüsnü Can Başer"], None),
]

def add_line(anchor, segments, tabstops):
    p = anchor.insert_paragraph_before(style="Body Text")
    if tabstops:
        for pos, align in tabstops:
            p.paragraph_format.tab_stops.add_tab_stop(Inches(pos), align)
    for seg in segments:
        if seg == "\t":
            p.add_run().add_tab()
        else:
            p.add_run(seg)
    return p

for segs, tabs in LINES:
    add_line(sect_para, segs, tabs)

doc.save(THESIS)
print(f"SAVED — inserted {len(LINES)} approval-page paragraphs before the section break.")
