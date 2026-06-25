"""Write Ch. V Discussion body — themed paragraphs on evaluation findings.

Each theme = one paragraph under its own H2 subsection (§5.1, §5.2, …).
Wipes everything between 'CHAPTER V' and 'References' and re-inserts from
the THEMES list below. If Ch. V doesn't exist yet, creates the CHAPTER V H1
(style 944) immediately before the 'References' heading first.

Idempotent: re-runs wipe-and-replace cleanly. Add new themes to THEMES and
re-run.

Conventions (see docs/research/THESIS_EDITING.md):
- Clone Heading 1 from CHAPTER IV (we added it ourselves; safe).
- Clone Heading 2 from §4.1 Evaluation Setup (we added it ourselves; safe).
- Clone Body Text from the first non-empty style-943 paragraph in the
  document (Ch. III body, written by us).
- Never clone original-template paragraphs that carry <w:drawing>.
- Skip any paragraph containing <w:sectPr> when wiping (pagination guard).
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"
QN_TBL = f"{{{W_NS}}}tbl"
QN_SECTPR = f"{{{W_NS}}}sectPr"
QN_DRAWING = f"{{{W_NS}}}drawing"


# Ch. V themes — each theme is one (H2 heading, body paragraph) pair.
# Add new themes to this list and re-run; the script wipes Ch. V body and
# re-inserts.
THEMES: list[tuple[str, str]] = [
    (
        "5.1 Evaluation Arm Disagreement",
        "Without revisiting the evaluation methodology in detail, we measured the configuration's performance using two different evaluation arms: the learned-scorer arm — the trained per-campaign predictor, which outputs a performance score from 0 to 1 for each ad copy — and the MAUVE arm (Pillutla et al., 2021), an embedding-overlap calculator that examines similarity between batches of generated campaigns rather than individual ones, also outputting scores from 0 to 1 related to similarity. Key tension arises when the learned-scorer results demonstrate the agent wrapper reducing composite scores consistently across the base model and the fine-tuned model, although marginally: the drop is −0.040 (95% CI [−0.061, −0.020], C → C_pipe) on the learned scorer, while MAUVE shows a lift of +0.122 (+21.9% relative) for the same comparison. The fact that this change has different signs across the two evaluation arms can be interpreted in different ways, but the answer lies in how each evaluation methodology works. The learned scorer is a pointwise trained model that answers the question: does this campaign look like the average gold-tier ad I was trained on? MAUVE is a distributional algorithm that answers the question: how much does this batch of generated ads semantically look like the reference gold batch in embedding space? This applies to the agent's score specifically, since an agent wrapper pulls in research and product specifics and market information and embeds them into the ad, giving them more variety, which allows the batch to have more overlap with good ads. The individual score reveals that the agent workflow makes each ad look less like the average ad the scorer was trained on. This remains fundamental to evaluating model performance and the effect of agent orchestration on output quality, since neither evaluation methodology is complete and neither is wrong. They simply reveal different information about the quality of the ads. The only real way to validate performance is real-world A/B deployment at scale, which is outside the scope of this thesis because of cost and time constraints.",
    ),
    (
        "5.2 Fine-Tuning–Agent Interaction",
        "Two base configurations can be wrapped by the agent workflow per the arm methodology: the base writer, which is off-the-shelf Qwen3-8B, and the fine-tuned writer, which is our fine-tuned Qwen3 with QLoRA (Dettmers et al., 2023). The intuitive expectation is that augmenting either model with agent capabilities would have independent improvements on the quality of the ad and should not disrupt fine-tuning performance. However, the data revealed a significant finding: both models, and especially the fine-tuned model, swing in the negative direction on the learned scorer — B → B_pipe loses 0.018 and C → C_pipe loses 0.040 (a 2.2× bigger drop on the FT writer) — while both swing in the positive direction on MAUVE by a comparable margin: base writer B → B_pipe gains 0.125 and fine-tuned writer C → C_pipe gains 0.122. The important finding is that although RAG mainly has complementary benefits to fine-tuning, there is some slight overlap where trade-offs must happen. Specifically, the retrieved context can shift the prompts away from the fine-tuning distribution, which lowers the effect of the fine-tuning on the output. This also explains why the shifts on the learned scorer are more dramatic in the case of the fine-tuned model than the base model, since the fine-tuned model is tied to a more specific mean vocabulary and format, which was used during training. This challenges the common assumption that fine-tuning and the agent are independent and contribute to output quality with no overlap. A possible mitigation is diversifying the fine-tuning training data to include RAG-like instructions, so the model learns to handle the context better during inference.",
    ),
    (
        "5.3 Two-Role Agent Architecture",
        "Most off-the-shelf models such as ChatGPT or Claude are general-purpose. They handle all tasks equally well, including research, planning, and writing in a single pass. They, however, lack the ability to have multi-step reasoning and agentic loops. They also lack the domain-specific training to generate successful advertising. In our agentic methodology, we combine the best of generalist models as well as our fine-tuned model. We use the general models as the orchestrator to perform general-purpose tasks such as research, tool calls, and answering questions. Our fine-tuned model sits with one responsibility, and that is to write ad creative, separating it from the main workflow. It is a strict split that enforces what each model should stick to in its domain: marketing domain skills live in the fine-tuned model weights, while general reasoning skills live in the orchestrator prompt. A specific architectural decision evident in this system is whether to use the fine-tuned model for orchestration tasks. The reasons we do not are control, and avoiding real performance degradation because of model size and fine-tuning constraints.",
    ),
    (
        "5.4 Domain Specialization Beats Frontier Scale",
        "The main premise of this thesis is to test whether a small (7–9B) fine-tuned open-source model can compete with a frontier proprietary LLM on a domain-specific task. The specific application of this thesis is to determine whether a marketing-specialized 7B model can produce copy at least as good as or better than proprietary large models. This comparison was made between the fine-tuned config (C: Draper-FT, Qwen3-8B + QLoRA, Dettmers et al., 2023) and the frontier model GPT-5.5 via API with no fine-tuning (config A). The findings in the thesis show that our fine-tuned model wins on every dimension our evaluation measures. Composite scores on our 215-brief test set show GOLD (real ad, ceiling) 0.684, C (Draper-FT) 0.651, B (Qwen3-8B base) 0.611, A (GPT-5.5) 0.603 — an absolute lift of +0.048 composite (8.0% relative) of C over A. More importantly, our model wins on all five platforms scored (Meta, TikTok, X, Google, Pinterest) as well as on all model heads (survivability, engagement_volume) except velocity, where it trails by a very small amount. Our model also reaches 95.2% of the GOLD ceiling. This result can be counterintuitive since these large provider models are hundreds of times larger than 8B and can indeed greatly surpass smaller models on general tasks, but when it comes to narrow domain-specific tasks, scale or parameter size isn't necessarily the best fit. Instead, specific domain knowledge can surpass their performance. This shows practical applications for real teams or researchers choosing between a fine-tuned model and prompting a frontier model depending on their use case.",
    ),
    (
        "5.5 Why Backtranslation Over Other Fine-Tuning Methods",
        "Having chosen the base Qwen3-8B model for fine-tuning, one important question was how it would be fine-tuned. The most critical point is that our collected ad corpus already contains classified successful ads, which are to be used as construction examples for fine-tuning. These key points eliminate methods such as human-annotated SFT, where real marketers write brief and ad pairs. It also eliminates forward-direction synthetic SFT, where an LLM writes brief and copy from its own training, omitting what we have already established as high-performance ads. Other methods such as continued pretraining on raw ad text, public instruction tuning, and preference tuning have also been considered. The important pivot to the back-translation method (Li et al., 2024, Humpback) comes from the fact that our data is grounded in high-performance ads and that briefs are absent from our dataset, making back-translating synthetically using a teacher LLM (claude-sonnet-4-6, gpt-5.4, gemini-3.1-pro-preview) to generate plausible briefs anchored to the established high performers the only viable option for proper fine-tuning training data. It works especially well since back-translation preserves the verbatim ad copy on the response side, and the LLM's synthetic behavior does not poison the training data. The takeaway is that when output data is available and interpretable but the input is private, back-translation is the number one option. This applies to marketing copy, legal motions, medical notes, and sales notes.",
    ),
]


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def find_paragraph_optional(doc, predicate):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    return None


def find_template_by_style(doc, style_id, label, prefix=None):
    """First paragraph with given style_id, no <w:drawing>, non-empty text.

    If prefix is provided, also require the paragraph text to start with it
    (used to disambiguate against TOC entries that share the style id).
    """
    for p in doc.paragraphs:
        if (p.style.style_id if p.style else "") != style_id:
            continue
        if p._element.find(f".//{QN_DRAWING}") is not None:
            continue
        ts = p._element.findall(f".//{QN_T}")
        text = "".join((t.text or "") for t in ts).strip()
        if not text:
            continue
        if prefix is not None and not text.startswith(prefix):
            continue
        return p._element
    raise RuntimeError(f"no usable style-{style_id} template found ({label})")


def set_paragraph_text(p_elem, new_text):
    """Set the first <w:t> in the paragraph to new_text; clear the rest.

    Preserves <w:rPr> and run structure so the paragraph style survives.
    """
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def ensure_chapter_v(doc):
    """Insert CHAPTER V H1 before References if missing.

    Returns (ch5_elem, refs_elem) — the CHAPTER V H1 element and the
    References heading element (the bounds of the body wipe range).
    """
    refs = find_paragraph(
        doc,
        lambda p: p.text.strip() == "References"
        and (p.style.style_id if p.style else "") == "945",
        "References heading",
    )

    ch5 = find_paragraph_optional(
        doc,
        lambda p: p.text.strip().upper() == "CHAPTER V"
        and (p.style.style_id if p.style else "") == "944",
    )
    if ch5 is not None:
        return ch5, refs

    ch4_h1 = find_paragraph(
        doc,
        lambda p: p.text.strip().upper() == "CHAPTER IV"
        and (p.style.style_id if p.style else "") == "944",
        "CHAPTER IV H1 template",
    )
    if ch4_h1.find(f".//{QN_DRAWING}") is not None:
        raise RuntimeError(
            "CHAPTER IV H1 template contains <w:drawing>; not safe to clone"
        )

    ch5 = deepcopy(ch4_h1)
    set_paragraph_text(ch5, "CHAPTER V")
    refs.addprevious(ch5)
    print("INSERTED: CHAPTER V (H1, style 944)")

    return ch5, refs


def write_section_body(doc):
    ch5, refs = ensure_chapter_v(doc)
    h2_tpl = find_template_by_style(doc, "945", "Heading 2", prefix="4.1 Evaluation")
    body_tpl = find_template_by_style(doc, "943", "Body Text")

    body = ch5.getparent()
    children = list(body)
    i_ch5 = children.index(ch5)
    i_refs = children.index(refs)

    to_delete = []
    skipped_sectpr = 0
    for c in children[i_ch5 + 1 : i_refs]:
        if c.tag == QN_P:
            if c.find(f".//{QN_SECTPR}") is not None:
                skipped_sectpr += 1
                continue
            to_delete.append(c)
        elif c.tag == QN_TBL:
            to_delete.append(c)

    print(f"DELETING: {len(to_delete)} elements between 'CHAPTER V' and 'References'")
    if skipped_sectpr:
        print(f"  (skipped {skipped_sectpr} paragraph(s) holding <w:sectPr>)")

    for c in to_delete:
        body.remove(c)

    insert_after = ch5
    for heading_text, body_text in THEMES:
        h = deepcopy(h2_tpl)
        set_paragraph_text(h, heading_text)
        insert_after.addnext(h)
        insert_after = h

        b = deepcopy(body_tpl)
        set_paragraph_text(b, body_text)
        insert_after.addnext(b)
        insert_after = b

        print(f"INSERTED: {heading_text!r}")
        print(f"          → {body_text[:90]}…")


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))
    write_section_body(doc)
    doc.save(str(path))
    print("\nDONE.")


if __name__ == "__main__":
    main()
