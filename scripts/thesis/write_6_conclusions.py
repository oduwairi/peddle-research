"""Write Ch. VI Conclusions and Recommendations body.

Each subsection = one paragraph under its own H2 (§6.1, §6.2, …). Wipes
everything between 'CHAPTER VI' and 'References' and re-inserts from the
SECTIONS list below. If Ch. VI doesn't exist yet, creates the CHAPTER VI H1
(style 944) immediately before the 'References' heading first.

Idempotent: re-runs wipe-and-replace cleanly. Add new sections to SECTIONS
and re-run.

Conventions (see docs/research/THESIS_EDITING.md):
- Clone Heading 1 from CHAPTER V (we added it ourselves; safe).
- Clone Heading 2 from §5.1 Evaluation Arm Disagreement (we added; safe).
- Clone Body Text from the first non-empty style-943 paragraph in the
  document (Ch. III body, written by us).
- Never clone original-template paragraphs that carry <w:drawing>.
- Skip any paragraph containing <w:sectPr> when wiping (pagination guard).
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"
QN_TBL = f"{{{W_NS}}}tbl"
QN_SECTPR = f"{{{W_NS}}}sectPr"
QN_DRAWING = f"{{{W_NS}}}drawing"


# Ch. VI subsections — each entry is one (H2 heading, body paragraph) pair.
# Add new entries here and re-run; the script wipes Ch. VI body and re-inserts.
SECTIONS: list[tuple[str, str]] = [
    (
        "6.1 Conclusion",
        "The aim of this thesis was to build a small fine-tuned model in the range of 7–9B based on an open-source model, wrapped in an agent workflow with live RAG, to see if it can compete with frontier models on a domain-specialized task (marketing and advertising). During the work, three main research questions guided the work. RQ1 asked whether a small 7–9B fine-tuned model can compete with large frontier proprietary models in generating successful ad campaigns. The fine-tuned writer (C) scored 0.651 composite, the frontier model (A, GPT-5.4-mini) scored 0.603, and the base model (B, Qwen3-8B) scored 0.611 — placing C at 95.2% of the GOLD ceiling (0.684); this key result highlights that, if done correctly, a small fine-tuned model can outperform a large model in a narrow domain task. RQ2 asked whether wrapping a fine-tuned model in an agent workflow with live RAG makes the output better. It was seen that the two evaluation arms — the learned scorer (pointwise) and MAUVE (distributional) — disagreed on this answer. What is going on? The learned scorer showed a reduction in overall performance for both the fine-tuned and the base model (B → B_pipe −0.025, C → C_pipe −0.040), while the MAUVE distribution showed a significant improvement in the score (C → C_pipe +0.133, a 46% relative lift). This means the answer can be interpreted in two ways. Augmenting the fine-tuned model with RAG can shift the output shape away from the learned format during fine-tuning based on the retrieved context, while making the output more varying and grounded in facts. RQ3 asked whether our custom proxy scoring system can generate training labels reliably when validated against ground-truth datasets. This is the most critical question, since publicly available and labeled advertising datasets are almost completely absent, which provides the necessity to develop a scoring system that uses proxy signals such as engagement or survivability to infer the scoring labels. The answer to this question was validated when our v3 scorer (Kaplan–Meier survival combined with engagement signals) was validated against the IRA Facebook ads corpus, where it yielded Spearman ρ = 0.749 (n=2,390) against real engagement (clicks).",
    ),
    (
        "6.2 Contributions",
        "Amongst the many contributions of this thesis, this section provides a summary on the contributions. Most notably, a deployed marketing LLM in the 7–9B parameter range (Qwen3-8B fine-tuned with QLoRA, served via Modal vLLM), with a live endpoint and API, a multi-platform ad corpus across a variety of verticals and commercial ads with proxy score labels, a custom hybrid two-arm evaluation methodology with a custom-trained learned scorer (DeBERTa-v3-base regressor) as well as a MAUVE algorithm comparator (Pillutla et al., 2021). This research also provides the learned scorer as a reusable model for any ad text and provides performance predictions. The frontend serves a full custom agent workflow with an orchestrator and writer architecture capable of performing research, writing angles, and emitting full campaigns to the user.",
    ),
    (
        "6.3 Limitations",
        "In this thesis, several limitations have been encountered that must be acknowledged. First, it is important to indicate that ad copy generation in this thesis is text-only and does not include the image or asset creative that often accounts for the other part of a campaign's success metrics; a plan for future integration of asset-brief generation into the fine-tuned model as a separate skill is recommended. The absence of explicitly labeled marketing ad datasets in the public domain provides the main limitation. Our approach was to design a proxy scorer for labeling the ads which, although accurate, remains an approximation of real market outcomes. A real-life deployment or an A/B test for different configurations remains the only true way to verify the performance of a campaign, which was omitted for this thesis due to time and cost constraints. During training, some limitations such as no hyperparameter tuning or base-model bake-off are also discussed. In our methodology, we picked a single quantized model, most appropriate for our fine-tuning task. We also picked the hyperparameters most known to perform in the existing literature instead of exploring optimal setups. Alternative, more thorough approaches would be hyperparameter search as well as base-model bake-offs between models such as Qwen, Gemma, or Mistral, however significantly increasing the cost and time of developing the project.",
    ),
    (
        "6.4 Recommendations for Future Work",
        "This work opens the door for many opportunities for future research and businesses to pursue further opportunities in the space of fine-tuning, marketing, and advertising. One of the most critical next steps is to expand the scope of the fine-tuned skills from ad copy text to a wider range of skills such as positioning, diagnostics, and asset generation. It will follow a similar workflow to the fine-tuning examples for the copywriting skill, expanding it to the wider range of skills for a full marketing agent. Future expansion has already been set up in the current work but has been deferred outside the scope of this thesis. Beyond this, a major encountered limitation in the current work was the lack of coordination between the fine-tuning examples and the augmenting RAG context during agentic inference. The most logical mitigation for this is to augment the training examples with RAG context as well as tool-call flows and agentic examples (retrieval-augmented fine-tuning; Zhang et al., 2024) so that the fine-tuned model learns to perform in similar conditions to the deployed inference case. This would greatly mitigate the encountered issue where RAG shifted the output distribution away from the fine-tuned model and resulted in a slightly weaker fine-tuning effect and a pivot back to the plain, generic AI style. However, such a strategy would require great amounts of complexity, construction, generation, and training costs, since responses will be much larger and more complex to design.",
    ),
]


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def find_paragraph_optional(doc, predicate):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    return None


def find_template_by_style(doc, style_id, label, prefix=None):
    """First paragraph with given style_id, no <w:drawing>, non-empty text.

    If prefix is provided, also require the paragraph text to start with it
    (used to disambiguate against TOC entries that share the style id).
    """
    for p in doc.paragraphs:
        if (p.style.style_id if p.style else "") != style_id:
            continue
        if p._element.find(f".//{QN_DRAWING}") is not None:
            continue
        ts = p._element.findall(f".//{QN_T}")
        text = "".join((t.text or "") for t in ts).strip()
        if not text:
            continue
        if prefix is not None and not text.startswith(prefix):
            continue
        return p._element
    raise RuntimeError(f"no usable style-{style_id} template found ({label})")


def set_paragraph_text(p_elem, new_text):
    """Set the first <w:t> in the paragraph to new_text; clear the rest.

    Preserves <w:rPr> and run structure so the paragraph style survives.
    """
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def ensure_chapter_vi(doc):
    """Insert CHAPTER VI H1 before References if missing.

    Returns (ch6_elem, refs_elem) — the CHAPTER VI H1 element and the
    References heading element (the bounds of the body wipe range).
    """
    refs = find_paragraph(
        doc,
        lambda p: p.text.strip() == "References"
        and (p.style.style_id if p.style else "") == "945",
        "References heading",
    )

    ch6 = find_paragraph_optional(
        doc,
        lambda p: p.text.strip().upper() == "CHAPTER VI"
        and (p.style.style_id if p.style else "") == "944",
    )
    if ch6 is not None:
        return ch6, refs

    ch5_h1 = find_paragraph(
        doc,
        lambda p: p.text.strip().upper() == "CHAPTER V"
        and (p.style.style_id if p.style else "") == "944",
        "CHAPTER V H1 template",
    )
    if ch5_h1.find(f".//{QN_DRAWING}") is not None:
        raise RuntimeError(
            "CHAPTER V H1 template contains <w:drawing>; not safe to clone"
        )

    ch6 = deepcopy(ch5_h1)
    set_paragraph_text(ch6, "CHAPTER VI")
    refs.addprevious(ch6)
    print("INSERTED: CHAPTER VI (H1, style 944)")

    return ch6, refs


def write_section_body(doc):
    ch6, refs = ensure_chapter_vi(doc)
    h2_tpl = find_template_by_style(doc, "945", "Heading 2", prefix="5.1 Evaluation")
    body_tpl = find_template_by_style(doc, "943", "Body Text")

    body = ch6.getparent()
    children = list(body)
    i_ch6 = children.index(ch6)
    i_refs = children.index(refs)

    to_delete = []
    skipped_sectpr = 0
    for c in children[i_ch6 + 1 : i_refs]:
        if c.tag == QN_P:
            if c.find(f".//{QN_SECTPR}") is not None:
                skipped_sectpr += 1
                continue
            to_delete.append(c)
        elif c.tag == QN_TBL:
            to_delete.append(c)

    print(f"DELETING: {len(to_delete)} elements between 'CHAPTER VI' and 'References'")
    if skipped_sectpr:
        print(f"  (skipped {skipped_sectpr} paragraph(s) holding <w:sectPr>)")

    for c in to_delete:
        body.remove(c)

    insert_after = ch6
    for heading_text, body_text in SECTIONS:
        h = deepcopy(h2_tpl)
        set_paragraph_text(h, heading_text)
        insert_after.addnext(h)
        insert_after = h

        b = deepcopy(body_tpl)
        set_paragraph_text(b, body_text)
        insert_after.addnext(b)
        insert_after = b

        print(f"INSERTED: {heading_text!r}")
        print(f"          → {body_text[:90]}…")


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))
    write_section_body(doc)
    doc.save(str(path))
    print("\nDONE.")


if __name__ == "__main__":
    main()
