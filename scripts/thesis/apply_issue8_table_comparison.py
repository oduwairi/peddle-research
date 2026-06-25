"""Reviewer round 2026-06-09 — issue #8: Table 2.6 deep-comparison paragraph.

Table 2.6 already lists 15 related works + citations; the reviewer wants an
interpretive comparison paragraph. This inserts the author's voiced paragraph
(Phase-5 polished — voice-to-text + typo fixes only, verbs/voice/order
preserved) as a NEW Body-Text paragraph immediately AFTER the Table 2.6 note,
at the end of Chapter II. No new heading -> no TOC change; Table 2.6 unchanged
-> no LoT change.

New-paragraph styling follows docs/research/THESIS_EDITING.md: the paragraph
copies the <w:pPr> (style + indent/spacing) and run <w:rPr> from a sibling
Body-Text paragraph (§2.9 closing synthesis) so it matches surrounding prose,
since a freshly added paragraph does NOT inherit direct formatting.

Run:  uv run python scripts/thesis/apply_issue8_table_comparison.py --dry-run
      uv run python scripts/thesis/apply_issue8_table_comparison.py

Idempotent; anchor-located; aborts without saving on any failure.
"""

from __future__ import annotations

import sys
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn

DOCX = "docs/research/THESIS.docx"

_SKIP_STYLES = ("toc", "table of figures", "table of contents", "table of tables")

# Anchors
NOTE_ANCHOR = "downloadable model weights"  # Table 2.6 Note paragraph (after table)
SIB_ANCHOR = "This survey has examined the research landscape"  # §2.9 closing prose
MARKER = "The reviewed works are split into two main categories"

PARA = (
    "The reviewed works are split into two main categories. Rows 1–4 feature "
    "open small models used as domain-specialized agents but not in the marketing "
    "domain — in finance, medicine, and others, such as Alpaca, Med42, "
    "AlpaCare, and FinGPT. The remaining rows in the table are marketing-domain "
    "specific; however, they are either closed, proprietary, or not fine-tunes. "
    "The thesis positions itself as the only sub-10B, fine-tuned, "
    "marketing-specialized model grounded in real ad performance. The thesis also "
    "features the first labeled ad data from proxy signals for commercial ads. "
    "The others use either different data or different languages and metrics."
)


def find_para(paras, anchor: str):
    for p in paras:
        style = (p.style.name if p.style else "").lower()
        if any(s in style for s in _SKIP_STYLES):
            continue
        if anchor in p.text:
            return p
    return None


def main() -> int:
    dry = "--dry-run" in sys.argv
    doc = Document(DOCX)
    paras = doc.paragraphs

    if any(MARKER in p.text for p in paras):
        print("skip: Table 2.6 comparison paragraph already present.")
        return 0

    note = find_para(paras, NOTE_ANCHOR)
    sib = find_para(paras, SIB_ANCHOR)
    if note is None:
        print(f"FAIL: Table 2.6 note not found -> {NOTE_ANCHOR!r}")
        return 1
    if sib is None:
        print(f"FAIL: sibling Body-Text para not found -> {SIB_ANCHOR!r}")
        return 1
    if not sib.runs:
        print("FAIL: sibling has no runs to copy run formatting from")
        return 1

    print(f"{'DRY-RUN ' if dry else ''}insert after note: {note.text[:60]!r}…")
    print(f"  copying formatting from sibling: {sib.text[:60]!r}…")
    print(f"  new paragraph ({len(PARA)} chars): {PARA[:70]!r}…")

    if dry:
        print("DRY-RUN: not saving.")
        return 0

    # Build a fresh empty paragraph at end of doc, then relocate it.
    new_para = doc.add_paragraph()
    new_p = new_para._p

    # Replace its pPr with a copy of the sibling's (carries pStyle + indent/spacing).
    cur_pPr = new_p.find(qn("w:pPr"))
    if cur_pPr is not None:
        new_p.remove(cur_pPr)
    sib_pPr = sib._p.find(qn("w:pPr"))
    if sib_pPr is not None:
        new_p.insert(0, deepcopy(sib_pPr))

    # Add a single run carrying the comparison text + the sibling run's rPr.
    run = new_para.add_run(PARA)
    sib_rPr = sib.runs[0]._r.find(qn("w:rPr"))
    if sib_rPr is not None:
        run._r.insert(0, deepcopy(sib_rPr))

    # Move the new <w:p> to immediately after the Table 2.6 note.
    note._p.addnext(new_p)

    doc.save(DOCX)
    print(f"saved {DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
