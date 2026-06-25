"""Consolidate reviewer-flagged repetition in THESIS.docx (reviewer point #4).

Delete-and-cross-reference only — every surviving sentence stays verbatim.
Three edits, each located by unique text (never by paragraph index, which
shifts). The script aborts rather than clobber if any find-string is absent
(meaning the author has since changed that paragraph) or if OnlyOffice holds
the file.

Edits:
  A  §5.4  — cut the duplicated composite/per-platform/per-head/95.2% recital
            (already reported in §4.2); keep premise + interpretation, add
            a (§4.2) cross-reference. The Ch IV<->V fix.
  B  §2.4  — cut the duplicated proxy-signal enumeration + the non-sequitur
            trailing sentence; forward-reference §3.3.1 (its methodological home).
  C  §3.8  — trim the re-posed RQ1 to "answers RQ1 (§1.3)".

GPT model names are left untouched (config A = GPT-5.5 is correct; the §6.1
mislabel is handled separately, pending author confirmation).
"""

from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document

DOCX = Path("docs/research/THESIS.docx")
LOCK = Path("docs/research/.~lock.THESIS.docx#")
BACKUP = Path("docs/research/THESIS.docx.pre-redundancy.bak")

# Each edit: an anchor substring (locates the unique paragraph) and a list of
# (find, replace) operations applied in order to that paragraph's live text.
EDITS = [
    {
        "name": "A  §5.4  cut Ch IV<->V results re-tabulation",
        "anchor": "The main premise of this thesis is to test whether a small",
        "ops": [
            # delete the three recited-results sentences (all already in §4.2)
            (
                "Composite scores on our 215-brief test set show GOLD (real ad, "
                "ceiling) 0.684, C (Draper-FT) 0.651, B (Qwen3-8B base) 0.611, "
                "A (GPT-5.5) 0.603 — an absolute lift of +0.048 composite "
                "(8.0% relative) of C over A. ",
                "",
            ),
            (
                "More importantly, our model wins on all five platforms scored "
                "(Meta, TikTok, X, Google, Pinterest) as well as on all model "
                "heads (survivability, engagement_volume) except velocity, where "
                "it trails by a very small amount. ",
                "",
            ),
            ("Our model also reaches 95.2% of the GOLD ceiling. ", ""),
            # add the cross-reference into the surviving premise sentence
            (
                "our fine-tuned model wins on every dimension our evaluation "
                "measures. This result",
                "our fine-tuned model wins on every dimension our evaluation "
                "measures (the full per-platform and per-head breakdown is "
                "reported in §4.2). This result",
            ),
        ],
    },
    {
        "name": "B  §2.4  cut duplicated proxy rationale, forward-ref §3.3.1",
        "anchor": "Our proposed system draws from these approaches and aims to deduce performance metrics",
        "ops": [
            (
                "These signals include: ad longevity (the number of days an ad "
                "remains active), engagement volume (likes, comments if "
                "available), engagement velocity (engagement over a duration) "
                "early termination (ads removed within a short period of time) "
                "and more depending on the availability of signals in data, this "
                "approach aggregates a collection of signals rather than relying "
                "on a single one, it also addresses the lack of available data "
                "on marketing campaigns with engagement metrics such as click "
                "rate, conversions, and so on. Combining these capabilities "
                "presents the foundation of domain specialized agents capable of "
                "reasoning and acting depending on the model training and "
                "provided tools.",
                "The full set of proxy signals and the rationale for using them "
                "in place of unavailable conversion labels are detailed in "
                "§3.3.1.",
            ),
        ],
    },
    {
        "name": "C  §3.8  trim re-posed RQ1 to a cross-reference",
        "anchor": "Model evaluation remains as the final and most critical step",
        "ops": [
            (
                "most critically answers the research question about whether an "
                "eight-billion-parameter domain-specialized model can outperform "
                "a frontier LLM on marketing domain-specific tasks.",
                "most critically answers RQ1 (§1.3).",
            ),
        ],
    },
]


def fail(msg: str) -> None:
    print(f"ABORT: {msg}")
    sys.exit(1)


def main() -> None:
    if LOCK.exists():
        fail("OnlyOffice lock present (.~lock.THESIS.docx#) — close the editor first.")
    if not DOCX.exists():
        fail(f"{DOCX} not found")

    shutil.copy2(DOCX, BACKUP)
    print(f"backup: {BACKUP}")

    doc = Document(str(DOCX))

    # Resolve each edit to exactly one paragraph and pre-validate every op.
    plan = []
    for edit in EDITS:
        matches = [p for p in doc.paragraphs if edit["anchor"] in p.text]
        if len(matches) != 1:
            fail(f"{edit['name']}: anchor matched {len(matches)} paragraphs (need 1)")
        para = matches[0]
        text = para.text
        for find, _ in edit["ops"]:
            if find not in text:
                fail(
                    f"{edit['name']}: find-string not present (paragraph changed "
                    f"since audit?) -> {find[:60]!r}..."
                )
            text = text.replace(find, "", 1) if _ == "" else text.replace(find, _, 1)
        plan.append((edit, para, text))

    # Apply: collapse into run[0] (all runs share one format; verified) and
    # clear the rest, preserving run[0]'s rPr.
    for edit, para, new_text in plan:
        runs = para.runs
        if not runs:
            fail(f"{edit['name']}: paragraph has no runs")
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
        print(f"applied: {edit['name']}")

    doc.save(str(DOCX))
    print(f"saved: {DOCX}")

    # Re-open and verify.
    print("\n=== VERIFICATION (reloaded from disk) ===")
    doc2 = Document(str(DOCX))
    checks = [
        ("§4.2", "cross-ref §4.2 present (A)"),
        ("detailed in §3.3.1", "cross-ref §3.3.1 present (B)"),
        ("answers RQ1 (§1.3)", "cross-ref §1.3 present (C)"),
    ]
    body = "\n".join(p.text for p in doc2.paragraphs)
    for needle, label in checks:
        print(f"  [{'OK' if needle in body else 'MISSING'}] {label}")
    gone = [
        "Composite scores on our 215-brief test set show GOLD",
        "These signals include: ad longevity",
        "answers the research question about whether an eight-billion-parameter",
    ]
    for needle in gone:
        print(f"  [{'OK-removed' if needle not in body else 'STILL-PRESENT'}] {needle[:50]}...")
    # GPT-5.5 in §5.4 must survive untouched
    print(f"  [{'OK' if 'frontier model GPT-5.5 via API' in body else 'CHANGED'}] §5.4 GPT-5.5 label preserved")


if __name__ == "__main__":
    main()
