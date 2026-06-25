"""Replace the Özet placeholder with a Turkish translation of the English
Abstract, mirroring the Abstract's structure: a single justified Body-Text
paragraph plus a (non-bold) "Anahtar Kelimeler:" line.

The translation faithfully renders the author's existing English abstract;
technical anchors (model names, MAUVE, QLoRA, Kaplan-Meier, DeBERTa, RAG, RQ1)
are kept verbatim. Author to verify fluency/terminology.
"""
from __future__ import annotations

import os
import sys
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

THESIS = "docs/research/THESIS.docx"
if os.path.exists("docs/research/.~lock.THESIS.docx#"):
    sys.exit("ABORT: editor lock present.")

OZET = (
    "Yapay zekâ teknolojisinin ve alana özelleşmiş ajanların hızla gelişmesiyle "
    "birlikte, bu sistemler pazarlama ve reklamcılık alanlarında geniş bir uygulama "
    "alanı bulmuştur. Ancak kullanıma hazır çoğu büyük dil modeli (LLM) ve model, "
    "gerçek dünya dağıtımlarında düşük performans gösterme eğiliminde olan genel "
    "reklam metinleri ve kampanyalar üretmektedir. Mevcut araştırmalar, ince ayar "
    "yapılmış 8B'lik bir modelin pazarlama alanı dikeyinde büyük tescilli modellerle "
    "eşleşip eşleşemeyeceğini ya da onları aşıp aşamayacağını; bunların bir ajan iş "
    "akışı ve canlı RAG ile sarmalanmasının üretilen reklamların performansını artırıp "
    "artıramayacağını ortaya koymamaktadır. Bu araştırma, söz konusu boşlukları, alana "
    "özelleşmiş 8B'lik bir modelin reklam metni üretiminde bir öncü (frontier) modeli "
    "geçip geçemeyeceğini sorarak gidermeyi ve ayrıca ince ayar ile canlı RAG'in "
    "birikimli etkisini ve her bir bileşenin ayrı ayrı etkisini anlamayı "
    "amaçlamaktadır. Bu araştırmada, beş platformdan (Facebook, TikTok, Pinterest, "
    "Twitter, Reddit) AdFlex aracılığıyla toplanan 55.000 reklamlık bir derlem "
    "kullanılmakta ve her reklam, vekil (proxy) puanlama sistemimizle "
    "puanlanmaktadır: kampanya ömrü üzerinde platform bazlı Kaplan-Meier sağkalım "
    "eğrileri ile sürekli etkileşim sinyallerinin (beğeni, paylaşım, yorum) "
    "birleşimi. Derlemdeki en yüksek performanslı reklamlar eğitim materyali olarak "
    "kullanılmaktadır. Eğitim oluşturma hattı, öğrenci modelin taklit etmeyi "
    "öğreneceği brief-yanıt çiftlerinin öğretmen modeller kullanılarak "
    "oluşturulmasını içerir ve talimat geri-çevirisi (instruction backtranslation) "
    "tekniğini kullanır. Öğrenci model, QLoRA ile ince ayar yapılan Qwen3-8B'dir. "
    "Ortaya çıkan ince ayarlı yazar, genel sorguların ve araç çağrılarının bir "
    "orkestratör tarafından yönetildiği, yaratıcı görevler söz konusu olduğunda ise "
    "ince ayarlı modelin yazar olarak çağrıldığı bir ajan iş akışıyla sarmalanmıştır. "
    "Değerlendirme metodolojisi, oluşturma kümesinden ayrılan ve platformlar arasında "
    "katmanlandırılmış 215 brief'lik bir test kümesi kullanır. Beş yapılandırma ile "
    "birlikte, kümedeki gerçek yüksek performanslı reklamlar olan GOLD reklamları "
    "karşılaştırılmıştır; yapılandırmalar bir öncü modelin yanı sıra ince ayar ve RAG "
    "kombinasyonlarını (dört kombinasyon) içerir. İki değerlendirme kolu "
    "tasarlanmıştır: etkileşim etiketleri üzerinde eğitilmiş bir DeBERTa-v3-base "
    "regresör puanlayıcı modeli ve gömme (embedding) örtüşmesine dayalı, derlem "
    "düzeyinde bir MAUVE dağılım karşılaştırması. Tüm değerler %95 bootstrap güven "
    "aralıklarıyla raporlanmaktadır. Sonuçlar, reklam başına bileşik puanları şöyle "
    "göstermektedir: GOLD 0,684, C 0,651, B 0,611, C_pipe 0,607, A 0,603, "
    "B_pipe 0,586. İnce ayarlı model diğer dört yapılandırmadan daha iyi performans "
    "göstermekte, ancak GOLD referansının gerisinde kalmaktadır; bu da RQ1'e olumlu "
    "bir yanıt verildiğini göstermektedir. Derlem düzeyindeki MAUVE sonuçları ise "
    "şöyledir: GOLD 0,462, C_pipe 0,420, B_pipe 0,302, C 0,287, B 0,183, A 0,180. "
    "Burada da ajanla sarmalanmış ince ayarlı modelimiz, diğer tüm yapılandırmaları "
    "belirgin bir farkla geçmektedir. İki değerlendirme kolu, ajanla sarmalanmış iş "
    "akışlarının etkisi konusunda çelişkili sonuçlar üretmektedir. Bu araştırmanın "
    "katkıları arasında etkileşim ve sağkalabilirlik sinyallerine dayalı bir vekil "
    "puanlama sistemi, başarılı reklam verileri üzerinde eğitilmiş, ince ayarlı açık "
    "kaynaklı bir model ve uçtan uca tam olarak dağıtılmış bir pazarlama ajanı yer "
    "almaktadır."
)

KEYWORDS = (
    "Anahtar Kelimeler: alana özelleşmiş LLM, reklam metni üretimi, talimat "
    "geri-çevirisi, QLoRA ince ayarı, geri-getirmeli artırılmış üretim (RAG), "
    "etkileşim temelli puanlama, pazarlama ajanı."
)

doc = Document(THESIS)
paras = doc.paragraphs

# locate Özet heading and its placeholder
oz_i = next(i for i, p in enumerate(paras) if p.text.strip() == "Özet" and p.style.style_id == "945")
ph_i = next(i for i in range(oz_i + 1, len(paras)) if paras[i].text.strip().startswith("[Özet"))
placeholder = paras[ph_i]
print(f"Özet heading=[{oz_i}] placeholder=[{ph_i}]")


def set_style_943(p):
    pPr = p._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = p._element.makeelement(qn("w:pPr"), {})
        p._element.insert(0, pPr)
    ps = pPr.find(qn("w:pStyle"))
    if ps is None:
        ps = pPr.makeelement(qn("w:pStyle"), {})
        pPr.insert(0, ps)
    ps.set(qn("w:val"), "943")


# 1) convert placeholder -> Özet body (Body Text 943, justified, single run)
for r in list(placeholder.runs):
    r._element.getparent().remove(r._element)
set_style_943(placeholder)
placeholder.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
placeholder.add_run(OZET)

# 2) keywords paragraph immediately after, same Body Text style (non-bold)
kw_el = deepcopy(placeholder._element)
for r in kw_el.findall(qn("w:r")):
    kw_el.remove(r)
placeholder._element.addnext(kw_el)
from docx.text.paragraph import Paragraph  # noqa: E402

kw = Paragraph(kw_el, placeholder._parent)
kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw.add_run(KEYWORDS)

doc.save(THESIS)
print("SAVED — Özet body + keywords written.")
