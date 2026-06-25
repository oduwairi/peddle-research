"""Embed the PRISMA literature-flow diagram into THESIS.docx as Figure 1.1.

Placement: Chapter I -> "Survey Methodology" subsection. The figure is
inserted immediately after the survey-methodology body paragraph (the one
starting "To conduct this survey, an AI-powered systematic search ...").
Because that subsection sits inside Chapter I (before the "CHAPTER II"
heading), the figure is numbered **1.1** per the document's chapter-based
numbering scheme (Fig 2.2 in ch.2, Fig 3.x in ch.3, ...).

The figure + caption are cloned from the existing Figure 3.4 image/caption
pair (paragraphs holding `image_*` drawing and the `Fig_3_4` bookmark) so
the inline-image formatting, the `Fig_1_1` bookmark, and the hidden
`TC "..." \\f F \\l 1` field that feeds the List of Figures all match the
document's existing convention exactly.

After running, refresh fields in OnlyOffice/Word (updateFields) so the
List of Figures picks up the new TC entry and its page number.

Idempotent: skips if a `Fig_1_1` bookmark is already present.
"""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches

THESIS = Path("docs/research/THESIS.docx")
LOCK = Path("docs/research/.~lock.THESIS.docx#")
BACKUP = Path("docs/research/THESIS.docx.pre-prisma.bak")
PRISMA_PNG = Path("docs/research/figures/fig-peerj-prisma-flow.png")

ANCHOR_PREFIX = "To conduct this survey"
NEW_BOOKMARK = "Fig_1_1"
NEW_BOOKMARK_ID = "10169"  # max existing id (10168) + 1
CAPTION = (
    "Figure 1.1: PRISMA-style flow of the literature identification, "
    "screening, and inclusion process."
)
TC_FIELD = f' TC "{CAPTION}" \\f F \\l 1 '
WIDTH_IN = 4.0  # 1024x1536 px -> 4.0in wide = 6.0in tall (fits one page)

QN_T = qn("w:t")
QN_R = qn("w:r")
QN_DRAWING = qn("w:drawing")
QN_INSTR = qn("w:instrText")
QN_BMS = qn("w:bookmarkStart")
QN_BME = qn("w:bookmarkEnd")


def ptext(elem) -> str:
    return "".join((t.text or "") for t in elem.findall(f".//{QN_T}"))


def main() -> int:
    if LOCK.exists():
        print("ABORT: OnlyOffice lock present — close the editor first.", file=sys.stderr)
        return 1
    if not THESIS.exists():
        print(f"ABORT: {THESIS} not found", file=sys.stderr)
        return 1
    if not PRISMA_PNG.exists():
        print(f"ABORT: {PRISMA_PNG} not found", file=sys.stderr)
        return 1

    doc = Document(str(THESIS))
    body = doc.element.body

    # Idempotency.
    existing = [b.get(qn("w:name")) for b in body.iter(QN_BMS)]
    if NEW_BOOKMARK in existing:
        print(f"SKIP: {NEW_BOOKMARK} bookmark already present; nothing to do.")
        return 0

    shutil.copy(THESIS, BACKUP)

    # --- Locate templates: Fig 3.4 image para + caption para -----------------
    img_tpl = None  # paragraph element holding a drawing (Fig 3.4 image)
    cap_tpl = None  # caption element carrying the Fig_3_4 bookmark + TC field
    for p in doc.paragraphs:
        el = p._element
        if cap_tpl is None and any(
            b.get(qn("w:name")) == "Fig_3_4" for b in el.findall(f".//{QN_BMS}")
        ):
            cap_tpl = el
            # The image holder is the immediately preceding paragraph with a drawing.
            prev = el.getprevious()
            while prev is not None and prev.tag == qn("w:p"):
                if prev.find(f".//{QN_DRAWING}") is not None:
                    img_tpl = prev
                    break
                prev = prev.getprevious()
            break
    if cap_tpl is None or img_tpl is None:
        print("ABORT: could not locate Fig 3.4 image/caption template.", file=sys.stderr)
        return 1

    # --- Locate the anchor (survey-methodology body paragraph) ---------------
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip().startswith(ANCHOR_PREFIX):
            anchor = p._element
            break
    if anchor is None:
        print(f"ABORT: anchor paragraph {ANCHOR_PREFIX!r} not found.", file=sys.stderr)
        return 1

    # --- Build the image paragraph (clone Fig 3.4 holder, strip its drawing) -
    img_el = deepcopy(img_tpl)
    for r in img_el.findall(QN_R):
        img_el.remove(r)  # drop the old drawing run(s); keep pPr
    anchor.addnext(img_el)

    new_img_para = next(p for p in doc.paragraphs if p._element is img_el)
    new_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    new_img_para.add_run().add_picture(str(PRISMA_PNG), width=Inches(WIDTH_IN))

    # --- Build the caption paragraph (clone Fig 3.4 caption, re-key) ---------
    cap_el = deepcopy(cap_tpl)
    for b in cap_el.findall(f".//{QN_BMS}"):
        if b.get(qn("w:name")) == "Fig_3_4":
            b.set(qn("w:name"), NEW_BOOKMARK)
            b.set(qn("w:id"), NEW_BOOKMARK_ID)
    for b in cap_el.findall(f".//{QN_BME}"):
        b.set(qn("w:id"), NEW_BOOKMARK_ID)
    for t in cap_el.findall(f".//{QN_T}"):
        if (t.text or "").startswith("Figure 3.4"):
            t.text = CAPTION
    for instr in cap_el.findall(f".//{QN_INSTR}"):
        if "TC " in (instr.text or ""):
            instr.text = TC_FIELD
    img_el.addnext(cap_el)

    doc.save(str(THESIS))

    print("INSERTIONS (after anchor 'To conduct this survey ...'):")
    print(f"  1. inline image   : {PRISMA_PNG} @ {WIDTH_IN}in wide (centered, style 943)")
    print(f"  2. caption + field: {CAPTION!r}")
    print(f"     bookmark={NEW_BOOKMARK} id={NEW_BOOKMARK_ID}  TC field -> List of Figures")
    print(f"  backup: {BACKUP}")
    print("NEXT: open in OnlyOffice and refresh fields (updateFields) to populate the LoF entry.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
