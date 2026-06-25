"""Reviewer feedback #7 — expand §2.8 Ethical Considerations in place.

Inserts three NEW body paragraphs (no new headings) into the existing
§2.8 structure, plus three interleaved References entries:

  - privacy passage      -> after §2.8 intro body  (data privacy in scraped corpus)
  - hallucination passage -> after §2.8.1 body      (hallucinated claims)
  - brand-voice passage   -> after §2.8.2 body      (brand voice manipulation)

References added (APA, plain Body Text, alphabetical interleave):
  - European Parliament & Council of the European Union. (2016).  (after Efron, before Fan)
  - Huang, L., et al. (2025).                                     (after Hu, L., before Jeong)
  - Karpinska-Krakowiak, M., & Eisend, M. (2025).                 (after Karami, before Karpukhin)

Idempotent: each insertion is skipped if its marker text is already present.
Anchored by visible text (never hardcoded indices). Clones <w:pPr>/<w:rPr>
from a safe sibling Body Text paragraph (no <w:drawing> involved).
"""

from __future__ import annotations

import copy
import sys

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCX = "docs/research/THESIS.docx"

# --- new body paragraphs (Phase 6 enriched) -------------------------------

PRIVACY = (
    "The only data source for this research is from public ad libraries such as "
    "Meta, Google, and TikTok. This means that advertisers willing to post on these "
    "platforms have consented to publish information such as the advertiser name, the "
    "ad's creative content, the engagement, the dates, and other metadata used by the "
    "data corpus, and since this information is publicly available, the corpus does not "
    "use any sensitive or private user information without consent. This includes "
    "information such as names, user IDs, contact details, or other sensitive information. "
    "The saved data is also never used to link back to user identities. This keeps the "
    "corpus building and usage ethical by design, consistent with the data minimisation "
    "principle of the General Data Protection Regulation (European Parliament & Council of "
    "the European Union, 2016), since it only draws on publicly shared information with "
    "prior consent by the advertisers. Despite this, there is still some minor risk that a "
    "published ad may still contain sensitive information of a real person, by some image "
    "or sensitive text the advertiser chose to use."
)

HALLUCINATION = (
    "One of the most significant risks in LLM generative systems, especially in advertising "
    "domains, is that the model could hallucinate information or product facts that are "
    "untrue (Huang et al., 2025). This includes false testimonies, fake features, and "
    "fabricated numbers. These can have serious legal repercussions if used in real "
    "advertising settings. Minimizing this risk includes different strategies, most "
    "importantly making sure the training data is grounded in the real ad it is derived "
    "from and avoids any fabrication or paraphrasing by the LLM. This keeps the construction "
    "examples which the model learns from grounded and away from hallucination risks. "
    "Additionally, augmenting the model with an agent workflow that can fact-check or revise "
    "the campaign before submission can also further reduce the chances of hallucinations. "
    "Despite these mitigations, some amount of hallucination may still happen in the real "
    "world. Prior warning and consent to the users must be given, in that AI-generated "
    "content can be misleading or false."
)

BRAND = (
    "An additional risk is that a marketing model can imitate a real brand voice without "
    "explicit instruction, and may emit information that the brand may not have said or "
    "warranted (Karpinska-Krakowiak & Eisend, 2025). This can have serious legal risks if "
    "not safeguarded properly. For our model, the specific use case is to be grounded in the "
    "user's product or brand and never infer product or brand information outside of it. This "
    "is enforced in the training example generation as well as in the inference prompts, for "
    "the model to avoid bringing outside facts into the information presented to the user. "
    "Additionally, the system output should never be attributed to any third-party brand. "
    "Any brand voice or visual identity used by the model is only used when the user "
    "explicitly provides this information to the model, and is restricted to the brand the "
    "user presented, never from the outside."
)

# (marker prefix to locate, anchor-paragraph prefix to insert AFTER, new text)
BODY_INSERTS = [
    (
        "The only data source for this research is from public ad libraries",
        "As for all contemporary AI systems, ethical and legal concerns remain",
        PRIVACY,
    ),
    (
        "One of the most significant risks in LLM generative systems",
        "A central concern in AI-generated advertising is misleading",
        HALLUCINATION,
    ),
    (
        "An additional risk is that a marketing model can imitate a real brand voice",
        "As for all generative AI models, the chance of models learning biases",
        BRAND,
    ),
]

# --- new references (APA, plain) ------------------------------------------

REF_EU = (
    "European Parliament & Council of the European Union. (2016). Regulation (EU) "
    "2016/679 of the European Parliament and of the Council of 27 April 2016 on the "
    "protection of natural persons with regard to the processing of personal data and on "
    "the free movement of such data, and repealing Directive 95/46/EC (General Data "
    "Protection Regulation). Official Journal of the European Union, L 119, 1–88. "
    "https://eur-lex.europa.eu/eli/reg/2016/679/oj/eng"
)
REF_HUANG = (
    "Huang, L., Yu, W., Ma, W., Zhong, W., Feng, Z., Wang, H., Chen, Q., Peng, W., Feng, X., "
    "Qin, B., & Liu, T. (2025). A survey on hallucination in large language models: "
    "Principles, taxonomy, challenges, and open questions. ACM Transactions on Information "
    "Systems, 43(2). https://doi.org/10.1145/3703155"
)
REF_KARP = (
    "Karpinska-Krakowiak, M., & Eisend, M. (2025). Realistic portrayals of untrue "
    "information: The effects of deepfaked ads and different types of disclosures. Journal "
    "of Advertising, 54(3), 432–442. https://doi.org/10.1080/00913367.2024.2306415"
)

# (marker prefix, successor-entry prefix to insert BEFORE, ref text)
REF_INSERTS = [
    ("European Parliament & Council of the European Union. (2016).", "Fan, W., Ding, Y., Ning, L.", REF_EU),
    ("Huang, L., Yu, W., Ma, W.", "Jeong, S., Baek, J., Cho, S.", REF_HUANG),
    ("Karpinska-Krakowiak, M., & Eisend, M. (2025).", "Karpukhin, V., Oguz, B., Min, S.", REF_KARP),
]


def make_para_like(model_p, text: str):
    """Build a fresh <w:p> cloning the model paragraph's pPr and first-run rPr."""
    if model_p.find(qn("w:drawing")) is not None:
        raise RuntimeError("refusing to clone a paragraph containing a drawing")
    new_p = OxmlElement("w:p")
    ppr = model_p.find(qn("w:pPr"))
    if ppr is not None:
        new_p.append(copy.deepcopy(ppr))
    r = OxmlElement("w:r")
    first_r = model_p.find(qn("w:r"))
    if first_r is not None:
        rpr = first_r.find(qn("w:rPr"))
        if rpr is not None:
            r.append(copy.deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    new_p.append(r)
    return new_p


def find_para(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def present(doc, marker):
    return any(marker in p.text for p in doc.paragraphs)


def main() -> int:
    doc = docx.Document(DOCX)
    summary = []

    # body paragraphs: insert AFTER anchor
    for marker, anchor_prefix, text in BODY_INSERTS:
        if present(doc, marker):
            summary.append(f"SKIP (present): {marker[:45]}...")
            continue
        anchor = find_para(doc, anchor_prefix)
        if anchor is None:
            print(f"ERROR: anchor not found: {anchor_prefix!r}")
            return 1
        anchor._p.addnext(make_para_like(anchor._p, text))
        summary.append(f"INSERT body after [{anchor_prefix[:40]}...] -> {marker[:45]}...")

    # references: insert BEFORE successor
    for marker, succ_prefix, text in REF_INSERTS:
        if present(doc, marker):
            summary.append(f"SKIP (present): {marker[:45]}...")
            continue
        succ = find_para(doc, succ_prefix)
        if succ is None:
            print(f"ERROR: ref successor not found: {succ_prefix!r}")
            return 1
        succ._p.addprevious(make_para_like(succ._p, text))
        summary.append(f"INSERT ref before [{succ_prefix[:30]}...] -> {marker[:45]}...")

    doc.save(DOCX)
    print("INSERTIONS:")
    for line in summary:
        print(f"  {line}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
