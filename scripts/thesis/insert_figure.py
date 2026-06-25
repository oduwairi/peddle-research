"""Insert a captioned figure into THESIS.docx after a given anchor paragraph.

Idempotent: if a paragraph immediately after the anchor already contains an
inline image (a `<w:drawing>` element), the script does nothing. This lets the
caller rerun without stacking duplicate figures.

Usage:
    uv run python scripts/thesis/insert_figure.py \
        --anchor "Following the ad collection and sweeping" \
        --image  docs/research/figures/fig-3-2-adflex-collection.png \
        --caption "Figure 3.1: AdFlex-centred ad collection system." \
        --width-inches 5.5

The anchor is matched by `startswith` against the paragraph's visible text
(after `.strip()`). Pick a unique prefix from the target paragraph.
"""
from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"
QN_DRAWING = f"{{{W_NS}}}drawing"
QN_INSTR = f"{{{W_NS}}}instrText"
QN_BMS = f"{{{W_NS}}}bookmarkStart"
QN_BME = f"{{{W_NS}}}bookmarkEnd"
QN_PPR = f"{{{W_NS}}}pPr"
QN_NAME = f"{{{W_NS}}}name"
QN_ID = f"{{{W_NS}}}id"


def find_anchor(doc, prefix: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise RuntimeError(f"Anchor paragraph not found (prefix: {prefix!r})")


def has_inline_image(p_elem) -> bool:
    return p_elem.find(f".//{QN_DRAWING}") is not None


def caption_template(doc):
    """Find a style-943 (Body Text) paragraph with text to clone for the caption."""
    for p in doc.paragraphs:
        sid = p.style.style_id if p.style else ""
        if sid == "943" and p.text.strip() and p._element.findall(f".//{QN_T}"):
            return p._element
    raise RuntimeError("No style-943 Body Text template paragraph found")


def set_paragraph_text(p_elem, new_text: str):
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def tc_caption_template(doc):
    """Find an existing caption paragraph carrying a `TC "..." \\f F` field.

    These captions feed the live List of Figures; cloning one keeps the new
    figure consistent (justified body style + a TC field + a Fig bookmark)
    instead of the plain centred caption the default path produces.
    """
    for p in doc.paragraphs:
        instrs = p._element.findall(f".//{QN_INSTR}")
        if any((it.text or "").lstrip().startswith('TC "') for it in instrs) and any(
            "\\f F" in (it.text or "") for it in instrs
        ):
            return p._element
    raise RuntimeError("No TC-field caption template paragraph found")


def blank_para_from(tpl_elem):
    """Clone a paragraph keeping only its <w:pPr> (to hold an inline image)."""
    new = deepcopy(tpl_elem)
    for child in list(new):
        if child.tag != QN_PPR:
            new.remove(child)
    return new


def build_tc_caption(tpl_elem, caption: str, bookmark: str, bm_id: int):
    """Clone a TC caption template, swapping bookmark, visible text and TC text."""
    new = deepcopy(tpl_elem)
    for bs in new.findall(f".//{QN_BMS}"):
        bs.set(QN_NAME, bookmark)
        bs.set(QN_ID, str(bm_id))
    for be in new.findall(f".//{QN_BME}"):
        be.set(QN_ID, str(bm_id))
    ts = new.findall(f".//{QN_T}")  # visible runs only (field uses instrText)
    if ts:
        ts[0].text = caption
        for t in ts[1:]:
            t.text = ""
    instrs = new.findall(f".//{QN_INSTR}")
    if instrs:
        instrs[0].text = f' TC "{caption}" \\f F \\l 1 '
        for it in instrs[1:]:
            it.text = ""
    return new


def _para_for(doc, elem):
    for p in doc.paragraphs:
        if p._element is elem:
            return p
    raise RuntimeError("paragraph not re-locatable after insertion")


def insert_figure(
    doc_path: Path,
    anchor_prefix: str,
    image_path: Path,
    caption: str,
    width_inches: float,
    bookmark: str | None = None,
    bm_id: int = 30000,
) -> None:
    if not image_path.exists():
        raise FileNotFoundError(image_path)

    doc = Document(str(doc_path))

    anchor = find_anchor(doc, anchor_prefix)
    anchor_elem = anchor._element

    # Idempotency: scan the next two paragraphs for an inline image whose
    # following sibling paragraph starts with the caption.
    body = anchor_elem.getparent()
    children = list(body)
    idx = children.index(anchor_elem)
    for offset in (1, 2):
        if idx + offset >= len(children):
            break
        nxt = children[idx + offset]
        if nxt.tag != QN_P:
            continue
        if has_inline_image(nxt):
            # Check the following paragraph for the caption — if it matches,
            # this figure is already present.
            cap_p = children[idx + offset + 1] if idx + offset + 1 < len(children) else None
            cap_text = ""
            if cap_p is not None and cap_p.tag == QN_P:
                cap_text = "".join(t.text or "" for t in cap_p.findall(f".//{QN_T}"))
            if cap_text.strip().startswith(caption.split(":")[0]):
                print(f"SKIP: figure already present after anchor {anchor_prefix!r}")
                return

    if bookmark:
        # TC mode: clone a real TC caption so the figure matches the doc's
        # existing figures (justified body style + TC field for the LoF).
        tpl = tc_caption_template(doc)
        img_p = blank_para_from(tpl)
        anchor_elem.addnext(img_p)
        new_para = _para_for(doc, img_p)
        new_para.add_run().add_picture(str(image_path), width=Inches(width_inches))
        cap_p = build_tc_caption(tpl, caption, bookmark, bm_id)
        img_p.addnext(cap_p)
        doc.save(str(doc_path))
        print(f"INSERTED (TC): figure + caption after anchor {anchor_prefix!r}")
        print(f"  image: {image_path}")
        print(f"  caption: {caption}")
        print(f"  bookmark: {bookmark} (id={bm_id})")
        return

    # Default mode: centred image + plain (no-TC) caption.
    tpl = caption_template(doc)

    img_p = deepcopy(tpl)
    set_paragraph_text(img_p, "")  # blank, will receive the inline drawing
    anchor_elem.addnext(img_p)

    new_para = _para_for(doc, img_p)
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    # Insert the caption paragraph after the image.
    cap_p = deepcopy(tpl)
    set_paragraph_text(cap_p, caption)
    img_p.addnext(cap_p)
    for p in doc.paragraphs:
        if p._element is cap_p:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break

    doc.save(str(doc_path))
    print(f"INSERTED: figure + caption after anchor {anchor_prefix!r}")
    print(f"  image: {image_path}")
    print(f"  caption: {caption}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--anchor", required=True, help="Anchor paragraph prefix")
    parser.add_argument("--image", required=True, type=Path)
    parser.add_argument("--caption", required=True)
    parser.add_argument("--width-inches", type=float, default=5.5)
    parser.add_argument("--doc", type=Path, default=Path("docs/research/THESIS.docx"))
    parser.add_argument(
        "--bookmark",
        default=None,
        help="Bookmark name (e.g. Fig_4_4_1). When set, emits a TC-field caption "
        "(justified, listed in the LoF) instead of a plain centred caption.",
    )
    parser.add_argument("--bm-id", type=int, default=30000, help="Unique bookmark id")
    args = parser.parse_args()

    insert_figure(
        doc_path=args.doc,
        anchor_prefix=args.anchor,
        image_path=args.image,
        caption=args.caption,
        width_inches=args.width_inches,
        bookmark=args.bookmark,
        bm_id=args.bm_id,
    )


if __name__ == "__main__":
    main()
