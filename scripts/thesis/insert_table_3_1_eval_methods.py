"""Insert Table 3.1 (evaluation methods comparison) into §3.8.

Places a caption paragraph + a 7x3 comparison table at the tail of §3.8,
immediately before the CHAPTER IV heading. Idempotent: skip if the
"Table 3.1" caption is already present inside §3.8.

write_3_8_evaluation.py only wipes <w:p> elements between the §3.8 heading
and CHAPTER IV, so the <w:tbl> + its caption (re-inserted here) survive
re-runs of that script.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"
QN_TBL = f"{{{W_NS}}}tbl"


CAPTION = "Table 3.1: Comparison of the two evaluation methods."

HEADER: tuple[str, str, str] = (
    "Aspect",
    "Per-ad absolute scoring",
    "Corpus distribution matching (MAUVE)",
)

ROWS: list[tuple[str, str, str]] = [
    ("Granularity", "Per ad", "Per model (whole-corpus)"),
    (
        "Output",
        "Four scores per ad: composite, survivability, engagement volume, engagement velocity",
        "Single similarity score in [0, 1] per model",
    ),
    (
        "Reference signal",
        "Real engagement labels from §3.3 (predictor trained on these)",
        "Real high-tier ads from the v3 corpus, with held-out test ads excluded",
    ),
    (
        "Embedding / model",
        "DeBERTa-v3-base regressor (§3.7)",
        "GPT-2 Large encoder",
    ),
    (
        "What it answers",
        "“How well would each ad have performed?”",
        "“Does the overall output resemble real winning ads?”",
    ),
    (
        "Reference",
        "This work, §3.7",
        "Pillutla et al. (2021)",
    ),
]


def _para_text(p_elem) -> str:
    return "".join(t.text or "" for t in p_elem.findall(f".//{QN_T}")).strip()


def main() -> None:
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    # Locate §3.8 heading and CHAPTER IV heading
    h_3_8 = None
    h_ch4 = None
    for p in doc.paragraphs:
        txt = p.text.strip()
        sid = p.style.style_id if p.style else ""
        if h_3_8 is None and sid == "945" and txt.startswith("3.8 Evaluation Methodology"):
            h_3_8 = p._element
        elif h_ch4 is None and sid == "944" and txt.upper() == "CHAPTER IV":
            h_ch4 = p._element
            break
    if h_3_8 is None or h_ch4 is None:
        raise RuntimeError("§3.8 or CHAPTER IV heading not found")

    body = h_ch4.getparent()
    children = list(body)
    i_h_3_8 = children.index(h_3_8)
    i_h_ch4 = children.index(h_ch4)

    # Idempotency: skip if Table 3.1 caption is already inside §3.8
    for c in children[i_h_3_8:i_h_ch4]:
        if c.tag == QN_P and _para_text(c).startswith("Table 3.1"):
            print("SKIP: Table 3.1 caption already present in §3.8")
            return

    # Body-text template (style 943) for the caption paragraph
    body_tpl = None
    for p in doc.paragraphs:
        sid = p.style.style_id if p.style else ""
        if sid == "943" and p.text.strip() and p._element.findall(f".//{QN_T}"):
            body_tpl = p._element
            break
    if body_tpl is None:
        raise RuntimeError("No style-943 Body Text template paragraph found")

    # Build table at end of doc, then relocate
    table = doc.add_table(rows=len(ROWS) + 1, cols=3)
    try:
        table.style = "Table Grid"
    except KeyError:
        # Style not defined in this docx — keep default; OnlyOffice will render fine
        pass

    # Header row
    for ci, val in enumerate(HEADER):
        cell = table.rows[0].cells[ci]
        cell.text = val
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Body rows
    for ri, row in enumerate(ROWS, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val

    # Detach the newly-created table from end-of-document
    tbl_elem = table._element
    tbl_elem.getparent().remove(tbl_elem)

    # Build the caption paragraph from the body-text template
    caption_p = deepcopy(body_tpl)
    ts = caption_p.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("Caption template paragraph has no <w:t>")
    ts[0].text = CAPTION
    for t in ts[1:]:
        t.text = ""

    # Insert caption then table immediately before CHAPTER IV
    h_ch4.addprevious(caption_p)
    h_ch4.addprevious(tbl_elem)

    doc.save(str(path))
    print(f"INSERTED: '{CAPTION}' + {len(ROWS) + 1}x3 table before CHAPTER IV")


if __name__ == "__main__":
    main()
