"""Reviewer round 2026-06-09 — issue #2: fill the Acknowledgments page.

Swaps the placeholder body ("The author has not declared any acknowledgements.")
for the author's drafted text (Phase-4) after a light Phase-5 polish. Heading
untouched; Body-Text style/formatting preserved (run-text replace only) → no
TOC change. Idempotent; anchor-located; aborts without saving on failure.

Run:  uv run python scripts/thesis/insert_acknowledgments.py --dry-run
      uv run python scripts/thesis/insert_acknowledgments.py
"""

from __future__ import annotations

import sys

from docx import Document

DOCX = "docs/research/THESIS.docx"

OLD = "The author has not declared any acknowledgements."
NEW = (
    "I would like to thank my supervisor Prof. Dr. Fadi Al-Turjman for his "
    "support on this project, as well as the committee of the Department of "
    "Artificial Intelligence Engineering at Near East University. Additionally, "
    "I would like to thank my family and friends for their continued support."
)
MARKER = "I would like to thank my supervisor Prof. Dr. Fadi Al-Turjman"


def replace_in_paragraph(p, old: str, new: str) -> bool:
    runs = p.runs
    texts = [r.text for r in runs]
    full = "".join(texts)
    idx = full.find(old)
    if idx == -1:
        return False
    start, end = idx, idx + len(old)
    pos = 0
    spans = []
    for t in texts:
        spans.append((pos, pos + len(t)))
        pos += len(t)
    out = list(texts)
    first = True
    for k, (lo, hi) in enumerate(spans):
        if hi <= start or lo >= end:
            continue
        a = max(start, lo) - lo
        b = min(end, hi) - lo
        if first:
            out[k] = texts[k][:a] + new + texts[k][b:]
            first = False
        else:
            out[k] = texts[k][:a] + texts[k][b:]
    for r, t in zip(runs, out):
        if r.text != t:
            r.text = t
    return True


def main() -> int:
    dry = "--dry-run" in sys.argv
    doc = Document(DOCX)
    for p in doc.paragraphs:
        if MARKER in p.text:
            print("skip: acknowledgments already filled.")
            return 0
    target = next((p for p in doc.paragraphs if OLD in p.text), None)
    if target is None:
        print(f"FAIL: placeholder not found -> {OLD!r}")
        return 1
    if not dry:
        if not replace_in_paragraph(target, OLD, NEW):
            print("FAIL: replacement did not fire")
            return 1
    print(f"OK: {'(dry-run) ' if dry else ''}acknowledgments swapped.")
    if dry:
        print("DRY-RUN: not saving.")
        return 0
    doc.save(DOCX)
    print(f"saved {DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
