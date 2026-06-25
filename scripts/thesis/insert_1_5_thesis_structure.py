"""Phase 7 insertion for the new §1.5 "Thesis Structure" section.

Replaces the bracketed placeholder body left by
`restructure_1_6_thesis_structure.py` with the author's Phase-4 draft after
Phase-5 polish + Phase-6 light enrichment (terminology precision only; this is
a roadmap section, so enrichment is deliberately minimal — no new claims, no
citations). Author voice, verbs, and clause order preserved.

Idempotent: locates the body paragraph by the placeholder token first; on a
re-run (placeholder already gone) it falls back to the first non-empty Body
Text paragraph between the "Thesis Structure" Heading-2 and the next heading,
and re-writes it to FINAL_TEXT. No paragraph indices are hardcoded.

TOC is a live Word field -> not touched here; refresh in OnlyOffice.
"""

from __future__ import annotations

import sys

from docx import Document
from docx.oxml.ns import qn

DOCX = "docs/research/THESIS.docx"
PLACEHOLDER_TOKEN = "THESIS STRUCTURE"  # matches the bracketed placeholder marker

FINAL_TEXT = (
    "This thesis is organized into six chapters in total. Each chapter "
    "represents a logical stage in the evolution of the thesis. Chapter one "
    "already covered the problem that this is trying to solve. Additionally, "
    "it has discussed the research questions aimed to be addressed by this "
    "research as well as the contributions of this thesis in the research "
    "field. Chapter two will review existing work and research in the "
    "literature related to marketing-specialized agents, including "
    "open-source foundation models, retrieval-augmented generation (RAG) "
    "techniques, AI in marketing and advertising, and other related "
    "literature. The chapter ends by stating open research gaps and future "
    "directions that this thesis hopes to address. The thesis continues with "
    "Chapter three where we propose our system methodology for our "
    "marketing-specialized agent, fine-tuned on successful marketing data. "
    "The section includes detailed information about the design and "
    "implementation of all critical steps in the pipeline. Following that, "
    "Chapter four presents the results of the evaluation of the obtained "
    "model from Chapter three. Specifically, it clarifies the multiple-arm "
    "evaluation setup, including the learned-scorer and MAUVE "
    "distribution-matching arms. It also presents the fine-tuning and agent "
    "ablation results. It concludes with the synthesis and limitations of "
    "the obtained results. Chapter five focuses on the discussion of the "
    "obtained results in the previous chapter, such as evaluation-arm "
    "agreement or disagreement. The intersection between fine-tuning and "
    "agent RAG architecture is also discussed, alongside the implications of "
    "domain-specialized small models over frontier models. The thesis "
    "concludes with Chapter six where we present our conclusion of answers to "
    "the proposed research questions and the contributions obtained at the "
    "end. Furthermore, we present the limitations of the presented work as "
    "well as recommendations for future work and research."
)


def norm(s: str) -> str:
    return " ".join(s.split()).strip()


def set_flat_text(p_el, text: str) -> None:
    ts = p_el.findall(".//" + qn("w:t"))
    if not ts:
        raise SystemExit("no <w:t> in target paragraph")
    ts[0].text = text
    ts[0].set(qn("xml:space"), "preserve")
    for t in ts[1:]:
        t.text = ""


def main() -> None:
    d = Document(DOCX)
    ps = list(d.paragraphs)

    # locate the "Thesis Structure" Heading-2
    h_idx = None
    for i, p in enumerate(ps):
        if p.style.name == "Heading 2" and norm(p.text) == "Thesis Structure":
            h_idx = i
            break
    if h_idx is None:
        raise SystemExit("'Thesis Structure' heading not found")

    # next heading bounds the section
    end = len(ps)
    for j in range(h_idx + 1, len(ps)):
        if ps[j].style.name in ("Heading 1", "Heading 2"):
            end = j
            break

    # prefer the placeholder paragraph; else first non-empty Body Text in range
    target = None
    for j in range(h_idx + 1, end):
        if PLACEHOLDER_TOKEN in ps[j].text:
            target = ps[j]
            break
    if target is None:
        for j in range(h_idx + 1, end):
            if norm(ps[j].text):
                target = ps[j]
                break
    if target is None:
        raise SystemExit("no body paragraph found under 'Thesis Structure'")

    before = norm(target.text)[:60]
    set_flat_text(target._p, FINAL_TEXT)
    d.save(DOCX)

    print("INSERTION SUMMARY")
    print(f"  - section: §1.5 Thesis Structure (heading style: {ps[h_idx].style.name})")
    print(f"  - body paragraph style: {target.style.name}")
    print(f"  - replaced: {before!r} ...")
    print(f"  - wrote {len(FINAL_TEXT)} chars, {FINAL_TEXT.count('. ') + 1} sentences")

    # verify
    d2 = Document(DOCX)
    leftover = [norm(p.text)[:40] for p in d2.paragraphs if PLACEHOLDER_TOKEN in p.text
                and norm(p.text) != "Thesis Structure"]
    print(f"\nVERIFY  placeholder remaining: {leftover if leftover else 'none'}")
    # confirm final text landed
    hit = any(norm(p.text) == norm(FINAL_TEXT) for p in d2.paragraphs)
    print(f"VERIFY  final prose present as one paragraph: {hit}")


if __name__ == "__main__":
    sys.exit(main())
