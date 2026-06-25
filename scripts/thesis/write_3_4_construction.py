"""Write §3.4 Training-Data Construction via Backtranslation body.

Same pattern as write_3_3_scoring.py: wipes every paragraph between the §3.4
heading and the §3.5 heading, then re-inserts the current PARAGRAPHS list
using a Body Text (style 943) template paragraph from CH3. New bibliography
entries (if any) are inserted in NEW_REFERENCES at their alphabetical
positions before APPENDIX, idempotent.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"


PARAGRAPHS: list[str] = [
    # ¶1 — Why we need construction (Group A)
    "Following data collection, the final AdFlex corpus obtained consists of 55,000 ads ranging from low to high tiers and spanning different verticals and platforms. The construction phase follows this in order to create training examples (supervised fine-tuning pairs) from which the student model can learn the style of high-performing ads. The format of the training examples follows a brief-to-assistant-response format (a standard chat-completion template) that can be used to train or fine-tune LLMs. The collected data contains real ads with engagement scores. However, they are not formatted in a way that can be used directly in the training pipeline. Since writing briefs by hand is expensive and slow, we turn to instruction backtranslation (Li et al., 2024) — example synthesis using teacher LLM models.",
    # ¶2 — Selection filters and caps (Group B)
    "Before feeding data from the collected corpus directly to the teacher for construction, several phases and filters are first cleared to ensure the data quality is sufficient for fine-tuning the model. Specifically, only ads with composite score over 0.70 are kept. This is to ensure truly the high-performing ads are the ones representing the ad set. Then additional small cascade filters such as English-only ads, character minimum thresholds of 60 characters (headline, body, description, and CTA combined), are used to ensure expected ads are being passed. Additionally, a few smaller filters such as training quality and content safety rating data columns, which are filled out by an LLM, are used to assess the quality of the training examples beyond rule-based deterministic features. Hard caps are also applied to each industry, advertiser, and vertical (maximum 400 ads per vertical, 50 per advertiser, minimum 30 ads per vertical) to ensure not one type dominates the training dataset.",
    # ¶3 — Teacher prompt: unlabeled input + rule sheet + register + models (Group C + D)
    "Following the filters and finalization of the ad set to be used for training, the raw ads are sent to the teacher as clean, unlabeled text, without labels such as platform, vertical, or score, to keep the model totally unbiased and get a clear text input: “This ad ran for [advertiser]. The ad copy reads:” More importantly, a fixed rule set is attached to each prompt or teacher bundle that focuses on two main points. First, the brief must be purely factual about the product only, what it is, and what it does, and clean of any creative information. It closely resembles real-world operation where the model is prompted on a product without any prior knowledge of any creative work or advertising direction. The second most critical rule is that the assistant response must reproduce the golden ad reference verbatim, since this is the core value of the training instruction and should not be paraphrased or meddled with by the LLM. The LLM is encouraged to add a short rationale grounded in specifics present in the ad itself (a word, a comma, a sequence, a line break), not inferred. Additional rules such as output schema requirements to ensure parsing of the results are also attached to the teacher bundle. To increase variety of the responses from LLMs, each instruction is also assigned a conversation style — either conversational, structured, or imperative. This roll is randomly distributed across source ads and deterministically derived from source ad IDs. Teacher models chosen for instruction generation are chosen to balance cost and quality. Three main chosen providers are Claude Haiku 4.5 at 40% of the total ad count, GPT-5.5 for 35%, and Gemini 3 Flash Preview at 25% of the total. Total planned number of generations is approximately 5,625 (a 4,500-example target with a 1.25× overgeneration buffer). Generation is done through the OpenAI, Anthropic, and Gemini batch APIs, which are 50% cheaper than normal sync calls. Completion time ranges between 1 and 24 hours, making it perfectly suitable for this application. The batch generation pipeline allows resumability and retry logic for rejected ads, to ensure smooth operation if model outputs violate any rules.",
    # ¶4 — Quality gates: tags, fidelity, schema-leak, dedup (Group E)
    "Before ingesting the teacher's responses, a few quality gates exist to ensure the ingested examples are of sufficient quality. First, making sure the system tags are present; malformed or unstructured responses are rejected straight away. A secondary fidelity gate is used to ensure the exact words are appearing in the response verbatim, indicating the model has adhered to the rule that the golden ad must be preserved word-for-word and not paraphrased. Models who fabricate responses or ad copies are rejected. Additionally, schema leak filters are also applied to filter out responses that contain literal headlines, body labels, or sponsor labels in case the models try to copy the ad structure instead of mentioning it naturally as part of the response. Finally, deduplication using TF-IDF cosine vectors (threshold 0.80) is applied on responses to ensure duplicates or near-duplicates are rejected, maintaining diversity.",
    # ¶5 — Final dataset: three-message chat, splits, metadata (Group F)
    "The final training dataset example structure follows a three-message chat. The first is a system prompt which is fixed across all examples and is used during model inference: “You are an ad copywriter. When a user describes a product or campaign, you write ad copy and a short rationale explaining why the execution works.” The user prompt or teacher's brief is the model's proposed plausible user prompt that can result in the high-performing ad as part of the response. And finally, the assistant response includes the corresponding real ad as part of a natural LLM response as well as a rationale on why this ad should be chosen. The final dataset is split into 85% train, 7.5% validation, and 7.5% test (2,442 / 215 / 215 examples, 2,872 total), stratified by platform so that each split gets an equal representation of all platforms in the corpus. Each training example is tagged with metadata fields such as the source tier, the platform, the vertical, the model, etc., to enhance debugging and iteration. All these fields are protected from leaking into the student messages, which are strictly limited to the three chat components discussed earlier.",
]


# (alphabetical-anchor-prefix, full reference text)
# Empty for §3.4 — Li et al. (2024) is already in the bibliography from §2.7.3.
NEW_REFERENCES: list[tuple[str, str]] = []


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def set_paragraph_text(p_elem, new_text):
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def write_section_body(doc):
    h_3_4 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("3.4 Training-Data Construction"),
        "§3.4 heading",
    )
    h_3_5 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("3.5 Fine-Tuning"),
        "§3.5 heading",
    )

    # Body Text template: first style-943 paragraph inside CH3 (excluding
    # §3.4's own body, which we're about to wipe).
    body_tpl = None
    in_ch3 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.upper() == "CHAPTER III":
            in_ch3 = True
            continue
        if t.upper() == "CHAPTER IV":
            break
        if t.startswith("3.4 Training-Data Construction"):
            in_ch3 = "in-34"
            continue
        if in_ch3 == "in-34" and t.startswith("3.5 "):
            in_ch3 = True
            continue
        if in_ch3 == "in-34":
            continue
        sid = p.style.style_id if p.style else ""
        if in_ch3 and sid == "943" and t and p._element.findall(f".//{QN_T}"):
            body_tpl = p._element
            break
    if body_tpl is None:
        raise RuntimeError("No style-943 Body Text paragraph found inside CH3 (outside §3.4)")

    body = h_3_4.getparent()
    children = list(body)
    i_start = children.index(h_3_4) + 1
    i_end = children.index(h_3_5)

    to_delete = [c for c in children[i_start:i_end] if c.tag == QN_P]
    print(f"DELETING: {len(to_delete)} existing §3.4 body paragraphs")
    for c in to_delete:
        body.remove(c)

    insert_after = h_3_4
    for i, text in enumerate(PARAGRAPHS, 1):
        p = deepcopy(body_tpl)
        set_paragraph_text(p, text)
        insert_after.addnext(p)
        insert_after = p
        print(f"  ¶{i}: {text[:80]}…")

    print(f"\nWROTE: {len(PARAGRAPHS)} paragraph(s) into §3.4 Training-Data Construction")


def insert_references(doc):
    if not NEW_REFERENCES:
        print("(no new references for §3.4)")
        return

    refs_heading = None
    appendix_heading = None
    for p in doc.paragraphs:
        t = p.text.strip().upper()
        if t == "REFERENCES" and refs_heading is None:
            refs_heading = p._element
        elif t.startswith("APPENDIX") and appendix_heading is None:
            appendix_heading = p._element
    if refs_heading is None or appendix_heading is None:
        raise RuntimeError("References / APPENDIX bounds not found")

    body = refs_heading.getparent()
    children = list(body)
    i_refs = children.index(refs_heading)
    i_appendix = children.index(appendix_heading)
    ref_range = children[i_refs + 1 : i_appendix]

    ref_tpl = None
    for c in ref_range:
        if c.tag != QN_P:
            continue
        pPr = c.find(f"{{{W_NS}}}pPr")
        pStyle = pPr.find(f"{{{W_NS}}}pStyle") if pPr is not None else None
        sid = pStyle.get(f"{{{W_NS}}}val") if pStyle is not None else ""
        ts = c.findall(f".//{QN_T}")
        if sid == "943" and ts and any(t.text for t in ts):
            ref_tpl = c
            break
    if ref_tpl is None:
        raise RuntimeError("No style-943 reference template paragraph found")

    inserted = 0
    skipped = 0
    for anchor_prefix, ref_text in NEW_REFERENCES:
        first_words = ref_text.split("(")[0].strip().rstrip(",")
        already_present = any(
            c.tag == QN_P and "".join(t.text or "" for t in c.findall(f".//{QN_T}")).startswith(first_words)
            for c in ref_range
        )
        if already_present:
            print(f"  SKIP (already present): {ref_text[:80]}…")
            skipped += 1
            continue

        anchor = None
        for c in ref_range:
            if c.tag != QN_P:
                continue
            text = "".join(t.text or "" for t in c.findall(f".//{QN_T}"))
            if text.startswith(anchor_prefix):
                anchor = c
                break
        if anchor is None:
            raise RuntimeError(f"Alphabetical anchor not found: {anchor_prefix!r}")

        new_p = deepcopy(ref_tpl)
        set_paragraph_text(new_p, ref_text)
        anchor.addprevious(new_p)
        inserted += 1
        print(f"  INSERT before {anchor_prefix!r}: {ref_text[:80]}…")

        children = list(body)
        i_refs_new = children.index(refs_heading)
        i_appendix_new = children.index(appendix_heading)
        ref_range = children[i_refs_new + 1 : i_appendix_new]

    print(f"\nREFERENCES: inserted {inserted}, skipped {skipped} (already present)")


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    print("=== §3.4 body ===")
    write_section_body(doc)

    print("\n=== References ===")
    insert_references(doc)

    doc.save(str(path))
    print(f"\nSAVED: {path}")


if __name__ == "__main__":
    main()
