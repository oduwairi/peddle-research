"""Rebuild THESIS.docx List of Tables / List of Figures as borderless column
tables that match the reference exemplar's format exactly.

Reference (`SURE NEU THESIS FORMAT abduallah inal PDF 12.docx`) builds both lists
as BORDERLESS tables (every <w:tblBorders> entry is val=none), not the Word-native
dot-leader Table-of-Figures field we currently use:

  - List of Tables  -> 2 cols  [Table N: caption | page]   widths 5214/2005 twips
  - List of Figures -> 3 cols  [Figure N | description | page] widths 1182/6159/881

We keep the existing PAGEREF page fields (moved into the page cell) so the page
column auto-fills the real page on F9 refresh — visually identical to the
exemplar's static numbers, but correct (the field cache currently reads 0). Pass
--static to instead drop a literal placeholder you fill by hand.

Idempotent: if a table already sits immediately after the heading, that list is
skipped. Anchored by heading text + style name, never by index. python-docx +
lxml only; no byte-slicing.

Run:  uv run python scripts/thesis/rebuild_lof_lot_as_tables.py            # dry-run
      uv run python scripts/thesis/rebuild_lof_lot_as_tables.py --apply
      uv run python scripts/thesis/rebuild_lof_lot_as_tables.py --apply --static
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

DOCX = Path("docs/research/THESIS.docx")
FONT = "Times New Roman"
PT = 12

ENTRY_STYLE = "table of figures"  # style name of the current dot-leader entries
HEADING_STYLE = "Heading 2"

FIG_LABEL_RE = re.compile(r"^(Figure\s+\d+(?:\.\d+)*)\s*[:.]\s*(.*)$", re.S)
PAGEREF_RE = re.compile(r"PAGEREF\s+(\S+)")


# ----------------------------------------------------------------------------- helpers


def set_run_font(run) -> None:
    run.font.name = FONT
    run.font.size = Pt(PT)
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)


def no_borders(tbl) -> None:
    """Force every table border to val=none (the exemplar's borderless look)."""
    tblPr = tbl._tbl.tblPr
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")
        borders.append(e)
    tblPr.append(borders)


def set_grid(tbl, widths: list[int]) -> None:
    """Fixed layout with explicit column widths (twips)."""
    tblPr = tbl._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl._tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    # tblGrid must follow tblPr
    tblPr.addnext(grid)
    # also set each cell width to match
    for row in tbl.rows:
        for cell, w in zip(row.cells, widths):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(w))


def write_cell(cell, text: str, *, align: str | None, bold: bool) -> None:
    p = cell.paragraphs[0]
    p.style = cell.part.document.styles["Body Text"]
    # clear default spacing so rows are tight like the exemplar
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    if align == "center":
        p.alignment = 1
    elif align == "right":
        p.alignment = 2
    else:
        p.alignment = 0
    run = p.add_run(text)
    run.bold = bold
    set_run_font(run)


def write_page_cell(cell, bookmark: str, static: bool) -> None:
    p = cell.paragraphs[0]
    p.style = cell.part.document.styles["Body Text"]
    p.alignment = 2  # right
    if static:
        run = p.add_run("00")  # placeholder for author to fill
        set_run_font(run)
        return
    # live PAGEREF field via fldSimple, cached "0" (refreshes on F9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f" PAGEREF {bookmark} \\h ")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(PT * 2))
    rpr.append(sz)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = "0"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


# ----------------------------------------------------------------------------- core


def parse_entry(p) -> tuple[str, str]:
    """Return (caption_text, bookmark_name) for a dot-leader list entry."""
    caption = p.text.split("\t")[0].strip()
    bookmark = ""
    for instr in p._p.iter(qn("w:instrText")):
        m = PAGEREF_RE.search(instr.text or "")
        if m:
            bookmark = m.group(1)
            break
    return caption, bookmark


def find_heading(doc, title: str):
    for p in doc.paragraphs:
        name = (p.style.name if p.style else "").lower()
        if name == HEADING_STYLE.lower() and p.text.strip() == title:
            return p
    return None


def collect_entries(heading_p):
    """Consecutive 'table of figures' entries following the heading."""
    entries = []
    el = heading_p._p.getnext()
    from docx.text.paragraph import Paragraph

    doc = heading_p.part.document
    while el is not None and el.tag == qn("w:p"):
        p = Paragraph(el, doc)
        if (p.style.name if p.style else "").lower() == ENTRY_STYLE:
            entries.append(p)
            el = el.getnext()
        else:
            break
    return entries


def rebuild(doc, title: str, kind: str, widths: list[int], static: bool, dry: bool) -> str:
    heading = find_heading(doc, title)
    if heading is None:
        return f"FAIL  {title}: heading not found"

    nxt = heading._p.getnext()
    if nxt is not None and nxt.tag == qn("w:tbl"):
        return f"skip  {title}: a table already follows the heading (idempotent)"

    entries = collect_entries(heading)
    if not entries:
        return f"FAIL  {title}: no '{ENTRY_STYLE}' entries found after heading"

    rows = []
    for p in entries:
        caption, bookmark = parse_entry(p)
        if kind == "fig":
            m = FIG_LABEL_RE.match(caption)
            if m:
                label, desc = m.group(1).strip(), m.group(2).strip()
            else:
                label, desc = "", caption
            rows.append((label, desc, bookmark))
        else:  # tables: keep full caption in one cell
            rows.append((caption, "", bookmark))

    if dry:
        out = [f"OK    {title}: would build {len(rows)}-row "
               f"{'3' if kind == 'fig' else '2'}-col borderless table"]
        for r in rows[:4]:
            out.append(f"        {r}")
        if len(rows) > 4:
            out.append(f"        … (+{len(rows) - 4} more)")
        bad = [r for r in rows if not r[-1]]
        if bad:
            out.append(f"        WARN: {len(bad)} entries missing a PAGEREF bookmark")
        return "\n".join(out)

    # Build the table at end, then relocate after the heading.
    ncols = 3 if kind == "fig" else 2
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.autofit = False
    no_borders(tbl)
    set_grid(tbl, widths)

    for i, r in enumerate(rows):
        cells = tbl.rows[i].cells
        if kind == "fig":
            label, desc, bookmark = r
            write_cell(cells[0], label, align="center", bold=True)
            write_cell(cells[1], desc, align="left", bold=False)
            write_page_cell(cells[2], bookmark, static)
        else:
            caption, _, bookmark = r
            write_cell(cells[0], caption, align="left", bold=False)
            write_page_cell(cells[1], bookmark, static)

    # Remove old entry paragraphs, then move the table directly after the heading.
    for p in entries:
        p._p.getparent().remove(p._p)
    heading._p.addnext(tbl._tbl)
    return f"OK    {title}: built {len(rows)}-row table, removed {len(entries)} old entries"


def main() -> int:
    apply = "--apply" in sys.argv
    static = "--static" in sys.argv
    dry = not apply

    if not DOCX.exists():
        print(f"ERROR: {DOCX} not found", file=sys.stderr)
        return 1

    lock = DOCX.parent / f".~lock.{DOCX.name}#"
    if apply and lock.exists():
        print(f"ERROR: {lock} present — THESIS.docx is open in OnlyOffice/LibreOffice.\n"
              "Close it first (scripted saves get clobbered otherwise).", file=sys.stderr)
        return 1

    doc = Document(str(DOCX))
    print(f"{'DRY-RUN' if dry else 'APPLY'}{' (static pages)' if static else ''} on {DOCX}")
    print("-" * 70)

    results = [
        rebuild(doc, "List of Tables", "tbl", [5214, 2005], static, dry),
        rebuild(doc, "List of Figures", "fig", [1182, 6159, 881], static, dry),
    ]
    for r in results:
        print(r)
    print("-" * 70)

    if any(r.startswith("FAIL") for r in results):
        print("ABORT: not saving (a rebuild failed).")
        return 1
    if dry:
        print("DRY-RUN: no changes written. Re-run with --apply (close OnlyOffice first).")
        return 0

    doc.save(str(DOCX))
    print(f"saved {DOCX}")
    print("FOLLOW-UP in OnlyOffice: Ctrl+A → F9 to refresh page-number fields "
          "(they currently read 0 until refreshed).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
