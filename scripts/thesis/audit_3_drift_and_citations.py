"""Chapter III audit pass — drift fixes + citation strengthening + figure renumber.

Single idempotent script that does precise in-place text substitutions inside
Chapter III paragraphs and inserts new bibliography entries alphabetically.

Unlike the per-section ``write_3_N_*.py`` scripts (which wipe-and-replace the
whole §3.x body and would clobber figure-caption paragraphs sitting between
body paragraphs), this script only modifies the runs whose text matches each
substitution rule. It is safe to re-run.

What it does:
 1. SUBSTITUTIONS — exact in-paragraph text replacements (drift fixes,
    inline citation insertions, figure renumbering). Each rule names a
    paragraph anchor prefix (first N chars of the target paragraph) plus a
    find/replace pair. Idempotent: a rule that has already been applied is
    skipped silently.
 2. NEW_REFERENCES — APA bibliography entries inserted before the first
    paragraph in the References section whose text starts with the given
    alphabetical anchor. Idempotent: skips if already present (matches on
    first author+year prefix).

Run with:  uv run python scripts/thesis/audit_3_drift_and_citations.py
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"


# (paragraph_anchor_prefix, find_text, replace_text, label)
SUBSTITUTIONS: list[tuple[str, str, str, str]] = [
    # ----- §3.3 Engagement-Based Scoring -----
    (
        "The main limitation of AdFlex ads",
        "indirect measurement of ad performance,",
        "indirect measurement of ad performance (Aral & Walker, 2011),",
        "§3.3 ¶381: Aral & Walker (engagement-as-proxy anchor)",
    ),
    (
        "As for lifespan, we use a Kaplan",
        "(Lee, Moon, & Salamatian, 2012).",
        "(Lee, Moon, & Salamatian, 2012; Cha et al., 2007).",
        "§3.3 ¶383: Cha et al. (online-content lifecycle)",
    ),
    (
        "A developed scoring system uses proxies",
        "Spearman rank correlation with bootstrap 95% confidence intervals",
        "Spearman rank correlation (Spearman, 1904) with bootstrap 95% confidence intervals (Efron, 1979)",
        "§3.3 ¶385: Spearman 1904 + Efron 1979",
    ),
    (
        "A developed scoring system uses proxies",
        "The Kruskal–Wallis H-test",
        "The Kruskal–Wallis H-test (Kruskal & Wallis, 1952)",
        "§3.3 ¶385: Kruskal & Wallis 1952",
    ),
    # ----- §3.4 Training-Data Construction -----
    (
        # Caption-suffix anchor (substring) — stable across the renumber.
        "Instruction backtranslation",
        "Figure 3.2: Instruction backtranslation",
        "Figure 3.3: Instruction backtranslation",
        "§3.4 ¶389: Figure 3.2 → 3.3 (renumber to break collision with §3.2 figure)",
    ),
    (
        "Before ingesting the teacher's responses",
        "deduplication using TF-IDF cosine vectors (threshold 0.80)",
        "deduplication using TF-IDF cosine vectors (Manning, Raghavan, & Schütze, 2008; threshold 0.80)",
        "§3.4 ¶392: Manning et al. (TF-IDF / IR textbook anchor)",
    ),
    # ----- §3.5 Fine-Tuning -----
    (
        "The setup phase includes shrinking",
        "rsLoRA (rank-stabilized LoRA),",
        "rsLoRA (rank-stabilized LoRA; Kalajdzievski, 2023),",
        "§3.5 ¶396: Kalajdzievski (rsLoRA)",
    ),
    (
        "During training, each training example",
        "(this is known as assistant-only loss masking)",
        "(this is known as assistant-only loss masking; Ouyang et al., 2022)",
        "§3.5 ¶397: Ouyang et al. (assistant-only SFT loss)",
    ),
    (
        "The training is set to run for three epochs",
        "Weight updates use an eight-bit version of the AdamW optimizer in order to save memory.",
        "Weight updates use an eight-bit version of the AdamW optimizer (Loshchilov & Hutter, 2019; Dettmers et al., 2022) in order to save memory.",
        "§3.5 ¶398: AdamW + 8-bit (Loshchilov 2019 + Dettmers 2022)",
    ),
    (
        "The training is set to run for three epochs",
        "The learning rate is set to follow a cosine schedule,",
        "The learning rate is set to follow a cosine schedule (Loshchilov & Hutter, 2017),",
        "§3.5 ¶398: Loshchilov & Hutter 2017 (SGDR / cosine)",
    ),
    # ----- §3.6 Agent Architecture -----
    (
        "The agent system consists of two main models",
        "an orchestrator model and a writer model.",
        "an orchestrator model and a writer model (Wu et al., 2023).",
        "§3.6 ¶402: Wu et al. AutoGen (two-role orchestrator-worker)",
    ),
    (
        # Caption-suffix anchor (substring) — stable across the renumber.
        "Draper.ai agent architecture",
        "Figure 3.3: Draper.ai agent architecture",
        "Figure 3.4: Draper.ai agent architecture",
        "§3.6 ¶405: Figure 3.3 → 3.4 (renumber cascade from §3.4)",
    ),
    (
        "Communication between the orchestrator and the writer",
        "The proposed architecture fires multiple requests in parallel with different temperatures (six draws at temperatures 0.5 to 1.0) and scores them with the predictor,",
        "On ask_draper, the proposed architecture fires multiple requests in parallel with different temperatures (six draws at temperatures 0.5 to 1.0) and scores them with the predictor (Cobbe et al., 2021),",
        "§3.6 ¶407: best-of-N scope fix + Cobbe et al. (verifier reranking)",
    ),
    (
        "Communication between the orchestrator and the writer",
        "calls the OpenAI gpt-image-1.5 model for visual generation",
        "calls the OpenAI image API (gpt-image-1.5 by default) for visual generation",
        "§3.6 ¶407: image model id (env-overridable)",
    ),
    # ----- §3.7 Scoring Predictor -----
    (
        "The idea is to train a small text regressor",
        "(a fine-tuned DeBERTa-v3-base)",
        "(a fine-tuned DeBERTa-v3-base; He, Gao, & Chen, 2021)",
        "§3.7 ¶410: He, Gao & Chen (DeBERTa-v3)",
    ),
    (
        "To train the model we are using the same 55,000-ad corpus",
        "The model is trained to predict all four labels instead of relying only on the composite",
        "The model is trained to predict all four labels instead of relying only on the composite (multi-head shared-encoder regression; Caruana, 1997)",
        "§3.7 ¶411: Caruana (multitask learning)",
    ),
    (
        "We use a small pretrained text model with four prediction outputs",
        "the four prediction heads get larger learning rates (1e-3)",
        "the four prediction heads get larger learning rates (1e-4)",
        "§3.7 ¶412: DRIFT FIX head LR 1e-3 → 1e-4",
    ),
    (
        "As for evaluation of the trained regressor",
        "AUC for the bottom and top tier ad sets",
        "AUC (Fawcett, 2006) for the bottom and top tier ad sets",
        "§3.7 ¶414: Fawcett (ROC-AUC)",
    ),
    # ----- §3.8 Evaluation Methodology -----
    (
        "Model evaluation remains as the final",
        "For this reason, we use multiple evaluation arms,",
        "For this reason, we use multiple evaluation arms (Sai, Mohankumar, & Khapra, 2022),",
        "§3.8 ¶416: Sai et al. (NLG eval taxonomy)",
    ),
    (
        "For the evaluation, the held-out test split",
        "to ensure fair comparison.",
        "to ensure fair comparison (Magar & Schwartz, 2022).",
        "§3.8 ¶417: Magar & Schwartz (data contamination)",
    ),
]


# (alphabetical-anchor-prefix, full reference text)
# Same-anchor entries are inserted in list order; reverse-order placement
# (last inserted = closest to anchor) is accounted for so the final
# alphabetical sequence is correct.
NEW_REFERENCES: list[tuple[str, str]] = [
    # ===== A =====
    (
        "Aralim",
        "Aral, S., & Walker, D. (2011). Creating social contagion through viral product design: A randomized trial of peer influence in networks. Management Science, 57(9), 1623–1639. https://doi.org/10.1287/mnsc.1110.1421",
    ),
    # ===== C =====
    (
        "Casper",
        "Caruana, R. (1997). Multitask learning. Machine Learning, 28, 41–75. https://doi.org/10.1023/A:1007379606734",
    ),
    (
        "Chiang",
        "Cha, M., Kwak, H., Rodriguez, P., Ahn, Y.-Y., & Moon, S. (2007). I Tube, You Tube, everybody tubes: Analyzing the world's largest user-generated content video system. In Proceedings of the 7th ACM SIGCOMM Conference on Internet Measurement (pp. 1–14).",
    ),
    # Cobbe + Dettmers 2022 both anchor on Dettmers 2023; Cobbe first so it
    # lands further from the anchor (correct alphabetical: ... Chung → Cobbe →
    # Dettmers 2022 → Dettmers 2023 → Dubois).
    (
        "Dettmers, T., Pagnoni",
        "Cobbe, K., Kosaraju, V., Bavarian, M., Chen, M., Jun, H., Kaiser, L., Plappert, M., Tworek, J., Hilton, J., Nakano, R., Hesse, C., & Schulman, J. (2021). Training verifiers to solve math word problems (arXiv:2110.14168). https://arxiv.org/abs/2110.14168",
    ),
    (
        "Dettmers, T., Pagnoni",
        "Dettmers, T., Lewis, M., Shleifer, S., & Zettlemoyer, L. (2022). 8-bit optimizers via block-wise quantization. 10th International Conference on Learning Representations. https://arxiv.org/abs/2110.02861",
    ),
    # ===== E / F =====
    (
        "Fan",
        "Efron, B. (1979). Bootstrap methods: Another look at the jackknife. The Annals of Statistics, 7(1), 1–26. https://doi.org/10.1214/aos/1176344552",
    ),
    (
        "G. Team",
        "Fawcett, T. (2006). An introduction to ROC analysis. Pattern Recognition Letters, 27(8), 861–874. https://doi.org/10.1016/j.patrec.2005.10.010",
    ),
    # ===== H =====
    (
        "Helberger",
        "He, P., Gao, J., & Chen, W. (2021). DeBERTaV3: Improving DeBERTa using ELECTRA-style pre-training with gradient-disentangled embedding sharing (arXiv:2111.09543). https://arxiv.org/abs/2111.09543",
    ),
    # ===== K =====
    (
        "Kaplan",
        "Kalajdzievski, D. (2023). A rank stabilization scaling factor for fine-tuning with LoRA (arXiv:2312.03732). https://arxiv.org/abs/2312.03732",
    ),
    (
        "Kurihara",
        "Kruskal, W. H., & Wallis, W. A. (1952). Use of ranks in one-criterion variance analysis. Journal of the American Statistical Association, 47(260), 583–621. https://doi.org/10.1080/01621459.1952.10483441",
    ),
    # ===== L =====
    # Loshchilov 2017 first so 2019 lands closer to anchor (correct order:
    # ... Lops → Loshchilov 2017 → Loshchilov 2019 → Lu).
    (
        "Lu",
        "Loshchilov, I., & Hutter, F. (2017). SGDR: Stochastic gradient descent with warm restarts. 5th International Conference on Learning Representations. https://arxiv.org/abs/1608.03983",
    ),
    (
        "Lu",
        "Loshchilov, I., & Hutter, F. (2019). Decoupled weight decay regularization. 7th International Conference on Learning Representations. https://arxiv.org/abs/1711.05101",
    ),
    # ===== M =====
    # Magar first so Manning lands closer to anchor (correct order:
    # ... Lu → Magar → Manning → Matias).
    (
        "Matias",
        "Magar, I., & Schwartz, R. (2022). Data contamination: From memorization to exploitation. In Proceedings of the 60th Annual Meeting of the Association for Computational Linguistics (Volume 2: Short Papers) (pp. 157–165). Association for Computational Linguistics. https://aclanthology.org/2022.acl-short.18",
    ),
    (
        "Matias",
        "Manning, C. D., Raghavan, P., & Schütze, H. (2008). Introduction to information retrieval. Cambridge University Press.",
    ),
    # ===== O =====
    (
        "Papineni",
        "Ouyang, L., Wu, J., Jiang, X., Almeida, D., Wainwright, C., Mishkin, P., Zhang, C., Agarwal, S., Slama, K., Ray, A., Schulman, J., Hilton, J., Kelton, F., Miller, L., Simens, M., Askell, A., Welinder, P., Christiano, P., Leike, J., & Lowe, R. (2022). Training language models to follow instructions with human feedback. Advances in Neural Information Processing Systems, 35, 27730–27744. https://arxiv.org/abs/2203.02155",
    ),
    # ===== S =====
    (
        "Sanh",
        "Sai, A. B., Mohankumar, A. K., & Khapra, M. M. (2022). A survey of evaluation metrics used for NLG systems. ACM Computing Surveys, 55(2), Article 39. https://doi.org/10.1145/3485766",
    ),
    (
        "Susser",
        "Spearman, C. (1904). The proof and measurement of association between two things. The American Journal of Psychology, 15, 72–101.",
    ),
    # ===== W =====
    (
        "Xi",
        "Wu, Q., Bansal, G., Zhang, J., Wu, Y., Li, B., Zhu, E., Jiang, L., Zhang, X., Zhang, S., Liu, J., Awadallah, A. H., White, R. W., Burger, D., & Wang, C. (2023). AutoGen: Enabling next-generation LLM applications via multi-agent conversation (arXiv:2308.08155). https://arxiv.org/abs/2308.08155",
    ),
]


def paragraph_text(p_elem) -> str:
    return "".join(t.text or "" for t in p_elem.findall(f".//{QN_T}"))


def find_ch3_paragraph(doc, anchor: str):
    """Return the first paragraph inside CHAPTER III whose text contains anchor as a substring.

    Substring matching lets figure-renumber rules survive their own
    substitution (the figure number changes but a stable caption suffix
    stays). Body-paragraph anchors use unique full-sentence openings, so
    substring matching can't cross-match them.
    """
    in_ch3 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        upper = t.upper()
        if upper == "CHAPTER III":
            in_ch3 = True
            continue
        if upper == "CHAPTER IV":
            break
        if in_ch3 and anchor in t:
            return p._element
    return None


def substitute_in_paragraph(p_elem, find_str: str, replace_str: str) -> tuple[bool, bool]:
    """Replace first occurrence of find_str in the paragraph's concatenated text.

    Returns (changed, already_applied):
      - (True, False)  — substitution applied
      - (False, True)  — replace_str already in paragraph; idempotent skip
      - (False, False) — find_str not in paragraph; rule didn't match
    """
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        return (False, False)
    full = "".join(t.text or "" for t in ts)
    if replace_str in full:
        return (False, True)
    if find_str not in full:
        return (False, False)
    new = full.replace(find_str, replace_str, 1)
    ts[0].text = new
    for t in ts[1:]:
        t.text = ""
    return (True, False)


def apply_substitutions(doc) -> None:
    applied = 0
    skipped = 0
    missing = 0
    for anchor, find, replace, label in SUBSTITUTIONS:
        p = find_ch3_paragraph(doc, anchor)
        if p is None:
            print(f"  MISS (anchor not found): {label}")
            missing += 1
            continue
        changed, already = substitute_in_paragraph(p, find, replace)
        if changed:
            print(f"  APPLY: {label}")
            applied += 1
        elif already:
            print(f"  SKIP (idempotent): {label}")
            skipped += 1
        else:
            print(f"  MISS (find string not in paragraph): {label}")
            missing += 1
    print(f"\nSUBSTITUTIONS: applied {applied}, skipped {skipped}, missing {missing}")
    if missing:
        raise RuntimeError(
            f"{missing} substitution(s) did not match — refusing to save. "
            "Likely the docx prose has drifted from the anchor strings; "
            "either re-run the per-section write_3_N script first or update "
            "the anchors in this script."
        )


def insert_references(doc) -> None:
    refs_heading = None
    appendix_heading = None
    for p in doc.paragraphs:
        t = p.text.strip().upper()
        if t == "REFERENCES" and refs_heading is None:
            refs_heading = p._element
        elif t.startswith("APPENDIX") and appendix_heading is None:
            appendix_heading = p._element
    if refs_heading is None or appendix_heading is None:
        raise RuntimeError("References / APPENDIX bounds not found")

    body = refs_heading.getparent()

    def current_ref_range():
        children = list(body)
        i_refs = children.index(refs_heading)
        i_app = children.index(appendix_heading)
        return children[i_refs + 1 : i_app]

    # Body Text template = first existing reference entry (style 943).
    ref_tpl = None
    for c in current_ref_range():
        if c.tag != QN_P:
            continue
        pPr = c.find(f"{{{W_NS}}}pPr")
        pStyle = pPr.find(f"{{{W_NS}}}pStyle") if pPr is not None else None
        sid = pStyle.get(f"{{{W_NS}}}val") if pStyle is not None else ""
        ts = c.findall(f".//{QN_T}")
        if sid == "943" and ts and any((t.text or "") for t in ts):
            ref_tpl = c
            break
    if ref_tpl is None:
        raise RuntimeError("No style-943 reference template paragraph found")

    inserted = 0
    skipped = 0
    for anchor_prefix, ref_text in NEW_REFERENCES:
        # Idempotency: skip if author+year prefix already present.
        first_words = ref_text.split("(")[0].strip().rstrip(",")
        already_present = any(
            c.tag == QN_P and paragraph_text(c).startswith(first_words)
            for c in current_ref_range()
        )
        if already_present:
            print(f"  SKIP (already present): {ref_text[:80]}…")
            skipped += 1
            continue

        anchor = None
        for c in current_ref_range():
            if c.tag != QN_P:
                continue
            if paragraph_text(c).startswith(anchor_prefix):
                anchor = c
                break
        if anchor is None:
            raise RuntimeError(f"Alphabetical anchor not found: {anchor_prefix!r}")

        new_p = deepcopy(ref_tpl)
        ts = new_p.findall(f".//{QN_T}")
        ts[0].text = ref_text
        for t in ts[1:]:
            t.text = ""
        anchor.addprevious(new_p)
        inserted += 1
        print(f"  INSERT before {anchor_prefix!r}: {ref_text[:80]}…")

    print(f"\nREFERENCES: inserted {inserted}, skipped {skipped} (already present)")


def main() -> None:
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    print("=== Substitutions (drift fixes + inline cites + figure renumber) ===")
    apply_substitutions(doc)

    print("\n=== Reference insertions ===")
    insert_references(doc)

    doc.save(str(path))
    print(f"\nSAVED: {path}")


if __name__ == "__main__":
    main()
