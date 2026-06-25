"""Scaffold Chapter IV (Results).

Replaces the template author's brain-tumour CH4 content with our 5-subsection
results skeleton. Each subsection gets a numbered Heading 2 + a placeholder
body paragraph. The "Result and discussion" H2 anchor is preserved (the same
way scaffold_chapter_iii.py preserved the "Methodology" H2 inside CHAPTER III).

Conventions (see docs/research/THESIS_EDITING.md):
- Clone Heading 2 from §2.7.3 (plain inline, no <w:drawing>), never from
  original-template headings.
- Clone Body Text from the first style-943 paragraph before CHAPTER IV with
  non-empty text. This is robust to write_3_X.py rewrites of earlier
  templates (the original "The methodology of the research" anchor used by
  the CH3 scaffold may no longer exist by the time this script runs).
- Delete `<w:p>` and `<w:tbl>` between "Result and discussion" and "CHAPTER V";
  skip any paragraph containing `<w:sectPr>` (pagination guard).
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_TBL = f"{{{W_NS}}}tbl"
QN_T = f"{{{W_NS}}}t"
QN_SECTPR = f"{{{W_NS}}}sectPr"
QN_DRAWING = f"{{{W_NS}}}drawing"


INTRO = (
    "[Chapter intro — placeholder. To be drafted: framing of the evaluation "
    "around the two absolute-scoring arms (learned-scorer and MAUVE "
    "distribution matching), the test-set definition (n=215 held-out briefs "
    "across the A/B/C/B_pipe/C_pipe/GOLD configs introduced in §3.8), and a "
    "note that the pairwise LLM-judge tournament arm is deferred — early "
    "results were not robust enough to report.]"
)

SUBSECTIONS: list[tuple[str, str]] = [
    (
        "4.1 Evaluation Setup",
        "[Section body — placeholder. To be drafted: the n=215 brief test "
        "split drawn from the held-out shard of the v2 construction corpus, "
        "the six evaluated configs (A: base orchestrator only; B: base "
        "orchestrator + untuned writer; C: base orchestrator + fine-tuned "
        "Draper writer; B_pipe / C_pipe: the same writers driven by the "
        "freeform agent loop; GOLD: the real ad from the source corpus), and "
        "the normalize → score pipeline (regex pre-cleaner + Haiku-based "
        "ad-copy extractor before any scoring arm).]",
    ),
    (
        "4.2 Learned-Scorer Absolute Scores",
        "[Section body — placeholder. To be drafted: per-config composite + "
        "four-head means (survivability, engagement_volume, "
        "engagement_velocity, composite) from the DeBERTa-v3-base predictor; "
        "per-platform breakdown (Meta / TikTok / X / Google / Pinterest / "
        "Reddit); paired contrasts (C vs A, C vs B, GOLD as ceiling); "
        "calibration and confidence intervals; reference to "
        "SCORING_PREDICTOR_PHASE2_RESULTS_2026-05.md.]",
    ),
    (
        "4.3 MAUVE Distribution Matching",
        "[Section body — placeholder. To be drafted: corpus-level MAUVE "
        "scores for A/B/C/GOLD against the v3 high-tier reference corpus; "
        "per-platform MAUVE table; interpretation of distribution distance "
        "as a complement to the learned-scorer's per-ad scoring; sanity "
        "bounds (GOLD highest, untuned base lowest); reference to run "
        "2026-05-18-mauve-initial.]",
    ),
    (
        "4.4 Fine-Tuning and Agent Ablation",
        "[Section body — placeholder. To be drafted: the 2×2 ablation "
        "(B / C / B_pipe / C_pipe) isolating the fine-tuning effect (+0.036 "
        "composite, 95% CI [+0.017, +0.055]) from the agent-loop effect; the "
        "negative interaction where the freeform agent on terse briefs hurts "
        "the fine-tuned writer (−0.040 composite); discussion of why a "
        "text-only learned predictor likely underestimates the agent's "
        "multimodal value; reference to RQ2_OFFLINE_2x2_RESULTS_2026-05.md.]",
    ),
    (
        "4.5 Synthesis and Limitations",
        "[Section body — placeholder. To be drafted: what each arm reveals "
        "and conceals; whether the learned-scorer and MAUVE arms agree on "
        "the ordering across configs; absence of a pairwise LLM-judge "
        "tournament in this iteration; known limitations (text-only "
        "evaluation, single test split, predictor in-distribution bias "
        "toward AdFlex-style ads); open questions deferred to a live A/B "
        "test (RQ1) and downstream user studies.]",
    ),
]


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def find_body_template(doc, ch4_elem):
    """First style-943 paragraph before CHAPTER IV that is safe to clone.

    Must (a) have non-empty <w:t> sitting OUTSIDE any <w:drawing>, and
    (b) carry no <w:drawing> at all. Otherwise cloning it produces a
    sidebar text-frame instead of an inline body paragraph.

    Robust to write_3_X.py rewrites: any §3.x body that the author has filled
    in will do, as long as it sits before the CHAPTER IV anchor and is a
    plain inline paragraph.
    """
    for p in doc.paragraphs:
        if p._element is ch4_elem:
            break
        if (p.style.style_id if p.style else "") != "943":
            continue
        if p._element.find(f".//{QN_DRAWING}") is not None:
            continue
        ts = p._element.findall(f".//{QN_T}")
        if any((t.text or "").strip() for t in ts):
            return p._element
    raise RuntimeError("no usable style-943 body template found before CHAPTER IV")


def set_paragraph_text(p_elem, new_text):
    """Set the first <w:t> in the paragraph to new_text; clear the rest.

    Preserves <w:rPr> and run structure so the heading/body style survives.
    """
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("paragraph template has no <w:t> to write into")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    # Anchors
    rd = find_paragraph(
        doc,
        lambda p: p.text.strip() == "Result and discussion"
        and (p.style.style_id if p.style else "") == "945",
        "'Result and discussion' heading (style 945)",
    )
    ch5 = find_paragraph(
        doc,
        lambda p: p.text.strip().upper() == "CHAPTER V",
        "CHAPTER V heading",
    )
    ch4 = find_paragraph(
        doc,
        lambda p: p.text.strip().upper() == "CHAPTER IV",
        "CHAPTER IV heading",
    )

    # Templates (deep-copied later)
    h2_tpl = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("2.7.3")
        and (p.style.style_id if p.style else "") == "945",
        "§2.7.3 heading template",
    )
    if h2_tpl.find(f".//{QN_DRAWING}") is not None:
        raise RuntimeError("§2.7.3 template paragraph unexpectedly contains <w:drawing>")
    body_tpl = find_body_template(doc, ch4)

    # Range to wipe
    body = rd.getparent()
    children = list(body)
    i_rd = children.index(rd)
    i_ch5 = children.index(ch5)

    to_delete = []
    skipped_sectpr = 0
    for c in children[i_rd + 1 : i_ch5]:
        if c.tag == QN_P:
            if c.find(f".//{QN_SECTPR}") is not None:
                skipped_sectpr += 1
                continue
            to_delete.append(c)
        elif c.tag == QN_TBL:
            to_delete.append(c)

    print(f"DELETING: {len(to_delete)} elements between 'Result and discussion' and CHAPTER V")
    print(f"  (skipped {skipped_sectpr} paragraph(s) holding <w:sectPr> for pagination)")

    for c in to_delete:
        body.remove(c)

    # Insertion
    inserted = []
    insert_after = rd

    intro_p = deepcopy(body_tpl)
    set_paragraph_text(intro_p, INTRO)
    insert_after.addnext(intro_p)
    insert_after = intro_p
    inserted.append(("intro", INTRO[:80]))

    for title, body_text in SUBSECTIONS:
        h = deepcopy(h2_tpl)
        set_paragraph_text(h, title)
        insert_after.addnext(h)
        insert_after = h

        b = deepcopy(body_tpl)
        set_paragraph_text(b, body_text)
        insert_after.addnext(b)
        insert_after = b

        inserted.append((title, body_text[:80]))

    doc.save(str(path))

    print("\nINSERTIONS:")
    for title, snippet in inserted:
        print(f"  {title!r}")
        print(f"    → {snippet}…")


if __name__ == "__main__":
    main()
