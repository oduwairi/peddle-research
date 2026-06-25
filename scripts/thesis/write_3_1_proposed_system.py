"""Write §3.1 Proposed System body — running list of paragraphs.

Each polished paragraph is appended to PARAGRAPHS; rerun wipes the body of §3.1
(every paragraph between the §3.1 heading and §3.2 heading) and re-inserts the
current list using the Body Text template paragraph from §3.7.3-era prose.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"


PARAGRAPHS: list[str] = [
    # ¶1 — Goal & framing
    "The proposed marketing AI agent system takes a marketing brief as input — including product information, platform constraints, or anything a user prompt might include — and produces a full campaign output in that platform's native shape (Meta, TikTok, X, Google Responsive Search Ads, Pinterest, or Reddit). This makes the system distinct from normal chat assistants in that it is domain-specific and takes action autonomously to provide the full campaign output. A core part of the system lies in an open-source fine-tuned model in the 8-billion-parameter range, as opposed to using a frontier model with a system prompt. This improves long-term cost and deployability — the writer fits on a single L4 GPU under vLLM — as well as fine-tuning the model's weights for the marketing ad-generation domain itself, in line with the taste-by-weights argument from §2.7.",
    # ¶2 — Phases of the system
    "The proposed system is mainly developed over three main phases. The first and most important phase is data corpus construction (§3.2–§3.4), where real ads are scraped, cleaned, or labeled to be used in downstream instruction generation. Phase two involves training the model itself using QLoRA fine-tuning on a small open-source model (§3.5), which produces a fine-tuned version of the base model. Phase three involves deployment and inference of the developed model on the cloud (§3.6–§3.7), as well as performing evaluation of the model's performance (§3.8).",
    # ¶3 — Two-role agent topology
    "The proposed agent system architecture involves a two-role agent system, not a single model. The main orchestrator uses a frontier general-purpose LLM (gpt-5.4-mini by default) and is responsible for coordination, tool calls, and routing requests. A general-purpose model is better suited to this role since its broad training provides a more balanced approach for general-intelligence applications. Our small fine-tuned model serves as the writer, which the orchestrator calls for all creative and writing tasks. This division provides the right balance between what our fine-tuned model does best — writing good ad copy — and what a generalist model handles, such as coordination and operational tasks. The full tool surface and call-routing logic are detailed in §3.6, and a third component — a learned scoring model that auto-evaluates generated copy — is described in §3.7.",
]


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def set_paragraph_text(p_elem, new_text):
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    h_3_1 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("3.1 Proposed System"),
        "§3.1 heading",
    )
    h_3_2 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("3.2 Data Acquisition"),
        "§3.2 heading",
    )

    # Body Text template: first style-943 paragraph inside CH3 (robust against
    # placeholder removal — uses an existing §3.1 body paragraph if present).
    body_tpl = None
    in_ch3 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.upper() == "CHAPTER III":
            in_ch3 = True
            continue
        if t.upper() == "CHAPTER IV":
            break
        sid = p.style.style_id if p.style else ""
        if in_ch3 and sid == "943" and t and p._element.findall(f".//{QN_T}"):
            body_tpl = p._element
            break
    if body_tpl is None:
        raise RuntimeError("No style-943 Body Text paragraph found inside CH3")

    body = h_3_1.getparent()
    children = list(body)
    i_start = children.index(h_3_1) + 1
    i_end = children.index(h_3_2)

    to_delete = [c for c in children[i_start:i_end] if c.tag == QN_P]
    print(f"DELETING: {len(to_delete)} existing §3.1 body paragraphs")
    for c in to_delete:
        body.remove(c)

    insert_after = h_3_1
    for i, text in enumerate(PARAGRAPHS, 1):
        p = deepcopy(body_tpl)
        set_paragraph_text(p, text)
        insert_after.addnext(p)
        insert_after = p
        print(f"  ¶{i}: {text[:80]}…")

    doc.save(str(path))
    print(f"\nWROTE: {len(PARAGRAPHS)} paragraph(s) into §3.1 Proposed System")


if __name__ == "__main__":
    main()
