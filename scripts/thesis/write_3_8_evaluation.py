"""Write §3.8 Evaluation Methodology body.

Same pattern as write_3_7_scoring_predictor.py: wipes every paragraph between
the §3.8 heading and the CHAPTER IV heading, then re-inserts the current
PARAGRAPHS list using a Body Text (style 943) template paragraph from CH3.
New bibliography entries are inserted in NEW_REFERENCES at their alphabetical
positions before APPENDIX, idempotent.

Idempotent: re-run safely as more paragraphs land in PARAGRAPHS.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"


PARAGRAPHS: list[str] = [
    # ¶1 — Why this section exists (motivation)
    "Model evaluation remains as the final and most critical step in this research. We want to be able to evaluate our fine-tuned model, in an absolute sense as well as relative to the base model without fine-tuning as well as the frontier LLMs. This provides insight into the fine-tuned model performance and most critically answers the research question about whether an eight-billion-parameter domain-specialized model can outperform a frontier LLM on marketing domain-specific tasks. It's essential to ensure that evaluation is grounded in real-world signals and not in opinions. For this reason, we use multiple evaluation arms, such as absolute scoring using our custom-trained regressor head (§3.7) as well as embedding distribution matching against the raw corpus (MAUVE; Pillutla et al., 2021).",
    # ¶2 — The test set and what we compare
    "For the evaluation, the held-out test split of the construction dataset serves as the foundation for evaluation, since neither the predictor nor the fine-tuned model have seen this data split during training, to ensure fair comparison. The held-out split carries both the brief as well as the response containing the real ad. These are a gold standard, similar to the training instructions themselves. We evaluate the fine-tuned Draper model against the frontier baseline as well as the untuned base model that was picked for fine-tuning (§3.5), to see the effect of fine-tuning on the base model. Additionally, we evaluate agent-integrated versions of the models (§3.6), in order to evaluate the effect of fine-tuning and agent RAG on the output quality, and what each — fine-tuning and RAG, or both — contribute to model performance. In all cases, scoring the held-out high-performance ads serves as the higher ceiling that the other models are compared to.",
    # ¶3 — Cleanup before any arm reads anything
    "Before feeding the text into the scorers, we have to ensure that the ad copy text is all that's being passed and scored by these scorers. And since different models wrap their responses differently, some use markdown notation, others use headers. We need to make sure we unify the structure and feed only the ad copy, not the entire response, into these scorers for fair comparison. For this, we devise a two-stage cleanup on every inference. First, a regex-based cleanup for deterministic checks, then a small LLM (claude-haiku-4-5) acts as an extractor and extracts the ad out of each response, which may be arbitrary. Empty or whitespace-only outputs from any model are automatically excluded from the results, since they are broken and shouldn't be used. Clean inferences are saved to be fed to the downstream evaluation pipelines.",
    # ¶4 — Per-ad absolute scoring (predictor from §3.7)
    "With our obtained scoring predictor model, every clean ad text goes through it, and the output is the predictor's prediction about the ad's performance, including composite score as well as individual scores for engagement volume, velocity, and survivability. Results are aggregated per model and per platform, which is used to detect models that perform well in certain categories and not others. Since the predictor is trained on real ground-truth engagement labels, the score directly reflects the performance score of the ad, and it's not grounded in opinion. The held-out test ads are fed to the model for reference as a top ceiling for performance.",
    # ¶5 — Corpus-level distribution matching (MAUVE)
    "As a second technique to evaluate the model's performance, we use MAUVE, which essentially embeds every piece of copy into GPT-2 Large embeddings and compares embedding distributions for similarity, to get a number between zero and one representing the similarity in the vector-representation space. We dig out all contestant models' inferences on the held-out test set, and we'll exclude those from the main v3 corpus to ensure no cross-leakage, and ask the question whether the inferences and the real ads cluster look alike and possess the same kind of writing. The gold ad inferences serve as the baseline and the highest-yielding the other models are trying to chase. This has the benefit of the score: it takes distributions rather than single-point scoring and comparison. It gives an overall idea of the model's output shape compared to the real winning batch.",
    # ¶6 — Sanity bounds on the distribution-matching arm
    "Gold ads compared against the baseline should have the highest scores, but since the corpus is large and has many types of variance and fat tails, the similarity won't go up beyond certain points. However, this still serves as the baseline that the other models are expected to try and chase. Random text compared against the reference pool should score the lowest. If these boundaries fail, then the embedding model needs re-evaluation since the baseline is broken.",
]


# (alphabetical-anchor-prefix, full reference text)
NEW_REFERENCES: list[tuple[str, str]] = [
    (
        "Qin",
        "Pillutla, K., Swayamdipta, S., Zellers, R., Welleck, S., Choi, Y., & Harchaoui, Z. (2021). MAUVE: Measuring the gap between neural text and human text using divergence frontiers. Advances in Neural Information Processing Systems, 34.",
    ),
]


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def set_paragraph_text(p_elem, new_text):
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def write_section_body(doc):
    h_3_8 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("3.8 Evaluation Methodology"),
        "§3.8 heading",
    )
    h_ch4 = find_paragraph(
        doc,
        lambda p: p.text.strip().upper() == "CHAPTER IV"
        and (p.style.style_id if p.style else "") == "944",
        "CHAPTER IV heading",
    )

    # Body Text template: first style-943 paragraph inside CH3 (excluding
    # §3.8's own body, which we're about to wipe).
    body_tpl = None
    in_ch3 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.upper() == "CHAPTER III":
            in_ch3 = True
            continue
        if t.upper() == "CHAPTER IV":
            break
        if t.startswith("3.8 Evaluation Methodology"):
            in_ch3 = "in-38"
            continue
        if in_ch3 == "in-38":
            # everything after §3.8 heading inside CH3 is body we're wiping
            continue
        sid = p.style.style_id if p.style else ""
        if in_ch3 and sid == "943" and t and p._element.findall(f".//{QN_T}"):
            body_tpl = p._element
            break
    if body_tpl is None:
        raise RuntimeError("No style-943 Body Text paragraph found inside CH3 (outside §3.8)")

    body = h_3_8.getparent()
    children = list(body)
    i_start = children.index(h_3_8) + 1
    i_end = children.index(h_ch4)

    to_delete = [c for c in children[i_start:i_end] if c.tag == QN_P]
    print(f"DELETING: {len(to_delete)} existing §3.8 body paragraphs")
    for c in to_delete:
        body.remove(c)

    insert_after = h_3_8
    for i, text in enumerate(PARAGRAPHS, 1):
        p = deepcopy(body_tpl)
        set_paragraph_text(p, text)
        insert_after.addnext(p)
        insert_after = p
        print(f"  ¶{i}: {text[:80]}…")

    print(f"\nWROTE: {len(PARAGRAPHS)} paragraph(s) into §3.8 Evaluation Methodology")


def insert_references(doc):
    if not NEW_REFERENCES:
        print("(no new references for §3.8)")
        return

    refs_heading = None
    appendix_heading = None
    for p in doc.paragraphs:
        t = p.text.strip().upper()
        if t == "REFERENCES" and refs_heading is None:
            refs_heading = p._element
        elif t.startswith("APPENDIX") and appendix_heading is None:
            appendix_heading = p._element
    if refs_heading is None or appendix_heading is None:
        raise RuntimeError("References / APPENDIX bounds not found")

    body = refs_heading.getparent()
    children = list(body)
    i_refs = children.index(refs_heading)
    i_appendix = children.index(appendix_heading)
    ref_range = children[i_refs + 1 : i_appendix]

    ref_tpl = None
    for c in ref_range:
        if c.tag != QN_P:
            continue
        pPr = c.find(f"{{{W_NS}}}pPr")
        pStyle = pPr.find(f"{{{W_NS}}}pStyle") if pPr is not None else None
        sid = pStyle.get(f"{{{W_NS}}}val") if pStyle is not None else ""
        ts = c.findall(f".//{QN_T}")
        if sid == "943" and ts and any(t.text for t in ts):
            ref_tpl = c
            break
    if ref_tpl is None:
        raise RuntimeError("No style-943 reference template paragraph found")

    inserted = 0
    skipped = 0
    for anchor_prefix, ref_text in NEW_REFERENCES:
        first_words = ref_text.split("(")[0].strip().rstrip(",")
        already_present = any(
            c.tag == QN_P and "".join(t.text or "" for t in c.findall(f".//{QN_T}")).startswith(first_words)
            for c in ref_range
        )
        if already_present:
            print(f"  SKIP (already present): {ref_text[:80]}…")
            skipped += 1
            continue

        anchor = None
        for c in ref_range:
            if c.tag != QN_P:
                continue
            text = "".join(t.text or "" for t in c.findall(f".//{QN_T}"))
            if text.startswith(anchor_prefix):
                anchor = c
                break
        if anchor is None:
            raise RuntimeError(f"Alphabetical anchor not found: {anchor_prefix!r}")

        new_p = deepcopy(ref_tpl)
        set_paragraph_text(new_p, ref_text)
        anchor.addprevious(new_p)
        inserted += 1
        print(f"  INSERT before {anchor_prefix!r}: {ref_text[:80]}…")

        children = list(body)
        i_refs_new = children.index(refs_heading)
        i_appendix_new = children.index(appendix_heading)
        ref_range = children[i_refs_new + 1 : i_appendix_new]

    print(f"\nREFERENCES: inserted {inserted}, skipped {skipped} (already present)")


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    print("=== §3.8 body ===")
    write_section_body(doc)

    print("\n=== References ===")
    insert_references(doc)

    doc.save(str(path))
    print(f"\nSAVED: {path}")


if __name__ == "__main__":
    main()
