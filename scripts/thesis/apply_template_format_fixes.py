"""Template structure/format compliance fixes for THESIS.docx.

Reviewer round (2026-06-09): "follow the structure and format of the SURE NEU
thesis template carefully." We cross-referenced THESIS.docx against
`docs/research/SURE NEU THESIS FORMAT abduallah inal PDF 12.docx` chapter by
chapter and apply the safe, mechanical divergences here. Prose-level gaps
(missing Ch I subsections) are left for the author by decision — not touched.

Operations (each idempotent, anchored by text/style_id, never by index):

  1. FONT      — rewrite every run/style font "Times Roman" -> "Times New Roman"
                 (the malformed legacy name font-substitutes on render). Leaves
                 "Courier New" (code blocks) and "Times New Roman" untouched.
  2. NUMBERING — strip the leading decimal prefix (2.1 / 2.1.1 ...) from every
                 Heading-2 (945) / Heading-3 (882) section heading, to mirror the
                 template's bare descriptive titles. Chapter headings (944),
                 appendix letter-numbering (A.1/B.1 — start with a letter, so the
                 digit-anchored regex skips them) and figure/table caption numbers
                 (not heading styles) are preserved. Also trims stray leading
                 whitespace in headings.
  3. ALIGN     — centre the chapter (CHAPTER X) headings and the front-matter page
                 titles (template centres them); body section headings untouched.
  4. REORDER   — move the List of Figures block to AFTER the List of Tables block
                 (template puts List of Tables first). Both are hardcoded text.
  5. ABSTRACT  — insert the required title + degree line above the Abstract body,
                 and the Turkish title + author + degree line above the Özet body
                 (template front-matter requirement). Author must finalise the
                 bracketed placeholders ([NN] page count; Turkish title).
  6. XREF SCAN — report (never rewrite) in-text "Section X.Y" / "§X.Y" references
                 that go stale once section numbers are stripped, for the author.

Run:  uv run python scripts/thesis/apply_template_format_fixes.py --dry-run
      uv run python scripts/thesis/apply_template_format_fixes.py

After applying: open in OnlyOffice, Ctrl+A -> F9 to refresh the TOC live field
(picks up the de-numbered headings + new list order) and repaginate the
front matter. Per docs/research/THESIS_EDITING.md.
"""

from __future__ import annotations

import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DOCX = "docs/research/THESIS.docx"

# Filter headings by stable style NAME, not style_id: OnlyOffice/Word reassign
# numeric style_ids on every save (944 -> 964 -> … across sessions), so ids
# drift and any id-keyed filter silently no-ops on a re-saved document.
H1, H2, H3 = "Heading 1", "Heading 2", "Heading 3"
BODY = "Body Text"

# ---------------------------------------------------------------------------
# run-level text replacement (preserves run formatting) — from
# scripts/thesis/apply_reviewer6_fixes.py
# ---------------------------------------------------------------------------


def replace_in_paragraph(p, old: str, new: str) -> bool:
    runs = p.runs
    texts = [r.text for r in runs]
    full = "".join(texts)
    idx = full.find(old)
    if idx == -1:
        return False
    start, end = idx, idx + len(old)
    pos = 0
    spans = []
    for t in texts:
        spans.append((pos, pos + len(t)))
        pos += len(t)
    out = list(texts)
    first = True
    for k, (lo, hi) in enumerate(spans):
        if hi <= start or lo >= end:
            continue
        a = max(start, lo) - lo
        b = min(end, hi) - lo
        if first:
            out[k] = texts[k][:a] + new + texts[k][b:]
            first = False
        else:
            out[k] = texts[k][:a] + texts[k][b:]
    for r, t in zip(runs, out):
        if r.text != t:
            r.text = t
    return True


def sid(p) -> str | None:
    """Style NAME (stable across editor re-saves), despite the legacy alias."""
    return p.style.name if (p is not None and p.style) else None


# ---------------------------------------------------------------------------
# 1. font normalization
# ---------------------------------------------------------------------------

_FONT_ATTRS = (qn("w:ascii"), qn("w:hAnsi"), qn("w:cs"), qn("w:eastAsia"))


def _fix_fonts_in(root) -> int:
    n = 0
    for rf in root.iter(qn("w:rFonts")):
        for a in _FONT_ATTRS:
            if rf.get(a) == "Times Roman":
                rf.set(a, "Times New Roman")
                n += 1
    return n


def fix_fonts(doc) -> int:
    roots = [doc.element, doc.styles.element]
    for sec in doc.sections:
        for hf in (sec.header, sec.footer):
            try:
                roots.append(hf._element)
            except Exception:
                pass
    return sum(_fix_fonts_in(r) for r in roots)


# ---------------------------------------------------------------------------
# 2. strip leading section numbers / whitespace from headings
# ---------------------------------------------------------------------------

_NUM = re.compile(r"^(\s*\d+(?:\.\d+)*\.?\s+)")
_LEAD = re.compile(r"^(\s+)")


def _heading_prefix(text: str) -> str:
    m = _NUM.match(text)
    if m:
        return m.group(1)
    m = _LEAD.match(text)
    if m:
        return m.group(1)
    return ""


def strip_numbering(doc, dry: bool) -> int:
    n = 0
    for p in doc.paragraphs:
        if sid(p) not in (H2, H3):
            continue
        full = "".join(r.text for r in p.runs)
        prefix = _heading_prefix(full)
        if not prefix:
            continue
        new_text = full[len(prefix) :]
        print(f"    strip  {full!r} -> {new_text!r}")
        if not dry:
            replace_in_paragraph(p, prefix, "")
        n += 1
    return n


# ---------------------------------------------------------------------------
# 3. centre chapter + front-matter page headings
# ---------------------------------------------------------------------------

_FRONT_MATTER_TITLES = {
    "Approval",
    "Declaration of Ethical Principles",
    "Acknowledgments",
    "Abstract",
    "Özet",
    "Table of Contents",
    "List of Figures",
    "List of Tables",
    "List of Abbreviations",
}


def centre_headings(doc, dry: bool) -> int:
    n = 0
    for p in doc.paragraphs:
        s = sid(p)
        t = p.text.strip()
        is_chapter = s == H1 and t.startswith("CHAPTER ")
        is_fm = s == H2 and t in _FRONT_MATTER_TITLES
        if not (is_chapter or is_fm):
            continue
        if p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            continue
        print(f"    centre {t!r}")
        if not dry:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        n += 1
    return n


# ---------------------------------------------------------------------------
# 4. reorder: List of Figures block AFTER List of Tables block
# ---------------------------------------------------------------------------


def _heading_index(paras, text: str) -> int | None:
    for i, p in enumerate(paras):
        if sid(p) == H2 and p.text.strip() == text:
            return i
    return None


def reorder_lists(doc, dry: bool) -> int:
    paras = doc.paragraphs
    i_lof = _heading_index(paras, "List of Figures")
    i_lot = _heading_index(paras, "List of Tables")
    i_abbr = _heading_index(paras, "List of Abbreviations")
    if i_lof is None or i_lot is None or i_abbr is None:
        print("    SKIP reorder: could not locate all three list headings")
        return 0
    if i_lot < i_lof:
        print("    skip reorder: List of Tables already precedes List of Figures")
        return 0
    if not (i_lof < i_lot < i_abbr):
        print(f"    SKIP reorder: unexpected order (lof={i_lof} lot={i_lot} abbr={i_abbr})")
        return 0
    # LoF block = [i_lof .. i_lot-1]; move it to just before List of Abbreviations.
    block = [paras[j]._element for j in range(i_lof, i_lot)]
    ref = paras[i_abbr]._element
    parent = ref.getparent()
    if any(el.getparent() is not parent for el in block) or ref.getparent() is None:
        print("    SKIP reorder: block/anchor not sibling-aligned in body tree")
        return 0
    print(
        f"    reorder: move LoF block ({len(block)} paras) to after LoT, "
        f"before 'List of Abbreviations'"
    )
    if not dry:
        for el in block:  # in order -> preserves block order before ref
            ref.addprevious(el)
    return 1


# ---------------------------------------------------------------------------
# 5. abstract / özet title + degree-line blocks
# ---------------------------------------------------------------------------

TITLE_EN = (
    "Domain Specialized Marketing Agent for High Performance Campaign "
    "Generation: Fine-Tuning and RAG-Based Approach"
)
DEGREE_EN = (
    "M.Sc., Department of Artificial Intelligence Engineering, May, 2026, "
    "[NN] pages"
)
AUTHOR = "OSAMA DUWAIRI"
TITLE_TR_PLACEHOLDER = "[Türkçe tez başlığı — yazar tarafından eklenecek]"
DEGREE_TR = (
    "Yapay Zeka Mühendisliği Bölümü, Yüksek Lisans, Mayıs, 2026, [NN] sayfa"
)

_DEGREE_MARK = "Department of Artificial Intelligence Engineering, May"
_DEGREE_TR_MARK = "Yapay Zeka Mühendisliği Bölümü, Yüksek Lisans"


def _section_window(doc, heading_text: str):
    """Return (heading_idx, first_body_idx, window_texts) for a front-matter
    section. first_body is the first non-empty paragraph after the heading;
    window_texts is the text of the next ~8 paragraphs (for presence checks)."""
    paras = doc.paragraphs
    h = None
    for i, p in enumerate(paras):
        if sid(p) == H2 and p.text.strip() == heading_text:
            h = i
            break
    if h is None:
        return None, None, []
    first_body = None
    window = []
    for j in range(h + 1, min(h + 9, len(paras))):
        window.append(paras[j].text)
        if first_body is None and paras[j].text.strip():
            first_body = j
    return h, first_body, window


def _insert_before(anchor_para, text: str, style_obj, *, bold: bool):
    """Insert a centred paragraph before anchor, styled as body text (kept OUT
    of the TOC outline). Bold the run for the title line."""
    new = anchor_para.insert_paragraph_before(text, style=style_obj)
    new.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if bold:
        for r in new.runs:
            r.bold = True
    return new


def add_abstract_blocks(doc, dry: bool) -> int:
    n = 0
    # Body Text style object (943) — title/degree lines are centred body text,
    # NOT headings, so the live TOC field does not absorb them as entries.
    body_style = None
    for p in doc.paragraphs:
        if sid(p) == BODY and p.text.strip():
            body_style = p.style
            break
    if body_style is None:
        print("    SKIP abstract blocks: could not resolve body style")
        return 0

    # --- English Abstract: title + degree line ---
    _, body_i, window = _section_window(doc, "Abstract")
    if body_i is None:
        print("    SKIP abstract block: heading/body not found")
    elif any(_DEGREE_MARK in t for t in window):
        print("    skip abstract block: already present")
    else:
        print("    abstract: insert title + degree line (centred body text)")
        if not dry:
            body = doc.paragraphs[body_i]
            _insert_before(body, TITLE_EN, body_style, bold=True)
            _insert_before(body, DEGREE_EN, body_style, bold=False)
        n += 1

    # --- Özet: Turkish title + author + degree line ---
    _, body_i, window = _section_window(doc, "Özet")
    if body_i is None:
        print("    SKIP özet block: heading/body not found")
    elif any(_DEGREE_TR_MARK in t for t in window):
        print("    skip özet block: already present")
    else:
        print("    özet: insert Turkish title + author + degree line (centred body text)")
        if not dry:
            body = doc.paragraphs[body_i]
            _insert_before(body, TITLE_TR_PLACEHOLDER, body_style, bold=True)
            _insert_before(body, AUTHOR, body_style, bold=True)
            _insert_before(body, DEGREE_TR, body_style, bold=False)
        n += 1
    return n


# ---------------------------------------------------------------------------
# 6. in-text section cross-reference scan (report only)
# ---------------------------------------------------------------------------

_XREF = re.compile(r"(?:Sections?|§)\s*\d+(?:\.\d+)*", re.IGNORECASE)


def scan_section_xrefs(doc) -> list[tuple[int, str]]:
    hits = []
    for i, p in enumerate(doc.paragraphs):
        if sid(p) in (H1, H2, H3):
            continue
        for m in _XREF.finditer(p.text):
            ctx = p.text[max(0, m.start() - 25) : m.end() + 25]
            hits.append((i, ctx.strip()))
    return hits


# --- deterministic cross-reference resolver (number -> heading name) --------

# a whole reference cluster: "Section 2.1", "Sections 2.1, 2.2", "§3.2–§3.4",
# "§4.2 and §4.3", "§3.6–§3.7"
_XREF_CLUSTER = re.compile(
    r"(?P<word>Sections?|§)\s*"
    r"(?P<body>\d+(?:\.\d+)*"
    r"(?:\s*(?:,|–|-|and|&|to|\bthrough\b)\s*§?\s*\d+(?:\.\d+)*)*)",
    re.IGNORECASE,
)
_NUMTOK = re.compile(r"\d+(?:\.\d+)*")


def build_section_name_map(doc) -> dict[str, str]:
    """number -> heading name, from the *currently numbered* H2/H3 headings."""
    m = {}
    for p in doc.paragraphs:
        if sid(p) not in (H2, H3):
            continue
        full = "".join(r.text for r in p.runs)
        mm = _NUM.match(full)
        if not mm:
            continue
        num = mm.group(1).strip().rstrip(".").strip()
        name = full[len(mm.group(1)) :].strip()
        if num and name:
            m[num] = name
    return m


def _ref_phrase(tokens, names, body, sentence_initial: bool) -> str:
    art = "The" if sentence_initial else "the"
    if len(names) == 1:
        return f"{art} {names[0]} section"
    is_range = "–" in body or "-" in body or "to" in body.lower() or "through" in body.lower()
    if is_range:
        return f"{art} {names[0]} through {names[-1]} sections"
    if len(names) == 2:
        joined = f"{names[0]} and {names[1]}"
    else:
        joined = ", ".join(names[:-1]) + f", and {names[-1]}"
    return f"{art} {joined} sections"


def resolve_section_xrefs(doc, name_map, dry: bool):
    """Rewrite in-text 'Section X.Y'/'§X.Y' clusters to the referenced heading
    name(s). Returns (resolved, unresolved) lists of (para_idx, before, after)."""
    resolved, unresolved = [], []
    for i, p in enumerate(doc.paragraphs):
        if sid(p) in (H1, H2, H3):
            continue
        text = p.text
        if not _XREF_CLUSTER.search(text):
            continue
        edits = []  # (old, new) for this paragraph
        for m in _XREF_CLUSTER.finditer(text):
            tokens = _NUMTOK.findall(m.group("body"))
            names = [name_map.get(t) for t in tokens]
            before = m.group(0)
            if any(n is None for n in names):
                missing = [t for t, n in zip(tokens, names) if n is None]
                unresolved.append((i, before, f"unmapped: {missing}"))
                continue
            prev = text[max(0, m.start() - 2) : m.start()]
            sent_init = m.start() == 0 or prev.strip().endswith((".", "?", ":"))
            after = _ref_phrase(tokens, names, m.group("body"), sent_init)
            edits.append((before, after))
            resolved.append((i, before, after))
        if not dry:
            for old, new in edits:
                replace_in_paragraph(p, old, new)
    return resolved, unresolved


# ---------------------------------------------------------------------------


def main() -> int:
    dry = "--dry-run" in sys.argv
    doc = Document(DOCX)
    tag = "DRY-RUN " if dry else ""
    print(f"{tag}template-format fixes on {DOCX}\n" + "=" * 64)

    print("[1] font: Times Roman -> Times New Roman")
    n_font = fix_fonts(doc)
    print(f"    fixed {n_font} font references")

    strip_num = "--strip-numbers" in sys.argv
    name_map = build_section_name_map(doc)  # built BEFORE stripping
    print("[2] strip section numbering (Heading 2/3)")
    if not strip_num:
        n_num = 0
        print("    SKIPPED (default): heading numbers left intact (pass --strip-numbers)")
    else:
        n_num = strip_numbering(doc, dry)
        print(f"    stripped {n_num} headings")
        resolved, unresolved = resolve_section_xrefs(doc, name_map, dry)
        print(f"    resolved {len(resolved)} in-text refs to section names:")
        for i, before, after in resolved:
            print(f"      para {i}: {before!r} -> {after!r}")
        if unresolved:
            print(f"    {len(unresolved)} UNRESOLVED (no numbered heading to map — needs manual fix):")
            for i, before, why in unresolved:
                print(f"      para {i}: {before!r}  [{why}]")

    print("[3] centre chapter + front-matter headings")
    n_ctr = centre_headings(doc, dry)
    print(f"    centred {n_ctr} headings")

    print("[4] reorder List of Tables before List of Figures")
    n_ord = reorder_lists(doc, dry)

    print("[5] abstract / özet title + degree blocks")
    n_abs = add_abstract_blocks(doc, dry)

    hits = scan_section_xrefs(doc)
    if not strip_num:
        print("[6] scan in-text 'Section X.Y' references (report only)")
        print(f"    {len(hits)} in-text section-number reference(s) would dangle if de-numbered:")
        for i, ctx in hits:
            print(f"      para {i}: …{ctx}…")

    print("=" * 64)
    print(
        f"summary: font={n_font} numbering={n_num} centre={n_ctr} "
        f"reorder={n_ord} abstract_blocks={n_abs} xref_hits={len(hits)}"
    )
    if dry:
        print("DRY-RUN: not saving.")
        return 0
    doc.save(DOCX)
    print(f"saved {DOCX}")
    print(
        "\nFOLLOW-UP (author, in OnlyOffice):\n"
        "  - Ctrl+A -> F9 to refresh the TOC live field (de-numbered headings + new list order).\n"
        "  - Repaginate front matter so each list/abstract page starts cleanly.\n"
        "  - Fill the bracketed placeholders: Abstract/Özet '[NN]' page count; "
        "Özet '[Türkçe tez başlığı …]'.\n"
        "  - Review the in-text 'Section X.Y' references listed above (now stale "
        "after numbers were stripped)."
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
