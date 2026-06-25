"""Chapter IV audit pass — drift fixes + citation enrichment + bib correction.

Single idempotent script that does precise in-place text substitutions inside
Chapter IV paragraphs, inserts new bibliography entries alphabetically, and
corrects one drift in an existing References entry.

Modelled on ``scripts/thesis/audit_3_drift_and_citations.py``: same paragraph
substitution + reference-insertion machinery, retargeted from Ch III to Ch IV,
plus a small ``BIB_FIXES`` extension for in-bibliography substitutions.

Citation enrichment follows APA first-mention convention: each ref is cited
at its first body-text mention in Ch IV (§4.1.3 P472 for GPT-2 and bootstrap;
§4.2.5 P489 for ECE and AUC; §4.4.1 P503 for RAG; etc.), not at every
subsequent re-mention.

What it does:
 1. SUBSTITUTIONS — exact in-paragraph text replacements inside Ch IV:
    drift fixes (§4.2.3 per-platform n counts for Config C), one Phase-5
    polish (§4.4.2 "simply...simply" redundancy), and 10 parenthetical
    citation enrichments. Idempotent: an already-applied rule is skipped.
 2. BIB_FIXES — exact in-paragraph substitution inside the References
    section. Currently: restore co-author Thickstun to the Pillutla et al.
    2021 MAUVE entry. Idempotent in the same way.
 3. NEW_REFERENCES — six APA entries newly cited by §4. Lewis et al. 2020
    (RAG) and Zheng et al. 2023 (LLM-as-Judge) are already in the
    bibliography from prior passes and are NOT re-inserted; the script
    only adds the genuinely missing entries.

Run with:  uv run python scripts/thesis/audit_4_drift_and_citations.py
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"


# (paragraph_anchor_prefix, find_text, replace_text, label)
SUBSTITUTIONS: list[tuple[str, str, str, str]] = [
    # ----- §4.1.2 Configurations — Qwen3 citation -----
    (
        "The evaluation pipeline uses five main configurations",
        "the base 8B model (Qwen3-8B) without our fine-tuning.",
        "the base 8B model (Qwen3-8B; Qwen Team, 2025) without our fine-tuning.",
        "§4.1.2 ¶470: Qwen3-8B first-mention cite (Qwen Team, 2025)",
    ),
    # ----- §4.1.3 Evaluation Arms — GPT-2 + bootstrap first body mentions -----
    (
        "In our methodology, we have two evaluation arms",
        "embedding representation (GPT-2 Large), and focuses less",
        "embedding representation (GPT-2 Large; Radford et al., 2019), and focuses less",
        "§4.1.3 ¶472: GPT-2 Large first-mention cite (Radford et al., 2019)",
    ),
    (
        "In our methodology, we have two evaluation arms",
        "all reported numbers come with 95% confidence intervals (via bootstrap resampling).",
        "all reported numbers come with 95% confidence intervals (via bootstrap resampling; Efron, 1979).",
        "§4.1.3 ¶472: bootstrap first-mention cite (Efron, 1979)",
    ),
    # ----- §4.2.3 Per-Platform Scores — drift fixes -----
    # P481 reports Config C per-platform composites with the n counts of GOLD/B
    # (no extraction failures on those configs) rather than Config C's actual
    # attrited n. The composites and ordering are correct; only 3 of the 5 n
    # values need fixing (reddit=39 and tiktok=28 are unchanged because no
    # Config-C attrition on those platforms; facebook 93→86, pinterest 28→25,
    # twitter 27→26).
    (
        "Now focusing on our fine-tuned model's performance",
        "with n: facebook 93,",
        "with n: facebook 86,",
        "§4.2.3 ¶481: P481 n facebook 93 → 86 (Config C attrition)",
    ),
    (
        "Now focusing on our fine-tuned model's performance",
        "pinterest 28, reddit 39,",
        "pinterest 25, reddit 39,",
        "§4.2.3 ¶481: P481 n pinterest 28 → 25 (Config C attrition)",
    ),
    (
        "Now focusing on our fine-tuned model's performance",
        "tiktok 28, twitter 27.",
        "tiktok 28, twitter 26.",
        "§4.2.3 ¶481: P481 n twitter 27 → 26 (Config C attrition)",
    ),
    # ----- §4.2.5 Predictor Reliability — ECE + AUC citations -----
    (
        "Given these results, it is important to address the reliability",
        "Calibration error (ECE) = 0.0074",
        "Calibration error (ECE; Guo et al., 2017) = 0.0074",
        "§4.2.5 ¶489: ECE first-mention cite (Guo et al., 2017)",
    ),
    (
        "Given these results, it is important to address the reliability",
        "The model's top-tier AUC = 0.865",
        "The model's top-tier AUC (Fawcett, 2006) = 0.865",
        "§4.2.5 ¶489: AUC first-mention cite (Fawcett, 2006)",
    ),
    # ----- §4.3 Figure 4.3.1 caption — UMAP citation -----
    # Anchor must be a substring that survives the substitution. "UMAP
    # projection" is disrupted by inserting "(McInnes, Healy, & Melville,
    # 2018)" between the two words, so anchor on the stable suffix instead.
    (
        "GPT-2 Large embeddings per configuration",
        "Figure 4.3.1: UMAP projection of GPT-2 Large embeddings",
        "Figure 4.3.1: UMAP (McInnes, Healy, & Melville, 2018) projection of GPT-2 Large embeddings",
        "§4.3 ¶498 caption: UMAP first-mention cite (McInnes et al., 2018)",
    ),
    # ----- §4.4.1 2×2 Design — RAG first-mention citation -----
    (
        "This section reports the 2×2 ablation results for RQ2.",
        "Retrieval-Augmented Generation on the quality of output",
        "Retrieval-Augmented Generation (Lewis et al., 2020) on the quality of output",
        "§4.4.1 ¶503: RAG first-mention cite (Lewis et al., 2020)",
    ),
    # ----- §4.4.2 Paired Contrasts — Phase-5 polish + Cohen's d citation -----
    (
        "For paired contrast, we simply cannot",
        "we simply cannot just subtract cell scores, simply because",
        "we simply cannot just subtract cell scores, because",
        "§4.4.2 ¶505: remove redundant 'simply' (Phase-5 polish)",
    ),
    (
        "For paired contrast, we simply cannot",
        "at Cohen's dz ≈ 0.41",
        "at Cohen's dz (Cohen, 1988) ≈ 0.41",
        "§4.4.2 ¶505: Cohen's d first-mention cite (Cohen, 1988)",
    ),
    # ----- §4.5.3 Limitations — LLM-as-judge + self-preference-bias citations -----
    (
        "Having discussed the interpretation of the results",
        "such as the LLM-as-a-judge tournament.",
        "such as the LLM-as-a-judge tournament (Zheng et al., 2023).",
        "§4.5.3 ¶520: LLM-as-judge first-mention cite (Zheng et al., 2023)",
    ),
    (
        "Having discussed the interpretation of the results",
        "would offer a great bias toward generative AI content.",
        "would offer a great bias toward generative AI content (Panickssery, Bowman, & Feng, 2024).",
        "§4.5.3 ¶520: self-preference bias cite (Panickssery et al., 2024)",
    ),
]


# In-bibliography fixes — substring substitutions inside the References section.
# (anchor_prefix, find_text, replace_text, label)
BIB_FIXES: list[tuple[str, str, str, str]] = [
    (
        "Pillutla, K.,",
        "Zellers, R., Welleck, S.,",
        "Zellers, R., Thickstun, J., Welleck, S.,",
        "References: Pillutla et al. 2021 — restore missing co-author Thickstun",
    ),
]


# (alphabetical-anchor-prefix, full reference text)
# Same-anchor entries are inserted in list order; reverse-order placement
# (last inserted = closest to anchor) is accounted for so the final
# alphabetical sequence is correct.
NEW_REFERENCES: list[tuple[str, str]] = [
    # ===== C =====
    # Cobbe → Cohen → Dettmers 2022. Anchor on the first Dettmers entry.
    (
        "Dettmers, T., Lewis",
        "Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Lawrence Erlbaum Associates.",
    ),
    # ===== G =====
    # Grigsby → Guo → Guu. Anchor on Guu.
    (
        "Guu, K.",
        "Guo, C., Pleiss, G., Sun, Y., & Weinberger, K. Q. (2017). On calibration of modern neural networks. Proceedings of the 34th International Conference on Machine Learning, 70, 1321–1330. https://proceedings.mlr.press/v70/guo17a.html",
    ),
    # ===== M =====
    # Matz 2024 → McInnes → Mita. Anchor on Mita.
    (
        "Mita",
        "McInnes, L., Healy, J., & Melville, J. (2018). UMAP: Uniform manifold approximation and projection for dimension reduction (arXiv:1802.03426). https://arxiv.org/abs/1802.03426",
    ),
    # ===== P =====
    # Ouyang → Panickssery → Papineni. Anchor on Papineni.
    (
        "Papineni",
        "Panickssery, A., Bowman, S. R., & Feng, S. (2024). LLM evaluators recognize and favor their own generations. Advances in Neural Information Processing Systems, 37. https://arxiv.org/abs/2404.13076",
    ),
    # ===== Q / R =====
    # Qin → Qwen Team → Radford → Ratner. Both anchor on Ratner. Insert order
    # matters: Qwen Team first (lands further from anchor), Radford second
    # (lands immediately before Ratner). Final sequence: ... Qin → Qwen Team
    # → Radford → Ratner.
    (
        "Ratner, A. J.",
        "Qwen Team. (2025). Qwen3 technical report (arXiv:2505.09388). https://arxiv.org/abs/2505.09388",
    ),
    (
        "Ratner, A. J.",
        "Radford, A., Wu, J., Child, R., Luan, D., Amodei, D., & Sutskever, I. (2019). Language models are unsupervised multitask learners. OpenAI. https://cdn.openai.com/better-language-models/language_models_are_unsupervised_multitask_learners.pdf",
    ),
]


def paragraph_text(p_elem) -> str:
    return "".join(t.text or "" for t in p_elem.findall(f".//{QN_T}"))


def find_ch4_paragraph(doc, anchor: str):
    """Return the first paragraph inside CHAPTER IV whose text contains anchor.

    Bounds: starts after the "CHAPTER IV" heading, ends at the next chapter
    heading or at the REFERENCES heading (whichever comes first).
    """
    in_ch4 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        upper = t.upper()
        if upper == "CHAPTER IV":
            in_ch4 = True
            continue
        if upper in ("CHAPTER V", "REFERENCES"):
            break
        if in_ch4 and anchor in t:
            return p._element
    return None


def find_ref_paragraph(doc, anchor_prefix: str):
    """Return the first paragraph inside REFERENCES whose text starts with anchor_prefix."""
    in_refs = False
    for p in doc.paragraphs:
        t = p.text.strip()
        upper = t.upper()
        if upper == "REFERENCES":
            in_refs = True
            continue
        if upper.startswith("APPENDIX"):
            break
        if in_refs and t.startswith(anchor_prefix):
            return p._element
    return None


def substitute_in_paragraph(p_elem, find_str: str, replace_str: str) -> tuple[bool, bool]:
    """Replace first occurrence of find_str in the paragraph's concatenated text.

    Returns (changed, already_applied):
      - (True, False)  — substitution applied
      - (False, True)  — replace_str already in paragraph; idempotent skip
      - (False, False) — find_str not in paragraph; rule didn't match
    """
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        return (False, False)
    full = "".join(t.text or "" for t in ts)
    if replace_str in full:
        return (False, True)
    if find_str not in full:
        return (False, False)
    new = full.replace(find_str, replace_str, 1)
    ts[0].text = new
    for t in ts[1:]:
        t.text = ""
    return (True, False)


def apply_substitutions(doc) -> None:
    applied = 0
    skipped = 0
    missing = 0
    for anchor, find, replace, label in SUBSTITUTIONS:
        p = find_ch4_paragraph(doc, anchor)
        if p is None:
            print(f"  MISS (anchor not found): {label}")
            missing += 1
            continue
        changed, already = substitute_in_paragraph(p, find, replace)
        if changed:
            print(f"  APPLY: {label}")
            applied += 1
        elif already:
            print(f"  SKIP (idempotent): {label}")
            skipped += 1
        else:
            print(f"  MISS (find string not in paragraph): {label}")
            missing += 1
    print(f"\nSUBSTITUTIONS: applied {applied}, skipped {skipped}, missing {missing}")
    if missing:
        raise RuntimeError(
            f"{missing} Ch IV substitution(s) did not match — refusing to save. "
            "Likely the §4 prose has drifted from the anchor strings; "
            "verify the anchors against the current docx state."
        )


def apply_bib_fixes(doc) -> None:
    applied = 0
    skipped = 0
    missing = 0
    for anchor, find, replace, label in BIB_FIXES:
        p = find_ref_paragraph(doc, anchor)
        if p is None:
            print(f"  MISS (anchor not found): {label}")
            missing += 1
            continue
        changed, already = substitute_in_paragraph(p, find, replace)
        if changed:
            print(f"  APPLY: {label}")
            applied += 1
        elif already:
            print(f"  SKIP (idempotent): {label}")
            skipped += 1
        else:
            print(f"  MISS (find string not in paragraph): {label}")
            missing += 1
    print(f"\nBIB_FIXES: applied {applied}, skipped {skipped}, missing {missing}")
    if missing:
        raise RuntimeError(
            f"{missing} bibliography fix(es) did not match — refusing to save."
        )


def insert_references(doc) -> None:
    refs_heading = None
    appendix_heading = None
    for p in doc.paragraphs:
        t = p.text.strip().upper()
        if t == "REFERENCES" and refs_heading is None:
            refs_heading = p._element
        elif t.startswith("APPENDIX") and appendix_heading is None:
            appendix_heading = p._element
    if refs_heading is None or appendix_heading is None:
        raise RuntimeError("References / APPENDIX bounds not found")

    body = refs_heading.getparent()

    def current_ref_range():
        children = list(body)
        i_refs = children.index(refs_heading)
        i_app = children.index(appendix_heading)
        return children[i_refs + 1 : i_app]

    # Body Text template = first existing reference entry (style 943).
    ref_tpl = None
    for c in current_ref_range():
        if c.tag != QN_P:
            continue
        pPr = c.find(f"{{{W_NS}}}pPr")
        pStyle = pPr.find(f"{{{W_NS}}}pStyle") if pPr is not None else None
        sid = pStyle.get(f"{{{W_NS}}}val") if pStyle is not None else ""
        ts = c.findall(f".//{QN_T}")
        if sid == "943" and ts and any((t.text or "") for t in ts):
            ref_tpl = c
            break
    if ref_tpl is None:
        raise RuntimeError("No style-943 reference template paragraph found")

    inserted = 0
    skipped = 0
    for anchor_prefix, ref_text in NEW_REFERENCES:
        # Idempotency: skip if author+year prefix already present.
        first_words = ref_text.split("(")[0].strip().rstrip(",")
        already_present = any(
            c.tag == QN_P and paragraph_text(c).startswith(first_words)
            for c in current_ref_range()
        )
        if already_present:
            print(f"  SKIP (already present): {ref_text[:80]}…")
            skipped += 1
            continue

        anchor = None
        for c in current_ref_range():
            if c.tag != QN_P:
                continue
            if paragraph_text(c).startswith(anchor_prefix):
                anchor = c
                break
        if anchor is None:
            raise RuntimeError(f"Alphabetical anchor not found: {anchor_prefix!r}")

        new_p = deepcopy(ref_tpl)
        ts = new_p.findall(f".//{QN_T}")
        ts[0].text = ref_text
        for t in ts[1:]:
            t.text = ""
        anchor.addprevious(new_p)
        inserted += 1
        print(f"  INSERT before {anchor_prefix!r}: {ref_text[:80]}…")

    print(f"\nREFERENCES: inserted {inserted}, skipped {skipped} (already present)")


def main() -> None:
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    print("=== §4 Substitutions (drift fixes + polish + citation enrichment) ===")
    apply_substitutions(doc)

    print("\n=== References-section fixes ===")
    apply_bib_fixes(doc)

    print("\n=== Reference insertions ===")
    insert_references(doc)

    doc.save(str(path))
    print(f"\nSAVED: {path}")


if __name__ == "__main__":
    main()
