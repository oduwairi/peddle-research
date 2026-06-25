"""Write §3.3 Engagement-Based Scoring body + insert two references.

Same pattern as write_3_2_data_acquisition.py: wipes every paragraph between
the §3.3 heading and the §3.4 heading, then re-inserts the current PARAGRAPHS
list using a Body Text (style 943) template paragraph from CH3. Additionally
inserts two new bibliography entries at their alphabetical positions before
the APPENDIX heading (idempotent — skips if already present).
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
QN_P = f"{{{W_NS}}}p"
QN_T = f"{{{W_NS}}}t"


PARAGRAPHS: list[str] = [
    # ¶1 — Motivation + signal families
    "The main limitation of AdFlex ads is the fact that they don't contain explicit performance labels such as click-through rate (CTR) or return on ad spend (ROAS). This creates the need for a proxy scoring system that uses proxy signals such as engagement volume and engagement velocity as well as survivability as an indirect measurement of ad performance, which are readily available in the API response.",
    # ¶2 — Engagement volume + velocity
    "Engagement volume refers to the total bulk volume of the engagement a single ad has received, summed across likes, shares, comments, reactions, and views with quality weights (5:3:1:1:0.1 for shares, comments, likes, reactions, and views respectively), depending on which fields the platform exposes. Velocity points to the speed at which this engagement was achieved (engagement per day the ad has been running), providing a direct proxy for viral ads that catch attention fast.",
    # ¶3 — Survivability + KM rationale + prior-art anchor
    "As for lifespan, we use a Kaplan–Meier survival curve (Kaplan & Meier, 1958) — a non-parametric estimator that handles right-censored lifespan data, since ads still running at scrape time contribute a known minimum lifespan but no terminal event (a 7-day censoring window distinguishes live ads from observed deaths) — to estimate, for each platform (or globally where the per-platform cohort is below 20 ads), how likely an ad is to stay alive after X days based on those we've already seen die, normalized per platform since each platform possesses different standards for long- and short-running ads. Each ad's survivability score is the share of its platform peers it has outlasted at the observed duration. This is consistent with prior survival modelling of online-content lifespan (Lee, Moon, & Salamatian, 2012).",
    # ¶4 — Composite, tiering, downstream uses
    "These three signals are then combined into a single composite score via a weighted sum (survivability 0.50, engagement volume 0.25, engagement velocity 0.25; weights are renormalized over the available signals when one is missing) that offers an overall image of the performance of the ad. This scored dataset is then split into three tiers based on the percentile scores across the entire corpus: HIGH (top 20%), MEDIUM (middle 50%), and LOW (bottom 30%). The focus would be on the high-tier performing ads as they will be the primary training material for the instruction-backtranslation pipeline (§3.4). This full scored dataset is also used to train the learned scorer (§3.7) — a text-only model that learns to estimate this composite score from ad copy alone — which is used later on for evaluation using the corpus as ground truth (§3.8).",
    # ¶5 — Proxy validation methodology (addresses RQ3)
    "A developed scoring system uses proxies, not absolute signals. So the obvious research question at hand is how accurately these proxies represent actual ad performance and real metrics (RQ3). To validate this, we pick three main test sources in which the proxy scorer is validated against the datasets' ground-truth performance labels. Google Political Ads are the main stream, where ads are labeled with real spend data (in coarse buckets) which often correlates very closely with performance, but is not a direct performance signal. For this reason, a second more robust data stream is the IRA Facebook ads, which has real impressions and, more importantly, click labels that are a direct performance label. A third, less primary data stream is the Upworthy archive A/B dataset. However, this dataset lacks the inputs the scorer needs for meaningful scoring, which makes it a less viable option. Spearman rank correlation with bootstrap 95% confidence intervals is used to calculate how well our score aligns with ground-truth performance levels. The Kruskal–Wallis H-test is used to validate tier separation between HIGH, MEDIUM, and LOW in our scoring system.",
]


# (alphabetical-anchor-prefix, full reference text)
# Each entry is inserted BEFORE the first References-section paragraph whose
# text begins with the anchor prefix.
NEW_REFERENCES: list[tuple[str, str]] = [
    (
        "Karami",
        "Kaplan, E. L., & Meier, P. (1958). Nonparametric estimation from incomplete observations. Journal of the American Statistical Association, 53(282), 457–481.",
    ),
    (
        "Lewis, P.",
        "Lee, J. G., Moon, S., & Salamatian, K. (2012). Modeling and predicting the popularity of online contents with Cox proportional hazard regression model. Neurocomputing, 76(1), 134–145.",
    ),
]


def find_paragraph(doc, predicate, label):
    for p in doc.paragraphs:
        if predicate(p):
            return p._element
    raise RuntimeError(f"{label} not found")


def set_paragraph_text(p_elem, new_text):
    ts = p_elem.findall(f".//{QN_T}")
    if not ts:
        raise RuntimeError("template paragraph has no <w:t>")
    ts[0].text = new_text
    for t in ts[1:]:
        t.text = ""


def write_section_body(doc):
    h_3_3 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("3.3 Engagement-Based Scoring"),
        "§3.3 heading",
    )
    h_3_4 = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("3.4 Training-Data Construction"),
        "§3.4 heading",
    )

    # Body Text template: first style-943 paragraph inside CH3 (excluding
    # the §3.3 body itself, which we're about to wipe).
    body_tpl = None
    in_ch3 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.upper() == "CHAPTER III":
            in_ch3 = True
            continue
        if t.upper() == "CHAPTER IV":
            break
        # Skip §3.3's own body — we're about to delete it.
        if t.startswith("3.3 Engagement-Based Scoring"):
            in_ch3 = "in-33"
            continue
        if in_ch3 == "in-33" and t.startswith("3.4 "):
            in_ch3 = True
            continue
        if in_ch3 == "in-33":
            continue
        sid = p.style.style_id if p.style else ""
        if in_ch3 and sid == "943" and t and p._element.findall(f".//{QN_T}"):
            body_tpl = p._element
            break
    if body_tpl is None:
        raise RuntimeError("No style-943 Body Text paragraph found inside CH3 (outside §3.3)")

    body = h_3_3.getparent()
    children = list(body)
    i_start = children.index(h_3_3) + 1
    i_end = children.index(h_3_4)

    to_delete = [c for c in children[i_start:i_end] if c.tag == QN_P]
    print(f"DELETING: {len(to_delete)} existing §3.3 body paragraphs")
    for c in to_delete:
        body.remove(c)

    insert_after = h_3_3
    for i, text in enumerate(PARAGRAPHS, 1):
        p = deepcopy(body_tpl)
        set_paragraph_text(p, text)
        insert_after.addnext(p)
        insert_after = p
        print(f"  ¶{i}: {text[:80]}…")

    print(f"\nWROTE: {len(PARAGRAPHS)} paragraph(s) into §3.3 Engagement-Based Scoring")


def insert_references(doc):
    # Locate References + APPENDIX bounds.
    refs_heading = None
    appendix_heading = None
    for p in doc.paragraphs:
        t = p.text.strip().upper()
        if t == "REFERENCES" and refs_heading is None:
            refs_heading = p._element
        elif t.startswith("APPENDIX") and appendix_heading is None:
            appendix_heading = p._element
    if refs_heading is None or appendix_heading is None:
        raise RuntimeError("References / APPENDIX bounds not found")

    body = refs_heading.getparent()
    children = list(body)
    i_refs = children.index(refs_heading)
    i_appendix = children.index(appendix_heading)
    ref_range = children[i_refs + 1 : i_appendix]

    # Body Text template = first reference entry (style 943).
    ref_tpl = None
    for c in ref_range:
        if c.tag != QN_P:
            continue
        # Find any style-943 entry with real text.
        pPr = c.find(f"{{{W_NS}}}pPr")
        pStyle = pPr.find(f"{{{W_NS}}}pStyle") if pPr is not None else None
        sid = pStyle.get(f"{{{W_NS}}}val") if pStyle is not None else ""
        ts = c.findall(f".//{QN_T}")
        if sid == "943" and ts and any(t.text for t in ts):
            ref_tpl = c
            break
    if ref_tpl is None:
        raise RuntimeError("No style-943 reference template paragraph found")

    inserted = 0
    skipped = 0
    for anchor_prefix, ref_text in NEW_REFERENCES:
        # Idempotency: skip if already present (match on author + year prefix).
        first_words = ref_text.split("(")[0].strip().rstrip(",")  # e.g. "Kaplan, E. L., & Meier, P."
        already_present = any(
            c.tag == QN_P and "".join(t.text or "" for t in c.findall(f".//{QN_T}")).startswith(first_words)
            for c in ref_range
        )
        if already_present:
            print(f"  SKIP (already present): {ref_text[:80]}…")
            skipped += 1
            continue

        # Find anchor: first ref paragraph in [i_refs+1, i_appendix) starting with anchor_prefix.
        anchor = None
        for c in ref_range:
            if c.tag != QN_P:
                continue
            text = "".join(t.text or "" for t in c.findall(f".//{QN_T}"))
            if text.startswith(anchor_prefix):
                anchor = c
                break
        if anchor is None:
            raise RuntimeError(f"Alphabetical anchor not found: {anchor_prefix!r}")

        new_p = deepcopy(ref_tpl)
        set_paragraph_text(new_p, ref_text)
        anchor.addprevious(new_p)
        inserted += 1
        print(f"  INSERT before {anchor_prefix!r}: {ref_text[:80]}…")

        # Refresh ref_range because we just mutated the tree.
        children = list(body)
        i_refs_new = children.index(refs_heading)
        i_appendix_new = children.index(appendix_heading)
        ref_range = children[i_refs_new + 1 : i_appendix_new]

    print(f"\nREFERENCES: inserted {inserted}, skipped {skipped} (already present)")


def main():
    path = Path("docs/research/THESIS.docx")
    doc = Document(str(path))

    print("=== §3.3 body ===")
    write_section_body(doc)

    print("\n=== References ===")
    insert_references(doc)

    doc.save(str(path))
    print(f"\nSAVED: {path}")


if __name__ == "__main__":
    main()
