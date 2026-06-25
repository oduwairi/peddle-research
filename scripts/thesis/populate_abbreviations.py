"""Populate the List of Abbreviations section in THESIS.docx.

Replaces the placeholder paragraph at the end of the LOA section with a
verified alphabetical list of abbreviations that appear in the body text
(Chapter I through the end of Chapter IV; references are excluded).

Each entry is a Body Text paragraph: ACRONYM\\tFull Expansion.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


THESIS = Path("docs/research/THESIS.docx")

ABBREVIATIONS: list[tuple[str, str]] = [
    ("AdamW", "Adam with Decoupled Weight Decay"),
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("ARC", "AI2 Reasoning Challenge"),
    ("AUC", "Area Under the Curve"),
    ("BERT", "Bidirectional Encoder Representations from Transformers"),
    ("BLEU", "Bilingual Evaluation Understudy"),
    ("BLEURT", "Bilingual Evaluation Understudy with Representations from Transformers"),
    ("CI", "Confidence Interval"),
    ("COMET", "Crosslingual Optimized Metric for Evaluation of Translation"),
    ("CoT", "Chain-of-Thought"),
    ("CRAG", "Corrective Retrieval-Augmented Generation"),
    ("CTA", "Call to Action"),
    ("CTR", "Click-Through Rate"),
    ("DeBERTa", "Decoding-enhanced BERT with Disentangled Attention"),
    ("DoRA", "Weight-Decomposed Low-Rank Adaptation"),
    ("ECE", "Expected Calibration Error"),
    ("FT", "Fine-Tuning"),
    ("GDPR", "General Data Protection Regulation"),
    ("GPT", "Generative Pre-trained Transformer"),
    ("GPU", "Graphics Processing Unit"),
    ("IRA", "Internet Research Agency"),
    ("JSONL", "JSON Lines (newline-delimited JSON)"),
    ("LLaMA", "Large Language Model Meta AI"),
    ("LLM", "Large Language Model"),
    ("LoRA", "Low-Rank Adaptation"),
    ("MAUVE", "Measuring the Gap Between Neural and Human-Written Text"),
    ("MMLU", "Massive Multitask Language Understanding"),
    ("NLG", "Natural Language Generation"),
    ("NLP", "Natural Language Processing"),
    ("OOM", "Out of Memory"),
    ("PEFT", "Parameter-Efficient Fine-Tuning"),
    ("QA", "Question Answering"),
    ("QLoRA", "Quantized Low-Rank Adaptation"),
    ("RAFT", "Retrieval-Augmented Fine-Tuning"),
    ("RAG", "Retrieval-Augmented Generation"),
    ("RLHF", "Reinforcement Learning from Human Feedback"),
    ("ROAS", "Return on Ad Spend"),
    ("ROC", "Receiver Operating Characteristic"),
    ("ROUGE", "Recall-Oriented Understudy for Gisting Evaluation"),
    ("TF-IDF", "Term Frequency – Inverse Document Frequency"),
    ("UMAP", "Uniform Manifold Approximation and Projection"),
    ("VeRA", "Vector-based Random Matrix Adaptation"),
    ("VRAM", "Video Random-Access Memory"),
]

# Sort alphabetically, case-insensitive, by acronym.
ABBREVIATIONS.sort(key=lambda kv: kv[0].lower())


def find_placeholder_paragraph(doc) -> Paragraph:
    """Locate the '[List of Abbreviations — to be populated ...]' placeholder."""
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("[List of Abbreviations") and "to be populated" in text:
            return p
    raise RuntimeError("Placeholder paragraph not found in THESIS.docx")


def clear_paragraph(p: Paragraph) -> None:
    """Remove all runs from a paragraph, leaving it empty."""
    for r in list(p._p.findall(qn("w:r"))):
        p._p.remove(r)


def write_entry(p: Paragraph, acronym: str, expansion: str) -> None:
    """Write 'ACRONYM<TAB>Expansion' into an (already-cleared) paragraph."""
    run_acro = p.add_run(acronym)
    run_acro.bold = False
    p.add_run("\t")
    p.add_run(expansion)


def insert_body_text_paragraph_after(p: Paragraph, style_name: str) -> Paragraph:
    """Insert a new empty paragraph immediately after ``p`` with the given style."""
    new_p_el = OxmlElement("w:p")
    p._p.addnext(new_p_el)
    new_p = Paragraph(new_p_el, p._parent)
    new_p.style = p.part.document.styles[style_name]
    return new_p


def main() -> None:
    if not THESIS.exists():
        raise SystemExit(f"Thesis file not found: {THESIS}")

    doc = Document(str(THESIS))
    placeholder = find_placeholder_paragraph(doc)
    placeholder_style = "Body Text"  # match the rest of the front matter

    # Replace placeholder text with the first entry.
    placeholder.style = doc.styles[placeholder_style]
    clear_paragraph(placeholder)
    first_acro, first_exp = ABBREVIATIONS[0]
    write_entry(placeholder, first_acro, first_exp)

    # Insert the remaining entries after the (now-populated) placeholder, in order.
    anchor = placeholder
    for acro, exp in ABBREVIATIONS[1:]:
        new_para = insert_body_text_paragraph_after(anchor, placeholder_style)
        write_entry(new_para, acro, exp)
        anchor = new_para

    doc.save(str(THESIS))
    print(f"Wrote {len(ABBREVIATIONS)} abbreviation entries to {THESIS}")


if __name__ == "__main__":
    main()
