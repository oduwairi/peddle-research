"""Insert sec 3.8.6 "Reference-Overlap Metrics" (the BLEU/chrF/ROUGE-L/METEOR/
BERTScore arm) into Ch. III, after the MAUVE arm; renumber Gold Baseline 3.8.6
-> 3.8.7.

Phase-5/6 polished author prose (single body paragraph to match the sibling
sec 3.8.x subsections). Index-free anchors, idempotent. In-text citations are
included; their References entries + verification are the Phase-8 follow-up.
The TOC is a live field -- refresh fields in OnlyOffice after this edit.
"""

from __future__ import annotations

import docx

DOC = "docs/research/THESIS.docx"

HEADING = "3.8.6 Reference-Overlap Metrics"

BODY = (
    "This marketing-domain-specific evaluation framework has tried to avoid "
    "additional LLM evaluation methods and reference-overlap methods, including "
    "BLEU and chrF (shared words / characters), ROUGE-L (longest shared run of "
    "words), METEOR (allows word-form variation), and BERTScore (meaning in "
    "embedding space) (Papineni et al., 2002; Popović, 2015; Lin, 2004; "
    "Banerjee & Lavie, 2005; Zhang et al., 2020), since creative generation is "
    "an open-ended task with no clear one right answer. However, due to the "
    "presence of a golden, high-performing, ground-truth ad set, we can still "
    "apply these evaluation metrics in reference to the gold ads. This gives us "
    "a clear indication of how close each generated ad is to the real winning "
    "ad, measured on the wording or voice. This can be justified since, for an "
    "ad, the wording and style is exactly what drives performance; so although "
    "there isn't a single high-performing ad per brief, it can still be argued "
    "that being close to high-performing ads consistently indicates better "
    "performance. This evaluation arm employs two references per brief: the "
    "real winner, as well as a small pool of similar high-performing ads on the "
    "same platform. Taking a multi-ad reference helps mitigate the effect of "
    "being punished for generating a winning ad worded differently. "
    "Additionally, memorization flags are in place to flag winners that copy "
    "winning ads too literally instead of copying the style. The metrics being "
    "used are validated against real A/B outcomes from the Upworthy Research "
    "Archive (Matias et al., 2021), asking the question: do these metrics "
    "really classify winning ads? This evaluation arm is added as more of a "
    "supplementary or informational stage, to show the limitations of using "
    "traditional metrics for the creative marketing domain."
)


def main() -> None:
    d = docx.Document(DOC)

    def by_eq(s, style=None):
        for p in d.paragraphs:
            if p.text.strip() == s and (style is None or p.style.name == style):
                return p
        return None

    if by_eq(HEADING) is not None:
        print("already inserted; nothing to do")
        return

    gold = by_eq("3.8.6 Gold Baseline", style="Heading 3")
    if gold is None:
        raise SystemExit("'3.8.6 Gold Baseline' heading not found")

    h3_style = gold.style
    body_style = next(p for p in d.paragraphs if p.style.name == "Body Text").style

    h = d.add_paragraph()
    h.add_run(HEADING)
    h.style = h3_style

    b = d.add_paragraph()
    b.add_run(BODY)
    b.style = body_style

    gold._p.addprevious(h._p)
    gold._p.addprevious(b._p)

    # renumber Gold Baseline 3.8.6 -> 3.8.7
    for r in gold.runs:
        if "3.8.6" in r.text:
            r.text = r.text.replace("3.8.6", "3.8.7")
            break
    else:
        full = "".join(r.text for r in gold.runs)
        gold.runs[0].text = full.replace("3.8.6", "3.8.7")
        for r in gold.runs[1:]:
            r.text = ""

    d.save(DOC)
    print("INSERTED: §3.8.6 Reference-Overlap Metrics (1 heading + 1 body para)")
    print("RENUMBERED: Gold Baseline 3.8.6 -> 3.8.7")
    print("NOTE: refresh fields in OnlyOffice (live TOC) to sync the contents page")


if __name__ == "__main__":
    main()
