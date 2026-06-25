"""Convert Figures 2.2 and 2.7 from floating (anchored) to inline images,
matching the other 13 figures (inline, centered, Body Text style 943).

Fig 2.2: single floating image (image2.png); caption already a normal para.
Fig 2.7: floating image (image3.png) + floating text-box caption; we make the
image inline AND rebuild the caption as a normal Body Text paragraph.

Reuses the already-embedded image bytes (exact current pixels) and each image's
current extent (exact current size). Idempotent-ish: re-running re-embeds.
"""
from __future__ import annotations
import os, shutil, sys
from copy import deepcopy
from docx import Document
from docx.shared import Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

THESIS = "docs/research/THESIS.docx"
LOCK = "docs/research/.~lock.THESIS.docx#"
MEDIA = "/tmp/thx/word/media"  # exact embedded bytes, extracted earlier

if os.path.exists(LOCK):
    sys.exit("ABORT: OnlyOffice lock present — close the editor first.")
for src in (f"{MEDIA}/image2.png", f"{MEDIA}/image3.png"):
    if not os.path.exists(src):
        sys.exit(f"ABORT: embedded image source missing: {src}")

shutil.copy(THESIS, THESIS + ".bak")
doc = Document(THESIS)

def resolve(blip):
    rid = blip.get(qn("r:embed"))
    try:
        return doc.part.rels[rid].target_ref
    except KeyError:
        return ""

def find_img_para(name):
    for p in doc.paragraphs:
        for blip in p._element.iter(qn("a:blip")):
            if name in resolve(blip):
                return p
    return None

def extent_cx(p):
    ext = p._element.find(".//" + qn("wp:extent"))
    return int(ext.get("cx")) if ext is not None else None

def clear_body(p):
    for ch in list(p._element):
        if ch.tag != qn("w:pPr"):
            p._element.remove(ch)

def set_style_943(p):
    pPr = p._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = p._element.makeelement(qn("w:pPr"), {})
        p._element.insert(0, pPr)
    ps = pPr.find(qn("w:pStyle"))
    if ps is None:
        ps = pPr.makeelement(qn("w:pStyle"), {})
        pPr.insert(0, ps)
    ps.set(qn("w:val"), "943")

p22 = find_img_para("image2.png")
p27 = find_img_para("image3.png")
assert p22 is not None and p27 is not None, "could not locate figure paragraphs"
cx22, cx27 = extent_cx(p22), extent_cx(p27)
print(f"Fig2.2 para found, current width EMU={cx22} ({cx22/914400:.2f} in)")
print(f"Fig2.7 para found, current width EMU={cx27} ({cx27/914400:.2f} in)")

# --- Fig 2.2: image -> inline, centered, style 943 (caption already at next paras) ---
clear_body(p22); set_style_943(p22)
p22.add_run().add_picture(f"{MEDIA}/image2.png", width=Emu(cx22))
p22.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- Fig 2.7: image -> inline, centered, style 943; rebuild caption as a normal para ---
cap_text = ("Figure 2.7 — Overview of the Self-Instruct pipeline (Wang et al., 2023). "
            "Where a model starts from a small pool of human written samples then "
            "iteratively generates more samples to be used in instruction tuning "
            "tasks in low data applications.")
clear_body(p27); set_style_943(p27)
p27.add_run().add_picture(f"{MEDIA}/image3.png", width=Emu(cx27))
p27.alignment = WD_ALIGN_PARAGRAPH.CENTER

# clone an existing Body-Text caption paragraph (Fig 2.2's) as the template
templ = None
for p in doc.paragraphs:
    if p.text.strip().startswith("Figure 2.2 (Gao"):
        templ = p._element
        break
assert templ is not None, "caption template (Fig 2.2) not found"
newcap = deepcopy(templ)
# strip any drawing just in case, then set text
for dr in newcap.findall(".//" + qn("w:drawing")):
    dr.getparent().remove(dr)
ts = newcap.findall(".//" + qn("w:t"))
assert ts, "template caption has no <w:t>"
ts[0].text = cap_text
for t in ts[1:]:
    t.text = ""
# ensure caption is left-aligned/justified like others (clear any center jc)
capPr = newcap.find(qn("w:pPr"))
if capPr is not None:
    jc = capPr.find(qn("w:jc"))
    if jc is not None:
        capPr.remove(jc)
p27._element.addnext(newcap)

doc.save(THESIS)
print("SAVED.")
