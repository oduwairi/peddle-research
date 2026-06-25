"""Convert THESIS.docx Chapter II into the PeerJ literature-review template.

Pass 1: structural shell only.
- Title / author block: filled from thesis front matter.
- Abstract: kept as placeholder (mirrors thesis).
- Introduction: placeholder.
- Survey methodology: thesis para 268 (already drafted in Chapter I).
- Body: §2.1..§2.9 subheadings + body paragraphs.
- Conclusions: thesis Chapter II §Conclusion (paras 661..664).
- Acknowledgements: thesis acknowledgements text.
- References: placeholder.

Inline figures, tables, and bibliography migration are explicitly deferred to a
later pass.

Run:
    uv run python scripts/explore/convert_litreview_to_peerj.py
"""

from __future__ import annotations

import copy
import re
import shutil
from pathlib import Path

import docx
from docx.document import Document as _DocxDocument
from docx.oxml.ns import qn
from docx.shared import Emu, Pt

REPO = Path(__file__).resolve().parents[2]
THESIS_PATH = REPO / "docs" / "research" / "THESIS.docx"
TEMPLATE_PATH = REPO / "docs" / "research" / "PeerJ-literature-review-template.docx"
OUTPUT_PATH = REPO / "docs" / "research" / "literature-review-peerj.docx"

# --- Thesis content anchors ---------------------------------------------------
TITLE = (
    "DOMAIN SPECIALIZED MARKETING AGENT FOR HIGH PERFORMANCE CAMPAIGN GENERATION: "
    "FINE-TUNING AND RAG-BASED APPROACH"
)
AUTHOR_NAME = "Osama Duwairi"
AFFILIATION = (
    "Department of Artificial Intelligence Engineering, "
    "Near East University Institute of Graduate Studies, Nicosia, Cyprus"
)
CORRESPONDING_NAME = "Osama Duwairi"
CORRESPONDING_ADDRESS = "Near East University, Near East Boulevard, ZIP 99138, Nicosia, North Cyprus"
CORRESPONDING_EMAIL = "oduwairi@gmail.com"

ABSTRACT_TEXT = (
    "This literature survey aims to explore existing research and works on the topic "
    "of domain specialized AI agents in marketing. The survey starts by examining "
    "open-source foundation models and parameter-efficient fine-tuning techniques such "
    "as QLoRA and other LoRA variants, results show that these models although in "
    "7B-9B range can show competitive performance when compared to large proprietary "
    "models in domain-specific tasks. Next we examine retrieval-augmented generation "
    "architectures including naive to advanced and efficient RAG showing that these "
    "techniques help reduce hallucinations and provide more factual responses in "
    "fast-changing domains. The survey continues by exploring the applications of LLMs "
    "in marketing and advertising and the challenge of low-data domains for constructing "
    "fine-tuned models, mitigation techniques such as LLM data synthesis, and data "
    "labeling based on proxy signals. Agentic LLM architectures such as single and "
    "multi-agent systems are briefly discussed. The survey concludes by exploring "
    "evaluation methodologies for open-ended domains, and ethical/legal considerations "
    "in AI-generated marketing. The survey reveals key research gaps: while domains "
    "such as medicine have produced numerous fine-tuned LLMs (Med42, AlpaCare, "
    "Med-PaLM, among others), no existing work fine-tunes an open-source sub-10B model "
    "for marketing campaign generation (Table 2.3). Additionally, while several public "
    "marketing datasets exist, none combine English-language ad creative text with "
    "real-world performance labels suitable for fine-tuning content generation models "
    "(Table 2.4). Other gaps include the absence of established evaluation benchmarks "
    "for creative marketing generation as well as gaps in ethical and legal aspects in "
    "AI-driven advertising systems."
)
KEYWORDS_TEXT = (
    "Keywords: agentic AI, marketing campaign generation, parameter-efficient "
    "fine-tuning, retrieval-augmented generation, QLoRA, domain-specialized language "
    "models, LLM evaluation."
)
# Introduction is assembled from thesis Ch I §Background..§Paper Organization
# (paras 248..265). Survey Methodology (para 266) is excluded — it has its own slot.
INTRO_SUBSECTIONS = [
    (248, 251, "Background"),
    (251, 254, "Problem Statement"),
    (254, 260, "Research Questions"),
    (260, 263, "Contributions"),
    (263, 266, "Paper Organization"),
]

# PeerJ requires figures/tables uploaded separately, numbered in order of first
# in-text appearance. Thesis uses §-aligned numbering (Table 2.X / Figure 2.Y);
# normalize each to a single sequence.
FIGURE_RENUMBER = {
    "2.2": "1",  # §2.2.1 RAG architectures (Gao et al., 2024)
    "2.7": "2",  # §2.7.2 Self-Instruct pipeline (Wang et al., 2023)
}
TABLE_RENUMBER = {
    "2.1": "1",  # §2.1.1 Open-Source Foundation Models
    "2.2": "2",  # §2.1.2 PEFT Methods
    "2.3": "3",  # §2.1.3 Domain-Specialized Small Models
    "2.4": "4",  # §2.4.1 Marketing Datasets
    "2.5": "5",  # §2.6.4 Evaluation Frameworks
}

_FIG_REF_RE = re.compile(r"\b(Figure|figure)\s+(2\.\d)\b")
_TBL_REF_RE = re.compile(r"\b(Table|table)\s+(2\.\d)\b")


def renumber_refs(text: str) -> str:
    """Rewrite in-text Figure/Table references to PeerJ sequential numbering."""

    def fix_fig(m: re.Match[str]) -> str:
        old = m.group(2)
        new = FIGURE_RENUMBER.get(old)
        return f"Figure {new}" if new else m.group(0)

    def fix_tbl(m: re.Match[str]) -> str:
        old = m.group(2)
        new = TABLE_RENUMBER.get(old)
        return f"Table {new}" if new else m.group(0)

    text = _FIG_REF_RE.sub(fix_fig, text)
    text = _TBL_REF_RE.sub(fix_tbl, text)
    return text
ACKNOWLEDGEMENTS_TEXT = "The author has not declared any acknowledgements."
# References are migrated whole from the thesis bibliography (paras 944..1032).
THESIS_REFERENCES_START = 944  # paragraph immediately after "References" heading
THESIS_REFERENCES_END = 1033   # paragraph index of "APPENDIX" — exclusive

# --- Heading sizes mirroring the PeerJ template ------------------------------
MAJOR_HEADING_PT = Pt(14)
SUB_HEADING_PT = Pt(12)


def extract_thesis_chapter2() -> tuple[list[tuple[str, str]], list[str]]:
    """Return (body_blocks, conclusion_paragraphs).

    body_blocks: list of (kind, text) covering paras 273..660.
      kind ∈ {"h2_major", "h2_sub", "body"}.
    conclusion_paragraphs: plain paragraph texts from paras 662..664.

    Table-internal paragraphs (cell contents) are skipped — they keep the
    overall regex `<w:p>` numbering aligned with the heading map but are NOT
    emitted as body prose, since they belong to tables we strip in this pass.
    """
    import io
    import zipfile

    with open(THESIS_PATH, "rb") as fh:
        raw = fh.read()
    with zipfile.ZipFile(io.BytesIO(raw)) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")

    paras_with_pos = [(m.start(), m.group(0)[len("<w:p") :]) for m in re.finditer(r"<w:p\b[^>]*>.*?</w:p>", xml, flags=re.S)]
    # Re-find full match — the slice above was just for clarity. Use a single pass:
    para_matches = list(re.finditer(r"<w:p\b[^>]*>(.*?)</w:p>", xml, flags=re.S))
    paras = [m.group(1) for m in para_matches]
    para_positions = [m.start() for m in para_matches]

    # Compute table extents [start, end) so we can flag table-internal paragraphs.
    table_extents: list[tuple[int, int]] = []
    pos = 0
    while True:
        open_m = re.search(r"<w:tbl\b[^>]*>", xml[pos:])
        if not open_m:
            break
        open_pos = pos + open_m.start()
        close = xml.find("</w:tbl>", open_pos)
        if close == -1:
            break
        end = close + len("</w:tbl>")
        table_extents.append((open_pos, end))
        pos = end

    def in_table(p_pos: int) -> bool:
        for s, e in table_extents:
            if s <= p_pos < e:
                return True
            if p_pos < s:
                return False
        return False

    text_re = re.compile(r"<w:t[^>]*>([^<]*)</w:t>", re.S)
    heading_re = re.compile(r'<w:pStyle\s+w:val="(944|945)"', re.S)

    def para_text(p: str) -> str:
        import html

        return html.unescape("".join(text_re.findall(p))).strip()

    def classify(idx: int) -> tuple[str, str]:
        p = paras[idx]
        text = para_text(p)
        sm = heading_re.search(p)
        if sm:
            m = re.match(r"^(\d+)\.(\d+)(?:\.(\d+))?\s", text)
            if m:
                return ("h2_sub" if m.group(3) else "h2_major", text)
            return ("other_heading", text)
        return ("body", text)

    body_blocks: list[tuple[str, str]] = []
    last_body_text: str | None = None
    for i in range(273, 661):
        if in_table(para_positions[i]):
            continue
        kind, text = classify(i)
        if not text:
            continue
        if kind == "other_heading":
            continue
        text = renumber_refs(text)
        # Drop accidental consecutive duplicates (e.g. Figure 2.7 caption appears
        # twice in the thesis source — once as the figure-anchor para's text and
        # once in the following paragraph).
        if kind == "body" and text == last_body_text:
            continue
        body_blocks.append((kind, text))
        last_body_text = text if kind == "body" else None

    conclusion: list[str] = []
    for i in range(662, 665):
        if in_table(para_positions[i]):
            continue
        text = para_text(paras[i])
        if text:
            conclusion.append(renumber_refs(text))

    return body_blocks, conclusion


def extract_thesis_references() -> list[str]:
    """Return thesis bibliography entries as a list of unescaped paragraph texts.

    Each non-empty paragraph in [THESIS_REFERENCES_START, THESIS_REFERENCES_END)
    is treated as one reference. The leading "References" heading and the
    trailing "APPENDIX" header are excluded by construction.
    """
    import html
    import zipfile

    with zipfile.ZipFile(THESIS_PATH) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    paras = re.findall(r"<w:p\b[^>]*>(.*?)</w:p>", xml, flags=re.S)
    text_re = re.compile(r"<w:t[^>]*>([^<]*)</w:t>", re.S)

    refs: list[str] = []
    for i in range(THESIS_REFERENCES_START, THESIS_REFERENCES_END):
        text = "".join(text_re.findall(paras[i])).strip()
        if not text or text == "References":
            continue
        refs.append(html.unescape(text))
    return refs


def extract_thesis_introduction() -> list[tuple[str, str]]:
    """Return Ch I §Background..§Paper Organization as (kind, text) blocks.

    kind ∈ {"h_sub", "body"}. Subsection headings are emitted as h_sub.
    """
    import zipfile

    with zipfile.ZipFile(THESIS_PATH) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    paras = re.findall(r"<w:p\b[^>]*>(.*?)</w:p>", xml, flags=re.S)
    text_re = re.compile(r"<w:t[^>]*>([^<]*)</w:t>", re.S)

    def para_text(p: str) -> str:
        import html

        return html.unescape("".join(text_re.findall(p))).strip()

    blocks: list[tuple[str, str]] = []
    for start, end, label in INTRO_SUBSECTIONS:
        blocks.append(("h_sub", label))
        for i in range(start + 1, end):
            text = para_text(paras[i])
            if text:
                blocks.append(("body", renumber_refs(text)))
    return blocks


# --- docx helpers -------------------------------------------------------------

def replace_paragraph_text(p, text: str, *, bold: bool = False, size: Pt | None = None) -> None:
    """Replace all runs in *p* with a single run carrying *text*.

    Preserves the paragraph's pPr (style/numbering) but resets runs.
    """
    pPr = p._p.find(qn("w:pPr"))
    # remove all children except pPr
    for child in list(p._p):
        if child.tag != qn("w:pPr"):
            p._p.remove(child)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size is not None:
        run.font.size = size


def insert_paragraph_after(prev_p, text: str, *, bold: bool = False, size: Pt | None = None):
    """Insert a new paragraph right after *prev_p* and return it."""
    new_p = copy.deepcopy(prev_p._p)
    # Strip existing runs from the clone (keep pPr only)
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    prev_p._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    para = Paragraph(new_p, prev_p._parent)
    run = para.add_run(text)
    if bold:
        run.bold = True
    if size is not None:
        run.font.size = size
    return para


# --- Main conversion ----------------------------------------------------------

def convert() -> None:
    print(f"[1/4] Copying template -> {OUTPUT_PATH.name}")
    shutil.copyfile(TEMPLATE_PATH, OUTPUT_PATH)

    print("[2/4] Extracting thesis content")
    intro_blocks = extract_thesis_introduction()
    body_blocks, conclusion = extract_thesis_chapter2()
    references = extract_thesis_references()
    print(f"      Introduction: {len(intro_blocks)} blocks")
    print(f"      Body (Ch II): {len(body_blocks)} blocks (headings + paragraphs)")
    print(f"      Conclusion (Ch II): {len(conclusion)} paragraphs")
    print(f"      References:   {len(references)} entries")

    print("[3/4] Loading template and locating slots")
    doc: _DocxDocument = docx.Document(str(OUTPUT_PATH))

    # Index template slot paragraphs by current text-prefix.
    slots: dict[str, int] = {}
    for idx, p in enumerate(doc.paragraphs):
        t = p.text.strip().lower()
        if t == "manuscript title":
            slots["title"] = idx
        elif t.startswith("firstname middlename lastname1, firstname"):
            slots["authors"] = idx
        elif t.startswith("1 department name"):
            slots["aff1"] = idx
        elif t.startswith("2 department name"):
            slots["aff2"] = idx
        elif t.startswith("3 department name"):
            slots["aff3"] = idx
        elif t.startswith("4 department name"):
            slots["aff4"] = idx
        elif t.startswith("firstname lastname1"):
            slots["corresp_name"] = idx
        elif t.startswith("street address"):
            slots["corresp_addr"] = idx
        elif t.startswith("email address"):
            slots["corresp_email"] = idx
        elif t.startswith("add your abstract"):
            slots["abstract"] = idx
        elif t.startswith("add your introduction"):
            slots["intro"] = idx
        elif t.startswith("add your survey"):
            slots["survey"] = idx
        elif t.startswith("add any major subheadings"):
            slots["body"] = idx
        elif t.startswith("add your conclusions"):
            slots["conclusions"] = idx
        elif t.startswith("add your acknowledgements"):
            slots["ack"] = idx
        elif t.startswith("add your references"):
            slots["refs"] = idx

    missing = {
        "title",
        "authors",
        "aff1",
        "corresp_name",
        "corresp_addr",
        "corresp_email",
        "abstract",
        "intro",
        "survey",
        "body",
        "conclusions",
        "ack",
        "refs",
    } - set(slots)
    if missing:
        raise RuntimeError(f"Could not locate template slots: {sorted(missing)}")

    print(f"      Located slots: {sorted(slots)}")

    paragraphs = doc.paragraphs

    # Title
    replace_paragraph_text(paragraphs[slots["title"]], TITLE, bold=True, size=Emu(228600))

    # Authors line: one author only
    replace_paragraph_text(paragraphs[slots["authors"]], f"{AUTHOR_NAME}¹")

    # Affiliations: keep aff1, clear aff2..aff4
    replace_paragraph_text(paragraphs[slots["aff1"]], f"¹ {AFFILIATION}")
    for key in ("aff2", "aff3", "aff4"):
        if key in slots:
            replace_paragraph_text(paragraphs[slots[key]], "")

    # Corresponding author
    replace_paragraph_text(paragraphs[slots["corresp_name"]], CORRESPONDING_NAME)
    replace_paragraph_text(paragraphs[slots["corresp_addr"]], CORRESPONDING_ADDRESS)
    replace_paragraph_text(paragraphs[slots["corresp_email"]], f"Email address: {CORRESPONDING_EMAIL}")

    # Abstract + keywords line directly after
    replace_paragraph_text(paragraphs[slots["abstract"]], renumber_refs(ABSTRACT_TEXT))
    insert_paragraph_after(paragraphs[slots["abstract"]], KEYWORDS_TEXT)

    # Introduction: inject §Background..§Paper Organization after the slot.
    intro_slot = paragraphs[slots["intro"]]
    replace_paragraph_text(intro_slot, "")
    anchor = intro_slot
    for kind, text in intro_blocks:
        if kind == "h_sub":
            anchor = insert_paragraph_after(anchor, text, bold=True, size=SUB_HEADING_PT)
        else:
            anchor = insert_paragraph_after(anchor, text)

    # Survey methodology
    survey_text = _load_survey_methodology()
    replace_paragraph_text(paragraphs[slots["survey"]], survey_text)

    # Acknowledgements
    replace_paragraph_text(paragraphs[slots["ack"]], ACKNOWLEDGEMENTS_TEXT)

    # References: inject all thesis bibliography entries after the slot.
    refs_slot = paragraphs[slots["refs"]]
    if references:
        replace_paragraph_text(refs_slot, references[0])
        anchor = refs_slot
        for entry in references[1:]:
            anchor = insert_paragraph_after(anchor, entry)
    else:
        replace_paragraph_text(refs_slot, "[References — placeholder.]")

    # Body: inject body_blocks after the body slot, then clear the slot itself.
    body_slot = paragraphs[slots["body"]]
    replace_paragraph_text(body_slot, "")  # blank the placeholder
    anchor = body_slot
    for kind, text in body_blocks:
        if kind == "h2_major":
            anchor = insert_paragraph_after(anchor, text, bold=True, size=MAJOR_HEADING_PT)
        elif kind == "h2_sub":
            anchor = insert_paragraph_after(anchor, text, bold=True, size=SUB_HEADING_PT)
        else:  # body
            anchor = insert_paragraph_after(anchor, text)

    # Conclusions: inject after conclusions slot
    concl_slot = paragraphs[slots["conclusions"]]
    if conclusion:
        replace_paragraph_text(concl_slot, conclusion[0])
        anchor = concl_slot
        for text in conclusion[1:]:
            anchor = insert_paragraph_after(anchor, text)
    else:
        replace_paragraph_text(concl_slot, "[Conclusion — placeholder.]")

    print("[4/4] Saving")
    doc.save(str(OUTPUT_PATH))
    print(f"\nWrote: {OUTPUT_PATH.relative_to(REPO)}")


def _load_survey_methodology() -> str:
    import html
    import zipfile

    with zipfile.ZipFile(THESIS_PATH) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    paras = re.findall(r"<w:p\b[^>]*>(.*?)</w:p>", xml, flags=re.S)
    text_re = re.compile(r"<w:t[^>]*>([^<]*)</w:t>", re.S)
    return html.unescape("".join(text_re.findall(paras[268]))).strip()


if __name__ == "__main__":
    convert()
